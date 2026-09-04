"""The optional Cover Letter — the letter the customer portal shows AHEAD of the proposal.

It is a portal DOCUMENT page, not an email (Hanz, 2026-08-28): one letterhead template per
(work type, audience), filled from the same intake/estimate values the proposal uses, riding
inside the same `proposal_payload` so a sent revision pins the letter exactly as it pins the
prices.

Everything here EXECUTES the real writer and the real templates. A source-text assertion cannot
catch a token nothing fills, a numbering definition Word would reject, or a template that quietly
lost its letterhead — and those are the three ways a first-draft document set goes wrong.

Covers:
  (a) the templates themselves — one per (work type, audience), real letterhead, the floating
      date box, no raw token left behind;
  (b) GenerateIn/GenerateOut — off by default, and OFF means no download url at all;
  (c) /api/coverletter-template — the block model the document editor renders;
  (d) /api/admin/cover-letter-pdf — the portal's server-to-server render, its SERVICE_TOKEN gate,
      its revision pinning, and its named refusal when there is no letter;
  (e) PortalPublishIn.has_cover_letter — the omitted-means-nothing-forwarded contract.
"""
import io
import json
import os
import pathlib
import shutil
import subprocess
import zipfile

import docx
import pytest
from fastapi.testclient import TestClient

import cover_letter_writer as clw
import main
import prepare_cover_letter_templates as prep
import proposal_writer as pw

client = TestClient(main.app)

# The (work_type, audience) pairs that have a letter of their own.
VARIANTS = sorted(clw.TEMPLATE_PICKER, key=lambda k: (k[0], k[1] or ""))

BASE = {
    "work_type": "epoxy",
    "audience": "Direct",
    "values": {
        "job_name": "Cover Letter QA", "project_name": "Cover Letter QA",
        "city_state": "Olathe, KS", "bid_date_formatted": "8/26/26",
        "system_name": "Treadwell MACRO Flake", "texture": "Orange Peel",
        "epoxy_sf": "18,000", "cove_lf": "420", "polish_sf": "0",
        "schedule_notes": "~5 working days.", "estimator_name": "Kyle Loseke",
        "lump_sum": "$61,162.00",
    },
}

FULL_VALUES = dict(BASE["values"], **{
    "epoxy_system_name": "Treadwell MACRO Flake",
    "gyp_soft_thickness": '3/4"', "gyp_soft_sf": "9,000",
    "gyp_hard_thickness": '1"', "gyp_hard_sf": "2,500",
})


def _rendered(docx_bytes: bytes) -> str:
    """Every paragraph, TEXT BOXES INCLUDED. `d.paragraphs` alone would miss the floating date on
    the letter and essentially the whole front page of a proposal."""
    d = docx.Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in pw._iter_all_paragraphs(d) if p.text.strip())


def _generate(**extra):
    r = client.post("/api/generate", json={**BASE, **extra})
    assert r.status_code == 200, r.text
    return r.json()


# What the greeting paragraph says IN THE TEMPLATE, as opposed to in the rendered letter.
# Direct moved to "{{greeting}}" on 2026-09-03 (Will's wording); GC and Gyp still say "Hello,"
# literally. Five override tests below anchor on this paragraph to get a block id, and they broke
# together the last time the copy moved -- so the string lives in one place now.
_GREETING_SOURCE = {"Direct": "{{greeting}}", "GC": "Hello,", None: "Hello,"}


# Which letters print the job name in their own body copy.
#
# Until 2026-09-04 every variant did, because every variant opened with the red heading
# "<System> Proposal - {{job_name}}". Hanz took that heading off all seven, and the two copy sets
# behind them do not agree about what is left: the SHARED copy (GC and Gyp) still opens "Thanks
# for the opportunity to bid on this {{job_name}} project!", while Will's DIRECT copy says "Thank
# you for the opportunity to provide a quote for this project." and names nobody. So a Direct
# letter now identifies the job nowhere on its own page — the greeting names the PERSON — and the
# proposal stapled behind it is what carries the job name.
#
# That is a copy decision, not a bug, and it is recorded here rather than papered over: the
# assertion below is two-sided, so restoring the job name to Will's wording (or losing it from
# GC's on a copy pass) goes red here and has to be a deliberate edit.
_NAMES_THE_JOB = {"Direct": False, "GC": True, None: True}


def _greeting_block(body, audience="Direct"):
    """The block id the document editor would use for the greeting paragraph."""
    want = _GREETING_SOURCE[audience]
    for b in body["blocks"]:
        if b["text"].strip() == want:
            return b
    raise AssertionError(
        "no %r paragraph in the %s letter -- if the greeting copy moved, update "
        "_GREETING_SOURCE rather than loosening the match. Blocks: %r"
        % (want, audience, [b["text"][:30] for b in body["blocks"]][:12]))


def _template(work_type="epoxy", audience="Direct"):
    r = client.get("/api/coverletter-template?work_type=" + work_type
                   + "&audience=" + str(audience))
    assert r.status_code == 200, r.text
    return r.json()


# ── (a) the templates ────────────────────────────────────────────────────────
def test_the_picker_mirrors_the_proposal_s_shape():
    """Keyed on (work_type, audience) with audience-first folders, like
    `proposal_writer.TEMPLATE_PICKER` — including its asymmetry: gyp's audience is None there and
    here, because a gypsum bid reads the same to an owner and to a GC."""
    assert all(isinstance(k, tuple) and len(k) == 2 for k in clw.TEMPLATE_PICKER)
    assert clw.TEMPLATE_PICKER[("gyp", None)]
    assert ("gyp", "Direct") not in clw.TEMPLATE_PICKER
    assert {a for _, a in clw.TEMPLATE_PICKER} == {"Direct", "GC", None}
    for (_wt, aud), rel in clw.TEMPLATE_PICKER.items():
        assert rel.split("/")[1] == (aud or "Gyp"), rel


@pytest.mark.parametrize("key", VARIANTS)
def test_every_mapped_template_exists_on_disk(key):
    """`pick_template` falls back to Direct/Epoxy on purpose (an unmapped combination still gets a
    letter), which means a MISSING gyp file would send a gyp customer an epoxy letter and log a
    warning nobody reads. This is the assertion that turns that into a red test."""
    assert clw.has_template(*key), (
        str(key) + " has no cover-letter template on disk; pick_template would serve Direct/Epoxy")
    assert clw.pick_template(*key).is_file()


def test_the_generator_and_the_writer_name_the_same_files():
    """Two tables, one set of files. A generator that stops writing a variant the writer still
    picks is a 500 on a real send, and nothing else would notice."""
    assert {"CoverLetter/" + rel for _wt, _aud, rel in prep.VARIANTS} == set(clw.TEMPLATE_PICKER.values())
    assert {(wt, aud) for wt, aud, _ in prep.VARIANTS} == set(clw.TEMPLATE_PICKER)


def test_gc_combo_has_its_own_letter_unlike_the_proposal():
    """The proposal reuses the GC resinous DOCUMENT for a GC combo bid because Kyle never made a
    GC combo file. A letter that did the same would tell a combo customer, in prose, that the
    pages behind it are an epoxy proposal — half the scope missing from the one page written to
    describe the job."""
    assert clw.pick_template("combo", "GC") != clw.pick_template("epoxy", "GC")
    text = _rendered(clw.fill_cover_letter(work_type="combo", audience="GC", values=FULL_VALUES))
    assert "Polished Concrete" in text and "Epoxy" in text


def test_an_unmapped_work_type_still_produces_a_letter():
    """sealer/budget have no letter of their own. Falling back beats 500-ing a generate."""
    assert clw.pick_template("sealer", "GC").name == "Epoxy.docx"
    assert clw.pick_template("sealer", "GC").parent.name == "Direct"
    assert clw.has_template("sealer", "GC") is False


def test_gyp_reaches_its_one_file_from_either_audience():
    """The audience-agnostic rung of the ladder, the same one `proposal_writer.pick_template`
    uses for (gyp, None)."""
    one = clw.pick_template("gyp", None)
    assert clw.pick_template("gyp", "Direct") == one
    assert clw.pick_template("gyp", "GC") == one
    assert clw.has_template("gyp", "GC") is True


@pytest.mark.parametrize("key", VARIANTS)
def test_a_filled_letter_leaves_no_raw_token(key):
    """The rule the proposal already lives by: a customer-facing document never shows a literal
    {{token}}. Executed against the REAL template, so a token added to the copy without a value
    behind it fails here rather than on a customer's screen."""
    work_type, audience = key
    blob = clw.fill_cover_letter(work_type=work_type, audience=audience, values=FULL_VALUES)
    assert clw.unfilled_tokens(docx.Document(io.BytesIO(blob))) == set()
    text = _rendered(blob)
    assert "Kyle Loseke" in text              # {{estimator_name}} signed it
    assert ("Cover Letter QA" in text) is _NAMES_THE_JOB[audience]


@pytest.mark.parametrize("key", VARIANTS)
def test_the_letterhead_survives(key):
    """The branding is the whole point of the document. It is a full-page PNG anchored on page
    one; a template rebuilt without it is a blank sheet with three bullets on it."""
    path = clw.pick_template(*key)
    with zipfile.ZipFile(str(path)) as z:
        assert [n for n in z.namelist() if n.startswith("word/media/")], \
            path.name + " has no letterhead artwork"
    _, _, geometry = clw.describe_template(*key)
    assert geometry["images"], path.name + " draws no image on the page"


@pytest.mark.parametrize("key", VARIANTS)
def test_a_letter_is_pure_flow_with_no_floating_box(key):
    """Kyle's example letter floats the date over the artwork in a small anchored box, and these
    templates copied that box verbatim until 2026-09-04, when Hanz asked for the date off every
    format. Removing it took the ONLY text box out of these documents, so a letter is now nothing
    but flowing body paragraphs. Two things the editor reads off this response:

      * `in_block` is None everywhere — a letter has no priced/repeatable region, so every
        paragraph is freely editable and nothing is engine-owned;
      * NOTHING is positioned. The editor already had a no-box layout branch (a template with no
        boxes is a layout, not a failure — see `test_cover_letter_ui`); this is the assertion that
        says the letter takes it, and that a stray box re-appearing on a copy pass goes red here
        rather than on a customer's screen.
    """
    _, blocks, geometry = clw.describe_template(*key)
    assert len(blocks) > 5
    assert all(b["in_block"] is None for b in blocks)
    assert [b["text"] for b in blocks if b["in_txbx"]] == []
    assert geometry["boxes"] == []
    # The letterhead artwork is still drawn — it is an image, not a box.
    assert geometry["images"]


@pytest.mark.parametrize("key", VARIANTS)
def test_no_variant_prints_a_date_or_the_old_heading(key):
    """The two things Hanz took off, asserted at the FILE, for every format at once.

    Read from the raw XML rather than the paragraph text because the date lived inside
    `mc:AlternateContent` — stored twice, once as the modern `wps` shape and once as a VML
    fallback — and `d.paragraphs` sees neither copy. A half-removal that left the fallback behind
    would print the date on any consumer that reads VML, and every paragraph-level assertion in
    this file would stay green.
    """
    xml = zipfile.ZipFile(str(clw.pick_template(*key))).read("word/document.xml").decode("utf8")
    assert "{{proposal_date_short}}" not in xml, "the date box survived in this variant"
    assert "txbxContent" not in xml, "a text box survived in this variant"
    # The heading was "<System> Proposal - {{job_name}}". Match on the connective rather than one
    # variant's system name, so Polish and Combo are covered by the same line.
    assert "Proposal - " not in xml, "the red title line survived in this variant"


@pytest.mark.parametrize("key", VARIANTS)
def test_the_letterhead_artwork_is_the_only_drawing_left(key):
    """Two `wp:docPr` sharing an id is a file Word 'repairs' on open by dropping a shape — and the
    shape it drops is silent. There used to be two (the artwork and the date box); with the box
    gone there is exactly one, which is also a second, independent check that the box is really
    out of the file rather than merely emptied of its text."""
    d = docx.Document(str(clw.pick_template(*key)))
    ids = [e.get("id") for e in d.element.body.iter(pw.qn("wp:docPr"))]
    assert len(ids) == len(set(ids)) == 1, ids


def test_combo_numbering_restarts_for_the_second_system():
    """Combo carries two system sections and Hanz's Template 3 numbers each of them 1/2/3.

    This assertion is the EDITOR's view, and on its own it is not enough — it was green on
    2026-08-29 over a Word render that printed the Polished Concrete items as 4/5/6.
    `_ordered_markers` runs one counter per `(numId, ilvl)`, which is right for a list whose
    instance owns its count and blind to the case below. Kept because it is what the editor draws;
    paired with `test_the_second_system_s_numbering_really_restarts_in_word` because it is not
    what Word obeys."""
    _, blocks, _ = clw.describe_template("combo", "Direct")
    markers = [b["para"]["marker"] for b in blocks if b["para"]["marker"]]
    # DERIVED, not restated. Direct's first section gained an "Area:" item on 2026-09-03 (Will's
    # wording), so the two sections are no longer the same length and a hard-coded 3+3 was
    # asserting the copy rather than the numbering. The claim is that the count RESTARTS: two
    # runs, each beginning at 1 and counting up by one.
    runs, cur = [], []
    for m in markers:
        if m == "1." and cur:
            runs.append(cur)
            cur = []
        cur.append(m)
    if cur:
        runs.append(cur)
    assert len(runs) == 2, markers
    for run in runs:
        assert len(run) >= 3, markers        # an empty second run would satisfy "restarts"
        assert run == ["%d." % (i + 1) for i in range(len(run))], markers


@pytest.mark.parametrize("key", VARIANTS)
def test_the_second_system_s_numbering_really_restarts_in_word(key):
    """The rule Word actually applies, asserted on the real numbering part.

    A list's counter belongs to the `w:abstractNum`, NOT to the `w:num` that points at it. The
    generator gives the two Combo sections two `w:num` ids over one shared abstract definition and
    its comment claimed that was enough to restart the count; it is not, and Word rendered
    "4. Schedule:" under "Polished Concrete:". The only reset in OOXML is an explicit
    `w:lvlOverride`/`w:startOverride`.

    So: any two numbering instances this document uses that share an abstract must each carry a
    startOverride at the level they print. Checked on every variant, not just Combo — a
    single-section letter shares the same numbering part and would drift the same way if a second
    instance ever appeared."""
    d = docx.Document(str(clw.pick_template(*key)))
    numbering = d.part.numbering_part.element

    used = set()
    for p in d.element.body.iter(pw.qn("w:p")):
        for num in p.iter(pw.qn("w:numId")):
            used.add(num.get(pw.qn("w:val")))
    assert used, clw.pick_template(*key).name + " has no numbered list at all"

    by_abstract: dict = {}
    for num in numbering.findall(pw.qn("w:num")):
        if num.get(pw.qn("w:numId")) not in used:
            continue
        ref = num.find(pw.qn("w:abstractNumId"))
        by_abstract.setdefault(ref.get(pw.qn("w:val")), []).append(num)

    for abstract_id, instances in by_abstract.items():
        if len(instances) < 2:
            continue
        for num in instances:
            starts = [ov.find(pw.qn("w:startOverride"))
                      for ov in num.findall(pw.qn("w:lvlOverride"))
                      if ov.get(pw.qn("w:ilvl")) == "0"]
            starts = [s for s in starts if s is not None]
            assert starts and starts[0].get(pw.qn("w:val")) == "1", (
                "numId %s shares abstractNum %s with %d other instance(s) and has no "
                "startOverride, so Word continues the count instead of restarting it"
                % (num.get(pw.qn("w:numId")), abstract_id, len(instances) - 1))


# ── (a3) the letter is one page, and these are the four reasons it is ────────
# Combo stranded its entire sign-off on a second, LETTERHEAD-LESS page: the artwork is anchored
# to page one, so page two came out as five grey lines on blank paper. Rendered with Word COM and
# measured with PyMuPDF on 2026-08-29 — before, page one ended at 691.59pt against a 693.0pt
# text-area bottom with the tagline pushed over; after, at 677.55pt with 15.45pt to spare, and
# all seven variants are one page.
#
# The PAGE COUNT itself cannot be asserted here: there is no renderer in the test environment
# (no LibreOffice, no Word), and an estimator accurate to ±20pt cannot tell a 15pt margin from an
# overflow — it would be a green light with nothing behind it. What is asserted instead is each
# of the four measured causes, every one of them read off the real generated file.
_BODY_HALF_PT = int(round(prep.BODY_PT.pt * 2))
_SIG_HALF_PT = int(round(prep.SIG_PT.pt * 2))


def _generated_paragraphs(d):
    """The paragraphs this generator wrote: the direct `w:p` children of the body that follow the
    letterhead's `w:sdt`. Excludes Kyle's own two — the section-break paragraph before the `sdt`
    and the artwork host inside it — which are copied byte-for-byte and are not ours to size."""
    out, seen_sdt = [], False
    for child in d.element.body:
        if child.tag == pw.qn("w:sdt"):
            seen_sdt = True
        elif seen_sdt and child.tag == pw.qn("w:p"):
            out.append(child)
    assert out, "found no generated paragraphs after the letterhead sdt"
    return out


def _mark_half_pt(p):
    ppr = p.find(pw.qn("w:pPr"))
    rpr = ppr.find(pw.qn("w:rPr")) if ppr is not None else None
    sz = rpr.find(pw.qn("w:sz")) if rpr is not None else None
    return int(sz.get(pw.qn("w:val"))) if sz is not None else None


def _run_half_pts(p):
    out = []
    for r in p.findall(pw.qn("w:r")):
        sz = r.find(pw.qn("w:rPr") + "/" + pw.qn("w:sz"))
        if sz is not None:
            out.append(int(sz.get(pw.qn("w:val"))))
    return out


@pytest.mark.parametrize("key", VARIANTS)
def test_every_paragraph_mark_is_the_size_of_its_own_text(key):
    """Cause 1, and the one nothing on screen shows you.

    Word gives a line the height of the tallest thing on it and the invisible `¶` counts.
    `add_paragraph()` leaves that mark at the style default — 12pt here — so in an 11pt letter
    every paragraph's LAST line was 14.06pt tall against 12.94pt for its wrapped ones, and the
    blank line inside a 10pt signature was a 12pt blank line. Worth ~22pt on Combo. It is
    invisible in Word, invisible in the block model, and it is why the mark is set explicitly."""
    d = docx.Document(str(clw.pick_template(*key)))
    for p in _generated_paragraphs(d):
        mark = _mark_half_pt(p)
        text = "".join(t.text or "" for t in p.iter(pw.qn("w:t")))
        assert mark in (_BODY_HALF_PT, _SIG_HALF_PT), (
            "paragraph %r carries no explicit mark size, so its last line is the style's 12pt"
            % text[:40])
        runs = _run_half_pts(p)
        if runs:
            assert mark == max(runs), (
                "paragraph %r prints at %.1fpt but its mark is %.1fpt"
                % (text[:40], max(runs) / 2.0, mark / 2.0))


@pytest.mark.parametrize("key", VARIANTS)
def test_the_thank_you_runs_straight_into_the_line_that_introduces_the_proposal(key):
    """Cause 2. Example1 sets those two lines adjacent (measured at a 0.02pt gap in his render);
    the blank paragraph between them was this generator's invention and cost ~15pt."""
    d = docx.Document(str(clw.pick_template(*key)))
    texts = ["".join(t.text or "" for t in p.iter(pw.qn("w:t"))).strip()
             for p in _generated_paragraphs(d)]
    # Two wordings per line since Direct took Will's text on 2026-09-03. The claim under test
    # is ADJACENCY, not the copy -- so match either and assert nothing sits between them.
    thanks = ("Thanks for the opportunity", "Thank you for the opportunity")
    opens_the_list = ("The pages that follow", "A few things to note")
    i = next(i for i, t in enumerate(texts) if t.startswith(thanks))
    assert texts[i + 1].startswith(opens_the_list), texts[i:i + 3]


@pytest.mark.parametrize("key", VARIANTS)
def test_the_numbered_items_run_together_with_air_only_around_the_list(key):
    """Cause 3. Kyle's three items are contiguous — his `beforeAutospacing` puts ~14pt above the
    list and collapses to nothing between its rows. A `space_after` on every item spread the six
    Combo rows by ~36pt for no gain, and this asserts both halves of his shape: no space BETWEEN
    the rows, and real air ABOVE the row that opens a group."""
    d = docx.Document(str(clw.pick_template(*key)))
    numbered = [p for p in _generated_paragraphs(d)
                if p.find(pw.qn("w:pPr") + "/" + pw.qn("w:numPr")) is not None]
    assert len(numbered) >= 3, len(numbered)

    opens_a_group = []
    for p in numbered:
        spacing = p.find(pw.qn("w:pPr") + "/" + pw.qn("w:spacing"))
        after = spacing.get(pw.qn("w:after")) if spacing is not None else None
        assert after in (None, "0"), (
            "a numbered item carries %s twips of space_after; Kyle's run together" % after)
        before = spacing.get(pw.qn("w:before")) if spacing is not None else None
        opens_a_group.append(before not in (None, "0"))

    # Combo's second group opens on its heading, not on an item, so the count is 0 or 1 per
    # variant — what must not happen is that EVERY row is spaced, or that a group opens flush.
    heads = [p for p in _generated_paragraphs(d)
             if "".join(t.text or "" for t in p.iter(pw.qn("w:t"))).strip().endswith(":")
             and p.find(pw.qn("w:pPr") + "/" + pw.qn("w:spacing")) is not None]
    assert sum(opens_a_group) + len(heads) >= 1, "the numbered list opens flush against the intro"
    assert sum(opens_a_group) <= 1, "more than one numbered row is spaced off the one above it"


@pytest.mark.parametrize("key", VARIANTS)
def test_there_is_no_title_line_and_the_body_keeps_kyles_right_indent(key):
    """Every letter used to open with a red underlined heading — "Epoxy / Resinous Flooring
    Proposal - {{job_name}}" and its siblings. Hanz took it off every format on 2026-09-04: the
    proposal stapled behind the letter names the system and the job on its own front page, and a
    heading that repeats it is a second place for the job name to go stale. The greeting is now
    the first line on the page.

    The second half is the part that is easy to lose while deleting the first. Kyle's example
    carries `w:ind w:right="2340"` (117pt) on its BODY paragraphs — that is what keeps the text
    off the letterhead's right-hand artwork — and it was the TITLE that was the exception. Delete
    the title carelessly and the indent goes with it, which no reader of the diff would notice
    and every printed letter would show.
    """
    d = docx.Document(str(clw.pick_template(*key)))
    paras = _generated_paragraphs(d)
    first = "".join(t.text or "" for t in paras[0].iter(pw.qn("w:t")))
    assert "Proposal - " not in first, "the title line is still the first thing on the page"
    assert first.strip() in ("{{greeting}}", "Hello,"), first

    def right_tw(p):
        ind = p.find(pw.qn("w:pPr") + "/" + pw.qn("w:ind"))
        return ind.get(pw.qn("w:right")) if ind is not None else None

    assert all(right_tw(p) == "2340" for p in paras),         "the body lost Kyle's 117pt right indent"


def test_the_letter_does_not_read_like_an_email():
    """The PDF these templates came from was written as outbound email. Hanz confirmed this is a
    portal document page, so the email framing had to go — and this is the assertion that stops it
    creeping back in on a copy pass."""
    text = _rendered(clw.fill_cover_letter(work_type="epoxy", audience="Direct",
                                           values=FULL_VALUES)).lower()
    assert "subject line" not in text
    assert "to this email" not in text
    assert "attached" not in text


def test_the_date_is_backfilled_from_the_bid_date_not_a_clock():
    """No letter PRINTS a date any more (the letterhead box came off on 2026-09-04), but
    `_ensure_cover_letter_values` still resolves `{{proposal_date_short}}` because the same values
    dict fills the PROPOSAL in the same request. This box runs ~13 hours ahead of Central, so a
    date taken off `datetime.now()` would be a day out; a replayed payload that predates the field
    must still resolve to a real date rather than to today."""
    values = {k: v for k, v in FULL_VALUES.items() if k != "proposal_date"}
    values["bid_date_formatted"] = "8/26/26"
    assert clw._ensure_cover_letter_values(values)["proposal_date_short"] == "8/26/26"


# ── (a2) the short date still has to PARSE, even though no letter prints it ──
# Kyle drew a 63pt date box on his letterhead and Word clips rather than grows, so a long-form
# date once printed as the single word "August" on all seven letters. The box is gone (2026-09-04)
# and with it the clipping, but `{{proposal_date_short}}` is still resolved from the same values
# dict — the PROPOSAL uses it — so the parsing ladder below still has to hold.


def test_the_long_date_is_not_narrowed_for_the_proposal_too():
    """`{{proposal_date}}` and `{{proposal_date_short}}` are two tokens on purpose. The same values
    dict fills the PROPOSAL in the same request and its header prints long form; narrowing the one
    token would silently re-date every proposal document to M/D/YY."""
    out = clw._ensure_cover_letter_values(dict(FULL_VALUES, proposal_date="August 27, 2026"))
    assert out["proposal_date"] == "August 27, 2026"
    # NOT "8/27/26": `bid_date_formatted` ("8/26/26" in FULL_VALUES) outranks `proposal_date`
    # for the letterhead box, because it is what the proposal's own header prints. `proposal_date`
    # is stamped fresh to "today" on every generate and would otherwise date the letterhead the
    # day someone clicked Generate rather than the day the bid was actually dated.
    assert out["proposal_date_short"] == "8/26/26"


def test_the_short_date_prefers_the_bid_date_over_todays_stamped_proposal_date():
    """The bug this guards: `proposal_date` is recomputed to `new Date()` on every Proposal Review
    load (`proposal-review.js`), so a bid entered 2026-08-20 and finalized/sent 2026-08-27 must
    still letterhead itself 8/20/26 — the same date the proposal's own header prints — not the day
    someone happened to click Generate."""
    out = clw._ensure_cover_letter_values(dict(
        FULL_VALUES,
        bid_date_formatted="8/20/26",
        proposal_date="August 27, 2026",
    ))
    assert out["proposal_date_short"] == "8/20/26"


@pytest.mark.parametrize("raw, expect", [
    ("8/26/26", "8/26/26"),            # Kyle's own spelling, and what the payload already carries
    ("08/26/2026", "8/26/26"),         # four-digit year
    ("2026-08-26", "8/26/26"),         # the ISO bid_date column
    ("August 26, 2026", "8/26/26"),    # what the Proposal Review screen stamps
    ("Aug 26, 2026", "8/26/26"),
    ("", None),
    (None, None),
    ("TBD", None),                     # not a date; the caller must not print it in a 63pt box
    ("next Tuesday", None),
])
def test_the_short_date_reads_every_shape_the_payload_arrives_in(raw, expect):
    """`_short_date` PARSES, it never clocks — this box runs ~13 hours ahead of Central and a
    `now()` here would date a letter sent Tuesday evening as Wednesday. `%y` is tried before `%Y`
    so "8/26/26" is 2026 and not the year 26."""
    assert clw._short_date(raw) == expect


def test_an_undatable_payload_resolves_to_blank_rather_than_to_garbage():
    """Empty beats a half-parsed date. The refusal is logged naming the value that failed to parse
    — see `_ensure_cover_letter_values` — and nothing else moves. Asserted at the values pass
    because the letter no longer prints this token; a caller that DOES print it must be handed a
    blank string rather than the literal "TBD"."""
    values = {k: v for k, v in FULL_VALUES.items()
              if k not in ("proposal_date", "bid_date_formatted", "bid_date", "site_visit_date")}
    values["proposal_date"] = "TBD"
    assert clw._ensure_cover_letter_values(values)["proposal_date_short"] == ""
    # And the letter itself is still clean — blank, never a raw {{token}}.
    d = docx.Document(io.BytesIO(
        clw.fill_cover_letter(work_type="epoxy", audience="Direct", values=values)))
    assert clw.unfilled_tokens(d) == set()


def test_filling_the_letter_does_not_mutate_the_caller_s_values():
    """The SAME dict is handed to fill_proposal in the same request. A writer that grows keys on
    its caller's data is how the proposal starts printing something nobody typed."""
    values = dict(FULL_VALUES)
    before = dict(values)
    clw.fill_cover_letter(work_type="epoxy", audience="Direct", values=values)
    assert values == before


# ── (b) GenerateIn / GenerateOut ─────────────────────────────────────────────
def test_the_cover_letter_is_off_by_default():
    """Every draft saved before this feature carries none of the three keys, and
    `GenerateIn(**proposal_payload)` is how the portal PDF, the revision replay and the
    To-Dropbox re-upload rebuild those payloads."""
    gi = main.GenerateIn(**{"work_type": "epoxy", "values": {}})
    assert gi.cover_letter_enabled is False
    assert gi.cover_letter_paragraph_overrides == {}
    assert gi.cover_letter_template_version == ""


def test_disabled_means_no_download_url_at_all():
    """Not an empty string, not a url that 404s. The Done page and the portal both branch on
    whether this field is set."""
    out = _generate()
    assert out["cover_letter_download_url"] is None
    assert out["docx_download_url"]          # the proposal itself is unaffected


def test_enabled_returns_a_real_downloadable_letter():
    out = _generate(cover_letter_enabled=True)
    url = out["cover_letter_download_url"]
    assert url and url != out["docx_download_url"], (
        "the cover letter url must be its own cache token, not the proposal's")
    text = _rendered(client.get(url).content)
    # BASE is a Direct bid, and a Direct letter names neither the job nor the system since the
    # heading came off (see _NAMES_THE_JOB) — so assert on what a Direct letter DOES carry: the
    # estimator's signature and a resolved {{cover_system_line}}.
    assert "Kyle Loseke | Estimator" in text
    assert "Materials / System: Treadwell MACRO Flake" in text
    # And it is a DIFFERENT document from the proposal, not a second copy of it.
    assert "TERMS AND CONDITIONS" not in text.upper()


def test_a_broken_letter_fails_the_live_generate_loudly(monkeypatch):
    """The default half of the `want_cover_letter` gate, and the reason it defaults to True: an
    estimator who ticked the box must not be handed a silent proposal-only send. They would publish
    a portal that promises a letter and shows nothing, and only the customer would find out."""
    def boom(*a, **k):
        raise RuntimeError("template is corrupt")
    monkeypatch.setattr(clw, "fill_cover_letter", boom)
    r = client.post("/api/generate", json=dict(BASE, cover_letter_enabled=True))
    assert r.status_code >= 400, "a failed letter was swallowed on the live generate path"


def test_a_broken_letter_does_not_take_down_the_proposal_only_callers(monkeypatch):
    """`want_cover_letter=False`. /api/admin/proposal-pdf, the To-Dropbox re-file and the revision
    replay all rebuild a stored payload with `GenerateIn(**pp)` — so a draft that had the box
    ticked replays with it ticked — and NONE of them ever serves the letter. Before the gate, a
    cover-letter fault 500'd a proposal PDF the portal was waiting on and blocked a Dropbox filing
    that has nothing to do with the letter."""
    import inspect
    monkeypatch.setattr(clw, "fill_cover_letter",
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("template is corrupt")))
    gi = main.GenerateIn(**dict(BASE, cover_letter_enabled=True))
    # The signature is the contract these two call sites depend on.
    assert "want_cover_letter" in inspect.signature(main._generate).parameters
    out = main._generate(gi, _req(), persist=False, want_cover_letter=False)
    assert out.docx_download_url, "the proposal was lost to a cover-letter fault"
    assert out.cover_letter_download_url is None


def _req():
    """The minimum Request `_generate` touches (it reads headers for the base url)."""
    from starlette.requests import Request
    return Request({"type": "http", "method": "POST", "path": "/api/generate",
                    "headers": [(b"host", b"testserver")], "query_string": b"",
                    "scheme": "http", "server": ("testserver", 80), "client": ("test", 1)})


def test_every_proposal_only_call_site_still_passes_the_gate():
    """A reader of `main.py`, not of a mock: the gate is only worth having if the call sites it was
    added for actually pass it. THREE of them now — the customer proposal PDF, the To-Dropbox
    re-file, and the revision replay (/api/draft/{id}/revisions/{n}/files). If someone adds a
    fourth proposal-only replay this test does not catch it — but it does catch the known ones
    silently losing the argument in a refactor.

    Each has its own executed test asserting the value at the call site; this one exists because
    those use a stubbed `_generate` and so cannot notice a site that stopped calling it at all."""
    import inspect
    import re as _re
    src = inspect.getsource(main)
    # CALLS only. Counting the bare string would also count the comment above each call site and
    # the parameter's own docstring, which is how this test first went red against correct code.
    # Not `[^)]*` — two of the three call sites are `_generate(GenerateIn(**pp), ...)`, whose own
    # closing paren ends the class before the argument is reached.
    calls = _re.findall(r"_generate\(.{0,120}?want_cover_letter=False", src, _re.S)
    assert len(calls) == 3, (
        "a proposal-only caller stopped opting out of the letter (or a fourth one appeared "
        "without being reviewed): found %d" % len(calls))


@pytest.mark.parametrize("audience", ["Direct", "GC"])
def test_the_audience_on_the_generate_body_picks_the_letter(monkeypatch, audience):
    """`audience` already rides every generate for the proposal. The letter reads it from that
    same field rather than adding a second one that could disagree with it.

    Asserted at the seam, with the REAL writer still doing the work, because Direct/Epoxy.docx and
    GC/Epoxy.docx are byte-identical today (the source PDF has no GC copy — see the CoverLetter
    README). Comparing the produced documents would therefore pass with the audience hard-wired to
    None, and go on passing right up until the copy pass diverges them and a GC customer gets the
    owner's letter."""
    seen = {}
    real = clw.fill_cover_letter

    def spy(**kw):
        seen["audience"] = kw.get("audience")
        seen["path"] = clw.pick_template(kw["work_type"], kw.get("audience"))
        return real(**kw)

    monkeypatch.setattr(main.cover_letter_writer, "fill_cover_letter", spy)
    out = _generate(audience=audience, cover_letter_enabled=True)
    assert seen["audience"] == audience
    assert seen["path"].parent.name == audience
    text = _rendered(client.get(out["cover_letter_download_url"]).content)
    assert "Kyle Loseke" in text
    assert ("Cover Letter QA" in text) is _NAMES_THE_JOB[audience]


def test_the_letter_and_the_proposal_agree_on_the_job():
    """Built from the same `values` dict, after the same alias/backfill pass. Two documents that
    disagree about the job name is what a second values path would produce.

    The comparand is no longer the job name, and the reason is worth writing down because the
    obvious substitutes are both wrong. Since the title line came off (2026-09-04) a DIRECT letter
    names the job nowhere, so `{{job_name}}` cannot be checked against a Direct proposal; and
    swapping the whole test to GC — whose letter DOES still say "this {{job_name}} project" — only
    moves the hole, because the GC proposal template carries its job name inside the page artwork
    and no text assertion can see it. So the check runs on two values that both documents really
    print from the same dict: the system name and the estimator. A second values path would still
    be caught, which is the whole point of the test.
    """
    out = _generate(cover_letter_enabled=True)
    letter = _rendered(client.get(out["cover_letter_download_url"]).content)
    proposal = _rendered(client.get(out["docx_download_url"]).content)
    for shared in ("Treadwell MACRO Flake", "Kyle Loseke"):
        assert shared in letter and shared in proposal, shared
    # And the job name still round-trips into the proposal, which is now the only document in the
    # pair that prints it on a Direct bid.
    assert "Cover Letter QA" in proposal


def test_a_paragraph_override_reaches_the_letter():
    """The document editor's channel, keyed by the block id /api/coverletter-template hands out."""
    greeting = _greeting_block(_template())
    out = _generate(cover_letter_enabled=True,
                    cover_letter_paragraph_overrides={str(greeting["id"]): {"text": "Good morning,"}})
    text = _rendered(client.get(out["cover_letter_download_url"]).content)
    assert "Good morning," in text and "Hello," not in text


def test_a_stale_template_version_drops_the_overrides():
    """A block id is a position in a walk over ONE file. A regenerated template shifts every id
    after the changed paragraph, so a draft captured against the old one would rewrite the wrong
    sentence — silently, in a document a customer reads."""
    greeting = _greeting_block(_template())
    out = _generate(cover_letter_enabled=True,
                    cover_letter_paragraph_overrides={str(greeting["id"]): {"text": "Good morning,"}},
                    cover_letter_template_version="not-the-current-version")
    text = _rendered(client.get(out["cover_letter_download_url"]).content)
    assert "Hello," in text and "Good morning," not in text


def test_the_current_template_version_keeps_the_overrides():
    """The other half of the guard: matching versions must NOT drop an edit the estimator made."""
    body = _template()
    greeting = _greeting_block(body)
    out = _generate(cover_letter_enabled=True,
                    cover_letter_paragraph_overrides={str(greeting["id"]): {"text": "Good morning,"}},
                    cover_letter_template_version=body["template_version"])
    assert "Good morning," in _rendered(client.get(out["cover_letter_download_url"]).content)


@pytest.fixture
def identical_mtimes():
    """Force Direct/Epoxy.docx and GC/Epoxy.docx to the same mtime, and put them back after.

    Not a contrived condition: all seven letters are written by one generator run, and a `git
    checkout` or a Docker `COPY` stamps a whole tree at once. The test below would pass by luck on
    a box where the two files happened to land a few hundred nanoseconds apart."""
    paths = [clw.pick_template("epoxy", "Direct"), clw.pick_template("epoxy", "GC")]
    saved = [(p, p.stat().st_atime_ns, p.stat().st_mtime_ns) for p in paths]
    stamp = saved[0][2]
    for p in paths:
        os.utime(p, ns=(stamp, stamp))
    yield
    for p, atime, mtime in saved:
        os.utime(p, ns=(atime, mtime))


def test_an_override_captured_on_another_variant_is_dropped(identical_mtimes):
    """THE REASON `template_version` IS NOT A BARE MTIME HERE.

    With the two files stamped identically — see the fixture — mtime alone cannot tell Direct from
    GC, so this payload (a Direct edit replayed on a GC generate) would sail through the guard and
    rewrite whichever GC sentence happened to sit at that index, in a document a customer reads.
    The variant prefix is the server-side half of proposal-review.js's per-template
    `overrideKey(wt, audience)` store."""
    direct = _template("epoxy", "Direct")
    gc = _template("epoxy", "GC")
    assert direct["template_version"].split("@")[1] == gc["template_version"].split("@")[1], (
        "the fixture did not equalise the mtimes; this test would then pass for the wrong reason")
    assert direct["template_version"] != gc["template_version"]
    greeting = _greeting_block(direct)
    out = _generate(audience="GC", cover_letter_enabled=True,
                    cover_letter_paragraph_overrides={str(greeting["id"]): {"text": "Good morning,"}},
                    cover_letter_template_version=direct["template_version"])
    text = _rendered(client.get(out["cover_letter_download_url"]).content)
    assert "Hello," in text and "Good morning," not in text


def test_the_version_stamp_names_the_variant_actually_opened():
    """Built from the RESOLVED key, so the epoxy fallback and the audience-agnostic gyp entry
    stamp the file that was opened rather than the one that was asked for."""
    assert main._cover_letter_template_version("sealer", "GC").startswith("epoxy:Direct@")
    assert main._cover_letter_template_version("gyp", "GC").startswith("gyp:@")
    assert (main._cover_letter_template_version("gyp", "Direct")
            == main._cover_letter_template_version("gyp", "GC"))


@pytest.mark.parametrize("hostile", [
    {"not-an-int": {"text": "x"}},
    {"0": "a string, not a dict"},
    {"0": {"text": 17}},
    [],                      # the proposal's list shape, sent to the dict channel
    "nonsense",
])
def test_a_malformed_override_payload_cannot_break_a_generate(hostile):
    """A stale draft or a hand-built body must never 500 the one action that produces the
    customer's documents."""
    r = client.post("/api/generate", json={**BASE, "cover_letter_enabled": True,
                                           "cover_letter_paragraph_overrides": hostile})
    # A non-dict is refused by the model itself (422) — never accepted and half-applied.
    assert r.status_code in (200, 422), r.text
    if r.status_code == 200:
        assert r.json()["cover_letter_download_url"]


def test_overrides_are_applied_lowest_id_first():
    """A draft round-trips its dict keys in whatever order it was stored. Applying them in that
    order would make the result depend on JSON serialization, which is not a property anybody can
    reason about."""
    staged = main._sanitize_cover_letter_overrides({"7": {"text": "b"}, "2": {"text": "a"}})
    assert [o["id"] for o in staged] == [2, 7]


# ── (c) /api/coverletter-template ────────────────────────────────────────────
def test_the_template_endpoint_serves_the_block_model():
    body = _template("polish", "GC")
    assert body["work_type"] == "polish" and body["audience"] == "GC"
    # Folder AND file: "Polish.docx" alone does not say which of two documents this is.
    assert body["template_name"] == "GC/Polish.docx"
    assert body["template_version"].startswith("polish:GC@")
    assert body["geometry"]["page"]["w_pt"] == 612.0
    ids = [b["id"] for b in body["blocks"]]
    assert ids == list(range(len(ids))), "block ids must be the walk's positions, in order"
    # The fidelity metadata the editor renders from, on the same keys the proposal endpoint uses.
    for key in ("runs", "para", "align", "list", "price_flat", "style", "in_txbx", "txbx"):
        assert key in body["blocks"][0], key + " missing — the editor renders both documents"


def test_the_two_audiences_are_two_different_documents():
    assert _template("epoxy", "Direct")["template_name"] == "Direct/Epoxy.docx"
    assert _template("epoxy", "GC")["template_name"] == "GC/Epoxy.docx"


def test_the_template_endpoint_revalidates_cheaply():
    """Same ETag contract as /api/proposal-template — the editor re-fetches on every open."""
    first = client.get("/api/coverletter-template?work_type=epoxy&audience=Direct")
    again = client.get("/api/coverletter-template?work_type=epoxy&audience=Direct",
                       headers={"If-None-Match": first.headers["etag"]})
    assert again.status_code == 304


def test_the_etag_distinguishes_the_audiences():
    """One cached response serving both variants would show the estimator the wrong document and,
    worse, hand out ids captured against it."""
    direct = client.get("/api/coverletter-template?work_type=epoxy&audience=Direct")
    gc = client.get("/api/coverletter-template?work_type=epoxy&audience=GC",
                    headers={"If-None-Match": direct.headers["etag"]})
    assert gc.status_code == 200
    assert gc.headers["etag"] != direct.headers["etag"]


def test_an_unknown_work_type_serves_the_fallback_rather_than_an_empty_editor():
    assert _template("sealer", "GC")["template_name"] == "Direct/Epoxy.docx"


def test_the_media_route_only_serves_this_package_s_own_parts():
    """Whitelisted against the package's own listing — the same rule as the proposal's media
    route, because the name arrives in a query string."""
    _, _, geometry = clw.describe_template("epoxy", "GC")
    base = "/api/coverletter-template/media?work_type=epoxy&audience=GC&name="
    ok = client.get(base + geometry["images"][0]["name"])
    assert ok.status_code == 200 and ok.content[:4] == b"\x89PNG"
    for bad in ("../word/document.xml", "word/document.xml", "nope.png"):
        assert client.get(base + bad).status_code == 404


# ── (d) /api/admin/cover-letter-pdf ──────────────────────────────────────────
# A GET with query params, matching /api/admin/proposal-pdf exactly — the portal proxies straight
# through to it.
URL = "/api/admin/cover-letter-pdf?draft_id=d1"


def _pinned(monkeypatch, payload, live=None):
    monkeypatch.setitem(os.environ, "SERVICE_TOKEN", "svc-test")
    monkeypatch.setattr(main.drafts, "get_revision",
                        lambda did, no: {"data": {"proposal_payload": payload}})
    monkeypatch.setattr(main.drafts, "load_draft",
                        lambda did: {"data": {"proposal_payload":
                                              live if live is not None else payload}})


def _stub_generate(monkeypatch, seen, token="cl-tok"):
    def fake_generate(gi, request, *, persist=True, want_cover_letter=True):
        seen["name"] = gi.values.get("project_name")
        seen["persist"] = persist
        seen["want_cover_letter"] = want_cover_letter
        seen["enabled"] = gi.cover_letter_enabled
        return main.GenerateOut(
            work_type="epoxy", audience="Direct",
            xlsx_download_url="/api/file/x", docx_download_url="/api/file/d",
            pdf_download_url="/api/file/d/pdf", totals={},
            cover_letter_download_url="/api/file/" + token)
    monkeypatch.setattr(main, "_generate", fake_generate)
    main._FILE_CACHE[token] = {"content": b"docx", "_pdf": b"%PDF-1.4"}


def test_the_pdf_route_is_a_get_with_query_params():
    """The portal's proxy is built against this signature. A POST-with-a-body twin would be a
    second contract to keep in step with /api/admin/proposal-pdf."""
    route = next(r for r in main.app.routes
                 if getattr(r, "path", "") == "/api/admin/cover-letter-pdf")
    assert set(route.methods) == {"GET"}
    assert {"draft_id", "revision_no"} <= {p.name for p in route.dependant.query_params}


def test_the_pdf_route_is_service_token_gated(monkeypatch):
    """It is in _AUTH_PUBLIC_PATHS so the portal can reach it without a Google session — which is
    required, or the auth gate rejects the server-to-server call before this handler runs. The
    only thing then standing in front of a customer's document is this header."""
    monkeypatch.setitem(os.environ, "SERVICE_TOKEN", "svc-test")
    assert "/api/admin/cover-letter-pdf" in main._AUTH_PUBLIC_PATHS
    for headers in ({}, {"X-Service-Token": ""}, {"X-Service-Token": "svc-tes"},
                    {"X-Service-Token": "svc-test-extra"}):
        assert client.get(URL, headers=headers).status_code == 401, headers


def test_an_unset_service_token_refuses_everything(monkeypatch):
    """An unconfigured deploy must be closed, not open. `not token_env` comes FIRST so an empty
    env var can never compare equal to an empty header."""
    monkeypatch.delenv("SERVICE_TOKEN", raising=False)
    assert client.get(URL, headers={"X-Service-Token": ""}).status_code == 401


def test_the_pdf_renders_the_pinned_revision_not_the_live_draft(monkeypatch):
    """Same reason as the proposal PDF: a letter a customer opens must not disagree with the
    proposal below it because the estimator has since re-saved."""
    seen = {}
    _pinned(monkeypatch,
            {"values": {"project_name": "Snap"}, "cover_letter_enabled": True},
            live={"values": {"project_name": "LIVE"}, "cover_letter_enabled": True})
    _stub_generate(monkeypatch, seen)
    r = client.get(URL + "&revision_no=2", headers={"X-Service-Token": "svc-test"})
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith("application/pdf")
    assert seen["name"] == "Snap"
    assert seen["persist"] is False, "a customer's letter render wrote to the estimator's draft"


def test_a_project_without_a_cover_letter_is_404_by_name(monkeypatch):
    """404, NOT 500 — the portal passes it through to the customer as "no cover letter", and a
    500 there reads as a broken feature instead of a project that simply has none. The body names
    the missing thing, because that sentence is what the portal shows."""
    seen = {}
    _pinned(monkeypatch, {"values": {"project_name": "No letter"}})
    _stub_generate(monkeypatch, seen)
    r = client.get(URL, headers={"X-Service-Token": "svc-test"})
    assert r.status_code == 404
    assert "cover letter" in r.json()["detail"].lower()
    assert seen == {}, "a project with no cover letter still ran a full generate"


def test_an_ungenerated_proposal_is_422(monkeypatch):
    monkeypatch.setitem(os.environ, "SERVICE_TOKEN", "svc-test")
    monkeypatch.setattr(main.drafts, "load_draft", lambda did: {"data": {"project_name": "x"}})
    assert client.get(URL, headers={"X-Service-Token": "svc-test"}).status_code == 422


def test_a_missing_revision_is_404(monkeypatch):
    monkeypatch.setitem(os.environ, "SERVICE_TOKEN", "svc-test")
    monkeypatch.setattr(main.drafts, "get_revision", lambda did, no: None)
    assert client.get(URL + "&revision_no=9",
                      headers={"X-Service-Token": "svc-test"}).status_code == 404


def test_a_render_failure_is_a_500_that_names_the_letter(monkeypatch):
    """It must not fall back to serving the proposal, and it must not return 200 with nothing.
    The estimator reading "sent" while the customer sees an empty viewer is the failure this whole
    file is written against."""
    seen = {}
    _pinned(monkeypatch, {"values": {"project_name": "Boom"}, "cover_letter_enabled": True})
    _stub_generate(monkeypatch, seen, token="cl-boom")
    main._FILE_CACHE["cl-boom"] = {"content": b"docx"}      # no memoized _pdf -> a real render

    def boom(_blob):
        raise RuntimeError("soffice exited 1")
    monkeypatch.setattr(main.pdf_writer, "docx_to_pdf", boom)
    r = client.get(URL, headers={"X-Service-Token": "svc-test"})
    assert r.status_code == 500
    assert "cover letter" in r.json()["detail"].lower()


# ── (e) PortalPublishIn.has_cover_letter ─────────────────────────────────────
def test_has_cover_letter_defaults_to_forwarding_nothing():
    """Same contract as require_deposit beside it: omitted means the portal keeps its stored
    value, so a re-send from an older page cannot switch a customer's letter off."""
    assert main.PortalPublishIn().has_cover_letter is None


@pytest.mark.parametrize("sent,expected", [
    (None, None),        # omitted -> absent from the forwarded body
    (True, True),
    (False, False),      # explicitly OFF must travel; it is not the same as omitted
])
def test_has_cover_letter_is_forwarded_only_when_chosen(monkeypatch, sent, expected):
    """Wired like test_portal_publish.py's `_wire`: the draft, the revision snapshot and the
    outbound call are all stubbed, so what is asserted is exactly the body the portal receives."""
    cap = {}
    monkeypatch.setattr(main.drafts, "load_draft", lambda i: {"id": i, "data": {}})
    monkeypatch.setattr(main.profiles, "get_by_email", lambda e: {"email": e, "role": "admin"})
    monkeypatch.setattr(main.drafts, "create_revision", lambda did, data, by=None: 1)
    monkeypatch.setattr(main.drafts, "log_event", lambda *a, **k: None)
    monkeypatch.setattr(main, "_portal",
                        lambda path, method="GET", body=None: cap.update(body=body) or {"ok": True})

    payload = {"emails": [], "assigned_estimator": "kyle@wetreadwell.com"}
    if sent is not None:
        payload["has_cover_letter"] = sent
    r = client.post("/api/portal/publish?draft_id=d1", json=payload)
    assert r.status_code == 200, r.text
    assert cap["body"].get("has_cover_letter") == expected
    if sent is None:
        assert "has_cover_letter" not in cap["body"], (
            "an omitted flag was forwarded anyway — that overwrites the portal's stored value")


# ── Will Buchanan's Direct wording (2026-09-03) ──────────────────────────────
def test_the_direct_letter_carries_wills_text_and_not_the_first_draft():
    """His email, verbatim: *"In addition to what Greg sent you for GC projects please use the
    text template below for direct projects. The highlighted text should be pulled from the
    intake form."* The three sentences he wrote out are fixed copy and must read exactly as he
    sent them -- a paraphrase is a different letter going to a customer."""
    d = docx.Document(str(clw.pick_template("epoxy", "Direct")))
    texts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    assert "Thank you for the opportunity to provide a quote for this project." in texts
    assert "A few things to note:" in texts
    assert ("Schedule: This price is based on all work taking place in 1 phase/mobilization. "
            "If this needs to be split into multiple phases and/or over weekends, please let me "
            "know.") in texts
    assert "Feel free to reach out, if you have questions." in texts
    assert "Looking forward to working with you!" in texts
    body = "\n".join(texts)
    assert "Materials / System: {{cover_system_line}}" in body
    assert "Area: {{work_areas}}" in body


def test_the_direct_letter_no_longer_asks_the_estimator_to_pick_a_thickness():
    """Both placeholders it shipped with are answered now: the intake form asks for the thickness
    and `{{cover_system_line}}` composes the cove height. An instruction left beside a filled
    value is how a customer receives a document with "[COVE HEIGHT - pick one]" in it."""
    for wt in ("epoxy", "combo"):
        text = "\n".join(p.text for p in docx.Document(str(clw.pick_template(wt, "Direct"))).paragraphs)
        assert "[THICKNESS" not in text and "[COVE HEIGHT" not in text, wt
        assert "[SCHEDULE" not in text, wt


def test_the_gc_letter_did_not_move():
    """GC waits on Greg's text, which has not reached this repo. The module's standing rule --
    inventing contractor-flavoured sentences puts wording nobody approved in front of a customer
    -- is why Direct diverged alone rather than both letters moving together."""
    text = "\n".join(p.text for p in docx.Document(str(clw.pick_template("epoxy", "GC"))).paragraphs)
    assert "Thanks for the opportunity" in text
    assert "The pages that follow" in text
    assert "[THICKNESS" in text, "the GC letter lost its placeholder without Greg's wording"
    assert "{{cover_system_line}}" not in text
    assert prep.spec_for("epoxy", "GC") is prep.COPY["epoxy"]
    assert prep.spec_for("epoxy", "Direct") is prep.DIRECT_COPY["epoxy"]
    assert prep.spec_for("gyp", None) is prep.COPY["gyp"]


def test_the_polish_direct_letter_keeps_its_own_system_wording():
    """Will's Materials / System line is an epoxy one and he wrote nothing about polished
    concrete. Aggregate exposure and sheen are real spec decisions; guessing them to make the two
    letters look alike would put a spec claim in a customer document."""
    text = "\n".join(p.text for p in docx.Document(str(clw.pick_template("polish", "Direct"))).paragraphs)
    assert "{{cover_system_line}}" not in text
    assert "[AGGREGATE EXPOSURE" in text and "[SHEEN" in text
    assert "A few things to note:" in text            # his structure, though
    assert "Area: {{work_areas}}" in text


def test_the_area_line_appears_once_on_a_combo_letter():
    """It describes what the floor covers across the whole job. Under "Polished Concrete:" as
    well, the same words read as a second, different area."""
    text = "\n".join(p.text for p in docx.Document(str(clw.pick_template("combo", "Direct"))).paragraphs)
    assert text.count("Area: {{work_areas}}") == 1


@pytest.mark.parametrize("values, expect", [
    # Will's own example, end to end.
    ({"contact_name": "Brandon Weller", "system_name": "MACRO Flake Single Broadcast",
      "system_thickness": '1/4"', "cove_height": "6", "cove_lf": "420"},
     '1/4" MACRO Flake Single Broadcast with 6" Integral Cove Base'),
    # Three of Kyle's fifteen names already state a thickness. Prepending a DISAGREEING pick
    # would print two contradictory thicknesses in a spec line, so the name wins.
    ({"system_name": '3/16" Urethne Cement With Color Fast (SLB)', "system_thickness": '1/4"',
      "cove_lf": "0"},
     '3/16" Urethne Cement With Color Fast (SLB)'),
    # No cove, no cove clause. The proposal body still prints "with 0 LF of integral cove base"
    # on a no-cove job; the letter does not inherit that.
    ({"system_name": "MACRO Flake", "system_thickness": '1/8"', "cove_lf": "0"},
     '1/8" MACRO Flake'),
    ({"system_name": "MACRO Flake", "cove_lf": "1,200", "cove_height": "4"},
     'MACRO Flake with 4" Integral Cove Base'),      # thousands separator still reads as > 0
    ({"system_name": "MACRO Flake", "cove_lf": "420"},
     'MACRO Flake with 6" Integral Cove Base'),      # 6" is the default height
    ({"system_name": "Treadwell Polished Concrete", "cove_lf": "0"},
     "Treadwell Polished Concrete"),                 # polish: no thickness, no cove
])
def test_the_system_line_composes_what_no_template_text_could(values, expect):
    assert clw._ensure_cover_letter_values(values)["cover_system_line"] == expect


@pytest.mark.parametrize("contact, expect", [
    ("Brandon Weller", "Brandon,"),
    ("Brandon", "Brandon,"),
    ("brandon weller", "Brandon,"),        # capitalised
    ("McDonald Reyes", "McDonald,"),       # mixed case left as typed, never "Mcdonald,"
    ("", "Hello,"),                        # a bare comma is not a greeting
    (None, "Hello,"),
    ("brandon@acme.com", "Hello,"),        # Kyle sometimes has only an email
    # Found by this table, not confirmed by it: both sides greeted "B.," until the guard learned
    # to measure a name with an initial's punctuation removed. Kyle types this shape whenever he
    # has a surname and an initial and no more.
    ("B. Weller", "Hello,"),               # a lone initial is not a first name
    ("J.R. Weller", "J.R.,"),              # two real letters, and what that person is called
])
def test_the_greeting_is_a_first_name_or_an_honest_fallback(contact, expect):
    """Will's letter opens on the contact's first name and nothing else. NOTE a known wart, left
    alone on purpose: a contact typed "Weller, Brandon" greets "Weller," -- special-casing the
    comma would misfire on "Brandon Weller, PE", which is the commoner shape in this database."""
    assert clw._ensure_cover_letter_values({"contact_name": contact})["greeting"] == expect


def test_a_blank_area_box_prints_the_square_footage_rather_than_a_bare_label():
    """The letter's list items are static paragraphs -- `cover_letter_writer` deliberately does
    not port the `{{#block}}` expansion that could drop one -- so "Area:" with nothing after it
    is what an empty box would ship."""
    out = clw._ensure_cover_letter_values({"area_description": "~18,000 sf of epoxy flooring"})
    assert out["work_areas"] == "~18,000 sf of epoxy flooring"
    typed = clw._ensure_cover_letter_values(
        {"work_areas": "Warehouse expansion and 4 offices",
         "area_description": "~18,000 sf of epoxy flooring"})
    assert typed["work_areas"] == "Warehouse expansion and 4 offices"


def test_the_letter_prints_no_raw_token_with_the_new_fields_supplied():
    """The end-to-end shape: what the intake form now collects, through the real writer."""
    values = dict(FULL_VALUES, contact_name="Brandon Weller",
                  work_areas="Warehouse expansion and 4 offices", system_thickness='1/4"',
                  cove_height="6")
    for key in VARIANTS:
        d = docx.Document(io.BytesIO(clw.fill_cover_letter(
            work_type=key[0], audience=key[1], values=values)))
        assert clw.unfilled_tokens(d) == set(), key
    text = _rendered(clw.fill_cover_letter(work_type="epoxy", audience="Direct", values=values))
    assert "Brandon," in text
    assert "Area: Warehouse expansion and 4 offices" in text
    assert '1/4" Treadwell MACRO Flake with 6" Integral Cove Base' in text


# ── the two resolutions must agree ───────────────────────────────────────────
_HARNESS = pathlib.Path(__file__).resolve().parent / "js" / "payload-sync-harness.js"
_FRONTEND = pathlib.Path(__file__).resolve().parents[2] / "frontend"


@pytest.fixture(scope="module")
def js_tokens():
    """What the BROWSER resolves, from the real `computeTokenValues` out of proposal-review.js.

    Executed, not read. A source assertion that both files "look the same" cannot catch a
    regex that behaves differently in the two languages, which is the whole risk here."""
    if shutil.which("node") is None:
        pytest.skip("node is not installed")
    proc = subprocess.run(["node", str(_HARNESS), str(_FRONTEND)],
                          capture_output=True, text=True, encoding="utf-8", timeout=180)
    assert proc.returncode == 0, (
        "the harness itself failed — read this before assuming a product bug:\n" + proc.stderr)
    cases = json.loads(proc.stdout.strip().splitlines()[-1])["coverLetterTokens"]
    assert len(cases) >= 10, "the case table shrank; a thin table is how a fork survives"
    return cases


def test_the_screen_and_the_document_resolve_the_same_three_tokens(js_tokens):
    """THE BUG THIS EXISTS FOR. `{{greeting}}`, `{{work_areas}}` and `{{cover_system_line}}` are
    derived in two places: `computeTokenValues` (proposal-review.js) for the cover-letter editor's
    on-screen preview, and `_ensure_cover_letter_values` here for generate and for the portal's
    server-side replay of a pinned revision. A token resolved on only one side previews as a raw
    `{{token}}` over a correct PDF — exactly what PR #431 fixed for `{{proposal_date_short}}` —
    and an estimator proofreading on screen cannot tell that from a broken document.

    Neither side's answers are written down here. The JS side is executed, its answers and the
    upstream values it read are both returned, and Python is handed those same upstream values.
    So this asserts AGREEMENT, which is the property that matters; the per-branch expectations
    live in the parametrized tests above, where a literal is the point."""
    disagreements = []
    for case in js_tokens:
        up = {k: v for k, v in case["upstream"].items() if v is not None}
        out = clw._ensure_cover_letter_values(up)
        for token in ("greeting", "work_areas", "cover_system_line"):
            if out[token] != case[token]:
                disagreements.append(
                    "%s / %s: screen %r, document %r" % (case["name"], token, case[token], out[token]))
    assert not disagreements, (
        "the preview and the PDF have forked:\n  " + "\n  ".join(disagreements))


def test_the_agreement_table_actually_exercises_every_branch(js_tokens):
    """A green agreement test proves nothing if every row is the happy path — two identical
    implementations agree on anything. These are the branches that can fork."""
    lines = {c["name"]: c["cover_system_line"] for c in js_tokens}
    greets = {c["name"]: c["greeting"] for c in js_tokens}
    # a thickness prepended, and a thickness correctly NOT prepended
    assert lines["wills_own_example"].startswith('1/4"')
    assert lines["name_already_states_a_thickness"].startswith('3/16"'), lines
    assert "1/4" not in lines["name_already_states_a_thickness"], (
        "the estimator's pick was prepended onto a name that already states a different "
        "thickness — the letter now claims two")
    # the cove clause present and absent
    assert "Integral Cove Base" in lines["wills_own_example"]
    assert "Integral Cove Base" not in lines["no_cove_drops_the_clause"]
    # both greeting outcomes
    assert greets["wills_own_example"] == "Brandon,"
    assert set(greets.values()) >= {"Brandon,", "Hello,", "McDonald,"}
    # the fallback, and a typed value beating it
    fallback = next(c for c in js_tokens if c["name"] == "blank_area_falls_back_to_the_sf_line")
    assert "sf of" in fallback["work_areas"] and fallback["upstream"]["work_areas"] == ""
