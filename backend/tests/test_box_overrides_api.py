"""`box_overrides` through /api/generate — the plumbing, not the XML.

test_box_resize.py proves the resize itself. This proves the request actually carries it, and
that the two ways it can go wrong are covered:

  * **Stale ids.** A box id is a position in `_iter_txbx` over one specific template file. A
    re-annotation shifts those positions, so a draft captured against the old template would
    resize a DIFFERENT box than the estimator dragged — silently, and only visible in the
    customer's document. The same `template_version` guard that protects paragraph_overrides
    has to protect these, which is why the guard was extended rather than duplicated.
  * **Hostile or stale input.** A dict arriving from a saved draft must never 500 a generate.
"""
import io
import zipfile

import docx
import pytest
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

BASE = {
    "work_type": "polish",
    "audience": "GC",
    "values": {"job_name": "Box Override API", "city_state": "Lenexa, KS",
               "polish_sf": "1000", "lump_sum": "$10,000"},
}


def _first_box_cy(docx_bytes):
    """The first text box's wp:extent/@cy, straight out of the generated file."""
    d = docx.Document(io.BytesIO(docx_bytes))
    txbx = next(iter(pw._iter_txbx(d)))
    anchor = pw._txbx_anchor(txbx)
    return int(anchor.find(pw.qn("wp:extent")).get("cy"))


def _first_box_pos(docx_bytes):
    """The first text box's anchor OFFSETS in points, straight out of the generated file."""
    d = docx.Document(io.BytesIO(docx_bytes))
    anchor = pw._txbx_anchor(next(iter(pw._iter_txbx(d))))
    return (pw._anchor_offset(anchor, "positionH")[0],
            pw._anchor_offset(anchor, "positionV")[0])


def _generate(**extra):
    r = client.post("/api/generate", json={**BASE, **extra})
    assert r.status_code == 200, r.text
    return client.get(r.json()["docx_download_url"]).content


def test_a_resize_reaches_the_generated_document():
    plain = _first_box_cy(_generate())
    resized = _first_box_cy(_generate(box_overrides={"0": {"h_pt": 500.0}}))
    assert resized != plain
    assert resized == int(round(500.0 * pw._EMU_PER_PT))


def test_no_box_overrides_changes_nothing():
    assert _first_box_cy(_generate()) == _first_box_cy(_generate(box_overrides={}))


def test_a_move_reaches_the_generated_document():
    """The other half of the wire format. `x_pt`/`y_pt` ride the same dict as the size, so the
    only thing that could go wrong between the browser and the .docx is a field name."""
    plain = _first_box_pos(_generate())
    # Box 0 is the DATE/JOB NAME header, 72 x 18pt at (18.35, 36) on Kyle's sheet — the corner is
    # a legal position for it, which the sanitiser would refuse for a full-width box.
    moved = _first_box_pos(_generate(box_overrides={"0": {"x_pt": 200.0, "y_pt": 300.0}}))
    assert moved != plain, "the move never left the request"
    dx, dy = moved[0] - plain[0], moved[1] - plain[1]
    assert (dx, dy) != (0.0, 0.0)
    # An OFFSET, not a page coordinate: 200pt across on a 90pt left margin is a 110pt offset.
    assert moved[0] == pytest.approx(110.0, abs=0.02)


def test_a_stale_template_version_drops_the_move_too():
    """Same guard as the resize, because a box id means the same thing for both: a position in the
    backend's walk over one specific .docx. A stale draft must not relocate whatever box now
    happens to sit at that index."""
    plain = _first_box_pos(_generate())
    stale = _first_box_pos(_generate(box_overrides={"0": {"x_pt": 200.0, "y_pt": 300.0}},
                                     template_version="STALE-NOPE"))
    assert stale == plain, "a stale template_version still moved the box"


def test_a_stale_template_version_drops_the_resize():
    """The important one. A stale draft must NOT resize whatever box now happens to sit at
    that index — better the design size than a confidently wrong document."""
    plain = _first_box_cy(_generate())
    stale = _first_box_cy(_generate(box_overrides={"0": {"h_pt": 500.0}},
                                    template_version="STALE-NOPE"))
    assert stale == plain, "a stale template_version still applied the box resize"


def test_a_matching_template_version_applies_the_resize():
    tv = client.get("/api/proposal-template?work_type=polish&audience=GC").json()["template_version"]
    got = _first_box_cy(_generate(box_overrides={"0": {"h_pt": 500.0}}, template_version=tv))
    assert got == int(round(500.0 * pw._EMU_PER_PT))


def test_an_empty_template_version_is_a_legacy_caller_and_applies():
    got = _first_box_cy(_generate(box_overrides={"0": {"h_pt": 500.0}}, template_version=""))
    assert got == int(round(500.0 * pw._EMU_PER_PT))


def test_the_guard_drops_paragraph_and_box_overrides_together():
    """They share one guard; if a later edit split them apart, one kind would keep applying
    against ids the other had already decided were untrustworthy."""
    src = pw.pick_template("polish", "GC")
    d = docx.Document(str(src))
    pid = next(i for i, _k, _p, in_block, text, _t in pw.iter_editable_blocks(d)
               if in_block is None and (text or "").strip())
    blob = _generate(paragraph_overrides=[{"id": pid, "text": "BOTH-GUARD-TEST"}],
                     box_overrides={"0": {"h_pt": 500.0}},
                     template_version="STALE-NOPE")
    with zipfile.ZipFile(io.BytesIO(blob)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    assert "BOTH-GUARD-TEST" not in xml
    assert _first_box_cy(blob) == _first_box_cy(_generate())


@pytest.mark.parametrize("bad", [
    {"0": {"h_pt": "tall"}},
    {"0": {"h_pt": None}},
    {"0": None},
    {"nope": {"h_pt": 300}},
    {"0": {"h_pt": 99999}},
    {"0": {"h_pt": -300}},
    {"999999": {"h_pt": 300}},
    {"0": {"unknown_field": 3}},
])
def test_malformed_box_overrides_still_generate_a_document(bad):
    """A saved draft from an older client, or a hand-built request, must not take
    /api/generate down."""
    r = client.post("/api/generate", json={**BASE, "box_overrides": bad})
    assert r.status_code == 200, r.text


def test_a_non_dict_box_overrides_is_refused_by_validation_not_a_500():
    r = client.post("/api/generate", json={**BASE, "box_overrides": ["not", "a", "dict"]})
    assert r.status_code in (200, 422), r.text
    assert r.status_code != 500
