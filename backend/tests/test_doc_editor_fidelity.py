"""Bold, italic, underline and font size have to survive a reload — and a reload must never
make the saved draft poorer than it found it.

THE BUG, live on production until 2026-08-21.

`restoreSavedOverrides` replayed a saved override with `el.textContent = o.text` and never read
`o.runs`; `runs` was only ever read off the API's template record, never off a saved override.
So after ANY re-init — F5, re-opening the draft, a trip to Done or Estimate Review and back, a
work-type or template switch that re-runs `initDocumentEditor` in place — the paragraph was one
plain text node again, with no `tw-fmt` class.

That is where it stops being a display bug. `collectOverrides` then re-serialised the flattened
paragraph: `serializeRuns` returned one plain run, `runsArePlain` agreed, `fmtChanged` was false,
and the stored entry DEGRADED from `{id, text, runs}` to `{id, text}`. The 800 ms
`schedulePersistOverrides` wrote that back over the good one. The estimator's formatting was not
merely hidden from them, it was destroyed in their draft.

Hanz, repeatedly: "whatever the update is in the proposal tool it should be one to one, spacing,
font size, indentation ETC." The font-size half was broken outright.

TWO SEPARATE PROPERTIES ARE FIXED, and both are asserted independently:

  (a) a reload rebuilds the RUNS, not just the text (`restoreSavedOverrides`);
  (b) a serialise can never downgrade a stored entry (`preserveRichOverrides`), asserted with
      the restore DELIBERATELY SKIPPED. A guard that only holds while (a) is correct is not a
      guard: the next edit to either function would put the data loss straight back, and the
      symptom is silent and permanent.

THE TOKEN QUESTION, decided rather than guessed. The same paragraphs carry `.tw-fill`
`data-token` spans whose text comes from the estimate. A stored run now records the token whose
value it still holds, so the restore re-substitutes the CURRENT value instead of freezing the
one that happened to be on screen when the formatting was applied — an estimator who corrects
the square footage must not find the old number welded into a paragraph they once bolded. The
run's stored TEXT stays the resolved value, so the .docx path is unchanged and there is no new
way for a raw `{{token}}` to reach a customer.

A run is tagged only when replaying it cannot destroy anything, which rules out two cases:

  1. the estimator typed somewhere in the paragraph — `{{scope_notes}}` renders as an editable
     fill and rewording the scope in place is a first-class use of this editor, so once the
     words differ from the pristine rendering they belong to the estimator and no run may be
     replaced by a sidebar value;
  2. the token appears in more than one run, because formatting HALF a value splits its fill —
     writing the whole value into each half would print it twice.

Where those two goals genuinely conflict (case 2, e.g. bolding one word inside a scope
narrative) the fix keeps the estimator's characters and lets the value freeze, which is exactly
what the old code did for every case. Refusing is never worse than today.

WHY THE FRONTEND HALF RUNS UNDER NODE. This bug IS a disagreement between two functions that
each look correct alone: one writes a shape, the other reads a different one. No source-text
assertion can see that, and on 2026-08-12 `STAGE_CREATED` shipped unbound with every source
assertion green and took the production board down. So `js/doc-editor-fidelity-harness.js` lifts
the shipped functions — including the real `schedulePersistOverrides`, the writer that did the
damage — gives them the smallest DOM they touch, and drives the whole round trip through them.
Then the runs that survive the round trip are put through the REAL writer and read back out of a
generated .docx, because "the bold reached the document" is a claim about the file.
"""
import io
import json
import pathlib
import shutil
import subprocess

import docx
import pytest
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

import proposal_writer as pw

FRONTEND = pathlib.Path(__file__).resolve().parents[2] / "frontend"
HARNESS = pathlib.Path(__file__).resolve().parent / "js" / "doc-editor-fidelity-harness.js"

_EPOXY_TEMPLATE = "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx"

# The bolded phrase, as the harness produces it. Stated once here so every assertion below
# compares against the same object rather than restating a run list five times.
BOLD_ROUND_TRIP = [
    {"text": "Scope:", "bold": True, "size_pt": 8},
    {"text": "  ", "bold": False, "size_pt": 8},
    {"text": "Grind", "bold": True, "size_pt": 8},
    {"text": " and coat.", "bold": False, "size_pt": 8},
]


@pytest.fixture(scope="module")
def ran():
    if shutil.which("node") is None:
        pytest.skip("node is not installed")
    proc = subprocess.run(["node", str(HARNESS), str(FRONTEND)],
                          capture_output=True, text=True, encoding="utf-8", timeout=300)
    assert proc.returncode == 0, (
        "the harness itself failed — read this before assuming a product bug:\n" + proc.stderr)
    return json.loads(proc.stdout.strip().splitlines()[-1])


# ══ the fixture is the real template ═════════════════════════════════════════════════════════
def test_the_fixture_block_is_the_one_the_editor_actually_gets(ran):
    """Every assertion below is about block 115 of the Direct epoxy template. If the harness's
    copy of its runs has drifted from the file, the round trip proves nothing about what Kyle
    sees. `/api/proposal-template` reports the WORK rows after `_normalize_work_label_formatting`
    (main.py does that so the preview matches the generated file), so those are the runs."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    pw._normalize_work_label_formatting(d)
    real = None
    for _idx, _kind, p_elem, _in_block, text, _txbx in pw.iter_editable_blocks(d):
        if text.startswith("Scope:"):
            real = pw._block_runs(p_elem, Paragraph(p_elem, d))
            break
    assert real is not None, "the Scope: row is gone from the template"
    assert ran["fixture"]["block115"] == real, (
        "the harness fixture has drifted from the template:\nfixture %r\nreal    %r"
        % (ran["fixture"]["block115"], real))


# ══ (a) a reload rebuilds the RUNS ═══════════════════════════════════════════════════════════
def test_bolding_a_phrase_reaches_the_payload(ran):
    """The baseline. Bold five characters in the middle of a paragraph and the runs split at
    exactly those characters, with the neighbours keeping what the template gave them — the
    explicit `bold: False` included, because absent would mean "inherit" and on a bold paragraph
    style that is the opposite instruction."""
    assert ran["roundTrip"]["pristine"]["runs"] == [
        {"text": "Scope:", "bold": True, "size_pt": 8},
        {"text": "  Grind and coat.", "bold": False, "size_pt": 8},
    ]
    assert ran["roundTrip"]["bolded"]["runs"] == BOLD_ROUND_TRIP
    assert ran["roundTrip"]["sent"] == [
        {"id": 115, "text": "Scope:  Grind and coat.", "runs": BOLD_ROUND_TRIP}]
    assert ran["roundTrip"]["stored"] == ran["roundTrip"]["sent"], (
        "the persist did not store what the collector sent")


def test_the_reload_puts_the_runs_back_not_just_the_words(ran):
    """THE FIX. A fresh page over the same draft has nothing in memory: every span has to come
    back out of the stored entry. Before this, `restored.runs` was a single plain run."""
    restored = ran["roundTrip"]["restored"]
    assert restored["runs"] == BOLD_ROUND_TRIP, (
        "the reload flattened the paragraph: %r" % (restored["runs"],))
    assert restored["text"] == "Scope:  Grind and coat.", "the words changed on the way back"
    assert restored["dirty"] is True
    assert restored["fmt"] is True, (
        "tw-fmt was not restored, so the very next keystroke would degrade the entry again — "
        "the paragraph would LOOK right and still lose the work")


def test_the_next_autosave_no_longer_overwrites_the_draft_with_plain_text(ran):
    """The destructive half, in order: restore, re-serialise, persist. This is the exact
    sequence that used to replace `{id, text, runs}` with `{id, text}` in the saved draft."""
    assert ran["roundTrip"]["resent"] == ran["roundTrip"]["sent"], (
        "a serialise after a restore changed what gets sent")
    assert ran["roundTrip"]["afterSecondPersist"] == ran["roundTrip"]["stored"], (
        "the autosave after a reload rewrote the stored entry")


def test_it_does_not_degrade_over_repeated_reloads(ran):
    """The old loss was progressive — each visit flattened a little more of what was left — so
    one round trip is not enough evidence. Restore, persist, restore again."""
    assert ran["roundTrip"]["thirdTrip"]["runs"] == BOLD_ROUND_TRIP
    assert ran["roundTrip"]["thirdTrip"]["fmt"] is True


def test_a_font_size_and_an_italic_survive_the_same_trip(ran):
    """The half the reported symptom hid. A size has no HTML tag behind it — it exists only as
    an inline style and a `w:sz` — so it is the switch most likely to be dropped by a rebuild,
    and it is the one Hanz asked for by name."""
    assert ran["sizeAndItalic"]["restored"]["runs"] == [
        {"text": "Scope:", "bold": True, "italic": True, "size_pt": 8},
        {"text": "  ", "bold": False, "size_pt": 8},
        {"text": "Grind", "bold": False, "size_pt": 14},
        {"text": " and coat.", "bold": False, "size_pt": 8},
    ]
    assert ran["sizeAndItalic"]["stored"][0]["runs"] == \
        ran["sizeAndItalic"]["restored"]["runs"]


# ══ (b) the guard, independent of (a) ════════════════════════════════════════════════════════
def test_a_page_that_never_restored_still_cannot_destroy_the_runs(ran):
    """THE GUARD, exercised against a page in exactly the state the bug left it in: the runs are
    stored, the DOM knows nothing about them, and a keystroke elsewhere fires the persist. The
    collector has to hand back the stored runs it cannot see."""
    good = ran["guard"]["good"]
    assert good[0]["runs"] == BOLD_ROUND_TRIP, "the fixture never had formatting to protect"
    rescued = [o for o in ran["guard"]["collected"] if o["id"] == 115]
    assert rescued and rescued[0]["runs"] == BOLD_ROUND_TRIP, (
        "collectOverrides dropped the stored runs: %r" % (ran["guard"]["collected"],))
    assert ran["guard"]["stored"] == ran["guard"]["collected"], (
        "the persist wrote a poorer entry than it found")
    # And the unrelated edit that triggered the persist still went through.
    assert {"id": 116, "text": "Area: 5,200 SF and cove"} in ran["guard"]["collected"]


def test_the_guard_still_lets_the_estimator_take_formatting_OFF(ran):
    """It has to be narrow or it becomes a different bug: formatting that cannot be removed.
    Reset sends ONE plain run (the `tw-fmt` class is never taken off, so the collector keeps
    using the runs branch) and emptying a paragraph sends `runs: []`. Both are arrays, so
    neither is mistaken for the signature of the bug."""
    assert ran["reset"]["stored"] == [
        {"id": 115, "text": "Scope:  Grind and coat.", "runs": [{"text": "Scope:  Grind and coat."}]}], (
        "Reset was undone by the guard")
    assert ran["reset"]["onScreen"]["runs"] == [{"text": "Scope:  Grind and coat."}]
    assert ran["emptied"]["stored"] == [{"id": 115, "text": "", "runs": []}], (
        "an emptied paragraph came back from the dead")
    assert ran["emptied"]["onScreen"]["empty"] is True


def test_an_entry_from_another_version_of_the_template_is_neither_replayed_nor_rescued(ran):
    """Paragraph ids belong to ONE template file. A stored entry captured against a different
    version describes different paragraphs, so replaying it would rewrite whichever paragraph
    happens to hold that number — and the guard must not smuggle it back in either."""
    assert ran["staleVersion"]["restored"]["dirty"] is False
    assert ran["staleVersion"]["restored"]["fmt"] is False
    assert ran["staleVersion"]["stored"] == [{"id": 116, "text": "Area: 5,200 SF and cove"}], (
        "the version-mismatched entry was rescued into the new store")


def test_an_untouched_document_still_ships_nothing(ran):
    """The generated .docx for a document nobody edited has to be the file it was before any of
    this existed."""
    assert ran["untouched"] == []


# ══ the token question ═══════════════════════════════════════════════════════════════════════
def test_a_token_whose_value_changed_shows_the_NEW_value(ran):
    """Bold the "Area:" label, then correct the square footage on Estimate Review and come back.
    The number the customer reads must be the new one. The stored run records the token, so the
    restore re-reads it instead of replaying "5,200" forever."""
    saved = ran["tokenFresh"]["savedRuns"]
    tagged = [r for r in saved if r.get("tok")]
    assert tagged == [{"text": "5,200", "bold": False, "size_pt": 8, "tok": "epoxy_sf"}], (
        "the estimate-sourced run was not tagged with its token: %r" % (saved,))
    restored = ran["tokenFresh"]["restored"]
    assert restored["text"] == "Area: 6,000 SF", (
        "a formatted paragraph froze the old square footage: %r" % (restored["text"],))
    assert restored["runs"][0] == {"text": "Area:", "bold": True, "size_pt": 8}, \
        "the bold label was lost while re-reading the token"
    assert ran["tokenFresh"]["stored"]["text"] == "Area: 6,000 SF"


def test_the_fill_span_is_rebuilt_exactly_once(ran):
    """The other half of the same question: restoring runs must not orphan the `.tw-fill` spans
    or duplicate them. One tagged run, one span, one copy of the value."""
    assert ran["tokenFresh"]["restored"]["fills"] == [["epoxy_sf", "6,000"]]
    assert ran["liveFill"]["after"]["fills"] == [["epoxy_sf", "6,000"]]


def test_text_the_estimator_typed_into_a_fill_is_never_replaced_by_the_sidebar(ran):
    """`{{scope_notes}}` renders as an editable fill, and rewording the scope straight in the
    document is a first-class use of this editor. Reloaded against a DIFFERENT sidebar value —
    the hostile case — their characters come back untouched, because the collector refuses to
    tag any run in a paragraph whose words have changed."""
    assert not any(r.get("tok") for r in ran["tokenTyped"]["savedRuns"]), (
        "a run was tagged in a paragraph the estimator had typed in")
    assert ran["tokenTyped"]["restored"]["text"] == "Scope:  Grind, patch and coat."
    assert ran["tokenTyped"]["restored"]["runs"] == [
        {"text": "Scope:", "bold": True, "size_pt": 8},
        {"text": "  Grind, ", "bold": False, "size_pt": 8},
        {"text": "patch", "bold": True, "size_pt": 8},
        {"text": " and coat.", "bold": False, "size_pt": 8},
    ]


def test_the_typing_rule_holds_even_when_the_fill_is_left_whole(ran):
    """The test above formats INSIDE the fill, which splits it — so the duplicate-token rule
    would have dropped the tag on its own and the test proves nothing about the typing rule.
    Here the estimator types in the scope fill and bolds the LABEL instead. The fill is still one
    run, so only "the words changed, so they belong to the estimator" can save their wording."""
    got = ran["tokenTypedFormatOutside"]
    assert got["fillsAtSave"] == [["scope_notes", "Grind, patch and coat."]], (
        "the fill was split or lost, so this case has stopped isolating the typing rule: %r"
        % (got["fillsAtSave"],))
    assert not any(r.get("tok") for r in got["savedRuns"]), (
        "a whole, unsplit fill run was tagged in a paragraph the estimator had typed in: %r"
        % (got["savedRuns"],))
    assert got["restored"]["text"] == "Scope:  Grind, patch and coat."
    assert got["restored"]["runs"] == [
        {"text": "Scope:", "bold": True, "italic": True, "size_pt": 8},
        {"text": "  Grind, patch and coat.", "bold": False, "size_pt": 8},
    ]


def test_a_value_split_in_half_by_formatting_freezes_rather_than_doubling(ran):
    """Bold "5," out of "5,200" and the fill becomes two spans. There is no honest way to spread
    a NEW value across them, so both keep their stored text: the number goes stale, exactly as
    it did before this change, instead of being printed twice. The conflict is real and this is
    the side that loses least — a duplicated figure is a wrong price in a contract."""
    stored = ran["splitFill"]["stored"]
    assert not any(r.get("tok") for r in stored["runs"])
    assert ran["splitFill"]["restored"]["text"] == "Area: 5,200 SF", (
        "the split value was rebuilt from the token: %r" % (ran["splitFill"]["restored"]["text"],))
    assert ran["splitFill"]["restored"]["runs"] == [
        {"text": "Area: ", "bold": False, "size_pt": 8},
        {"text": "5,", "bold": True, "size_pt": 8},
        {"text": "200 SF", "bold": False, "size_pt": 8},
    ]


def test_a_live_sidebar_change_reaches_a_block_that_was_only_formatted(ran):
    """A stale number is the same wrong number with or without a reload. `refreshDocumentFills`
    skips every `tw-dirty` block, which is right for one somebody typed in and wrong for one
    where only the styling changed — its words are still the template's, so its fills still owe
    the sidebar their live values. Updated span by span, because the innerHTML IS the
    formatting."""
    assert ran["liveFill"]["before"]["text"] == "Area: 5,200 SF"
    after = ran["liveFill"]["after"]
    assert after["text"] == "Area: 6,000 SF"
    assert after["runs"][0] == {"text": "Area:", "bold": True, "size_pt": 8}, \
        "the live re-fill erased the formatting it was supposed to preserve"
    assert ran["liveFill"]["stored"]["text"] == "Area: 6,000 SF", (
        "the fresh value never reached the draft, so the next generate would still be stale")


def test_the_live_re_fill_leaves_a_hand_typed_block_alone(ran):
    """The same pass must not touch a paragraph whose words are the estimator's."""
    assert ran["liveFillTyped"]["text"] == "Area: 5,200 SF plus cove"


def test_a_block_that_was_formatted_AND_typed_in_keeps_its_words(ran):
    """The case the `tw-fmt` gate lets through, so the baseline comparison inside
    `refreshFillsInPlace` is the only thing left protecting the estimator's sentence. Without
    the case above it, removing that comparison passes every other test in this file."""
    assert ran["liveFillFormattedThenTyped"]["text"] == "Area: 5,200 SF plus 60 LF cove", (
        "the live re-fill overwrote words the estimator typed into a formatted paragraph")
    assert ran["liveFillFormattedThenTyped"]["fills"] == [["epoxy_sf", "5,200"]], (
        "the fill span was gone, so this case proves nothing about the baseline comparison")
    # And the harsher shape: the highlighted number itself overtyped. "5,250" is a figure the
    # estimator chose; the sidebar does not get to overrule it behind their back.
    assert ran["liveFillOvertypedNumber"]["text"] == "Area: 5,250 SF", (
        "the live re-fill overwrote a number the estimator typed over: %r"
        % (ran["liveFillOvertypedNumber"]["text"],))


def test_the_live_re_fill_still_works_the_SECOND_time(ran):
    """`refreshFillsInPlace` moves the pristine baseline with the value it writes. Without that
    the block reads as hand-edited from then on: the next re-fill refuses, the token tag is
    dropped from the stored runs, and the number quietly stops tracking the estimate — the same
    frozen figure, arrived at one edit later."""
    assert ran["liveFillTwice"]["mid"]["text"] == "Area: 6,000 SF"
    assert any(r.get("tok") for r in ran["liveFillTwice"]["mid"]["runs"]), (
        "the token tag was dropped after the first live re-fill")
    assert ran["liveFillTwice"]["after"]["text"] == "Area: 7,500 SF", (
        "a second sidebar change never reached the paragraph")
    assert ran["liveFillTwice"]["stored"]["text"] == "Area: 7,500 SF"
    assert any(r.get("tok") for r in ran["liveFillTwice"]["stored"]["runs"])


# ══ the three environments the brief asked for ═══════════════════════════════════════════════
def test_two_tabs_on_one_draft_do_not_wipe_each_others_formatting(ran):
    """Two real editors over one store. The second tab loaded BEFORE the first formatted
    anything, so its DOM cannot see the runs — which is what used to make its next autosave
    destructive. It edits an unrelated paragraph, and both edits survive."""
    after = ran["twoTabs"]["afterB"]
    kept = [o for o in after if o["id"] == 115]
    assert kept and kept[0]["runs"] == BOLD_ROUND_TRIP, (
        "tab B's autosave destroyed tab A's formatting: %r" % (after,))
    assert {"id": 116, "text": "Area: 5,200 SF and cove"} in after, "tab B lost its own edit"
    assert ran["twoTabs"]["tabBSees"]["runs"] == BOLD_ROUND_TRIP, (
        "tab B could not see the formatting after reloading")


def test_a_template_switch_mid_session_keeps_each_templates_formatting(ran):
    """`reloadForWorkType` re-runs `initDocumentEditor` in place, with no page load — the case
    the per-template store exists for. Epoxy's runs must survive the polish visit, and polish
    must not inherit them: the ids mean different paragraphs in a different file."""
    assert ran["templateSwitch"]["epoxyStored"][0]["runs"] == BOLD_ROUND_TRIP
    polish = ran["templateSwitch"]["polishAfterRestore"]
    assert polish["dirty"] is False and polish["fmt"] is False, (
        "the epoxy override was replayed into the polish template")
    assert ran["templateSwitch"]["keys"] == ["epoxy:Direct", "polish:Direct"]
    assert ran["templateSwitch"]["polishItems"] == [
        {"id": 115, "text": "Polish scope:  Diamond grind."}]
    assert ran["templateSwitch"]["backToEpoxy"]["runs"] == BOLD_ROUND_TRIP, (
        "the epoxy formatting was gone after the round trip through polish")
    assert ran["templateSwitch"]["epoxyStillStored"][0]["runs"] == BOLD_ROUND_TRIP


@pytest.mark.parametrize("case,text", [("legacy", "Scope:  legacy text"),
                                       ("legacyFlat", "Scope:  older still")])
def test_a_draft_saved_before_this_feature_restores_exactly_as_it_does_today(ran, case, text):
    """Entries with NO `runs` key at all — both the keyed store and the older single slot. The
    text comes back and the block goes dirty, and nothing gains formatting it never had: no
    `tw-fmt`, no style spans, and the collector sends the same `{id, text}` it always did. A
    draft in flight right now looks exactly like this."""
    got = ran[case]
    assert got["restored"]["text"] == text
    assert got["restored"]["dirty"] is True
    assert got["restored"]["fmt"] is False, "a legacy override gained a formatting flag"
    assert got["restored"]["html"] == text, "a legacy override gained style spans"
    assert got["collected"] == [{"id": 115, "text": text}]


def test_a_legacy_draft_is_not_rewritten_into_the_richer_shape(ran):
    """Its stored entry must stay byte-identical: nothing to gain, and a payload change on a
    draft nobody touched is a change to a document nobody asked to change."""
    assert ran["legacy"]["stored"] == [{"id": 115, "text": "Scope:  legacy text"}]


# ══ and now the document itself ══════════════════════════════════════════════════════════════
def _scope_block_id():
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    for idx, _kind, _p, in_block, text, _txbx in pw.iter_editable_blocks(d):
        if in_block is None and text.startswith("Scope:"):
            return idx
    pytest.fail("the Scope: row is gone from the Direct epoxy template")


def _generated_scope_runs(runs):
    """Put `runs` through the WHOLE writer and read the Scope paragraph back out of the file.

    `fill_proposal`, not `_apply_paragraph_overrides` alone: `_normalize_work_label_formatting`
    re-bolds WORK rows up to their first colon AFTER overrides are applied, and
    `_shrink_overflowing_text_boxes` rewrites sizes after that. Both have clobbered this
    paragraph before, so the assertion has to be made on the finished document."""
    blob = pw.fill_proposal(
        work_type="epoxy", audience="Direct",
        values={"job_name": "Cedar Ridge Distribution Center", "scope_notes": "SCOPE"},
        paragraph_overrides=[{"id": _scope_block_id(), "text": "ignored when runs are present",
                              "runs": runs}],
    )
    out = docx.Document(io.BytesIO(blob))
    for _idx, _kind, p_elem, _in_block, text, _txbx in pw.iter_editable_blocks(out):
        if text.startswith("Scope:"):
            return [(r.text, r.bold, r.italic,
                     r.font.size.pt if r.font.size else None) for r in Paragraph(p_elem, out).runs]
    pytest.fail("the overridden paragraph is not in the generated document")


def test_the_bold_that_survived_the_round_trip_reaches_the_real_docx(ran):
    """The end of the chain, and the only assertion that is about the customer's file: take the
    runs the draft holds AFTER a reload and a re-save, generate the document, and find the bold
    on the phrase the estimator highlighted."""
    assert ran["docxRuns"] == BOLD_ROUND_TRIP, "the harness handed over the wrong runs"
    got = _generated_scope_runs(ran["docxRuns"])
    texts = [r[0] for r in got]
    assert "".join(texts) == "Scope:  Grind and coat.", (
        "the document text is not what the editor sent: %r" % (got,))
    by_text = {r[0]: r for r in got}
    assert by_text["Grind"][1] is True, (
        "the bold phrase is not bold in the generated document: %r" % (got,))


def test_the_bold_did_NOT_leak_into_the_neighbouring_runs(ran):
    """The failure that would look like success. A writer that collapsed the paragraph, or a
    restore that widened the range, bolds the words around the phrase too — and a proposal with
    a random bold sentence in it is worse than one with none, because nobody can tell it was a
    bug."""
    got = _generated_scope_runs(ran["docxRuns"])
    by_text = {r[0]: r for r in got}
    assert by_text[" and coat."][1] is False, (
        "the run after the bold phrase was bolded too: %r" % (got,))
    assert by_text["  "][1] is False, (
        "the run before the bold phrase was bolded too: %r" % (got,))
    # And every run keeps the template's own 8pt, which is what "one to one" means here.
    assert {r[3] for r in got} == {8.0}, (
        "the writer changed a size nobody asked it to change: %r" % (got,))


def test_a_size_the_estimator_chose_reaches_the_docx_on_one_run_only(ran):
    """The size half, through the writer. 14pt on the phrase, 8pt either side — measured on the
    finished file, because the overflow shrink runs last and has turned a chosen size into
    4.5pt before now."""
    got = _generated_scope_runs(ran["sizeAndItalic"]["stored"][0]["runs"])
    by_text = {r[0]: r for r in got}
    assert by_text["Grind"][3] == 14.0, "the chosen size did not reach the document: %r" % (got,)
    assert by_text[" and coat."][3] == 8.0, "the size leaked past the selection: %r" % (got,)
    assert by_text["Scope:"][2] is True, "the italic label lost its italic: %r" % (got,)
    assert by_text["Grind"][2] is not True, "the italic leaked out of the label: %r" % (got,)


def test_the_tagged_token_run_does_not_add_a_raw_token_to_the_document(ran):
    """The token tag rides along in the stored entry and the writer ignores it. What it must NOT
    have done is put a literal `{{epoxy_sf}}` in front of a customer — the run's stored text is
    still the resolved value, which is why this fix cannot regress the no-raw-token invariant."""
    runs = ran["tokenFresh"]["stored"]["runs"]
    assert any(r.get("tok") for r in runs), "the fixture has no tagged run to prove anything about"
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    idx = None
    for i, _kind, _p, in_block, text, _txbx in pw.iter_editable_blocks(d):
        if in_block is None and text.startswith("Scope:"):
            idx = i
            break
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": runs}])
    for _i, _k, p_elem, _ib, text, _t in pw.iter_editable_blocks(d):
        if "6,000" in (text or ""):
            assert "{{" not in text, "a raw token reached the paragraph: %r" % (text,)
            assert [r.text for r in Paragraph(p_elem, d).runs] == \
                ["Area:", " ", "6,000", " SF"], (
                "the tagged run was not written as its own run: %r"
                % ([r.text for r in Paragraph(p_elem, d).runs],))
            return
    pytest.fail("the override did not reach the document")


def test_bad_runs_from_a_stale_draft_still_cannot_500_the_generator(ran):
    """House rule for every sanitiser here. The stored shape grew a key, so the parade of
    malformed entries gets walked again with it."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    for bad in ({"id": 0, "runs": [{"text": "x", "tok": 5}]},
                {"id": 0, "runs": [{"text": "x", "tok": None}]},
                {"id": 0, "runs": [{"tok": "epoxy_sf"}]},
                {"id": 0, "runs": [{"text": "x", "tok": {"nope": 1}}]}):
        pw._apply_paragraph_overrides(d, [bad])        # must not raise


def test_the_writer_ignores_the_token_tag_entirely(ran):
    """`tok` is frontend bookkeeping. A run carrying it must produce the same `<w:r>` as the
    same run without it, or the extra key would be changing customer documents."""
    def sizes_and_bolds(runs):
        d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
        idx = None
        for i, _k, _p, in_block, text, _t in pw.iter_editable_blocks(d):
            if in_block is None and text.startswith("Scope:"):
                idx = i
                break
        pw._apply_paragraph_overrides(d, [{"id": idx, "runs": runs}])
        for _i, _k, p_elem, _ib, text, _t in pw.iter_editable_blocks(d):
            if text.startswith("Area:"):
                return [(r.text, r.bold, r.italic, r.underline,
                         r.font.size.pt if r.font.size else None)
                        for r in Paragraph(p_elem, d).runs]
        pytest.fail("the override did not land")

    tagged = ran["tokenFresh"]["stored"]["runs"]
    stripped = [{k: v for k, v in r.items() if k != "tok"} for r in tagged]
    assert sizes_and_bolds(tagged) == sizes_and_bolds(stripped)


# ══ the shape of the stored entry ════════════════════════════════════════════════════════════
def test_the_plain_edit_payload_did_not_grow(ran):
    """Most edits are plain, and the 500-override cap and the draft blob were sized for the
    small shape. A paragraph nobody formatted still ships `{id, text}` with no runs at all."""
    plain = [o for o in ran["guard"]["collected"] if o["id"] == 116]
    assert plain == [{"id": 116, "text": "Area: 5,200 SF and cove"}]
    assert ran["legacy"]["collected"] == [{"id": 115, "text": "Scope:  legacy text"}]


def test_the_token_tag_only_appears_on_runs_that_carry_a_token(ran):
    """It is not a per-run flag on everything: three of the four runs of the Area row have no
    token, and tagging them would make the restore rewrite ordinary words."""
    runs = ran["tokenFresh"]["savedRuns"]
    assert [bool(r.get("tok")) for r in runs] == [False, False, True, False]
