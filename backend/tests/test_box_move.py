"""MOVING a proposal text box, and whether the customer's PDF honours it.

Hanz, 2026-08-13: "Allow me to drag and resize the text box for the proposal please."

Resizing already worked (test_box_resize.py). Nothing had ever written a POSITION: `_pos_of_anchor`
read `wp:positionH`/`wp:positionV` for the editor and no code path wrote them back.

TWO PLACES AGAIN, and a different decoy. A box's position lives in the DrawingML anchor's
posOffset pair AND in the `mc:Fallback` VML twin's @style — but as `margin-left`/`margin-top`, not
`left`/`top`. `Invoice_Deposit.docx` actually ships `position:absolute;left:0;text-align:left;
margin-left:-2.1pt;…`, so writing `left` would edit a declaration that means nothing here and
leave the real offset alone, while looking exactly like a successful move. Same shape of trap as
the `a:extLst/a:ext` one on the size side.

THE PARAGRAPH-RELATIVE PROBLEM. Both offsets are measured from a frame, not from the page corner:
every box in all nine templates is `positionH relativeFrom="column"` and `positionV
relativeFrom="paragraph"`. The vertical one cannot be resolved from the XML at all — the rendered
height of the empty anchor paragraphs is a layout result — so `_pos_of_anchor` estimates it as
`margin.top + paragraph_index * _ANCHOR_LINE_H_PT + posOffset`.

The writer never tries to invert that. It reads where `_pos_of_anchor` thinks the box is NOW and
shifts the offset by `wanted - current`, so the estimate appears on both sides and cancels exactly.
Two consequences worth stating, both pinned below: a move is exact even where the estimate is not,
and applying the same override twice lands in the same place.

The alternative — rewriting `relativeFrom` to "page" and writing an absolute offset — was
rejected on purpose. A page-relative box stops travelling with the text above it, and Kyle's page
design is a baked-in full-page PNG, so a box that stops tracking its own frame art is not a
cosmetic difference.

MEASURED, in a container with LibreOffice (the dev box has none), on box 3 of the GC Resinous
template, asking for (120, 200) from a design (161.8, 153.2):

    which half is written    where MARK00 lands in the PDF
    neither                  (161.9, 152.6)
    DrawingML only           (120.1, 199.4)
    VML fallback only        (161.9, 152.6)      <- moves nothing
    both                     (120.1, 199.4)

So LibreOffice takes the DrawingML branch, lands within ~1pt of the request, and would ignore a
move written only into the fallback. Both are still written every time: which branch a renderer
believes is version-dependent, and a half-written move gives Word and the customer's PDF two
different documents.
"""
import io
import pathlib
import shutil

import docx
import pytest

import pdf_writer
import proposal_writer as pw

TEMPLATES = pathlib.Path(__file__).resolve().parents[1] / "templates"
GC_RESINOUS = TEMPLATES / "GC" / "xx TREADWELL RESINOUS PROPOSAL - xx.docx"

BOX = 3                      # the WORK box: big, framed, and the one that overflows in practice


def _open():
    return docx.Document(str(GC_RESINOUS))


def _geo(d, i=BOX):
    return pw.template_geometry(d)["boxes"][i]


def _box(d, i=BOX):
    return list(pw._iter_txbx(d))[i]


def _vml_pos(txbx):
    """(margin-left, margin-top) in points from the VML twin, or (None, None)."""
    got = {}
    for shape in pw._txbx_vml_twins(txbx):
        for part in (shape.get("style") or "").split(";"):
            k, _, v = part.partition(":")
            if k.strip().lower() in ("margin-left", "margin-top"):
                got[k.strip().lower()] = pw._vml_len_pt(v)
        break
    return got.get("margin-left"), got.get("margin-top")


# ── the templates really are anchored the awkward way ─────────────────────
@pytest.mark.parametrize("path", sorted(TEMPLATES.rglob("*.docx")))
def test_every_box_is_offset_anchored_so_a_move_has_something_to_shift(path):
    """If a template ever positioned a box by keyword (`wp:align`) instead, the writer would be
    converting the anchor rather than nudging it — worth knowing, not worth guessing at."""
    d = docx.Document(str(path))
    for i, txbx in enumerate(pw._iter_txbx(d)):
        anchor = pw._txbx_anchor(txbx)
        if anchor is None:
            continue
        for tag in ("positionH", "positionV"):
            p = anchor.find(pw.qn("wp:" + tag))
            assert p is not None, "box %d in %s has no %s" % (i, path.name, tag)
            assert p.find(pw.qn("wp:posOffset")) is not None, (
                "box %d in %s positions by keyword, not by offset" % (i, path.name))


def test_the_vertical_anchor_really_is_paragraph_relative():
    """The whole reason the writer works in differences. If this ever became page-relative the
    estimate would stop being an estimate and this file's approach could be simplified."""
    anchor = pw._txbx_anchor(_box(_open()))
    assert pw._anchor_offset(anchor, "positionV")[1] == "paragraph"
    assert pw._anchor_offset(anchor, "positionH")[1] == "column"


def test_no_template_ships_a_box_off_the_paper():
    """The move bound is the sheet. If a design box already hung off it, that bound would refuse
    to move that box at all — a fix that breaks the feature for one template."""
    for path in sorted(TEMPLATES.rglob("*.docx")):
        d = docx.Document(str(path))
        page_w, page_h = pw.page_size(d)
        for b in pw.template_geometry(d)["boxes"]:
            if b["x_pt"] is None:
                continue
            assert 0 <= b["x_pt"] and b["x_pt"] + b["w_pt"] <= page_w + 0.5, (
                "box %d in %s runs off the paper horizontally" % (b["id"], path.name))
            assert 0 <= b["y_pt"] and b["y_pt"] + b["h_pt"] <= page_h + 0.5, (
                "box %d in %s runs off the paper vertically" % (b["id"], path.name))


def test_every_template_puts_boxes_outside_the_printable_area():
    """Which is why the move bound is the sheet and not the printable area. Stated out loud
    because "keep it inside the margins" is the obvious rule and it would refuse every drag."""
    for path in sorted(TEMPLATES.rglob("*.docx")):
        d = docx.Document(str(path))
        page = pw._page_metrics(d)
        m = page["margin"]
        outside = [b["id"] for b in pw.template_geometry(d)["boxes"]
                   if b["x_pt"] is not None and (
                       b["x_pt"] < m["left"] or b["y_pt"] < m["top"]
                       or b["x_pt"] + b["w_pt"] > page["w_pt"] - m["right"]
                       or b["y_pt"] + b["h_pt"] > page["h_pt"] - m["bottom"])]
        assert outside, "%s keeps every box inside the margins — re-read the bound" % path.name


# ── the move lands exactly where it was asked to ──────────────────────────
def test_a_move_puts_the_box_where_the_editor_said():
    """Round-tripped through the SAME reader the editor was given its numbers from, which is what
    makes the paragraph-relative estimate cancel."""
    d = _open()
    before = _geo(d)
    assert pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0, "y_pt": 200.0}}) == 1
    after = _geo(d)
    assert (after["x_pt"], after["y_pt"]) == (pytest.approx(120.0), pytest.approx(200.0))
    assert (after["w_pt"], after["h_pt"]) == (before["w_pt"], before["h_pt"]), (
        "a move resized the box")


def test_a_move_is_exact_however_wrong_the_line_height_estimate_is():
    """The load-bearing property of doing this as a difference. Recalibrate the anchor line height
    — or get it wrong in the first place — and a drag of 40pt still moves the box 40pt.

    If this ever fails, the writer has started trusting the absolute estimate."""
    design_offset = pw._anchor_offset(pw._txbx_anchor(_box(_open())), "positionV")[0]

    def moved_offset(line_h):
        """Drag this box 40pt UP with the calibration set to `line_h`, and report the offset the
        writer wrote. Up rather than down: box 3 anchors at paragraph 5, so a big line height puts
        its estimated y near the bottom of the sheet and a downward drag would be refused by the
        paper bound — which would make this pass for the wrong reason."""
        d = _open()
        old = pw._ANCHOR_LINE_H_PT
        pw._ANCHOR_LINE_H_PT = line_h
        try:
            start = _geo(d)["y_pt"]
            assert pw._apply_box_overrides(d, {str(BOX): {"y_pt": start - 40.0}}) == 1
            return pw._anchor_offset(pw._txbx_anchor(_box(d)), "positionV")[0]
        finally:
            pw._ANCHOR_LINE_H_PT = old

    for line_h in (14.0, 20.0, 99.0):
        assert moved_offset(line_h) == pytest.approx(design_offset - 40.0, abs=0.01), (
            "at %gpt per anchor line the writer moved the box by something other than the 40pt it "
            "was dragged — it has started trusting the absolute estimate" % line_h)


def test_applying_the_same_move_twice_lands_in_the_same_place():
    """It is an absolute position, not an accumulating nudge. /api/generate is re-POSTed with the
    identical payload whenever a download token has expired (done.js retries on a 404)."""
    d = _open()
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0, "y_pt": 200.0}})
    once = _geo(d)
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0, "y_pt": 200.0}})
    assert _geo(d) == once


def test_one_axis_at_a_time_leaves_the_other_alone():
    """A width-only drag sends only w_pt; a sideways drag only x_pt. Filling in the missing axis
    with a default would jump the box on the other one."""
    d = _open()
    before = _geo(d)
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 100.0}})
    after = _geo(d)
    assert after["x_pt"] == pytest.approx(100.0)
    assert after["y_pt"] == pytest.approx(before["y_pt"])


def test_a_move_and_a_resize_in_one_request_both_land():
    d = _open()
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 100.0, "y_pt": 300.0,
                                           "w_pt": 300.0, "h_pt": 250.0}})
    got = _geo(d)
    assert (got["x_pt"], got["y_pt"], got["w_pt"], got["h_pt"]) == (
        pytest.approx(100.0), pytest.approx(300.0),
        pytest.approx(300.0), pytest.approx(250.0))


# ── both recorded places, and the decoy ───────────────────────────────────
def test_the_move_is_written_to_the_drawingml_and_the_vml_twin():
    """LibreOffice reads the DrawingML branch; other renderers read the fallback. Writing one
    leaves the other claiming the old position, which is Word and the PDF disagreeing about a box
    the estimator can see on screen."""
    d = _open()
    txbx = _box(d)
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0, "y_pt": 200.0}})
    anchor = pw._txbx_anchor(txbx)
    ox = pw._anchor_offset(anchor, "positionH")[0]
    oy = pw._anchor_offset(anchor, "positionV")[0]
    vml_x, vml_y = _vml_pos(txbx)
    assert vml_x == pytest.approx(ox, abs=0.02), "the VML twin still claims the old left offset"
    assert vml_y == pytest.approx(oy, abs=0.02), "the VML twin still claims the old top offset"


def test_the_vml_twin_gets_the_offset_not_the_page_coordinate():
    """Its `mso-position-*-relative:text` means it is measured from the same frame the DrawingML
    anchor is. Writing 120 there rather than the offset would put the box 120pt from the anchor
    paragraph, not from the page edge."""
    d = _open()
    txbx = _box(d)
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0}})
    vml_x, _ = _vml_pos(txbx)
    assert vml_x != pytest.approx(120.0), "the fallback was given a page coordinate"
    assert vml_x == pytest.approx(120.0 - pw._page_metrics(d)["margin"]["left"], abs=0.02)


def test_a_bare_left_declaration_is_never_mistaken_for_the_offset():
    """`Invoice_Deposit.docx` ships `position:absolute;left:0;text-align:left;margin-left:…`.
    Writing `left` would edit a decoy, leave the real offset, and look like it worked."""
    src = TEMPLATES / "Invoice_Deposit.docx"
    d = docx.Document(str(src))
    txbx = list(pw._iter_txbx(d))[0]
    shape = pw._txbx_vml_twins(txbx)[0]
    before = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    assert before.get("left") == "0", "the fixture no longer carries the decoy; pick another box"
    pw._move_txbx(txbx, x_off_pt=42.0, y_off_pt=-7.0)
    after = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    assert after["left"] == "0", "the bare `left` declaration was rewritten"
    assert after["text-align"] == "left"
    assert pw._vml_len_pt(after["margin-left"]) == pytest.approx(42.0)
    assert pw._vml_len_pt(after["margin-top"]) == pytest.approx(-7.0)


def test_the_vml_style_keeps_everything_else():
    """The style also carries size, z-index and visibility. Losing those is worse than the wrong
    position."""
    d = _open()
    txbx = _box(d)
    shape = pw._txbx_vml_twins(txbx)[0]
    before = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    pw._move_txbx(txbx, x_off_pt=10.0, y_off_pt=10.0)
    after = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    for key, val in before.items():
        if key.strip().lower() in ("margin-left", "margin-top"):
            continue
        assert after.get(key) == val, "the VML style lost or changed %r" % key
    assert after["width"] == before["width"] and after["height"] == before["height"]


def test_a_box_with_no_anchor_position_is_left_alone_in_BOTH_places():
    """Both or neither. A drawing whose anchor carries no position — a `wp:inline`, which no
    template of Kyle's uses but a future one might — has nothing for the fallback to agree WITH,
    so writing the fallback alone would produce exactly the disagreement between Word and the
    customer's PDF that writing two places exists to prevent.

    Built by stripping the position off a real box, because no fixture has one."""
    d = _open()
    txbx = _box(d)
    anchor = pw._txbx_anchor(txbx)
    for tag in ("positionH", "positionV"):
        anchor.remove(anchor.find(pw.qn("wp:" + tag)))
    before = pw._txbx_vml_twins(txbx)[0].get("style")
    assert pw._move_txbx(txbx, x_off_pt=10.0, y_off_pt=10.0) == 0
    assert pw._txbx_vml_twins(txbx)[0].get("style") == before, (
        "the VML fallback was moved on its own, with no DrawingML position to match it")


def test_the_page_is_measured_from_the_document_rather_than_assumed():
    """Every template Kyle ships is US Letter, which makes an assumed 612 x 792 look correct
    forever. A tabloid template would then be bounded to a Letter sheet and refuse two thirds of
    its own page.

    The section is rewritten in memory only — `docx.Document` was opened from bytes on disk and
    the template file is never saved."""
    from docx.shared import Pt as _Pt

    d = _open()
    d.sections[0].page_width = _Pt(1224)      # 17in, tabloid landscape
    d.sections[0].page_height = _Pt(792)
    assert pw.page_size(d) == pytest.approx((1224.0, 792.0))
    page = pw._page_metrics(d)
    assert page["w_pt"] == pytest.approx(1224.0)
    # And the consequence: a box may now be dragged past where a Letter sheet would have stopped.
    b = _geo(d)
    assert pw._apply_box_overrides(d, {str(BOX): {"x_pt": 700.0}}) == 1, (
        "a position 700pt across a 1224pt sheet was refused, so the bound is a Letter constant")
    assert _geo(d)["x_pt"] == pytest.approx(700.0)
    assert b["w_pt"] + 700.0 <= 1224.0, "the fixture's premise no longer holds"


def test_a_move_does_not_disturb_the_size_sites():
    d = _open()
    txbx = _box(d)
    anchor = pw._txbx_anchor(txbx)
    ext = anchor.find(pw.qn("wp:extent"))
    before = (ext.get("cx"), ext.get("cy"))
    pw._apply_box_overrides(d, {str(BOX): {"x_pt": 120.0, "y_pt": 200.0}})
    assert (ext.get("cx"), ext.get("cy")) == before


# ── which SECTION a floating object is measured from ──────────────────────
# A `relativeFrom="column"` offset is measured from the margin of the section the anchor sits
# in. `_pos_of_anchor` used to take that margin from `d.sections[0]`, which is only the same
# thing on a single-section document — and all nine proposal templates are single-section, so
# nothing could see it — those templates' reported box positions (30-165pt across) were right
# all along and are not what this section is about.
#
# The LETTERHEAD IMAGES are the ones that flattered the old sum: each carries a posOffset of
# exactly minus its own margin, so x came out 0, the right answer from the wrong arithmetic.
# The Cover Letter templates inherit the Treadwell letterhead's two sections (171pt, then
# 49.5pt with the artwork and every body paragraph in it), where minus-49.5 plus section 0's
# 171 put the full-bleed letterhead 121.5pt off to the right in the editor's page view.
#
# The same first-section read reached the WORDS too, through `_page_metrics` -> the served
# `geometry.page.margin`, which the editor uses as the padding of its flowing text column.

def _end_section_at_top(d, left_pt, top_pt=None):
    """Insert a leading empty paragraph that CLOSES a section with its own page margins.

    Word puts the `sectPr` of a section in the `w:pPr` of its LAST paragraph, so this makes the
    document two-section with everything already in it living in the SECOND — exactly the shape
    the Treadwell letterhead has. In memory only; the template on disk is untouched.

    `top_pt` moves the vertical margin too, and it is not decoration: both sections of every
    real fixture here happen to share a `top` of 117pt, so a document that varies only `left`
    cannot tell a correct y from one still read off section 0. Vary the top and the y-axis
    becomes falsifiable as well."""
    import copy

    body = d.element.body
    sect = copy.deepcopy(body.find(pw.qn("w:sectPr")))
    pg = sect.find(pw.qn("w:pgMar"))
    pg.set(pw.qn("w:left"), str(int(round(left_pt * 20))))
    if top_pt is not None:
        pg.set(pw.qn("w:top"), str(int(round(top_pt * 20))))
    p = body.makeelement(pw.qn("w:p"), {})
    ppr = p.makeelement(pw.qn("w:pPr"), {})
    ppr.append(sect)
    p.append(ppr)
    body.insert(0, p)
    return d


def test_a_box_is_measured_from_its_own_sections_margin_not_the_first():
    """Red before the fix by 300pt across and 200pt down: the box was reported most of a sheet
    from where Word puts it, so the editor drew it there and a drag from that phantom position
    wrote the offset back just as far out.

    The y arm carries `+ _ANCHOR_LINE_H_PT` because the fixture inserts a paragraph, which is
    exactly one more empty line above the anchor — the estimate `_pos_of_anchor` documents.
    Anything past that one step is section 0's top margin leaking in."""
    d = _open()
    m0 = pw._page_metrics(d)["margin"]
    before = _geo(d)
    _end_section_at_top(d, left_pt=m0["left"] + 300.0, top_pt=m0["top"] + 200.0)
    assert len(d.sections) == 2, "the fixture did not actually create a second section"
    assert d.sections[0].left_margin.pt == pytest.approx(m0["left"] + 300.0), "fixture premise"
    assert d.sections[0].top_margin.pt == pytest.approx(m0["top"] + 200.0), "fixture premise"
    after = _geo(d)
    assert after["x_pt"] == pytest.approx(before["x_pt"]), (
        "a section the box is not in moved the box sideways")
    assert after["y_pt"] == pytest.approx(before["y_pt"] + pw._ANCHOR_LINE_H_PT), (
        "a section the box is not in moved the box down the page")


def _top_level_of(body, el):
    """The `w:body` child that contains `el` — the unit a section break divides."""
    while el is not None and el.getparent() is not body:
        el = el.getparent()
    return el


def _end_section_after(d, after_index, left_pt, top_pt):
    """Close a section immediately AFTER one top-level child, stranding it in a section of its
    own with margins nothing else in the document shares. In memory only.

    The counterpart to `_end_section_at_top`, and the only shape that can falsify
    `_pos_of_anchor`'s own section lookup. Once `_page_metrics` also resolves the section (it
    must — it feeds the editor's text column), the `page` dict handed to `_pos_of_anchor`
    ALREADY carries the right margin for any anchor sitting with the bulk of the text, so a
    fixture where the box and the words share a section cannot tell the per-anchor lookup from
    no lookup at all. Both of the real two-section templates are that shape."""
    import copy

    body = d.element.body
    sect = copy.deepcopy(body.find(pw.qn("w:sectPr")))
    pg = sect.find(pw.qn("w:pgMar"))
    pg.set(pw.qn("w:left"), str(int(round(left_pt * 20))))
    pg.set(pw.qn("w:top"), str(int(round(top_pt * 20))))
    p = body.makeelement(pw.qn("w:p"), {})
    ppr = p.makeelement(pw.qn("w:pPr"), {})
    ppr.append(sect)
    p.append(ppr)
    body.insert(after_index + 1, p)
    return d


def test_a_box_follows_its_own_section_even_when_the_text_is_in_another():
    """The per-anchor half of the fix, on BOTH axes.

    Box 3 is stranded in a leading section with margins 300pt across and 200pt down from the
    one the other 39 text-bearing paragraphs use, so the box must be reported to have moved by
    exactly that and the editor's column must not have moved at all. A `_pos_of_anchor` reading
    the served page margin instead scores nothing on either axis."""
    d = _open()
    body = d.element.body
    before, was = _geo(d), pw._page_metrics(d)["margin"]
    anchored_in = _top_level_of(body, pw._txbx_anchor(_box(d)))
    _end_section_after(d, list(body).index(anchored_in),
                       left_pt=was["left"] + 300.0, top_pt=was["top"] + 200.0)

    assert len(d.sections) == 2, "the fixture did not create a second section"
    assert pw._section_ordinal(body, anchored_in) == 0, "the box is not in the stranded section"
    assert pw._body_section_ordinal(d) == 1, "the words did not stay in the other section"
    assert pw._page_metrics(d)["margin"] == pytest.approx(was), (
        "one stranded paragraph dragged the editor's whole text column with it")

    after = _geo(d)
    assert after["x_pt"] == pytest.approx(before["x_pt"] + 300.0), (
        "the box was measured from the text's margin, not from its own section's")
    assert after["y_pt"] == pytest.approx(before["y_pt"] + 200.0), (
        "the box's vertical was measured from the text's top margin, not its own section's")


def test_the_page_the_editor_is_told_about_is_the_section_the_words_are_in():
    """`geometry.page.margin` is not decoration — `proposal-review.js` uses it verbatim as the
    padding of the one flowing text column it lays the body out in, and `max_box` bounds a
    resize. Read off section 0, the letterhead served a 396pt column indented 171pt when the
    letter's own paragraphs live in a 522pt column at 49.5pt: every line of the customer's
    cover letter rendered in a squeezed, over-indented ribbon.

    The same bug as the artwork, moved from the picture to the words."""
    d = _end_section_at_top(_open(), left_pt=390.0, top_pt=300.0)
    was = pw._page_metrics(_open())["margin"]
    page = pw._page_metrics(d)
    assert page["margin"]["left"] == pytest.approx(was["left"]), (
        "the editor was handed the margin of the section holding one empty paragraph")
    assert page["margin"]["top"] == pytest.approx(was["top"])
    # …and the two bounds derived from the same page setup still agree with it.
    assert page["max_box"]["w_pt"] == pytest.approx(
        page["w_pt"] - page["margin"]["left"] - page["margin"]["right"])
    assert pw.box_size_limits(d) == pytest.approx(
        (page["max_box"]["w_pt"], page["max_box"]["h_pt"]))


def test_the_sheet_is_the_one_the_body_text_prints_on():
    """`page_size` bounds a drag, and Word lets each section have its own paper — a title page
    on tabloid and a body on Letter is an ordinary thing to do. Every fixture in this repo has
    uniform paper, so this is the only place the section walk in `page_size` is observable at
    all: without it a leading tabloid title section would let the estimator drag a box 600pt
    across a Letter body page, and the customer's PDF would clip whatever went past 612."""
    from docx.shared import Pt as _Pt

    d = _end_section_at_top(_open(), left_pt=390.0, top_pt=300.0)
    d.sections[0].page_width = _Pt(1224)          # 17in title section…
    assert pw._body_section_ordinal(d) == 1, "the words must be in the OTHER section"
    assert pw.page_size(d) == pytest.approx((612.0, 792.0)), (
        "the drag was bounded by a sheet the body text is not printed on")
    assert pw._page_metrics(d)["w_pt"] == pytest.approx(612.0)
    # …and the bound really refuses the position that only the wrong sheet would allow.
    assert pw._apply_box_overrides(d, {str(BOX): {"x_pt": 700.0}}) == 0


def test_the_paragraph_carrying_a_break_belongs_to_the_section_it_closes():
    """The half of Word's rule that is easy to get backwards. The `sectPr` sits in the LAST
    paragraph of its section, not the first of the next one — so resolving 'the next break at or
    after me' has to be inclusive. Off by one here and every anchor lands a section too late."""
    d = _end_section_at_top(_open(), left_pt=390.0, top_pt=300.0)
    body = d.element.body
    closer, following = list(body)[0], list(body)[1]
    assert closer.find(pw.qn("w:pPr") + "/" + pw.qn("w:sectPr")) is not None, "fixture premise"
    assert pw._section_ordinal(body, closer) == 0
    assert pw._section_ordinal(body, following) == 1
    assert pw._margins_of(d, pw._section_ordinal(body, closer))["left"] == pytest.approx(390.0)
    assert pw._margins_of(d, pw._section_ordinal(body, following))["left"] == pytest.approx(
        pw._page_metrics(_open())["margin"]["left"])


def _letterhead_pgMar(d):
    """The `w:pgMar` of the letterhead's BODY section — the one the words are in.

    Deliberately not one of the proposal templates: every one of those is 90/90/72/72, which is
    exactly `_DEFAULT_MARGIN_PT`, so a margin read that collapsed entirely to the fallback would
    return the right numbers anyway and neither test below could fail. The letterhead's
    49.5/40.5/117/99 differs on all four sides."""
    pg = d.element.body.find(pw.qn("w:sectPr")).find(pw.qn("w:pgMar"))
    assert pw._margins_of(d, 1) != pw._DEFAULT_MARGIN_PT, "fixture cannot distinguish a fallback"
    return pg


def test_a_universal_measure_margin_is_read_rather_than_dropped():
    """`w:pgMar/@w:left` is an ST_TwipsMeasure: "990" and "0.6875in" are the same margin, and a
    hand-rolled int() of the second drops the page setup back to the Letter guess — which on the
    letterhead would re-indent the whole letter by 40pt. python-docx's own Length conversion
    reads both."""
    d = docx.Document(str(COVER_LETTERS[0]))
    want = d.sections[1].left_margin.pt
    _letterhead_pgMar(d).set(pw.qn("w:left"), "%sin" % (want / 72.0))
    assert pw._margins_of(d, 1)["left"] == pytest.approx(want)


def test_a_single_unreadable_margin_costs_only_that_side():
    """The old hand-rolled read took all four sides in one try/except, so one bad attribute
    silently relaid the whole page instead of just its own edge."""
    d = docx.Document(str(COVER_LETTERS[0]))
    kept = {s: d.sections[1].__getattribute__(s + "_margin").pt
            for s in ("right", "top", "bottom")}
    _letterhead_pgMar(d).set(pw.qn("w:left"), "banana")
    got = pw._margins_of(d, 1)
    assert got["left"] == pw._DEFAULT_MARGIN_PT["left"], "a bad side should take the fallback"
    for side, want in kept.items():
        assert got[side] == pytest.approx(want), (
            "a bad `left` took `%s` down with it" % side)


COVER_LETTERS = sorted((TEMPLATES / "CoverLetter").rglob("*.docx"))
# Without this, a tree whose Cover Letter templates land in a different commit than this fix
# collapses the parametrization to zero cases, and pytest reports that as SKIP rather than as
# failure — the whole two-section guard would quietly stop guarding anything.
assert COVER_LETTERS, (
    "no Cover Letter templates under %s: the two-section regression tests below would cover "
    "nothing. Run backend/prepare_cover_letter_templates.py." % (TEMPLATES / "CoverLetter"))


@pytest.mark.parametrize("path", COVER_LETTERS)
def test_the_letterhead_artwork_sits_flush_with_the_left_edge_of_the_sheet(path):
    """The real two-section fixture, and the bug as the frontend agent hit it. The letterhead is
    full-bleed art — bison top-right, red bar along the bottom — anchored at posOffset -49.5pt
    against its own section's 49.5pt margin, so it starts at x=0. Reported at 121.5pt it renders
    off the page in the editor preview."""
    d = docx.Document(str(path))
    geo = pw.template_geometry(d)
    assert len(d.sections) == 2, "%s stopped being the two-section case this guards" % path.name
    assert geo["images"], "%s lost its letterhead artwork" % path.name
    for img in geo["images"]:
        assert img["x_pt"] == pytest.approx(0.0, abs=0.5), (
            "%s reports the letterhead at x=%.2fpt" % (path.name, img["x_pt"]))
    page_w, _page_h = pw.page_size(d)
    for b in geo["boxes"]:
        assert 0.0 <= b["x_pt"] and b["x_pt"] + b["w_pt"] <= page_w, (
            "%s reports box %d off the sheet at x=%.2f" % (path.name, b["id"], b["x_pt"]))


@pytest.mark.parametrize("path", COVER_LETTERS)
def test_the_letter_is_laid_out_in_the_section_its_paragraphs_live_in(path):
    """Section 0 of the letterhead holds ONE empty paragraph; every body paragraph, and the
    artwork, are in section 1. So the editor's column must be section 1's 522pt at 49.5pt."""
    d = docx.Document(str(path))
    words = pw._body_section_ordinal(d)
    assert words != 0, "%s: the words are in section 0 after all" % path.name
    margin = pw._page_metrics(d)["margin"]
    assert margin["left"] == pytest.approx(d.sections[words].left_margin.pt)
    assert margin["left"] == pytest.approx(49.5), "%s: not the letterhead setup" % path.name
    page_w, _h = pw.page_size(d)
    assert page_w - margin["left"] - margin["right"] == pytest.approx(522.0)


@pytest.mark.parametrize("path", sorted(TEMPLATES.rglob("*.docx")))
def test_a_single_section_template_reads_exactly_as_it_did_before(path):
    """Requirement one of the fix: the templates already in production must report exactly what
    they reported when these four functions said `d.sections[0]` and nothing else. Pinned
    against that literal expression rather than against a copied constant, so the check cannot
    drift along with the code it is checking."""
    d = docx.Document(str(path))
    if len(d.sections) != 1:
        pytest.skip("%s is multi-section; covered above" % path.name)
    sec = d.sections[0]
    assert pw._body_section_ordinal(d) == 0
    for child in d.element.body:
        assert pw._section_ordinal(d.element.body, child) == 0
    assert pw._page_metrics(d)["margin"] == pytest.approx({
        "top": sec.top_margin.pt, "left": sec.left_margin.pt,
        "right": sec.right_margin.pt, "bottom": sec.bottom_margin.pt})
    assert pw.page_size(d) == pytest.approx((sec.page_width.pt, sec.page_height.pt))
    assert pw.box_size_limits(d) == pytest.approx((
        sec.page_width.pt - sec.left_margin.pt - sec.right_margin.pt,
        sec.page_height.pt - sec.top_margin.pt - sec.bottom_margin.pt))


# ── the bounds ────────────────────────────────────────────────────────────
@pytest.mark.parametrize("spec", [
    {"x_pt": 600.0},          # corner on the paper, box hanging off the right
    {"y_pt": 780.0},          # …and off the bottom
    {"x_pt": -5.0},           # corner off the paper
    {"y_pt": -1.0},
    {"x_pt": 700.0},          # corner past the sheet entirely
    {"y_pt": 900.0},
])
def test_a_position_that_would_run_off_the_paper_is_refused(spec):
    """LibreOffice CLIPS a box that leaves the sheet rather than spilling: nothing errors, and the
    customer's proposal is quietly missing a paragraph. Refusing leaves the design position, which
    is a document that still reads correctly."""
    d = _open()
    before = _geo(d)
    assert pw._apply_box_overrides(d, {str(BOX): spec}) == 0
    assert _geo(d) == before


@pytest.mark.parametrize("which", [BOX, 1])
def test_the_furthest_legal_corner_is_accepted(which):
    """The bound must not be so cautious that it refuses the position the handle stops at.

    Box 1 is the 72 x 18pt DATE/JOB NAME box, and it is the case that separates the two bounds:
    its furthest legal corner is (540, 774), both past the 432 x 648 printable area. Bounding the
    corner by the printable area — the obvious mistake, since that IS the size bound — would
    refuse to drag a small box anywhere near the right or the bottom of the sheet.
    """
    d = _open()
    b = _geo(d, which)
    page_w, page_h = pw.page_size(d)
    spec = {"x_pt": page_w - b["w_pt"], "y_pt": page_h - b["h_pt"]}
    assert pw._apply_box_overrides(d, {str(which): spec}) == 1, (
        "the furthest legal corner (%.1f, %.1f) for a %.0f x %.0f box was refused"
        % (spec["x_pt"], spec["y_pt"], b["w_pt"], b["h_pt"]))
    got = _geo(d, which)
    assert (got["x_pt"], got["y_pt"]) == (pytest.approx(spec["x_pt"]), pytest.approx(spec["y_pt"]))


def test_a_move_plus_a_resize_is_bounded_against_the_NEW_size():
    """The pair has to be judged together, or a legal-looking move plus a legal-looking resize
    add up to a box off the paper.

    y=600 with the design 183.75pt height ends at 783.75pt, comfortably on a 792pt sheet. Made
    648pt tall in the same request it ends at 1248pt, and the bottom two thirds of the text would
    be clipped away silently.

    The RESIZE still applies, and that is not an oversight. `box_size_limits` bounds a height by
    the printable area and says out loud that it "bounds the impossible, not every overflow" —
    648pt at the design y already runs past the bottom of the sheet, which is pre-existing and
    deliberate (the render check's job, not the geometry's). What this test pins is that the new
    position bound does not INHERIT that looseness: the composite request ends up doing exactly
    what a resize-only request would have done, and no more. The handle is stricter again — it
    clamps a height to `page - current y` — so it never asks for this pair in the first place.
    """
    d = _open()
    before = _geo(d)
    assert 600.0 + before["h_pt"] <= 792.0, "the fixture's premise no longer holds"
    n = pw._apply_box_overrides(d, {str(BOX): {"y_pt": 600.0, "h_pt": 648.0}})
    after = _geo(d)
    assert after["h_pt"] == pytest.approx(648.0), "the resize should still apply"
    assert after["y_pt"] == pytest.approx(before["y_pt"]), (
        "y=600 with h=648 is 1248pt of a 792pt sheet and was accepted anyway")
    assert n == 1


@pytest.mark.parametrize("raw", [
    {"0": {"x_pt": "left"}}, {"0": {"x_pt": True}}, {"0": {"y_pt": float("nan")}},
    {"0": {"y_pt": float("inf")}}, {"0": {"x_pt": None}}, {"0": {"x_pt": -1}},
    {"0": {"y_pt": 1e9}},
    # A NUMERIC STRING is refused too, not coerced. The client sends JSON numbers; a quoted one
    # arrived from something that is not this editor, and quietly parsing whatever a stranger
    # sends is how a sanitiser stops being one. (h_pt/w_pt have always worked this way.)
    {"0": {"x_pt": "100"}}, {"0": {"y_pt": "100"}}, {"0": {"h_pt": "100"}},
    {"0": {"x_pt": [100]}}, {"0": {"y_pt": {"v": 100}}},
])
def test_malformed_positions_never_raise_and_never_apply(raw):
    """A stale draft or a hand-built request must not 500 /api/generate."""
    assert pw._sanitize_box_overrides(raw) == {}
    d = _open()
    assert pw._apply_box_overrides(d, raw) == 0


def test_zero_is_a_legal_position():
    """The top-left corner of the sheet. Easy to lose to a falsy check, and the clamp on the
    handle produces exactly this when somebody drags up and to the left."""
    assert pw._sanitize_box_overrides({"0": {"x_pt": 0, "y_pt": 0}}) == {
        "0": {"x_pt": 0.0, "y_pt": 0.0}}


def test_the_position_bound_is_the_sheet_of_the_template_it_is_given():
    d = _open()
    sec = d.sections[0]
    assert pw.page_size(d) == pytest.approx((sec.page_width.pt, sec.page_height.pt))
    assert pw.page_size(d) == pytest.approx((612.0, 792.0)), "Kyle's templates are US Letter"


def test_a_template_with_no_usable_page_falls_back_to_letter():
    class _Fake:
        sections = []
    assert pw.page_size(_Fake()) == (612.0, 792.0)


def test_the_fallback_page_and_the_fallback_size_limit_describe_one_sheet():
    """612-90-90 = 432 and 792-72-72 = 648. Two different guesses here would let the size limit
    and the margins disagree on the one template that ever reaches this branch."""
    class _Fake:
        sections = []
    page = pw._page_metrics(_Fake())
    m = page["margin"]
    assert page["w_pt"] - m["left"] - m["right"] == pytest.approx(page["max_box"]["w_pt"])
    assert page["h_pt"] - m["top"] - m["bottom"] == pytest.approx(page["max_box"]["h_pt"])


def test_an_override_for_a_box_that_does_not_exist_is_ignored():
    d = _open()
    n = len(list(pw._iter_txbx(d)))
    assert pw._apply_box_overrides(d, {str(n + 50): {"x_pt": 100.0}}) == 0


def test_a_refused_move_is_not_reported_as_a_change():
    """The count is what tells the log (and any future caller) whether the estimator's drag
    landed. `_resize_txbx` counts the sites it VISITED, so calling it with nothing to write
    reported a write that never happened."""
    d = _open()
    assert pw._apply_box_overrides(d, {str(BOX): {"x_pt": 700.0}}) == 0


# ── end to end through fill_proposal ──────────────────────────────────────
def test_fill_proposal_moves_the_box():
    blob = pw.fill_proposal(
        work_type="polish", audience="GC",
        values={"job_name": "Move Check", "city_state": "Lenexa, KS"},
        box_overrides={str(BOX): {"x_pt": 120.0, "y_pt": 220.0}},
    )
    got = _geo(docx.Document(io.BytesIO(blob)))
    assert (got["x_pt"], got["y_pt"]) == (pytest.approx(120.0), pytest.approx(220.0))


def test_generation_without_a_move_is_unchanged():
    """The feature must be inert when nobody used it."""
    common = dict(work_type="polish", audience="GC",
                  values={"job_name": "Inert Move", "city_state": "Lenexa, KS"})
    a = _geo(docx.Document(io.BytesIO(pw.fill_proposal(**common))))
    b = _geo(docx.Document(io.BytesIO(pw.fill_proposal(box_overrides={}, **common))))
    c = _geo(docx.Document(io.BytesIO(pw.fill_proposal(box_overrides=None, **common))))
    assert a == b == c


# ── does the customer's PDF honour it? ────────────────────────────────────
def _numbered_pdf(box_h_override=None, **override):
    """The real template with box 3 filled with numbered markers, optionally moved, as PDF."""
    import copy

    d = _open()
    txbx = _box(d)
    proto = next(iter(txbx.iter(pw.qn("w:p"))))
    for p in list(txbx.iter(pw.qn("w:p"))):
        p.getparent().remove(p)
    for i in range(6):
        p = copy.deepcopy(proto)
        pw._set_paragraph_text(p, "MARK%02d" % i)
        txbx.append(p)
    if override:
        assert pw._apply_box_overrides(d, {str(BOX): override}) == 1, (
            "the override was refused, so the render would prove nothing")
    buf = io.BytesIO()
    d.save(buf)
    return pdf_writer.docx_to_pdf(buf.getvalue())


def _mark_positions(pdf_bytes):
    try:
        import fitz
    except ImportError:
        pytest.skip("PyMuPDF not installed")
    got = {}
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            for x0, y0, x1, y1, word, *_ in page.get_text("words"):
                if word.startswith("MARK"):
                    got.setdefault(word, (x0, y0, x1, y1))
    return got


@pytest.mark.skipif(shutil.which("soffice") is None and shutil.which("libreoffice") is None,
                    reason="LibreOffice is not installed (it is in the Docker image)")
def test_the_pdf_really_honours_a_move():
    """THE load-bearing test for the move, for the reason its resize twin gives: LibreOffice, not
    Word, renders what the customer receives, and it ignores things Word honours (DrawingML
    autofit, most famously). A move that Word showed and the PDF dropped would be invisible to the
    estimator until a customer received the wrong document.

    Asserted as a DELTA rather than an absolute landing point, because the box's own top inset puts
    the first baseline about a point below the box edge and that is not what is under test. The
    delta is what the estimator dragged.

    Measured in the container: design (161.9, 152.6), moved (120.1, 199.4) for a requested
    (161.8, 153.2) -> (120.0, 200.0). Within 0.7pt on both axes, and every marker moves together.
    """
    d0 = _open()
    design = _geo(d0)
    want_x, want_y = 120.0, 200.0

    base = _mark_positions(_numbered_pdf())
    moved = _mark_positions(_numbered_pdf(x_pt=want_x, y_pt=want_y))
    assert base and set(base) == set(moved), (
        "the markers did not survive both renders: %s vs %s" % (sorted(base), sorted(moved)))

    dx_want = want_x - design["x_pt"]
    dy_want = want_y - design["y_pt"]
    for mark in sorted(base):
        dx = moved[mark][0] - base[mark][0]
        dy = moved[mark][1] - base[mark][1]
        assert dx == pytest.approx(dx_want, abs=1.0), (
            "%s moved %.1fpt sideways in the PDF, the drag asked for %.1f" % (mark, dx, dx_want))
        assert dy == pytest.approx(dy_want, abs=1.0), (
            "%s moved %.1fpt down in the PDF, the drag asked for %.1f" % (mark, dy, dy_want))


@pytest.mark.skipif(shutil.which("soffice") is None and shutil.which("libreoffice") is None,
                    reason="LibreOffice is not installed (it is in the Docker image)")
def test_the_pdf_keeps_a_box_moved_to_the_furthest_legal_corner_on_the_paper():
    """The promise the sheet bound makes. Measured: at (189, 608.25) the last marker's baseline
    sits 652.7pt down a 792pt sheet, and every marker is still in the text layer."""
    d0 = _open()
    design = _geo(d0)
    page_w, page_h = pw.page_size(d0)
    x, y = page_w - design["w_pt"], page_h - design["h_pt"]
    got = _mark_positions(_numbered_pdf(x_pt=x, y_pt=y))
    assert len(got) == 6, "text was clipped away at the furthest position the product allows"
    for mark, (x0, y0, x1, y1) in got.items():
        assert 0 <= x0 and x1 <= page_w, "%s runs off the side of the sheet" % mark
        assert 0 <= y0 and y1 <= page_h, "%s runs off the bottom of the sheet" % mark
