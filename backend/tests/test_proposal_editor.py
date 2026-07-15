"""Proposal Review's document editor (Google-Docs-like editing of the REAL
.docx template).

Covers:
  - GET /api/proposal-template: ordered blocks + correct `in_block` flags,
    for both the Direct Epoxy and Direct Combo templates.
  - `paragraph_overrides` landing in the generated .docx with the paragraph's
    original formatting (first-run rPr) preserved.
  - Overrides targeting a paragraph inside a repeatable/priced block region
    are rejected (that content is pricing-engine/template owned).
  - Malformed overrides (bad ids, non-dict entries, huge lists) never 500
    `/api/generate` or `proposal_writer._apply_paragraph_overrides` directly.

The endpoint's `id` is `proposal_writer.iter_editable_blocks`'s walk index —
tests look up ids by TEXT (not hardcoded magic numbers) so they stay valid if
Kyle's templates are re-annotated with different paragraph counts.
"""
import io
import re

import docx
from docx import Document
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"


def _rendered(docx_bytes):
    """Text of paragraphs Word actually renders (mc:Choice copy), excluding
    the legacy mc:Fallback duplicate — same helper used across the suite."""
    d = Document(io.BytesIO(docx_bytes))
    out = []
    for p in d.element.xpath("//w:p"):
        if any(True for _ in p.iterancestors(f"{_MC}Fallback")):
            continue
        t = "".join(x.text or "" for x in p.xpath(".//w:t")).strip()
        if t:
            out.append(t)
    return "\n".join(out)


def _blocks_via_walk(template_rel_path):
    """The SAME id-mapping walk the endpoint + override-apply use, run
    directly against a template file (no HTTP round trip) so tests can find a
    paragraph's id by its (pristine) text."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / template_rel_path))
    return [
        {"id": idx, "kind": kind, "text": text, "in_block": in_block, "in_txbx": in_txbx}
        for idx, kind, _p, in_block, text, in_txbx in pw.iter_editable_blocks(d)
    ]


def _find(blocks, predicate):
    hits = [b for b in blocks if predicate(b["text"])]
    assert hits, "no block matched predicate"
    return hits[0]


_EPOXY_TEMPLATE = "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx"
_COMBO_TEMPLATE = "Direct/xx.xx.xx TREADWELL COMBO PROPOSAL - CUSTMOER NAME.docx"


# ── GET /api/proposal-template ──────────────────────────────────────────
def test_template_endpoint_epoxy_blocks_ordered_with_in_block():
    r = client.get("/api/proposal-template?work_type=epoxy&audience=Direct")
    assert r.status_code == 200, r.text
    j = r.json()
    blocks = j["blocks"]
    assert blocks, "no blocks returned"

    # ids are the walk's own 0..N-1 index, in order, no gaps/dupes.
    ids = [b["id"] for b in blocks]
    assert ids == list(range(len(blocks)))

    # kind is always one of the two contracted shapes.
    assert {b["kind"] for b in blocks} <= {"p", "cell"}

    # Ordinary boilerplate text is freely editable (in_block is None).
    terms = _find(blocks, lambda t: t.strip() == "TERMS AND CONDITIONS")
    assert terms["in_block"] is None

    # Priced/repeatable regions are correctly tagged, including NESTED
    # sub-blocks (tax_breakout / remodel / has_options all live inside
    # single_bid — see proposal_writer._expand_named_block's docstring).
    single_bid_start = _find(blocks, lambda t: t.strip() == "{{#single_bid}}")
    assert single_bid_start["in_block"] == "single_bid"
    base_bid = _find(blocks, lambda t: t.strip() == "Base Bid")
    assert base_bid["in_block"] == "single_bid"
    sales_tax_row = _find(blocks, lambda t: "Material Sales Tax" in t)
    assert sales_tax_row["in_block"] == "tax_breakout"
    remodel_row = _find(blocks, lambda t: "Remodel Tax" in t and "alternate" not in t.lower())
    assert remodel_row["in_block"] == "remodel"
    options_label = _find(blocks, lambda t: t.strip() == "Options:")
    assert options_label["in_block"] == "has_options"
    price_line_row = _find(blocks, lambda t: "price_line.label" in t)
    assert price_line_row["in_block"] == "price_line"
    alt_row = _find(blocks, lambda t: "ALTERNATE SYSTEM" in t)
    assert alt_row["in_block"] == "alternate"

    # Document-order sanity: header -> WORK -> PRICE reads top to bottom.
    job_name = _find(blocks, lambda t: t.strip() == "{{job_name}}")
    system_start = _find(blocks, lambda t: t.strip() == "{{#system}}")
    assert job_name["id"] < system_start["id"] < single_bid_start["id"] < price_line_row["id"] < alt_row["id"]

    # Front-page vs body flag: the tokenized front page lives in floating
    # text boxes (in_txbx True); the Terms boilerplate is plain body text.
    assert job_name["in_txbx"] is True
    assert terms["in_txbx"] is False


def test_template_endpoint_combo_blocks_in_block_flags():
    r = client.get("/api/proposal-template?work_type=combo&audience=Direct")
    assert r.status_code == 200, r.text
    j = r.json()
    blocks = j["blocks"]
    assert j["template_name"] == "xx.xx.xx TREADWELL COMBO PROPOSAL - CUSTMOER NAME.docx"

    option1 = _find(blocks, lambda t: t.strip().startswith("Option 1:"))
    assert option1["in_block"] is None            # freely editable WORK copy
    single_bid_start = _find(blocks, lambda t: t.strip() == "{{#single_bid}}")
    assert single_bid_start["in_block"] == "single_bid"
    remodel_row = _find(blocks, lambda t: t.strip() == "{{#remodel}}")
    assert remodel_row["in_block"] == "remodel"    # nested inside single_bid


def test_template_endpoint_unknown_combo_falls_back_but_never_500():
    r = client.get("/api/proposal-template?work_type=bogus&audience=Bogus")
    assert r.status_code == 200
    # main.proposal_writer.pick_template falls back to (epoxy, Direct)
    assert "EPOXY" in r.json()["template_name"].upper()


# ── paragraph_overrides: apply against the pristine template ───────────────
_BASE_VALS = {
    "lump_sum_formatted": "$61,162.00", "tax_amount_formatted": "$2,639.00",
    "state_name": "Kansas", "total_formatted": "$63,801.00",
    "total_label": "$63,801.00 – Total", "system_name": "MACRO", "texture": "OP",
    "epoxy_sf": "12,000", "polish_sf": "8,000", "cove_lf": "250",
    "bid_date_formatted": "6/15/26", "job_name": "J", "city_state": "C",
    "area_description": "x", "disposal": "d", "site_visit_date": "6/15",
    "schedule_notes": "REPLACED-SCHEDULE-TOKEN", "work_description": "w",
    "scope_notes": "demo + install scope",
    "base_bid_formatted": "$58,523.00", "material_tax_formatted": "$2,639.00",
    "base_tax_phrase": "(material sales tax INCLUDED)",
    "site_visit_phrase": "per site visit on 6/15/26", "exclusions": "standard exclusions",
}


def test_paragraph_override_lands_in_docx_with_formatting_preserved():
    blocks = _blocks_via_walk(_EPOXY_TEMPLATE)
    scope = _find(blocks, lambda t: t.strip().startswith("Scope:"))
    assert scope["in_block"] is None

    # Confirm the ORIGINAL paragraph's bold-ness so we can prove it survives.
    d0 = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    orig_p = [p for i, k, p, ib, t, tb in pw.iter_editable_blocks(d0) if i == scope["id"]][0]
    orig_bold = any(r.bold for r in docx.text.paragraph.Paragraph(orig_p, d0).runs if r.bold is not None)

    new_text = "Scope:  A totally custom scope, still with a live token {{schedule_notes}} kept."
    out = pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS,
                           paragraph_overrides=[{"id": scope["id"], "text": new_text}])
    rendered = _rendered(out)
    assert "A totally custom scope, still with a live token REPLACED-SCHEDULE-TOKEN kept." in rendered
    # The {{schedule_notes}} the estimator deliberately left IN the override
    # text got filled by the normal flat pass (Phase 2), same as any other
    # token — it wasn't just quoted back out.
    assert "{{schedule_notes}}" not in rendered
    assert not re.search(r"\{\{[#/]", rendered), "leftover block marker"

    # Formatting preserved: re-walk the OUTPUT doc with the SAME shared helper
    # (not a naive raw xpath text join, which would match the text box's
    # ANCHOR paragraph too and misreport its formatting — see `_own_text`'s
    # docstring) and confirm the overridden paragraph's bold-ness matches.
    d1 = docx.Document(io.BytesIO(out))
    hit_elem = next(
        p for i, k, p, ib, t, tb in pw.iter_editable_blocks(d1)
        if "totally custom scope" in t
    )
    hit = docx.text.paragraph.Paragraph(hit_elem, d1)
    new_bold = any(r.bold for r in hit.runs if r.bold is not None)
    assert new_bold == orig_bold


def test_paragraph_override_inside_block_region_is_rejected():
    blocks = _blocks_via_walk(_EPOXY_TEMPLATE)
    base_bid = _find(blocks, lambda t: t.strip() == "Base Bid")
    assert base_bid["in_block"] == "single_bid"

    out = pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS,
                           paragraph_overrides=[{"id": base_bid["id"], "text": "HACKED BASE BID"}])
    rendered = _rendered(out)
    assert "HACKED BASE BID" not in rendered
    assert "Base Bid" in rendered   # untouched, template's own pricing content stands


def test_apply_paragraph_overrides_ignores_malformed_entries_directly():
    """Unit-level: proposal_writer._apply_paragraph_overrides must never raise,
    even called directly (bypassing main.py's sanitization) with garbage."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    n = pw._apply_paragraph_overrides(d, [
        "not-a-dict",
        {"id": "not-an-int", "text": "x"},
        {"id": 3.5, "text": "x"},        # float id -> rejected (not int)
        {"id": True, "text": "x"},       # bool id -> rejected (isinstance(bool, int) trap)
        {"id": 0, "text": 123},          # non-str text
        {"id": 0},                        # missing text
        {"id": 999999, "text": "out of range"},
    ])
    assert n == 0


def test_generate_endpoint_survives_malformed_paragraph_overrides():
    """Drives the real /api/generate route (not just the unit helper) with a
    garbage paragraph_overrides list mixed with one valid override — must
    return 200, never 500."""
    blocks = _blocks_via_walk(_EPOXY_TEMPLATE)
    scope = _find(blocks, lambda t: t.strip().startswith("Scope:"))
    malformed = [
        "not-a-dict",
        {"id": "nope"},
        {"no_id": True, "text": "x"},
        {"id": None, "text": "x"},
        {"id": scope["id"], "text": "Scope:  Overridden via HTTP"},
    ]
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "paragraph_overrides": malformed}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    f = client.get(r.json()["docx_download_url"])
    assert f.status_code == 200
    assert "Overridden via HTTP" in _rendered(f.content)


def test_paragraph_overrides_list_is_capped():
    """main._sanitize_paragraph_overrides caps the accepted list so a client
    can't force fill_proposal to walk/rewrite an unbounded number of ids."""
    over_cap = [{"id": i, "text": f"x{i}"} for i in range(600)]
    sanitized = main._sanitize_paragraph_overrides(over_cap)
    assert len(sanitized) == main._PARAGRAPH_OVERRIDES_MAX


def test_sanitize_paragraph_overrides_coerces_and_drops_bad_entries():
    out = main._sanitize_paragraph_overrides([
        {"id": "5", "text": "ok, string id coerces"},
        {"id": 6, "text": 42},            # non-str text coerces via str()
        {"id": None, "text": "bad id"},
        {"id": True, "text": "bool id"},   # bool would coerce to 1 — dropped
        "not-a-dict",
        {"id": 7},                          # missing text
    ])
    assert out == [
        {"id": 5, "text": "ok, string id coerces"},
        {"id": 6, "text": "42"},
    ]


# ── fidelity metadata: runs / list flags / geometry ─────────────────────
def test_template_endpoint_run_formatting_and_list_flags():
    r = client.get("/api/proposal-template?work_type=epoxy&audience=Direct")
    blocks = r.json()["blocks"]

    # Every block's run segments must re-join to exactly its text (the
    # frontend falls back to flat rendering when this doesn't hold — for
    # Kyle's templates it always must).
    for b in blocks:
        if b["runs"]:
            assert "".join(s["text"] for s in b["runs"]) == b["text"], f"run join broke on id {b['id']}"

    # Scope row: bold "Scope:" lead-in, then the token isolated as its own
    # segment carrying the value run's (non-bold) formatting — the same
    # attribution the docx fill uses (_sub_runs_preserving start-run rule).
    scope = _find(blocks, lambda t: t.strip().startswith("Scope:"))
    segs = scope["runs"]
    assert segs[0]["text"] == "Scope:" and segs[0]["bold"] is True
    token_seg = next(s for s in segs if s["text"] == "{{scope_notes}}")
    assert token_seg["bold"] is not True
    assert token_seg["font"] == "Zetta Serif Book"    # real template face
    assert token_seg["size_pt"] == 8.0                # true point size
    assert token_seg["color"] == "404040"

    # Real Word numbering drives the bullet flag; "Base Bid" is an indent-only
    # List Paragraph heading and must NOT read as a bullet.
    assert scope["list"] is True
    base_bid = _find(blocks, lambda t: t.strip() == "Base Bid")
    assert base_bid["list"] is False

    # Box membership pairs blocks with page geometry.
    assert scope["txbx"] is not None
    terms = _find(blocks, lambda t: t.strip() == "TERMS AND CONDITIONS")
    assert terms["txbx"] is None


def test_template_endpoint_geometry_page_boxes_artwork():
    r = client.get("/api/proposal-template?work_type=epoxy&audience=Direct")
    j = r.json()
    geo = j["geometry"]

    # Letter page, the template's real margins.
    assert geo["page"]["w_pt"] == 612.0 and geo["page"]["h_pt"] == 792.0
    assert geo["page"]["margin"]["left"] == 90.0 and geo["page"]["margin"]["top"] == 72.0

    # One geometry entry per text box, ids aligned with the blocks' txbx.
    box_ids = {b["id"] for b in geo["boxes"]}
    used = {b["txbx"] for b in j["blocks"] if b["txbx"] is not None}
    assert used <= box_ids
    for box in geo["boxes"]:
        assert all(isinstance(box[k], (int, float)) for k in ("x_pt", "y_pt", "w_pt", "h_pt"))

    # The WORK content box lands where the letterhead artwork draws its frame
    # (calibrated anchoring — allow a loose tolerance, the spec accepts
    # approximate placement).
    work_box = next(b for b in geo["boxes"] if b["w_pt"] > 400 and 100 < b["y_pt"] < 220)
    assert 140 <= work_box.get("x_pt") <= 180

    # The full-page letterhead artwork is listed and covers the page.
    art = [im for im in geo["images"] if im["w_pt"] > 600]
    assert art, "letterhead artwork missing from geometry"
    assert art[0]["para_index"] == 0                      # page-1 background
    assert any(im["para_index"] > 0 for im in art)        # terms-page background(s)


# ── media endpoint (letterhead artwork) ─────────────────────────────────
def test_media_endpoint_serves_whitelisted_artwork():
    r = client.get("/api/proposal-template?work_type=epoxy&audience=Direct")
    art_name = r.json()["geometry"]["images"][0]["name"]
    m = client.get(f"/api/proposal-template/media?work_type=epoxy&audience=Direct&name={art_name}")
    assert m.status_code == 200
    assert m.headers["content-type"].startswith("image/")
    assert m.content[:8] == b"\x89PNG\r\n\x1a\n"          # real PNG bytes


def test_media_endpoint_rejects_unknown_and_traversal_names():
    for bad in ("nope.png", "../document.xml", "word/media/image1.png", "..%2Fdocument.xml", ""):
        m = client.get("/api/proposal-template/media",
                       params={"work_type": "epoxy", "audience": "Direct", "name": bad})
        assert m.status_code == 404, f"{bad!r} should 404, got {m.status_code}"


# ── overrides must never delete anchored artwork / text boxes ───────────
def test_override_on_anchor_paragraph_keeps_drawings():
    """Kyle's templates anchor the letterhead art + every floating text box in
    runs of blank body paragraphs. An override targeting such a paragraph
    must write its text WITHOUT dropping the drawing run — losing it would
    delete the letterhead (or a whole text box) from the customer docx."""
    d0 = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    anchor_id = next(
        i for i, k, p, ib, t, tb in pw.iter_editable_blocks(d0)
        if tb is None and ib is None
        and p.find(".//" + "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") is not None
    )
    n_drawings_before = len(d0.element.body.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))

    out = pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS,
                           paragraph_overrides=[{"id": anchor_id, "text": "note above the letterhead"}])
    d1 = docx.Document(io.BytesIO(out))
    n_drawings_after = len(d1.element.body.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
    assert n_drawings_after == n_drawings_before, "override dropped an anchored drawing"
    assert "note above the letterhead" in _rendered(out)


# ── system_overrides: per-option WORK-row DISPLAY edits (epoxy only) ─────
def _epoxy_cells(**over):
    """cell_values that make _build_epoxy_systems pick one named system."""
    cells = {"Epoxy!A22": "Zeta Broadcast System", "Epoxy!E20": 12000, "Epoxy!E34": 0}
    cells.update(over)
    return cells


def test_system_override_applied_to_epoxy_docx():
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "cell_values": _epoxy_cells(),
            "system_overrides": [{"name": "Custom Display Name", "sqft": "9,999"}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    assert "Custom Display Name" in txt
    assert "Zeta Broadcast System" not in txt        # computed name replaced
    assert "9,999" in txt


def test_system_override_second_row_index_alignment():
    """A null hole in the list must NOT shift row 1's override onto row 0 —
    the sanitizer coerces it to {} in place, keeping option indices aligned."""
    cells = _epoxy_cells(**{"Epoxy!A26": "Beta Polish System", "Epoxy!E24": 8000, "Epoxy!E37": 0})
    cells["Epoxy!A22"] = "Alpha Epoxy System"
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "cell_values": cells,
            "system_overrides": [None, {"name": "Second Sys Renamed"}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    assert "Alpha Epoxy System" in txt                # option 1 keeps its computed name
    assert "Second Sys Renamed" in txt                # option 2 renamed
    assert "Beta Polish System" not in txt


def test_system_override_empty_string_reverts_to_computed():
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "cell_values": _epoxy_cells(),
            "system_overrides": [{"name": "", "texture": "   "}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    assert "Zeta Broadcast System" in txt             # blank override -> computed name shows


def test_system_override_ignored_for_non_epoxy():
    body = {"work_type": "polish", "audience": "Direct", "values": dict(_BASE_VALS),
            "system_overrides": [{"name": "ShouldNotAppearInPolish"}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    assert "ShouldNotAppearInPolish" not in txt


def test_generate_survives_malformed_system_overrides():
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "cell_values": _epoxy_cells(),
            "system_overrides": ["junk", 42, None, {"name": {"nested": 1}},
                                 {"sqft": ["x"]}, {"name": "ok-row"}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text


def test_sanitize_system_overrides_index_preserving_cap_and_coercion():
    out = main._sanitize_system_overrides([
        {"name": 42, "texture": "T", "sqft": "  1,000  "},   # int coerces, sqft strips
        "not-a-dict",                                         # -> {} placeholder, NOT dropped
        {"name": {"nested": 1}, "sqft": ["x"]},               # nested containers dropped
        {"name": "", "bogus": "ignored"},                     # blank dropped, unknown key ignored
        {"name": True},                                       # bool dropped
    ])
    assert out == [
        {"name": "42", "texture": "T", "sqft": "1,000"},
        {},                                                   # index preserved
        {},
        {},
        {},
    ]
    # length cap
    assert len(main._sanitize_system_overrides([{"name": "x"}] * 50)) == main._SYSTEM_OVERRIDES_MAX
    # per-field length cap
    long_one = main._sanitize_system_overrides([{"name": "z" * 999}])
    assert len(long_one[0]["name"]) == main._SYSTEM_OVERRIDE_FIELD_MAXLEN


def test_system_override_does_not_touch_estimate_cells():
    """Doc-editor system overrides are DISPLAY-only: the generated .xlsx must
    still carry the original estimate value in Epoxy!A22, unaffected by an
    override that renames the system in the proposal .docx."""
    from openpyxl import load_workbook
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_BASE_VALS),
            "cell_values": _epoxy_cells(),
            "system_overrides": [{"name": "Renamed In Proposal Only"}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    wb = load_workbook(io.BytesIO(client.get(r.json()["xlsx_download_url"]).content))
    assert wb["Epoxy"]["A22"].value == "Zeta Broadcast System"


def test_price_rows_carry_price_flat_flag():
    """PRICE-list rows (numId=3) carry price_flat=True so the on-screen document
    editor can render them with the red-square PRICE-list style (.tw-priceline),
    matching Kyle's template. WORK + Terms rows are not flagged (price_flat
    falsy) — they render via their own bullet/number style."""
    r = client.get("/api/proposal-template?work_type=epoxy&audience=Direct")
    assert r.status_code == 200
    blocks = r.json()["blocks"]

    def all_flat(sub):
        hits = [b for b in blocks if sub in b["text"]]
        assert hits, f"{sub!r} not found among blocks"
        return all(b.get("price_flat") for b in hits)

    for sub in ("base_bid_formatted", "material_tax_formatted",
                "price_line.amount_formatted", "total_formatted"):
        assert all_flat(sub), f"{sub} rows should be price_flat"

    # Terms (numId 5) must NOT be flagged — they keep their bullets.
    terms = [b for b in blocks if b["text"].strip() == "TERMS AND CONDITIONS"]
    assert terms and not any(b.get("price_flat") for b in terms)
