"""Resizing a proposal text box, and whether the customer's PDF honours it.

Hanz's third ask for the Word-like Proposals tab was "allow that textbox to be adjustable as
well". That matters more than a convenience: measuring the GC Resinous NOTES box showed the
delivered PDF was LOSING CONTENT — at the template's 184pt the box renders 21 of 60 lines, and
Kyle's boilerplate is longer than that. Making the box bigger is the actual cure; the overflow
notice shipped in #244 only reports the problem.

A box's size lives in THREE places, in two unit systems:

    wp:extent/@cx,@cy                  EMU, the anchor's size
    wps:spPr/a:xfrm/a:ext/@cx,@cy      EMU, the shape transform
    mc:Fallback v:shape/@style         points (sometimes inches) in a CSS-ish string

Write only the DrawingML pair and the legacy VML twin still claims the old size. Which branch
a renderer reads is version-dependent, so Word and the customer's PDF could disagree about the
same box — the worst kind of bug here, because the estimator's screen would look right.

Two traps, both found by probing the real template rather than assuming:

  * `wsp.iter(a:ext)` matches TWO elements. The second, under `a:extLst`, has no cx/cy —
    writing to it silently does nothing while looking like it worked.
  * VML lengths are not uniformly points: box 0 is `width:324.8pt`, box 1 is `width:1in`.

The PDF test at the bottom is the one that actually settles it, and it is the reason this
feature was safe to build: project memory records that LibreOffice ignores DrawingML autofit,
so it was an open question whether it honours an explicit extent either.
"""
import io
import pathlib
import shutil

import docx
import pytest

import pdf_writer
import proposal_writer as pw

TEMPLATES = pathlib.Path(__file__).resolve().parents[1] / "templates"
GC_RESINOUS = TEMPLATES / "GC" / "xx TREADWELL RESINOUS PROPOSAL - xx.docx"

_A = "{%s}" % pw._A_NS
_V = "{%s}" % pw._V_NS
_WPS = "{%s}" % pw._WPS_NS


def _sites(txbx):
    """Every recorded size for one box: (extent_cy_emu, xfrm_cy_emu, vml_height_pt)."""
    anchor = pw._txbx_anchor(txbx)
    ext = anchor.find(pw.qn("wp:extent")) if anchor is not None else None
    extent_cy = int(ext.get("cy")) if ext is not None else None

    xfrm_cy = None
    for wsp in txbx.iterancestors(_WPS + "wsp"):
        xfrm = next(iter(wsp.iter(_A + "xfrm")), None)
        e = xfrm.find(_A + "ext") if xfrm is not None else None
        if e is not None:
            xfrm_cy = int(e.get("cy"))
        break

    vml_h = None
    for shape in pw._txbx_vml_twins(txbx):
        for part in (shape.get("style") or "").split(";"):
            k, _, v = part.partition(":")
            if k.strip().lower() == "height":
                vml_h = pw._vml_len_pt(v)
        break
    return extent_cy, xfrm_cy, vml_h


def _box(d, i=0):
    return list(pw._iter_txbx(d))[i]


# ── the units the templates actually use ──────────────────────────────
@pytest.mark.parametrize("raw,expect", [
    ("99pt", 99.0), ("324.8pt", 324.8), ("1in", 72.0), ("18pt", 18.0),
    ("2pc", 24.0), ("96px", 72.0), ("50", 50.0), ("-36pt", -36.0),
    ("", None), ("auto", None), ("12em", None), ("50%", None),
])
def test_vml_lengths_are_read_in_their_own_unit(raw, expect):
    """Box 1 of the real GC template is `width:1in`, not points. Reading that as 1pt would
    collapse the box to a sliver in every VML-reading renderer."""
    assert pw._vml_len_pt(raw) == expect


def test_the_fixture_really_does_mix_units():
    """If the templates ever became uniformly pt, the unit handling above would stop being
    exercised by anything real and this file should say so out loud."""
    d = docx.Document(str(GC_RESINOUS))
    units = set()
    for txbx in pw._iter_txbx(d):
        for shape in pw._txbx_vml_twins(txbx):
            for part in (shape.get("style") or "").split(";"):
                k, _, v = part.partition(":")
                if k.strip().lower() in ("width", "height"):
                    units.add("".join(c for c in v if c.isalpha()))
    assert len(units) > 1, "expected mixed VML units in the real template, got %s" % units


# ── all three sites move together ─────────────────────────────────────
def test_a_resize_writes_every_site():
    d = docx.Document(str(GC_RESINOUS))
    txbx = _box(d)
    before = _sites(txbx)
    assert all(v is not None for v in before), "fixture box lacks a site: %s" % (before,)

    assert pw._resize_txbx(txbx, h_pt=400.0) >= 3
    extent_cy, xfrm_cy, vml_h = _sites(txbx)
    assert extent_cy == int(round(400.0 * pw._EMU_PER_PT))
    assert xfrm_cy == int(round(400.0 * pw._EMU_PER_PT))
    assert vml_h == 400.0, "the legacy VML twin still claims the old height"


def test_the_drawingml_and_vml_sizes_agree_after_a_resize():
    """They must not merely both change — they must land on the SAME size, or Word and the
    PDF disagree about a box the estimator can see."""
    d = docx.Document(str(GC_RESINOUS))
    txbx = _box(d)
    pw._resize_txbx(txbx, w_pt=300.0, h_pt=250.0)
    extent_cy, xfrm_cy, vml_h = _sites(txbx)
    assert extent_cy == xfrm_cy
    assert abs(vml_h - extent_cy / pw._EMU_PER_PT) < 0.02


def test_resizing_one_dimension_leaves_the_other_alone():
    d = docx.Document(str(GC_RESINOUS))
    txbx = _box(d)
    anchor = pw._txbx_anchor(txbx)
    cx_before = anchor.find(pw.qn("wp:extent")).get("cx")
    pw._resize_txbx(txbx, h_pt=333.0)
    assert anchor.find(pw.qn("wp:extent")).get("cx") == cx_before


def test_the_vml_style_keeps_its_other_declarations():
    """The style also carries position, z-index and visibility. Losing those would unpin the
    box from the page — a far worse outcome than the wrong height."""
    d = docx.Document(str(GC_RESINOUS))
    txbx = _box(d)
    shape = pw._txbx_vml_twins(txbx)[0]
    before = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    pw._resize_txbx(txbx, h_pt=280.0)
    after = dict(p.split(":", 1) for p in shape.get("style").split(";") if ":" in p)
    for key, val in before.items():
        if key.strip().lower() in ("width", "height"):
            continue
        assert after.get(key) == val, "the VML style lost or changed %r" % key
    assert "position" in after and "z-index" in after


def test_the_extlst_ext_is_never_mistaken_for_geometry():
    """`wsp.iter(a:ext)` also matches an a:extLst child with no cx/cy. Writing there is a
    silent no-op that looks like a successful resize."""
    d = docx.Document(str(GC_RESINOUS))
    txbx = _box(d)
    pw._resize_txbx(txbx, h_pt=321.0)
    for wsp in txbx.iterancestors(_WPS + "wsp"):
        for e in wsp.iter(_A + "ext"):
            parent = e.getparent().tag.split("}")[-1]
            if parent == "extLst":
                assert e.get("cx") is None, "geometry was written onto the a:extLst ext"
        break


# ── the sanitiser ─────────────────────────────────────────────────────
@pytest.mark.parametrize("raw", [
    None, [], "nope", 5, {"a": 1}, {"0": None}, {"0": "tall"}, {"0": []},
    {"0": {"h_pt": "tall"}}, {"0": {"h_pt": True}}, {"0": {"h_pt": float("nan")}},
    {"0": {"h_pt": float("inf")}}, {"-1": {"h_pt": 200}}, {"0": {}},
])
def test_malformed_box_overrides_never_raise_and_never_apply(raw):
    """House rule for every sanitiser here: a stale draft or hand-built request must not 500
    /api/generate."""
    assert pw._sanitize_box_overrides(raw) == {}
    d = docx.Document(str(GC_RESINOUS))
    assert pw._apply_box_overrides(d, raw) == 0          # must not raise


@pytest.mark.parametrize("h,kept", [
    (12.0, True), (200.0, True), (648.0, True),
    (11.9, False), (0, False), (-5, False), (649.0, False), (700.0, False),
    (1600.0, False), (99999, False),
])
def test_heights_outside_the_printable_area_are_refused_rather_than_clamped(h, kept):
    """A 4pt box is a corrupt draft, not an intention. Refusing leaves the design size, which is
    a document that still reads correctly.

    The ceiling used to be 1600pt — about two pages. Measured in the container: at 1600pt the
    render keeps 56 of 60 lines and puts the last one 3pt from the sheet edge. The four that do
    not fit are silently dropped, so the customer gets a proposal missing text and nothing
    anywhere reports a problem."""
    got = pw._sanitize_box_overrides({"0": {"h_pt": h}})
    assert ("0" in got) is kept


def test_a_box_may_not_be_made_taller_than_the_page_can_print():
    """The bug this closes: 700pt was accepted and put text past the bottom edge."""
    d = docx.Document(str(GC_RESINOUS))
    assert pw._apply_box_overrides(d, {"0": {"h_pt": 700.0}}) == 0, (
        "a height the page cannot print was applied anyway")


@pytest.mark.parametrize("path", sorted(TEMPLATES.rglob("*.docx")))
def test_no_template_ships_a_box_bigger_than_the_limit(path):
    """If any design box exceeded the printable area, the new ceiling would forbid resizing it at
    all — a fix that broke the feature for that template. Measured: the largest is 424x222."""
    d = docx.Document(str(path))
    max_w, max_h = pw.box_size_limits(d)
    for i, tb in enumerate(pw._iter_txbx(d)):
        anc = pw._txbx_anchor(tb)
        ext = None if anc is None else anc.find(pw.qn("wp:extent"))
        if ext is None:
            continue
        w = float(ext.get("cx") or 0) / 12700.0
        h = float(ext.get("cy") or 0) / 12700.0
        assert w <= max_w and h <= max_h, (
            "box %d in %s is %.0fx%.0f, past its own page limit of %.0fx%.0f"
            % (i, path.name, w, h, max_w, max_h))


def test_the_limit_is_the_printable_area_of_the_template_it_is_given():
    """Not a constant: a template with different margins is held to its own page."""
    d = docx.Document(str(GC_RESINOUS))
    sec = d.sections[0]
    max_w, max_h = pw.box_size_limits(d)
    assert max_h == pytest.approx(sec.page_height.pt - sec.top_margin.pt - sec.bottom_margin.pt)
    assert max_w == pytest.approx(sec.page_width.pt - sec.left_margin.pt - sec.right_margin.pt)
    assert (max_w, max_h) == pytest.approx((432.0, 648.0)), "Kyle's templates are Letter, 1in"


def test_a_template_with_no_usable_page_falls_back_instead_of_refusing_everything():
    """Absurd margins would otherwise compute a limit below the 12pt minimum, which would refuse
    every resize on that template rather than bound it."""
    class _Fake:
        sections = []
    assert pw.box_size_limits(_Fake()) == (432.0, 648.0)


def test_the_editor_is_told_the_limit_rather_than_deriving_it():
    """Two independent subtractions of the same margins is how a drag handle ends up letting
    somebody drag to a size the server then throws away — which reads as the drag not working."""
    d = docx.Document(str(GC_RESINOUS))
    page = pw.template_geometry(d)["page"]
    assert page["max_box"]["h_pt"] == pytest.approx(648.0)
    assert page["max_box"]["w_pt"] == pytest.approx(432.0)
    assert page["max_box"]["min_pt"] == pytest.approx(12.0)
    got = pw._sanitize_box_overrides({"0": {"h_pt": page["max_box"]["h_pt"]}},
                                     pw.box_size_limits(d))
    assert "0" in got, "the limit the editor is given is itself refused by the sanitiser"


def test_a_string_key_that_is_a_number_is_accepted():
    """The client sends JSON, so ids arrive as strings."""
    assert pw._sanitize_box_overrides({"3": {"h_pt": 150}}) == {"3": {"h_pt": 150.0}}


def test_an_override_for_a_box_that_does_not_exist_is_ignored():
    d = docx.Document(str(GC_RESINOUS))
    n = len(list(pw._iter_txbx(d)))
    assert pw._apply_box_overrides(d, {str(n + 50): {"h_pt": 300}}) == 0


# ── ordering: the shrink must stand down ──────────────────────────────
def test_a_resized_box_is_no_longer_shrunk():
    """THE point of the feature, and the reason the resize runs before the shrink.

    A box is enlarged precisely because its text was being cut off. If the shrink still ran
    against the OLD geometry it would scale the runs down anyway, and the estimator would see
    a bigger box full of smaller text — the problem they were trying to solve."""
    def sizes_after_shrink(box_h):
        d = docx.Document(str(GC_RESINOUS))
        txbx = _box(d)
        # Fill the box past its design height so the shrink has something to do.
        p = next(iter(txbx.iter(pw.qn("w:p"))))
        import copy
        for _ in range(40):
            txbx.append(copy.deepcopy(p))
        if box_h:
            pw._resize_txbx(txbx, h_pt=box_h)
        pw._shrink_overflowing_text_boxes(d)
        return [sz.get(pw.qn("w:val")) for sz in txbx.iter(pw.qn("w:sz"))]

    at_design = sizes_after_shrink(None)
    enlarged = sizes_after_shrink(1500.0)
    assert at_design, "no runs carry a size; this test would prove nothing"
    assert enlarged != at_design, (
        "the shrink treated the enlarged box exactly like the small one — the resize is not "
        "being applied before the shrink reads the geometry")


def test_fill_proposal_accepts_box_overrides_end_to_end():
    blob = pw.fill_proposal(
        work_type="polish", audience="GC",
        values={"job_name": "Resize Check", "city_state": "Lenexa, KS"},
        box_overrides={"0": {"h_pt": 420.0}},
    )
    out = docx.Document(io.BytesIO(blob))
    extent_cy, xfrm_cy, vml_h = _sites(_box(out))
    assert extent_cy == int(round(420.0 * pw._EMU_PER_PT))
    assert xfrm_cy == extent_cy
    assert vml_h == 420.0


def test_generation_without_box_overrides_is_unchanged():
    """The feature must be inert when nobody used it."""
    common = dict(work_type="polish", audience="GC",
                  values={"job_name": "Inert Check", "city_state": "Lenexa, KS"})
    a = docx.Document(io.BytesIO(pw.fill_proposal(**common)))
    b = docx.Document(io.BytesIO(pw.fill_proposal(box_overrides=None, **common)))
    c = docx.Document(io.BytesIO(pw.fill_proposal(box_overrides={}, **common)))
    assert _sites(_box(a)) == _sites(_box(b)) == _sites(_box(c))


# ── does the customer's PDF honour it? ────────────────────────────────
@pytest.mark.skipif(shutil.which("soffice") is None and shutil.which("libreoffice") is None,
                    reason="LibreOffice is not installed (it is in the Docker image)")
def test_the_pdf_really_honours_an_explicit_resize():
    """The load-bearing test. Project memory records that LibreOffice ignores DrawingML
    autofit, which left it genuinely open whether it honours an explicit extent — and if it
    did not, the estimator's resize would look right in Word and be lost in the PDF the
    customer actually receives.

    Measured by putting numbered lines in a box and counting how many survive into the PDF's
    text layer at the design height versus enlarged."""
    import copy
    import re as _re

    def lines_in_pdf(box_h):
        d = docx.Document(str(GC_RESINOUS))
        txbx = _box(d)
        proto = next(iter(txbx.iter(pw.qn("w:p"))))
        for p in list(txbx.iter(pw.qn("w:p"))):
            p.getparent().remove(p)
        for i in range(60):
            p = copy.deepcopy(proto)
            pw._set_paragraph_text(p, "MARK%02d" % i)
            txbx.append(p)
        if box_h:
            pw._resize_txbx(txbx, h_pt=box_h)
        buf = io.BytesIO()
        d.save(buf)
        pdf = pdf_writer.docx_to_pdf(buf.getvalue())
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF not installed")
        with fitz.open(stream=pdf, filetype="pdf") as doc:
            text = "".join(page.get_text() for page in doc)
            # How far down the paper the last surviving line sits, so the limit can be held to
            # its actual promise rather than to a line count.
            low, page_h = 0.0, doc[0].rect.height
            for page in doc:
                for x0, y0, x1, y1, word, *_ in page.get_text("words"):
                    if word.startswith("MARK"):
                        low = max(low, y1)
        return len(set(_re.findall(r"MARK\d\d", text))), low, page_h

    # The tallest the product will actually accept, not an arbitrary large number. This used to
    # measure at 700pt, which the sanitiser now refuses for running off the paper — proving the
    # mechanism at a size nobody can ask for would be proving the wrong thing.
    limit = pw.box_size_limits(docx.Document(str(GC_RESINOUS)))[1]
    small, _, _ = lines_in_pdf(None)
    big, lowest, page_h = lines_in_pdf(limit)
    assert small < 60, ("the design-height box already fits all 60 lines, so this test cannot "
                        "distinguish honoured from ignored — it needs a box that overflows")
    assert big > small, (
        "LibreOffice ignored the explicit resize: %d lines at design height, %d at the %.0fpt "
        "limit. If this ever fails, box resizing cannot be trusted to reach the customer's PDF."
        % (small, big, limit))
    # The point of the limit: at the tallest size the product accepts, the text is still inside
    # the printable area. Measured — 648pt puts the last line at 684pt of a 720pt margin; 700pt
    # puts it at 740pt, into the margin; 1600pt at 789pt, with four lines dropped altogether.
    #
    # Asserting merely "on the paper" would pass at all three, because LibreOffice clips rather
    # than spilling. That is the assertion I first wrote, and it would not have caught the bug
    # this limit exists for.
    bottom = page_h - docx.Document(str(GC_RESINOUS)).sections[0].bottom_margin.pt
    assert lowest <= bottom, (
        "at the permitted maximum the last line sits %.0fpt down, past the %.0fpt bottom margin "
        "— the limit does not keep the box inside the printable area" % (lowest, bottom))
