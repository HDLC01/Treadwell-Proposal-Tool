"""The 27 numbered TERMS AND CONDITIONS clauses: the editor shows their NUMBERS, and no channel
can empty one.

TWO HALVES OF ONE MISUNDERSTANDING ABOUT ORDERED LISTS.

FIX A — the editor drew every clause as a red square. Two fields of `/api/proposal-template`
described the same paragraph and disagreed: `list` was True (the paragraph carries `w:numPr`) and
`para.bullet` was False (numId 5's `w:numFmt` is decimal). `renderBlock` trusted `list`, so all 27
clauses got `.tw-li` and the Wingdings square, and nothing corrected it — `applyParaToEl` is only
reachable through `setParaState`, which refuses a locked paragraph. So the screen showed a
bulleted list where the signed contract prints "1." to "27.". The flag was never wrong; it was
answering a different question. `para.marker` answers the right one: what the level PRINTS, read
off `w:lvlText` and `w:start` rather than inferred from `w:numFmt`.

FIX B — an emptied clause printed a bare clause number. The renumbering guard added 2026-08-20
holds (the clause count never moves), but it holds by KEEPING the numbering on a paragraph whose
text was just deleted, so Word printed "1." followed by nothing, in a contract, with no warning:
the paragraph controls are hidden on that row (`para.locked`). Of the three ways out —

  * refuse the edit and say so;
  * drop the paragraph, which renumbers every clause below it;
  * keep the text and say nothing —

the second is the exact regression the guard exists for (measured at 26 clauses of 27) and the
third is the same screen-versus-document lie as the red squares, pointed the other way. So the
edit is refused, at the API and in the editor, and the editor says why.

FIX A IS WHY FIX B WAS INVISIBLE: `.tw-block.tw-empty.tw-li::before { content: none }` meant an
emptied clause read as a blank line on screen while the .docx printed a bare number. That rule's
comment claimed "the preview shows what prints", which was false for an ordered list.

EXECUTED, NOT READ. The marker set is taken from the live endpoint and the browser behaviour from
the shipped renderBlock / input handler under node. A source-text assertion cannot see an unbound
identifier — on 2026-08-12 `STAGE_CREATED` shipped unbound with every source assertion green and
took the production board down.
"""
import io
import json
import pathlib
import shutil
import subprocess

import docx
import pytest
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

FRONTEND = pathlib.Path(__file__).resolve().parents[2] / "frontend"
HARNESS = pathlib.Path(__file__).resolve().parent / "js" / "doc-editor-fidelity-harness.js"
CSS = (FRONTEND / "styles.css").read_text(encoding="utf-8")

# Every proposal template, and the clause count MEASURED in each. Gyp really does have 28 — that
# is why the number is per template here instead of a single 27 asserted everywhere.
TEMPLATES = [
    ("epoxy", "Direct", 27),
    ("polish", "Direct", 27),
    ("combo", "Direct", 27),
    ("epoxy", "GC", 27),
    ("polish", "GC", 27),
    ("gyp", "Gyp", 28),
]

BASE_VALUES = {
    "job_name": "Clause Numbers", "project_name": "Clause Numbers", "city_state": "Olathe, KS",
    "scope_notes": "surface prep and coat", "schedule_notes": "four working days",
    "exclusions": "striping", "work_notes": "keep off for 24 hours",
}


def _doc(work_type="epoxy", audience="Direct"):
    return docx.Document(str(pw.pick_template(work_type, audience)))


def _blocks(work_type="epoxy", audience="Direct"):
    """The endpoint's own answer, over HTTP, for one template."""
    r = client.get("/api/proposal-template",
                   params={"work_type": work_type, "audience": audience})
    assert r.status_code == 200, r.text
    return r.json()["blocks"]


def _clause_ids(d):
    """The ids of the numbered clauses, by re-reading the file rather than hardcoding."""
    return [idx for idx, _k, p, ib, _t, _x in pw.iter_editable_blocks(d)
            if ib is None and pw._para_ordered_list(d, p)]


def _clause_markers(d):
    """The ordinals the generated/opened document's ordered paragraphs print, in document order."""
    return [m for _p, m in
            sorted(pw._ordered_markers(d).values(), key=lambda pm: _doc_index(d, pm[0]))]


def _doc_index(d, p_elem):
    for i, p in enumerate(d.element.body.iter(qn("w:p"))):
        if p is p_elem:
            return i
    return -1


def _clause_texts(d):
    return [(t or "").strip() for _i, _k, p, ib, t, _x in pw.iter_editable_blocks(d)
            if ib is None and pw._para_ordered_list(d, p)]


def _generate(overrides=None, work_type="epoxy", audience="Direct", **kw):
    return docx.Document(io.BytesIO(pw.fill_proposal(
        work_type=work_type, audience=audience, values=dict(BASE_VALUES),
        paragraph_overrides=overrides, **kw)))


# ══ the template facts, re-derived from Kyle's file ═══════════════════════════
def test_the_terms_level_is_decimal_and_says_so_in_its_lvl_text():
    """THE FACT THE WHOLE FIX RESTS ON. `w:numFmt` says "decimal"; it is `w:lvlText` that says
    "%1." — the period included — and `w:start` that says the first ordinal is 1. Neither was
    being read, which is why nothing downstream could tell "1." from a red square."""
    d = _doc()
    lvl = pw._numbering_levels(d).get(("5", "0"))
    assert lvl == {"fmt": "decimal", "text": "%1.", "start": 1,
                   "ind": {"left": "540", "hanging": "360"}}, lvl


def test_the_clause_number_prints_bold_which_the_preview_has_to_match():
    """The level carries `w:rPr><w:b/`, so Word prints "1." in bold. `.tw-num::before` is bold for
    that reason and no other — if Kyle re-authors the level, this says so before the preview
    quietly starts contradicting the document."""
    d = _doc()
    root = d.part.numbering_part.element
    aid = next(n.find(qn("w:abstractNumId")).get(qn("w:val"))
               for n in root.iter(qn("w:num")) if n.get(qn("w:numId")) == "5")
    lvl = next(a for a in root.iter(qn("w:abstractNum"))
               if a.get(qn("w:abstractNumId")) == aid).findall(qn("w:lvl"))[0]
    rpr = lvl.find(qn("w:rPr"))
    assert rpr is not None and rpr.find(qn("w:b")) is not None, \
        "the Terms level is no longer bold, so `.tw-num::before { font-weight: 700 }` now lies"
    rule = CSS.split(".tw-num::before", 1)[1].split("}", 1)[0]
    assert "font-weight: 700" in rule, rule


def test_the_bullet_levels_print_a_glyph_no_browser_font_has():
    """Which is why a bullet row keeps the preview's OWN square (CSS `content: "\\25AA"`) and
    reports `marker: ""`: numId 1/3/4 print U+F0A7, a Wingdings private-use codepoint. Rendering
    the level's own character would give the estimator a missing-glyph box."""
    levels = pw._numbering_levels(_doc())
    for num_id in ("1", "3", "4"):
        lvl = levels[(num_id, "0")]
        assert lvl["fmt"] == "bullet", (num_id, lvl)
        assert lvl["text"] == "", (num_id, lvl)


# ══ FIX A: what the endpoint actually returns ════════════════════════════════
@pytest.mark.parametrize("work_type,audience,count", TEMPLATES)
def test_every_numbered_clause_reports_the_number_it_prints(work_type, audience, count):
    """THE ENDPOINT, EXECUTED, for all 27 ids (28 in Gyp). The markers must be the contiguous run
    the contract shows — not "a marker exists", which a per-paragraph guess would also satisfy."""
    blocks = _blocks(work_type, audience)
    marked = [b for b in blocks if b["para"]["marker"]]
    assert [b["para"]["marker"] for b in marked] == \
        ["%d." % n for n in range(1, count + 1)], \
        "%s/%s clause numbering is not 1..%d" % (work_type, audience, count)
    assert len(marked) == count


@pytest.mark.parametrize("work_type,audience,count", TEMPLATES)
def test_a_clause_never_reports_a_bullet_and_a_bullet_never_reports_a_number(
        work_type, audience, count):
    """The two fields are the two halves of "what shows in front of the text", so they are never
    both set and never both empty on a list paragraph. THIS is the assertion the old `list` flag
    could not make: it was True for every one of them."""
    blocks = _blocks(work_type, audience)
    for b in blocks:
        p = b["para"]
        assert not (p["bullet"] and p["marker"]), b["id"]
        if b["list"]:
            assert p["bullet"] or p["marker"], (
                "block %s carries numbering but reports neither a bullet nor a number, so the "
                "editor has nothing to draw" % b["id"])
        else:
            assert not p["bullet"] and not p["marker"], b["id"]
    assert sum(1 for b in blocks if b["list"] and b["para"]["bullet"]) > 0
    assert sum(1 for b in blocks if b["list"] and b["para"]["marker"]) == count


@pytest.mark.parametrize("work_type,audience,count", TEMPLATES)
def test_locked_and_numbered_are_the_same_paragraphs_in_these_templates(
        work_type, audience, count):
    """`locked` (may the editor touch it) and `marker` (what does it print) are different
    questions, and conflating them was the bug — but in Kyle's files they land on the same rows.
    So the editor never has to fall back, and if that ever stops being true this says which
    paragraph is the exception rather than leaving a silent square behind."""
    blocks = _blocks(work_type, audience)
    locked = {b["id"] for b in blocks if b["para"]["locked"]}
    marked = {b["id"] for b in blocks if b["para"]["marker"]}
    assert locked == marked, sorted(locked ^ marked)
    assert len(locked) == count


def test_the_clause_markers_survive_a_whole_generate():
    """The document the customer receives, not the template: after block expansion (which inserts
    and removes paragraphs) the ordinals are still 1..27. Also the only test that exercises the
    marker cache being INVALIDATED — it is keyed on the paragraph count for exactly this reason."""
    template_markers = _clause_markers(_doc())
    generated = _generate()
    assert _clause_markers(generated) == template_markers
    assert template_markers == ["%d." % n for n in range(1, 28)]


# ══ the ordinal algebra, on its own ══════════════════════════════════════════
@pytest.mark.parametrize("n,fmt,want", [
    (1, "decimal", "1"), (27, "decimal", "27"),
    (1, "lowerLetter", "a"), (26, "lowerLetter", "z"), (27, "lowerLetter", "aa"),
    (2, "upperLetter", "B"),
    (4, "lowerRoman", "iv"), (9, "lowerRoman", "ix"), (14, "upperRoman", "XIV"),
    # Nothing recognisable, and nothing spellable, still produce a number rather than nothing:
    # a clause with no marker at all is indistinguishable on screen from a bullet row.
    (3, "cardinalText", "3"), (3, None, "3"), (0, "lowerRoman", "0"), (0, "lowerLetter", "0"),
])
def test_one_ordinal_in_every_format_the_templates_define(n, fmt, want):
    assert pw._format_ordinal(n, fmt) == want


def test_the_marker_is_the_lvl_text_not_the_number_alone():
    """"%1." is a template decision — the period is Kyle's, not ours. A renderer that built "1."
    from `w:numFmt` would print "1)" wrongly the day he changes the level."""
    levels = {("9", "0"): {"fmt": "decimal", "text": "(%1)", "start": 1, "ind": {}}}
    assert pw._render_lvl_text("(%1)", "9", {("9", "0"): 4}, levels) == "(4)"
    # A nested level whose parent counter has not been reached yet still reads as its start.
    levels[("9", "1")] = {"fmt": "lowerLetter", "text": "%1.%2", "start": 1, "ind": {}}
    assert pw._render_lvl_text("%1.%2", "9", {("9", "1"): 2}, levels) == "1.b"


def test_the_document_markers_come_from_the_lvl_text_not_the_fallback():
    """`_ordered_markers` has a fallback for a level with no `w:lvlText`, and on THESE templates it
    produces the same "1." as the real path — so nothing here would notice the marker being built
    from `w:numFmt` alone. Re-punctuating the level is the only way to tell them apart: change the
    level to print "(1)" and the document has to say "(1)"."""
    d = _doc()
    levels = pw._numbering_levels(d)
    levels[("5", "0")] = dict(levels[("5", "0")], text="(%1)")
    d._tw_num_levels = levels
    d._tw_ordered_markers = None
    assert _clause_markers(d)[:3] == ["(1)", "(2)", "(3)"]


def test_a_start_other_than_one_is_honoured():
    """`w:start` is read rather than assumed. Every level in these templates starts at 1, so this
    is the only place the code path is visible at all."""
    d = _doc()
    levels = pw._numbering_levels(d)
    levels[("5", "0")] = dict(levels[("5", "0")], start=5)
    d._tw_num_levels = levels
    d._tw_ordered_markers = None
    assert _clause_markers(d)[:3] == ["5.", "6.", "7."]


def test_the_marker_map_is_rebuilt_when_the_document_changes():
    """The cache is keyed on the paragraph COUNT for this reason. The map is an answer about
    DOCUMENT ORDER, and block expansion inserts and removes paragraphs on the same object — a
    cache that outlived that would hand out ordinals for a document that no longer exists, and
    (worse) hold a pinned element that has been detached from the tree."""
    d = _doc()
    assert len(pw._ordered_markers(d)) == 27
    by_id = {i: p for i, _k, p, _b, _t, _x in pw.iter_editable_blocks(d)}
    first = by_id[_clause_ids(d)[0]]
    first.getparent().remove(first)
    assert len(pw._ordered_markers(d)) == 26, "the marker map survived a paragraph being removed"
    # And the survivors renumbered, exactly as Word will renumber them.
    assert _clause_markers(d)[:2] == ["1.", "2."]


def test_a_document_with_no_numbering_definitions_reports_no_markers():
    """The honest degradation, and the reason `renderBlock` still falls back to the bullet class:
    an unreadable level yields "" rather than a guessed number."""
    d = _doc()
    d._tw_num_levels = {}
    d._tw_ordered_markers = None
    assert pw._ordered_markers(d) == {}
    for _i, _k, p, _b, _t, _x in pw.iter_editable_blocks(d):
        assert pw._para_marker(d, p) == ""


def test_a_bullet_row_is_never_given_a_marker():
    """Mutation bait: dropping the `fmt in ("bullet", "none")` skip would number all 24 WORK and
    NOTES rows, which reads as plausible output right up to the customer's document."""
    d = _doc()
    for _i, _k, p, _b, _t, _x in pw.iter_editable_blocks(d):
        if pw._num_fmt(d, p) == "bullet":
            assert pw._para_marker(d, p) == ""


def test_the_marker_lookup_survives_a_fresh_element_proxy():
    """The lxml trap `_hand_formatted` documents, in the one place it would be silent. The map is
    keyed on `id()` of an element proxy, and lxml frees a proxy the moment the last Python
    reference goes — then hands out a new one at a possibly REUSED address. The map pins the
    element, so a caller that walks the document AFTERWARDS looks up the same object."""
    d = _doc()
    first = pw._ordered_markers(d)          # built from its own walk, then that walk is gone
    assert first
    seen = [pw._para_marker(d, p) for _i, _k, p, ib, _t, _x in pw.iter_editable_blocks(d)
            if ib is None and pw._para_ordered_list(d, p)]
    assert seen == ["%d." % n for n in range(1, 28)], (
        "the marker map went stale between two walks of the same document: %r" % (seen,))


# ══ FIX B: no channel can empty a clause ═════════════════════════════════════
# text, runs, para, a lone newline, and a box override alongside. Every shape the editor and a
# stale draft can produce, against the count AND against the words.
def _blank_channels(tid):
    return {
        "text": [{"id": tid, "text": ""}],
        "whitespace": [{"id": tid, "text": "   "}],
        "newline": [{"id": tid, "text": "\n"}],
        "runs": [{"id": tid, "text": "", "runs": [{"text": ""}]}],
        "empty_runs": [{"id": tid, "text": "", "runs": []}],
        "runs_whitespace": [{"id": tid, "text": " ", "runs": [{"text": " ", "bold": True}]}],
        "with_para": [{"id": tid, "text": "", "para": {"bullet": False, "indent": 0}}],
        "para_only": [{"id": tid, "para": {"bullet": False, "indent": 0}}],
    }


@pytest.mark.parametrize("channel", sorted(_blank_channels(0)))
def test_no_channel_can_empty_a_clause_or_move_a_clause_number(channel):
    """The clause keeps its words, the count does not move, and every ordinal is where it was."""
    d0 = _doc()
    ids = _clause_ids(d0)
    tid = ids[0]
    before_texts, before_markers = _clause_texts(d0), _clause_markers(d0)
    assert len(before_texts) == 27

    got = _generate(_blank_channels(tid)[channel])
    assert _clause_texts(got) == before_texts, "%s emptied a clause" % channel
    assert _clause_markers(got) == before_markers, "%s moved a clause number" % channel


def test_a_box_override_alongside_the_blank_changes_nothing_about_the_clauses():
    """The fifth channel. A box resize travels in the same request and rewrites drawing anchors,
    so it is the one that could plausibly reach a paragraph by accident."""
    d0 = _doc()
    tid = _clause_ids(d0)[0]
    got = _generate([{"id": tid, "text": ""}], box_overrides={"0": {"h_pt": 500.0}})
    assert _clause_texts(got) == _clause_texts(d0)
    assert _clause_markers(got) == _clause_markers(d0)


def test_the_writer_reports_a_refused_blank_as_nothing_applied():
    """`_apply_paragraph_overrides` counts what it APPLIED, and a refusal applied nothing — the
    same answer it already gives for an id it skipped. Silently counting it would make the log
    line say the estimator's edit went through."""
    d = _doc()
    tid = _clause_ids(d)[0]
    assert pw._apply_paragraph_overrides(d, [{"id": tid, "text": ""}]) == 0
    # …and a NON-blank edit to the same clause is still applied: this is a rule about emptying,
    # not a read-only contract.
    assert pw._apply_paragraph_overrides(d, [{"id": tid, "text": "Agreement.  Reworded."}]) == 1
    assert "Agreement.  Reworded." in _clause_texts(d)


def test_a_bulleted_row_can_still_be_emptied_and_still_loses_its_square():
    """The behaviour the guard must not catch. A WORK row carries no meaning that outlives its
    text, so emptying it is a legitimate edit and its bullet goes with the words."""
    d = _doc()
    rows = {}
    for idx, _k, p, ib, t, _x in pw.iter_editable_blocks(d):
        if ib is None and (t or "").strip().startswith("Schedule:"):
            rows[idx] = p
    assert rows, "no free Schedule row in the template"
    sched = sorted(rows)[0]
    assert pw._apply_paragraph_overrides(d, [{"id": sched, "text": ""}]) == 1
    assert pw._para_num_ref(rows[sched]) is None, "the emptied bullet row kept its numbering"


def test_the_blank_is_refused_through_the_whole_generate_endpoint():
    """End to end over HTTP, with the real `template_version` so the staleness guard cannot make
    this pass by dropping the override before it reaches the writer."""
    tv = client.get("/api/proposal-template?work_type=epoxy&audience=Direct"
                    ).json()["template_version"]
    d0 = _doc()
    tid = _clause_ids(d0)[0]
    r = client.post("/api/generate", json={
        "work_type": "epoxy", "audience": "Direct",
        "values": dict(BASE_VALUES), "template_version": tv,
        "paragraph_overrides": [{"id": tid, "text": ""}],
    })
    assert r.status_code == 200, r.text
    got = docx.Document(io.BytesIO(client.get(r.json()["docx_download_url"]).content))
    assert _clause_texts(got) == _clause_texts(d0)
    assert _clause_markers(got) == _clause_markers(d0)


def test_the_sanitizer_still_carries_a_blank_through_to_the_writer():
    """Where the refusal must NOT live. `main._sanitize_paragraph_overrides` has no document, so
    it cannot tell a clause from a WORK row — a guard there would either break emptying a bullet
    row or be a guess. It keeps passing the entry along, and the writer decides."""
    out = main._sanitize_paragraph_overrides([{"id": 51, "text": ""}])
    assert out == [{"id": 51, "text": ""}]


@pytest.mark.parametrize("val,blank", [
    ("", True), ("   ", True), ("\n", True), ("\r\n\t ", True), ("x", False),
    ([], True), ([{"text": ""}], True), ([{"text": " "}], True),
    ([{"text": "x"}], False), ([{"text": ""}, {"text": "x"}], False),
    (None, True), ([{}], True),
])
def test_what_counts_as_blank(val, blank):
    """One definition for both shapes the editor sends. A lone newline is what a browser leaves
    behind when the last character is deleted, and it prints as nothing."""
    assert pw._override_is_blank(val) is blank


# ══ the browser, executed ════════════════════════════════════════════════════
@pytest.fixture(scope="module")
def ran():
    if shutil.which("node") is None:
        pytest.skip("node is not installed")
    proc = subprocess.run(["node", str(HARNESS), str(FRONTEND)],
                          capture_output=True, text=True, encoding="utf-8", timeout=180)
    assert proc.returncode == 0, (
        "the harness itself failed — read this before assuming a product bug:\n" + proc.stderr)
    return json.loads(proc.stdout.strip().splitlines()[-1])


def test_the_fixture_clause_is_the_shape_the_endpoint_really_sends():
    """So scenario 14 cannot drift into testing a payload the page never receives."""
    by_id = {b["id"]: b for b in _blocks()}
    clause = by_id[51]
    assert clause["list"] is True
    # v7 geometry. `hanging` is what makes this clause hang its number: text at 540tw (27pt),
    # the "1." at 540-360 = 180tw (9pt). Those are exactly the 9pt/18pt that .tw-num already
    # hardcoded -- the numbered clauses were the ONE place the editor was already reading the
    # file correctly, and this pins that they still agree now the numbers come from the record
    # rather than from the stylesheet.
    #
    # `contextual` is True here and False on the WORK rows: these clauses inherit
    # contextualSpacing from their own style, which is why a numbered contract has no gaps
    # between consecutive clauses.
    assert clause["para"] == {
        "bullet": False, "indent": 540, "hanging": 360, "first_line": None,
        "locked": True, "marker": "1.",
        "spacing": {"before": None, "after": None, "line": None,
                    "line_rule": None, "contextual": True},
    }
    assert clause["align"] == "justify"
    assert by_id[52]["para"]["marker"] == "2."
    work = by_id[115]
    assert work["list"] is True
    # The contrast the test is drawing: an editable WORK row, bulleted by a numbering level
    # that owns BOTH halves of its indent (288/288 -- text at 14.4pt, marker at the margin),
    # against the locked clause above. Spacing is compared for shape: this row carries the
    # file's own line spacing and no gaps, and no contextualSpacing of its own.
    wsp = work["para"].pop("spacing")
    assert work["para"] == {
        "bullet": True, "indent": 288, "hanging": 288, "first_line": None,
        "locked": False, "marker": "",
    }
    assert wsp["before"] is None and wsp["after"] is None, wsp
    assert wsp["line_rule"] == "auto" and wsp["line"] in (276, 300), wsp


def test_a_clause_is_drawn_as_its_number_and_not_as_a_bullet(ran):
    """Kyle's contract on screen. `.tw-num` carries the number in `data-marker`; `.tw-li` (the red
    square) is gone from the clause and still on the WORK row."""
    got = ran["clauseRender"]
    assert got["clause"]["num"] is True and got["clause"]["li"] is False
    assert got["clause"]["marker"] == "1."
    assert got["second"]["marker"] == "2."
    assert got["workRow"]["li"] is True and got["workRow"]["num"] is False
    assert got["workRow"]["marker"] == ""


def test_a_pre_v6_cached_block_still_gets_the_old_square(ran):
    """The honest fallback. A browser replaying a v5 response has no `marker`, and a paragraph
    that carries numbering has to be drawn as SOMETHING — so it keeps exactly what it drew before.
    `_BLOCK_SCHEMA_VERSION` was bumped so this never actually happens."""
    assert ran["clauseRender"]["noMarker"] == {
        "li": True, "num": False, "marker": "", "kept": False, "title": "",
        "dirty": False, "empty": False, "text": "Taxes.  Customer shall pay all taxes."}


def test_emptying_a_clause_puts_the_wording_back_and_says_why(ran):
    """The refusal the estimator can SEE, at the moment it happens. The clause text returns, the
    notice goes up, and nothing is shipped for that paragraph — so the document keeps the clause
    whether or not the backend guard also holds."""
    got = ran["clauseEmptied"]
    kept = got["kept"]
    assert kept["text"] == "Agreement. The Proposal of Treadwell, LLC."
    assert kept["kept"] is True
    assert "cannot be emptied" in kept["title"]
    assert kept["num"] is True and kept["marker"] == "1."
    # Back to matching the template, so it is not dirty and not empty either.
    assert kept["dirty"] is False and kept["empty"] is False
    assert got["collected"] == [] and got["stored"] == []
    # The clause below it never moved.
    assert got["neighbour"]["marker"] == "2."
    assert got["neighbour"]["kept"] is False


def test_the_notice_clears_when_they_carry_on_typing(ran):
    """A badge that outlives the thing it describes is the stale-badge bug this repo has already
    paid for once. It goes on the next real edit, and that edit ships normally."""
    got = ran["clauseEmptied"]
    assert got["afterTyping"]["kept"] is False
    assert got["afterTyping"]["title"] == ""
    assert got["afterTyping"]["dirty"] is True
    assert got["afterTypingCollected"] == [
        {"id": 51, "text": "Agreement.  The Proposal of Treadwell, LLC, as amended."}]


def test_a_lone_newline_is_treated_as_empty_in_the_browser_too(ran):
    """What a browser actually leaves in a contenteditable when the last character goes. The
    backend counts it as blank; the editor has to agree or the two halves disagree about the same
    keystroke."""
    got = ran["clauseNewline"]
    assert got["kept"]["kept"] is True
    assert got["kept"]["text"] == "Agreement. The Proposal of Treadwell, LLC."
    assert got["collected"] == []


def test_a_work_row_can_still_be_emptied_in_the_browser(ran):
    """The guard is about clause numbers, not about making the editor read-only."""
    got = ran["workRowEmptied"]
    assert got["look"]["text"] == ""
    assert got["look"]["empty"] is True
    assert got["look"]["kept"] is False
    assert got["collected"] == [{"id": 115, "text": ""}]


def test_a_draft_saved_with_an_emptied_clause_is_not_replayed(ran):
    """Drafts saved while this was possible still carry the entry. Replaying it would show a blank
    clause on screen while the .docx (which refuses it) prints the wording — the same lie as the
    red squares, pointed the other way. The stale entry is dropped on the next persist, and the
    rest of that draft is untouched."""
    got = ran["clauseLegacyBlank"]
    assert got["restored"]["text"] == "Agreement. The Proposal of Treadwell, LLC."
    assert got["restored"]["dirty"] is False
    assert got["workRow"]["text"] == "Scope:  kept from the same draft"
    assert got["collected"] == [{"id": 115, "text": "Scope:  kept from the same draft"}]
    assert got["stored"] == [{"id": 115, "text": "Scope:  kept from the same draft"}]


def test_a_stale_blank_runs_entry_is_dropped_instead_of_re_sent_forever(ran):
    """`runs: [{text: ""}]` — a non-empty array of nothing, which is the shape that does not heal
    itself. The restore skips it, so the DOM never reports that id, and `preserveRichOverrides`
    then treats the array as formatting worth rescuing and pushes the whole entry back in. Left
    alone it is re-sent on every persist for the life of the project, and only the writer's own
    refusal keeps a bare clause number out of the customer's document."""
    got = ran["clauseBlankRuns"]
    assert got["restored"]["text"] == "Agreement. The Proposal of Treadwell, LLC."
    assert got["collected"] == []
    assert got["stored"] == [], (
        "a blank-runs override for a numbered clause is still in the draft: %r" % (got["stored"],))


def test_blank_runs_beside_live_text_are_judged_on_the_runs(ran):
    """The two halves of an override entry can disagree, and `restoreSavedOverrides` renders the
    RUNS — so the runs are what decide whether the clause comes back empty. An entry judged on its
    `text` alone reads as harmless here and blanks the clause on screen anyway."""
    got = ran["clauseRunsDisagree"]
    assert got["restored"]["text"] == "Agreement. The Proposal of Treadwell, LLC."
    assert got["restored"]["empty"] is False
    assert got["stored"] == []


# ══ the CSS claims ═══════════════════════════════════════════════════════════
def test_the_number_on_screen_comes_from_the_documents_own_marker():
    """`content: attr(data-marker)` — the number is the one the backend resolved off `w:lvlText`,
    not a CSS counter that would restart wherever the preview happens to paginate."""
    rule = CSS.split(".tw-num::before", 1)[1].split("}", 1)[0]
    assert "attr(data-marker)" in rule, rule
    box = CSS.split("\n.tw-num {", 1)[1].split("}", 1)[0]
    # The level's own geometry: w:ind left=540 hanging=360 puts the text 27pt in and the number
    # 18pt to its left, i.e. 9pt from the margin.
    assert "margin-left: 9pt" in box and "padding-left: 18pt" in box, box


def test_the_emptied_bullet_rule_no_longer_claims_to_cover_ordered_lists():
    """FIX A IS WHY FIX B WAS INVISIBLE. This rule hid the square on an emptied `.tw-li`, and a
    numbered clause was ALSO a `.tw-li`, so an emptied clause read as a blank line on screen while
    Word printed a bare number. The rule is right for bullets and its comment used to claim more
    than that."""
    marker = ".tw-block.tw-empty.tw-li::before"
    head = CSS[:CSS.index(marker)]
    comment = head[head.rindex("/*"):]
    assert "the preview shows what prints." not in comment, (
        "the comment still makes the general claim that was false for ordered lists")
    assert "tw-num" in comment, (
        "the comment does not say where an ordered list is handled instead: %r" % (comment,))
    # The rule itself is unchanged and still outranks `.tw-li::before`.
    assert CSS.index(".tw-li::before") < CSS.index(marker)


def test_the_clause_notice_has_no_em_dashes_anywhere_the_estimator_reads_it():
    """Repo rule for user-visible frontend copy — the tooltip and the badge both."""
    src = (FRONTEND / "js" / "proposal-review.js").read_text(encoding="utf-8")
    msg = src.split("const _CLAUSE_KEPT_MSG =", 1)[1].split(";", 1)[0]
    assert "cannot be emptied" in msg
    assert "—" not in msg, msg
    badge = CSS.split(".tw-block.tw-clause-kept::after", 1)[1].split("}", 1)[0]
    assert "content:" in badge
    assert "—" not in badge, badge
