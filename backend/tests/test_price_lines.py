"""Generalized repeatable blocks ({{#price_line}}, {{#alternate}}, {{#system}}),
the alternate-system bid on /api/price, and the alternate estimate tab.

The block engine is unit-tested on synthetic docs (no template dependency); the
alternate bid + tab are tested through the real pricing/estimate paths.
"""
import io

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook

import main
import proposal_writer as pw
import estimate_writer as ew

client = TestClient(main.app)


def _doc(lines):
    d = Document()
    for line in lines:
        d.add_paragraph(line)
    return d


def _texts(d):
    return [p.text for p in d.paragraphs]


# ── repeatable block engine ────────────────────────────────────────────
def test_price_line_block_expands_to_n_rows():
    for n in (0, 3, 30):
        items = [{"label": f"Opt {i}", "amount_formatted": f"${i*100}"} for i in range(n)]
        d = _doc(["{{#price_line}}",
                  "{{price_line.amount_formatted}} – {{price_line.label}}",
                  "{{/price_line}}"])
        pw._expand_all_blocks(d, {"price_line": items})
        rows = [t for t in _texts(d) if "–" in t]
        assert len(rows) == n, f"expected {n} rows, got {len(rows)}"
        # markers always stripped, even when n == 0
        assert not any("price_line" in t for t in _texts(d))
        if n:
            assert rows[0] == "$0 – Opt 0"


def test_alternate_block_renders_and_leaves_bare_state_name():
    d = _doc(["{{#alternate}}",
              "ALT: {{alternate.system_name}}",
              "{{alternate.total_formatted}} – {{state_name}} Remodel Tax",
              "{{/alternate}}"])
    pw._expand_all_blocks(d, {"alternate": [{"system_name": "Quartz DB",
                                             "total_formatted": "$80,000"}]})
    txt = "\n".join(_texts(d))
    assert "ALT: Quartz DB" in txt
    # bare {{state_name}} is NOT a per-item key -> left for the flat pass
    assert "$80,000 – {{state_name}} Remodel Tax" in txt
    # flat pass then resolves it
    for p in pw._iter_all_paragraphs(d):
        pw._replace_in_paragraph(p, {"state_name": "Kansas"})
    assert "Remodel Tax" in "\n".join(_texts(d))


def test_empty_block_strips_markers_and_template():
    d = _doc(["before", "{{#price_line}}", "X {{price_line.label}}", "{{/price_line}}", "after"])
    pw._expand_all_blocks(d, {"price_line": []})
    txt = _texts(d)
    assert "before" in txt and "after" in txt
    assert not any("price_line" in t or t.startswith("X ") for t in txt)


def test_two_different_blocks_in_one_container():
    d = _doc(["{{#price_line}}", "{{price_line.label}}", "{{/price_line}}",
              "{{#alternate}}", "ALT {{alternate.system_name}}", "{{/alternate}}"])
    pw._expand_all_blocks(d, {"price_line": [{"label": "Mockup"}],
                              "alternate": [{"system_name": "Quartz"}]})
    txt = "\n".join(_texts(d))
    assert "Mockup" in txt and "ALT Quartz" in txt
    assert "{{#" not in txt and "{{/" not in txt


def test_system_block_regression():
    d = _doc(["{{#system}}", "{{system.name}}: {{system.sf}} SF", "{{/system}}"])
    pw._expand_all_blocks(d, {"system": [{"name": "A", "sf": "100"},
                                         {"name": "B", "sf": "200"}]})
    txt = "\n".join(_texts(d))
    assert "A: 100 SF" in txt and "B: 200 SF" in txt


# ── /api/price alternate bid (conftest bypasses auth) ──────────────────
def test_price_returns_distinct_alternate_bid():
    r = client.post("/api/price", json={
        "systems": [{"name": "MACRO Flake Single Broadcast", "sf": 12000}],
        "full_bid": True,
        "alternate_systems": [{"name": "40-S (Q28) Quartz Double Broadcast", "sf": 12000}],
        "alternate_label": "Recommended Quartz",
    })
    assert r.status_code == 200
    j = r.json()
    assert j["full_bid"]["total_base_bid"] > 0
    assert j["alternate_full_bid"]["total_base_bid"] > 0
    assert j["full_bid"]["total_base_bid"] != j["alternate_full_bid"]["total_base_bid"]
    assert j["alternate"]["material_sub"] > 0
    assert j["alternate"]["label"] == "Recommended Quartz"


def test_price_without_alternate_is_backcompat():
    r = client.post("/api/price", json={
        "systems": [{"name": "MACRO Flake Single Broadcast", "sf": 12000}], "full_bid": True})
    assert r.status_code == 200
    assert "alternate_full_bid" not in r.json()


# ── alternate estimate tab ─────────────────────────────────────────────
def test_fill_estimate_writes_alternate_tab():
    data = ew.fill_estimate({"project_name": "X"},
                            alternate={"sf": 9000, "material_sub": 12345, "label": "Quartz DB"})
    wb = load_workbook(io.BytesIO(data))
    assert ew.ALT_TAB_RENAME in wb.sheetnames        # blank tab renamed
    ws = wb[ew.ALT_TAB_RENAME]
    assert ws[ew.ALT_SF_CELL].value == 9000
    r = ew.ALT_MATERIAL_ROW
    assert ws[f"C{r}"].value == 12345
    assert ws[f"A{r}"].value == "Quartz DB"
    assert "Epoxy" in wb.sheetnames                  # base tab intact


def test_fill_estimate_no_alternate_keeps_blank_tab():
    data = ew.fill_estimate({"project_name": "X"})
    wb = load_workbook(io.BytesIO(data))
    assert ew.ALT_TAB_NAME in wb.sheetnames          # not renamed
    assert ew.ALT_TAB_RENAME not in wb.sheetnames


# ── real annotated Direct templates render the blocks ──────────────────
_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_BASE_VALS = {
    "lump_sum_formatted": "$61,162.00", "tax_amount_formatted": "$2,639.00",
    "state_name": "Kansas", "total_formatted": "$63,801.00",
    "total_label": "$63,801.00 – Total", "system_name": "MACRO", "texture": "OP",
    "epoxy_sf": "12,000", "polish_sf": "8,000", "cove_lf": "250",
    "bid_date_formatted": "6/15/26", "job_name": "J", "city_state": "C",
    "area_description": "x", "disposal": "d", "site_visit_date": "6/15",
    "schedule_notes": "s", "work_description": "w", "scope_notes": "demo + install scope",
    # epoxy PRICE breakdown tokens
    "base_bid_formatted": "$58,523.00", "material_tax_formatted": "$2,639.00",
    "base_tax_phrase": "(material sales tax INCLUDED)",
    "site_visit_phrase": "per site visit on 6/15/26", "exclusions": "standard exclusions",
}


def _rendered(docx_bytes):
    d = Document(io.BytesIO(docx_bytes))
    out = []
    for p in d.element.xpath("//w:p"):
        if any(True for _ in p.iterancestors(f"{_MC}Fallback")):
            continue
        t = "".join(x.text or "" for x in p.xpath(".//w:t")).strip()
        if t:
            out.append(t)
    return "\n".join(out)


def test_direct_templates_render_price_lines_and_alternate():
    import re
    import proposal_writer as pw
    pls = [{"label": "Onsite mockup", "amount_formatted": "$4,200"}]
    alts = [{"system_name": "Urethane", "lump_sum_formatted": "$92,500.00",
             "remodel_tax": "$3,900.00", "total_formatted": "$96,400.00"}]
    for wt in ("epoxy", "combo", "polish"):
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct",
                                          values=_BASE_VALS, price_lines=pls, alternates=alts))
        assert "$63,801" in blob, f"{wt}: base total missing"
        assert "$4,200 – Onsite mockup" in blob, f"{wt}: price line missing"
        assert "ALTERNATE SYSTEM — Urethane" in blob, f"{wt}: alternate missing"
        assert "$96,400.00 – Total" in blob, f"{wt}: alternate total missing"
        assert not re.search(r"\{\{[#/]", blob), f"{wt}: leftover block marker"


def test_epoxy_price_breakdown_kansas_and_remodel_toggle():
    """Epoxy PRICE: Base Bid + Material Sales Tax always shown; Remodel Tax
    only when a remodel row is supplied; never 'Missouri', never 'INCLUDED'."""
    import re
    import proposal_writer as pw
    on = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct",
        values=_BASE_VALS, remodel=[{"amount_formatted": "$2,639.00"}]))
    assert "Base Bid" in on and "Options:" in on
    assert "$58,523.00 – Epoxy flooring as described above (material sales tax INCLUDED)" in on
    assert "$2,639.00 – Material Sales Tax" in on
    assert "Remodel Tax" in on
    assert "Missouri Remodel Tax" not in on
    assert not re.search(r"\{\{[#/]", on)
    off = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS))
    assert "Material Sales Tax" in off          # always shown (transparency)
    assert "Remodel Tax" not in off             # hidden when remodel is off
    assert not re.search(r"\{\{[#/]", off)


def test_direct_templates_backcompat_no_blocks():
    """No price_lines / alternates -> markers stripped, base intact, no stray text."""
    import re
    import proposal_writer as pw
    for wt in ("epoxy", "combo", "polish"):
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct", values=_BASE_VALS))
        assert "$63,801" in blob
        assert "Onsite mockup" not in blob and "ALTERNATE SYSTEM" not in blob
        assert not re.search(r"\{\{[#/]", blob)


def test_epoxy_work_lists_systems():
    """Epoxy WORK section repeats per system: 1 → "System:" (no Option label),
    2 → "Option 1/2:". Never leaves a literal {{#system}} marker."""
    import re
    import proposal_writer as pw
    one = [{"prefix": "System:", "name": "Macro Flake Single Broadcast", "texture": "Orange Peel",
            "sqft": "4,000", "lf_clause": " and 100 LF of epoxy base"}]
    two = [{"prefix": "Option 1:", "name": "Macro Flake Single Broadcast", "texture": "Orange Peel",
            "sqft": "4,000", "lf_clause": " and 100 LF of epoxy base"},
           {"prefix": "Option 2:", "name": "micro Flake Double Broadcast", "texture": "Orange Peel",
            "sqft": "3,000", "lf_clause": ""}]
    t1 = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS, systems=one))
    assert "System:" in t1 and "Macro Flake Single Broadcast" in t1
    assert "Option 1" not in t1 and not re.search(r"\{\{[#/]", t1)
    t2 = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=_BASE_VALS, systems=two))
    assert "Option 1:" in t2 and "Macro Flake Single Broadcast" in t2
    assert "Option 2:" in t2 and "micro Flake Double Broadcast" in t2
    assert "~3,000 SF of epoxy flooring" in t2
    assert not re.search(r"\{\{[#/]", t2)
