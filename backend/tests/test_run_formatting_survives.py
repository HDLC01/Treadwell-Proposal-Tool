"""Hand-applied formatting has to survive the trip into the .docx.

Hanz asked for the Proposals tab to "act like a microsoft word" — edit text and change its
size. Before this, it couldn't: the pipeline was text-only end to end.

  * `serializeBlock` walked INTO the style spans the preview renders, kept the text and threw
    the styling away — so a bold word never left the browser. And because the TEXT was
    unchanged, the block wasn't marked dirty, so the next `refreshDocumentFills()` rewrote its
    innerHTML and erased the edit on screen too. Lost twice over.
  * `_set_paragraph_text` collapses a paragraph to run[0]'s rPr. Measured on the real GC
    Resinous label row: 20 runs mixing 9pt and 8pt with italic and underline came out as ONE
    run. `_normalize_work_label_formatting` then re-split at the first colon, which is why
    "bold through the colon" LOOKED preserved — it was being reconstructed, not carried.
  * `_shrink_overflowing_text_boxes` rewrites every w:sz in a box it thinks overflows.
    Measured: an edited NOTES line came out at 4.5pt. That silently undoes a deliberate size
    on exactly the boxes somebody resizes BECAUSE they overflow.

The last one is the subtle one and the reason `_user_sized_paragraphs` exists: an automatic
shrink that overrides a human choice is worse than a box that overflows, because the person
who chose has no way to see what happened.

2026-08-21, two more of exactly that shape:

  * `_normalize_work_label_formatting` un-bolded every run after the first colon
    UNCONDITIONALLY. Right for the template's own text, wrong for an override: bold applied to
    a phrase inside a Scope / Schedule / Exclusions value showed on screen, survived the
    reload, travelled in the payload, was rebuilt faithfully by `_set_paragraph_runs` — and was
    flattened one pass later. Measured on Direct block 115 with bold on "3-coat system": the
    run SPLIT survived, `w:b` came out `val="0"`. The format toolbar's most-used button was a
    silent no-op in the box estimators edit most. Italic and underline were never touched by
    that pass, so they always survived — the fix is what makes bold behave like them.
  * The exemption register itself was a no-op in production. It held bare `id()`s, and lxml
    frees an element proxy the moment the last Python reference goes, then hands out a NEW
    proxy at a different address. Every walk here is a generator that keeps no references, so
    the ids went stale between recording and consulting them: measured 1 paragraph recorded, 0
    still matching one `iter_editable_blocks` later, and an 11pt choice reaching the .docx at
    7pt. The register now holds the ELEMENTS (`_hand_formatted`), which is what keeps their
    `id()` alive and unique. The older test below passes either way because it holds `p_elem`
    itself; `test_a_chosen_size_survives_a_full_generate` is the one that does not.
"""
import io
import pathlib

import docx
import pytest
from docx.text.paragraph import Paragraph

import proposal_writer as pw

TEMPLATES = pathlib.Path(__file__).resolve().parents[1] / "templates"
GC_RESINOUS = TEMPLATES / "GC" / "xx TREADWELL RESINOUS PROPOSAL - xx.docx"
DIRECT_EPOXY = TEMPLATES / "Direct" / "XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx"
GYP = TEMPLATES / "Gyp" / "xx TREADWELL UNDERLAYMENT PROPOSAL - xx.docx"


def _label_block(d):
    """The mixed-format GC label row — 20 runs of 9pt/8pt with italic and underline."""
    for idx, _kind, p_elem, in_block, text, _txbx in pw.iter_editable_blocks(d):
        if in_block is None and "Resinous Flooring" in (text or ""):
            return idx, p_elem
    pytest.skip("the GC label row moved; this test needs a mixed-format paragraph")


def _runs(p_elem, d):
    return [(r.text, r.bold, r.italic, r.underline,
             r.font.size.pt if r.font.size else None) for r in Paragraph(p_elem, d).runs]


def test_the_fixture_really_is_mixed_format():
    """If the template ever became uniform, every test below would pass without proving
    anything."""
    d = docx.Document(str(GC_RESINOUS))
    _, p_elem = _label_block(d)
    runs = _runs(p_elem, d)
    assert len(runs) > 5, "expected a many-run paragraph"
    assert len({r[4] for r in runs if r[4]}) > 1, "expected more than one size in the paragraph"


# ── the writer ────────────────────────────────────────────────────────
def test_each_run_keeps_its_own_formatting():
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [
        {"text": "Label:", "bold": True, "underline": True, "size_pt": 11.0},
        {"text": " body", "bold": False, "italic": False, "underline": False, "size_pt": 8.0},
        {"text": " aside", "italic": True, "bold": False, "size_pt": 8.0},
    ]}])
    got = _runs(p_elem, d)
    assert len(got) == 3, f"runs were collapsed: {got}"
    assert got[0][1] is True and got[0][3] is True and got[0][4] == 11.0
    assert got[1][1] is False and got[1][4] == 8.0
    assert got[2][2] is True, "the italic aside lost its italic"


def test_an_absent_key_inherits_rather_than_switching_off():
    """`bold: False` and "no bold key" are different instructions. Absent must leave the
    template's own rPr alone, or every unformatted run would strip Kyle's design."""
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    base = _runs(p_elem, d)[0]
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [{"text": "kept"}]}])
    got = _runs(p_elem, d)[0]
    assert got[1] == base[1], "an absent bold key changed the inherited weight"
    assert got[4] == base[4], "an absent size key changed the inherited size"


def test_an_explicit_false_writes_an_explicit_off():
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [{"text": "x", "bold": False}]}])
    assert _runs(p_elem, d)[0][1] is False


def test_a_plain_text_override_still_takes_the_simple_path():
    """Most edits are plain. They must keep working exactly as before — same shape, same
    writer, no behaviour change."""
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    n = pw._apply_paragraph_overrides(d, [{"id": idx, "text": "just text"}])
    assert n == 1
    assert Paragraph(p_elem, d).text == "just text"


# ── the 4.5pt bug ─────────────────────────────────────────────────────
def _overflow_the_box(d, p_elem):
    """Make the paragraph's box genuinely overflow, so the shrink actually engages.

    Without this the test is hollow: `_estimate_txbx_scale` returns ~1.0 for a box that fits,
    `_scale_txbx_runs` is never called, and the assertion passes whether or not the exemption
    exists. Verified by removing the exemption and watching the test still pass — which is why
    the padding is here and why the guard below checks the shrink really ran."""
    import copy as _copy
    box = p_elem.getparent()
    for _ in range(30):
        box.append(_copy.deepcopy(p_elem))


def test_a_chosen_size_survives_the_overflow_shrink():
    """THE regression. The shrink rewrites every w:sz in a box it thinks overflows; measured
    turning an edited NOTES line into 4.5pt. A size the estimator chose has to be exempt."""
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [
        {"text": "Deliberately eleven point", "size_pt": 11.0}]}])
    assert _runs(p_elem, d)[0][4] == 11.0
    _overflow_the_box(d, p_elem)

    # Prove the shrink ENGAGED, or the assertion below means nothing.
    box = p_elem.getparent()
    sizes_before = [sz.get(pw.qn("w:val")) for sz in box.iter(pw.qn("w:sz"))]
    pw._shrink_overflowing_text_boxes(d)
    sizes_after = [sz.get(pw.qn("w:val")) for sz in box.iter(pw.qn("w:sz"))]
    assert sizes_before != sizes_after, (
        "the shrink did not run on this box, so this test proves nothing about the exemption")

    assert _runs(p_elem, d)[0][4] == 11.0, (
        "the automatic shrink overrode a size the estimator chose")


def test_the_shrink_still_shrinks_everything_else():
    """The exemption must be surgical. If it accidentally exempted the whole box, long content
    would start overlapping the next section again — the bug the shrink exists to prevent."""
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    # Stuff a box so it definitely overflows, and size only ONE paragraph.
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [
        {"text": "sized " * 40, "size_pt": 11.0}]}])
    before = {}
    for i, txbx in enumerate(pw._iter_txbx(d)):
        before[i] = [sz.get(pw.qn("w:val")) for sz in txbx.iter(pw.qn("w:sz"))]
    pw._shrink_overflowing_text_boxes(d)
    after = {}
    for i, txbx in enumerate(pw._iter_txbx(d)):
        after[i] = [sz.get(pw.qn("w:val")) for sz in txbx.iter(pw.qn("w:sz"))]
    assert any(before[i] != after[i] for i in before), (
        "nothing shrank at all — the exemption is too broad")


# ── through the whole pipeline ─────────────────────────────────────────
def test_formatting_survives_a_full_generate():
    """Piece-by-piece passing is not enough: `_normalize_work_label_formatting` runs AFTER
    overrides and re-splits label rows at the first colon, and the shrink runs after that."""
    d0 = docx.Document(str(GC_RESINOUS))
    idx, _ = _label_block(d0)
    blob = pw.fill_proposal(
        work_type="polish", audience="GC",
        values={"job_name": "Test", "city_state": "Lenexa, KS"},
        paragraph_overrides=[{"id": idx, "runs": [
            {"text": "Label:", "bold": True, "underline": True, "size_pt": 11.0},
            {"text": " body at eleven", "bold": False, "italic": False,
             "underline": False, "size_pt": 11.0},
        ]}],
    )
    out = docx.Document(io.BytesIO(blob))
    for _i, _k, p_elem, _ib, text, _t in pw.iter_editable_blocks(out):
        if "Label:" in (text or ""):
            got = _runs(p_elem, out)
            assert got[0][1] is True and got[0][4] == 11.0, got
            assert any(r[1] is False for r in got[1:]), (
                "the body run was re-bolded — normalize_work_label_formatting clobbered it")
            return
    pytest.fail("the overridden paragraph is not in the generated document")


# ── never 500 on bad input ────────────────────────────────────────────
@pytest.mark.parametrize("bad", [
    {"id": 0, "runs": "nope"},
    {"id": 0, "runs": []},
    {"id": 0, "runs": [{"nope": 1}]},
    {"id": 0, "runs": [{"text": 5}]},
    {"id": 0, "runs": [{"text": "x", "size_pt": "big"}]},
    {"id": 0, "runs": [{"text": "x", "size_pt": 0}]},
    {"id": 0, "runs": [{"text": "x", "size_pt": 9999}]},
    {"id": 0, "runs": [{"text": "x", "bold": "yes"}]},
    {"id": 0, "runs": [None]},
])
def test_malformed_runs_never_raise(bad):
    """A stale draft or a hand-built request must not 500 /api/generate — the house rule for
    every sanitiser in this codebase."""
    d = docx.Document(str(GC_RESINOUS))
    pw._apply_paragraph_overrides(d, [bad])          # must not raise


def test_a_bad_size_falls_back_rather_than_writing_nonsense():
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    base_size = _runs(p_elem, d)[0][4]
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [
        {"text": "x", "size_pt": 9999}]}])
    got = _runs(p_elem, d)[0][4]
    assert got != 9999, "an absurd size was written straight through"
    assert got == base_size, "a rejected size should leave the inherited one"


# ── the WORK box: hand formatting vs the label normalizer ─────────────
#
# Every test in this section goes through the WHOLE of `fill_proposal`, because the bug lived
# in the ORDER of the passes: each piece was correct on its own and the last one undid the
# first. Asserting on the raw `w:b`/`w:i`/`w:u` rather than on python-docx's tri-state keeps
# the record of what actually reached the customer's document.
_WORK_VALUES = {"job_name": "Test", "city_state": "Lenexa, KS",
                "scope_notes": "template scope text",
                "schedule_notes": "template schedule text",
                "exclusions": "template exclusions text"}


def _work_row(path):
    """(id, text) of the first Scope/Schedule/Exclusions row in a template's WORK box."""
    d = docx.Document(str(path))
    for idx, _k, _p, in_block, text, _tx in pw.iter_editable_blocks(d):
        if in_block is None and pw._WORK_ANCHOR_RE.match((text or "").strip()):
            return idx, text
    pytest.skip(f"no WORK label row in {path.name}; this test needs one")


def _switches(p_elem, d):
    """[(text, w:b, w:i, w:u)] per run, straight off the rPr. "-" is absent (inherit)."""
    out = []
    for r in Paragraph(p_elem, d).runs:
        rpr = r._r.find(pw.qn("w:rPr"))

        def val(tag, rpr=rpr):
            if rpr is None:
                return "-"
            el = rpr.find(pw.qn(tag))
            return "-" if el is None else (el.get(pw.qn("w:val")) or "on")
        out.append((r.text, val("w:b"), val("w:i"), val("w:u")))
    return out


def _generated_row(path, work_type, audience, overrides, needle):
    blob = pw.fill_proposal(work_type=work_type, audience=audience,
                            values=_WORK_VALUES, paragraph_overrides=overrides)
    out = docx.Document(io.BytesIO(blob))
    for _i, _k, p_elem, _ib, text, _tx in pw.iter_editable_blocks(out):
        if needle in (text or ""):
            return _switches(p_elem, out)
    pytest.fail(f"{needle!r} is not in the generated document")


ON = ("on", "1")          # Word writes bold as a bare <w:b/> OR <w:b w:val="1"/>


@pytest.mark.parametrize("path,work_type,audience", [
    (DIRECT_EPOXY, "epoxy", "Direct"),
    (GC_RESINOUS, "polish", "GC"),
    (GYP, "gyp", "Gyp"),
])
def test_bold_inside_a_work_value_survives(path, work_type, audience):
    """THE regression. Bold on a phrase inside a Scope/Schedule/Exclusions value must reach the
    .docx. Measured before the fix on all three template families: `w:b val="0"`."""
    idx, text = _work_row(path)
    label = text.split(":")[0] + ": "
    got = _generated_row(path, work_type, audience, [{
        "id": idx, "text": label + "install a 3-coat system today",
        "runs": [{"text": label, "bold": True},
                 {"text": "install a ", "bold": False},
                 {"text": "3-coat system", "bold": True},
                 {"text": " today", "bold": False}]}], "3-coat system")
    phrase = [r for r in got if r[0] == "3-coat system"]
    assert phrase, f"the estimator's run split did not survive: {got}"
    assert phrase[0][1] in ON, f"the label normalizer flattened the estimator's bold: {got}"
    assert [r for r in got if r[0] == "install a "][0][1] == "0", (
        f"the surrounding text was re-bolded: {got}")


def test_the_label_is_still_bold_with_no_override_at_all():
    """The pass has to keep doing its job on the pristine template — that is what it is for."""
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct", None, "template scope text")
    assert got[0][1] in ON, f"the template's own label lost its bold: {got}"
    assert got[-1][1] == "0", f"the template's own value came out bold: {got}"


def test_a_retyped_plain_text_row_still_gets_its_label_bolded():
    """A plain-TEXT override registers no runs, so nothing is exempt and the row is normalized
    exactly as before. This is the case that keeps Kyle's label bold on a row he just retyped,
    and the reason the exemption is per RUN and not per paragraph."""
    idx, _ = _work_row(DIRECT_EPOXY)
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct",
                         [{"id": idx, "text": "Scope: retyped by hand entirely"}],
                         "retyped by hand")
    assert got[0][0] == "Scope:" and got[0][1] in ON, f"the label was not bolded: {got}"
    assert got[-1][1] == "0", f"the retyped value came out bold: {got}"


def test_an_explicitly_unbolded_label_stays_unbold():
    """The other half of "the toolbar works": un-bolding the label is as deliberate as bolding
    the value, and the normalizer must not put it back."""
    idx, _ = _work_row(DIRECT_EPOXY)
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct", [{
        "id": idx, "text": "Scope: deliberately unbold label",
        "runs": [{"text": "Scope: ", "bold": False},
                 {"text": "deliberately unbold label", "bold": False}]}],
        "deliberately unbold")
    assert all(r[1] == "0" for r in got), f"the normalizer re-bolded the label: {got}"


def test_a_weight_stated_on_part_of_the_label_holds_too():
    """A run that sits ENTIRELY before the colon takes a different branch of the pass than the
    run containing it, and it needs the same exemption. Reachable for real: the toolbar formats
    a selection, and a selection can be half a label. Without this the test suite passed with
    the before-colon branch un-exempted."""
    idx, _ = _work_row(DIRECT_EPOXY)
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct", [{
        "id": idx, "text": "Scope: partial label weight",
        "runs": [{"text": "Sco", "bold": False},
                 {"text": "pe: ", "bold": True},
                 {"text": "partial label weight", "bold": False}]}], "partial label weight")
    assert got[0][0] == "Sco" and got[0][1] == "0", (
        f"the un-bolded first half of the label was re-bolded: {got}")
    assert got[1][1] in ON, f"the bolded second half of the label lost its bold: {got}"


def test_a_runs_override_that_states_no_weight_still_gets_the_auto_label_bold():
    """Formatting something OTHER than the weight (a size, say) must not cost the row its
    automatic label bold — only an explicit weight outranks the pass."""
    idx, _ = _work_row(DIRECT_EPOXY)
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct", [{
        "id": idx, "text": "Scope: sized but unweighted",
        "runs": [{"text": "Scope: ", "size_pt": 9.0},
                 {"text": "sized but unweighted", "size_pt": 9.0}]}], "sized but unweighted")
    assert got[0][1] in ON, f"the label lost its automatic bold: {got}"
    assert got[-1][1] == "0", f"the value came out bold: {got}"


def test_italic_and_underline_inside_a_work_value_survive():
    """They always did — the pass only ever wrote `w:b`. Pinned so bold cannot regress away
    from them again, and so a future pass that normalizes a second switch has to face this."""
    idx, _ = _work_row(DIRECT_EPOXY)
    got = _generated_row(DIRECT_EPOXY, "epoxy", "Direct", [{
        "id": idx, "text": "Scope: an italic aside and an underlined one",
        "runs": [{"text": "Scope: "},
                 {"text": "an italic aside", "italic": True},
                 {"text": " and ", "underline": False},
                 {"text": "an underlined one", "underline": True}]}], "an italic aside")
    ital = [r for r in got if r[0] == "an italic aside"]
    under = [r for r in got if r[0] == "an underlined one"]
    assert ital and ital[0][2] == "1", f"the italic aside lost its italic: {got}"
    assert under and under[0][3] == "single", f"the underlined phrase lost its underline: {got}"
    assert got[0][1] in ON, f"the label lost its bold: {got}"


# ── the register has to still be valid when it is consulted ────────────
def test_the_hand_format_register_is_not_stale_after_the_walk():
    """A bare set of id()s was measured going 1-recorded / 0-matching across one re-walk,
    because lxml frees the proxy and reissues it at a new address. Holding the element is the
    whole mechanism, so assert the property directly rather than only its consequence."""
    d = docx.Document(str(GC_RESINOUS))
    idx, p_elem = _label_block(d)
    pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [
        {"text": "Label:", "bold": True, "size_pt": 11.0},
        {"text": " body", "bold": False}]}])
    sized, bolded = pw._user_sized_paragraphs(d), pw._user_bolded_runs(d)
    assert len(sized) == 1 and len(bolded) == 2, (sized, bolded)
    del p_elem                                   # drop every reference the caller was holding
    still_sized = sum(1 for _i, _k, p, _b, _t, _x in pw.iter_editable_blocks(d)
                      if id(p) in sized)
    still_bolded = sum(1 for r in d.element.body.iter(pw.qn("w:r")) if id(r) in bolded)
    assert still_sized == 1, "the sized-paragraph register went stale after the walk"
    assert still_bolded == 2, "the bolded-run register went stale after the walk"


def test_a_chosen_size_survives_a_full_generate():
    """The shrink exemption, with NO reference held to the paragraph — which is the situation
    inside `fill_proposal` and the reason the register looked fine in a unit test while doing
    nothing in production. Measured with the un-pinned register: 11pt in, `w:sz` 14 (7pt) out.
    The gyp WORK box is the one that genuinely overflows on long content."""
    idx, _ = _work_row(GYP)
    long_text = ("Exclusions: " + ("ELEVENPOINT a deliberately very long exclusions value that "
                 "has to overflow this fixed box so the shrink actually engages. ") * 6)
    blob = pw.fill_proposal(work_type="gyp", audience="Gyp", values=_WORK_VALUES,
                            paragraph_overrides=[{"id": idx, "text": long_text,
                                                  "runs": [{"text": long_text,
                                                            "size_pt": 11.0}]}])
    out = docx.Document(io.BytesIO(blob))
    for _i, _k, p_elem, _ib, text, _tx in pw.iter_editable_blocks(out):
        if "ELEVENPOINT" in (text or ""):
            box = p_elem.getparent()
            sizes = {sz.get(pw.qn("w:val")) for sz in box.iter(pw.qn("w:sz"))}
            assert sizes - {"22"}, "nothing else in the box shrank, so this proves nothing"
            got = [r.font.size.pt for r in Paragraph(p_elem, out).runs if r.font.size]
            assert got and all(s == 11.0 for s in got), (
                f"the automatic shrink overrode the estimator's 11pt: {got}")
            return
    pytest.fail("the overridden gyp row is not in the generated document")


def test_media_runs_are_never_removed():
    """Kyle's templates anchor the letterhead AND every floating text box in runs of otherwise
    blank paragraphs. Dropping those would delete the letterhead from a customer document."""
    d = docx.Document(str(GC_RESINOUS))
    before = len(list(d.element.body.iter(pw.qn("w:drawing"))))
    assert before, "no drawings in the fixture — this test would prove nothing"
    for idx, _k, p_elem, in_block, _t, _tx in pw.iter_editable_blocks(d):
        if in_block is None and next(p_elem.iter(pw.qn("w:drawing")), None) is not None:
            pw._apply_paragraph_overrides(d, [{"id": idx, "runs": [{"text": "replaced"}]}])
            break
    assert len(list(d.element.body.iter(pw.qn("w:drawing")))) == before
