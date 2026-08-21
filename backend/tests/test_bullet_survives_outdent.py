"""A bulleted row that gets outdented must still print its square.

Kyle, via Hanz: "There is indentation in this but I cant remove tat if I want to to be aligned on
the polished concrete?" One press of the outdent button on a WORK row is the exact flow the
paragraph controls were built for, and it used to delete that row's red square from the customer's
document.

THE MECHANISM, because the numbers are the only way to see this bug. A bulleted paragraph's square
prints at `w:left - w:hanging`. numId 4's level is `left=288 hanging=288`, so every untouched WORK
row prints its square at x=0, flush with the margin, with the text 288 twips in. The editor sends
`indent: 0` for one outdent press, and `indent` means the TEXT's left edge everywhere in this
feature (`_effective_left_tw` reads w:left). The old code wrote `w:left="0"` and, because it passed
a hanging of `min(288, 0) = 0` into a writer that stored the hanging under a plain truthiness test,
dropped `w:hanging` altogether. The paragraph kept its `w:numPr`, so the LEVEL's hanging of 288
applied to a left of 0:

    square x = 0 - 288 = -288 twips

which is out in the page margin, where Word draws nothing. Measured with the same harness this test
uses, before and after:

    BEFORE  left=0    hang=None  squareX=-288   <- one square that will not print
    AFTER   left=288  hang=288   squareX=0      <- in line with all of its neighbours

Meanwhile the editor kept drawing the square (applyParaToEl keeps the tw-li class) and
`para_props()` re-read from the generated file kept reporting `bullet: True`. Three readers, three
answers, and the only one the customer sees was the wrong one.

So the floor for a still-bulleted row is its own hanging. These tests pin the square's POSITION
rather than the XML, because the position is what the requirement is about: a row the estimator
outdented has to line up with the rows it sits between.
"""

import io
import re
import zipfile

import docx
import pytest

import main as M
import proposal_writer as PW

VALUES = {"job_name": "Outdent Probe", "city_state": "Olathe, KS",
          "system_name": "Epoxy", "sqft": "4000", "lump_sum": "$40,000"}
SYSTEMS = [{"system_name": "Epoxy", "texture": "Orange Peel"}]

# The four bulleted, editable, non-block WORK rows of the Direct Epoxy template.
WORK_ROW = 115


def _generate(overrides):
    """Generate through the REAL sanitizer, so the payload is the one /api/generate would see."""
    return PW.fill_proposal(
        work_type="epoxy", audience="Direct", values=VALUES, systems=SYSTEMS,
        paragraph_overrides=M._sanitize_paragraph_overrides(overrides))


def _paragraphs(docx_bytes):
    xml = zipfile.ZipFile(io.BytesIO(docx_bytes)).read("word/document.xml").decode("utf-8")
    # Split on paragraph STARTS. A text box lives inside a paragraph, so a non-greedy
    # <w:p>...</w:p> match would stop at the first inner close and truncate the outer one.
    out = []
    for m in re.finditer(r"<w:p[ >]", xml):
        out.append(xml[m.start():xml.find("</w:p>", m.start()) + 6])
    return out


def _level_geometry(docx_bytes, num_id="4", ilvl="0"):
    d = docx.Document(io.BytesIO(docx_bytes))
    lvl = (PW._numbering_levels(d).get((num_id, ilvl)) or {}).get("ind") or {}
    left = lvl.get("left", lvl.get("start"))
    return (int(left) if left is not None else None,
            int(lvl["hanging"]) if lvl.get("hanging") is not None else None)


def _square_positions(docx_bytes, num_id="4"):
    """Where each numId-`num_id` paragraph will actually draw its bullet, in twips.

    None means "nothing can be computed", which prints nothing just as surely as a negative.
    """
    lvl_left, lvl_hang = _level_geometry(docx_bytes, num_id)
    out = []
    for p in _paragraphs(docx_bytes):
        m = re.search(r'<w:numId w:val="(\d+)"', p)
        if not m or m.group(1) != num_id:
            continue
        ind = re.search(r"<w:ind([^>]*?)/?>", p)
        left = hang = None
        if ind:
            lm = re.search(r'w:(?:left|start)="(-?\d+)"', ind.group(1))
            hm = re.search(r'w:hanging="(-?\d+)"', ind.group(1))
            left = int(lm.group(1)) if lm else None
            hang = int(hm.group(1)) if hm else None
        left = left if left is not None else lvl_left
        hang = hang if hang is not None else lvl_hang
        out.append(None if left is None or hang is None else left - hang)
    return out


def test_the_untouched_template_prints_every_work_square_at_the_margin():
    """The baseline the outdented row has to match. If this ever changes, the numbers in the
    docstring above stop being the right ones and the rest of this file needs re-reading."""
    b = _generate([])
    left, hang = _level_geometry(b)
    assert (left, hang) == (288, 288), (left, hang)
    xs = _square_positions(b)
    assert xs, "no numId 4 paragraphs found at all"
    assert set(xs) == {0}, xs


@pytest.mark.parametrize("indent", [0, 1, 100, 144, 287])
def test_an_outdent_below_the_hanging_still_prints_the_square(indent):
    """Every indent the toolbar can ask for below the level's hanging. Zero is the one the outdent
    button actually sends, and the others are here because a floor that only holds at exactly 0 is
    not a floor."""
    b = _generate([{"id": WORK_ROW, "para": {"bullet": True, "indent": indent}}])
    xs = _square_positions(b)
    unprintable = [x for x in xs if x is None or x < 0]
    assert not unprintable, (
        "indent=%d leaves %d square(s) at a position Word will not draw: %r"
        % (indent, len(unprintable), xs))
    # and it must line up with its neighbours, not merely exist
    assert set(xs) == {0}, (
        "indent=%d moved one square out of line with the rows around it: %r" % (indent, xs))


def test_the_outdented_row_keeps_its_numbering_and_gains_a_real_hanging():
    """The XML, once, so a future reader can see what the position above is made of."""
    b = _generate([{"id": WORK_ROW, "para": {"bullet": True, "indent": 0}}])
    hit = [p for p in _paragraphs(b)
           if re.search(r'<w:numId w:val="4"', p) and re.search(r"<w:ind[^>]*w:left", p)]
    assert len(hit) == 1, "expected exactly one row to state its own indent, got %d" % len(hit)
    ind = re.search(r"<w:ind([^>]*?)/?>", hit[0]).group(1)
    assert re.search(r'w:left="288"', ind), ind
    assert re.search(r'w:hanging="288"', ind), (
        "w:hanging is missing, so the level's own hanging applies to a left this feature set and "
        "the square lands in the margin: %s" % ind)


@pytest.mark.parametrize("indent", [576, 720, 1440])
def test_an_indent_above_the_hanging_is_untouched_by_the_floor(indent):
    """The floor must not disturb the direction that already worked. The row moves right, its
    square moves with it, and its neighbours stay where they were."""
    b = _generate([{"id": WORK_ROW, "para": {"bullet": True, "indent": indent}}])
    xs = _square_positions(b)
    assert None not in xs and min(xs) >= 0, xs
    moved = [x for x in xs if x != 0]
    assert moved == [indent - 288], (
        "expected exactly one square at %d and the rest at 0, got %r" % (indent - 288, xs))


def test_a_row_that_loses_its_bullet_puts_its_text_at_the_margin():
    """The floor applies to BULLETED rows only. Turning the bullet off is how an estimator gets the
    text itself flush left, and it must still do that - otherwise the floor has taken away the one
    way to reach the margin."""
    b = _generate([{"id": WORK_ROW, "para": {"bullet": False, "indent": 0}}])
    fours = _square_positions(b)
    # one fewer numId 4 paragraph than the baseline: this row left the list
    assert len(fours) == len(_square_positions(_generate([]))) - 1, fours
    assert all(x == 0 for x in fours), fours

    # AND its text really is at zero. Asserting only the square count let a mutation that applied
    # the bulleted floor to un-bulleted rows as well pass unnoticed - the squares were all still
    # fine, and the estimator's flush-left row had quietly been pushed 288 twips right. The floor
    # exists to protect a square; a row with no square has nothing to protect.
    rows = [par for par in _paragraphs(b)
            if "Scope:" in "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", par))
            and not re.search(r"<w:numId", par)]
    assert rows, "the un-bulleted Scope row is gone entirely"
    ind = re.search(r"<w:ind([^>]*?)/?>", rows[0])
    assert ind, "the un-bulleted row states no indent at all: %r" % rows[0][:200]
    left = re.search(r'w:(?:left|start)="(-?\d+)"', ind.group(1))
    assert left and int(left.group(1)) == 0, (
        "a row with no bullet must reach the margin, got %s" % ind.group(1))


def test_the_terms_clauses_are_untouched_by_any_of_this():
    """27 numbered clauses, before and after an outdent on a WORK row. The renumbering guard is why
    the release this ships in exists, so it gets re-asserted from here too."""
    base = _generate([])
    after = _generate([{"id": WORK_ROW, "para": {"bullet": True, "indent": 0}}])
    n_base = len([p for p in _paragraphs(base) if re.search(r'<w:numId w:val="5"', p)])
    n_after = len([p for p in _paragraphs(after) if re.search(r'<w:numId w:val="5"', p)])
    assert n_base == n_after, (n_base, n_after)
    assert n_base >= 27, n_base


def test_a_zero_hanging_is_written_rather_than_silently_dropped():
    """The writer-level half of the bug, tested directly.

    `_write_left_indent` stored the hanging under `if hanging_tw:`, so a request for zero was
    indistinguishable in the output from "no hanging at all" - and those two mean opposite things
    on a paragraph that carries numbering. No caller asks for zero any more, but the writer must
    not be the thing that decides that.
    """
    from docx.oxml.ns import qn

    d = docx.Document()
    p = d.add_paragraph("x")
    ppr = PW._get_or_make_ppr(p._p)

    PW._write_left_indent(ppr, 400, 0)
    ind = ppr.find(qn("w:ind"))
    assert ind.get(qn("w:hanging")) == "0", (
        "a hanging of zero was dropped: %r" % (dict(ind.attrib),))

    PW._write_left_indent(ppr, 400, None)
    assert ind.get(qn("w:hanging")) is None, (
        "None must still clear the attribute, as the docstring promises: %r" % (dict(ind.attrib),))
