"""Gyp (Gypsum Underlayment) work type — template annotation + backend narrative.

The gyp underlayment proposal uses ONE template regardless of audience
(TEMPLATE_PICKER[("gyp", None)]). annotate_templates.py tokenizes its floating
text boxes; main._ensure_value_aliases seeds gyp-specific scope/exclusions +
the gyp-only tokens (SF buckets, thicknesses, mobilization line). Gyp is
mobilization-based — its NOTES carry a "$8,600 / $10,800 per mob" bullet, NOT
the "$4,500 per additional phase" bullet, and it has no phase cell.

Covers:
  (a) the annotated template has every gyp token exactly 2x (mc:Choice +
      mc:Fallback) and zero leftover placeholders / static price rows;
  (b) pick_template gyp -> the underlayment doc for all audiences;
  (c) _ensure_value_aliases gyp defaults (scope/exclusions/schedule + gyp
      SF/thickness/mobilization backfills + forced work_description);
  (d) _notes_for("gyp") = the 14-item box, no "$xxxx" phase placeholder;
  (e) /api/default-notes?work_type=gyp;
  (f) E2E /api/generate (gyp, Branson values) -> a docx with the 3 SFs, the job
      name, the total, the "$8,600" mob note, and ZERO raw {{tokens}} / "$x".
"""
import io
import re
import zipfile

import docx
from docx import Document
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_GYP_REL = "Gyp/xx TREADWELL UNDERLAYMENT PROPOSAL - xx.docx"

# Every {{token}} the annotated gyp template carries, each expected exactly 2x
# (the letterhead art means each text box exists twice: mc:Choice + mc:Fallback).
_GYP_TOKENS = [
    "{{base_bid_formatted}}", "{{bid_date_formatted}}", "{{city_state}}",
    "{{estimator_name}}", "{{exclusions}}", "{{gyp_soft_sf}}", "{{gyp_hard_sf}}",
    "{{gyp_corridor_sf}}", "{{gyp_soft_thickness}}", "{{gyp_hard_thickness}}",
    "{{gyp_corridor_thickness}}", "{{job_name}}", "{{material_tax_formatted}}",
    "{{mobilizations_line}}", "{{tax_amount_formatted}}", "{{total_formatted}}",
    "{{work_description}}", "{{#notes}}", "{{/notes}}", "{{notes.text}}",
    "{{#has_options}}", "{{/has_options}}", "{{#price_line}}", "{{/price_line}}",
    "{{price_line.amount_formatted}}", "{{price_line.label}}",
]

# Placeholders / static rows that MUST be gone after annotation.
_GYP_LEFTOVERS = [
    "Greg Ingebretson",   # estimator placeholder -> {{estimator_name}}
    "7/1/26",             # header date -> {{bid_date_formatted}}
    "xx, KS", "xx, MO",   # city lines -> {{city_state}} / dropped
    "xx sf",              # SF placeholder -> {{gyp_*_sf}}
]


def _rendered(docx_bytes):
    """Text of the paragraphs Word renders (mc:Choice), skipping the VML
    mc:Fallback duplicate."""
    d = Document(io.BytesIO(docx_bytes))
    out = []
    for p in d.element.xpath("//w:p"):
        if any(True for _ in p.iterancestors(f"{_MC}Fallback")):
            continue
        t = "".join(x.text or "" for x in p.xpath(".//w:t")).strip()
        if t:
            out.append(t)
    return "\n".join(out)


def _doc_xml(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        return z.read("word/document.xml").decode("utf-8")


def _template_xml(rel_path):
    with zipfile.ZipFile(str(pw.TEMPLATES_ROOT / rel_path)) as z:
        return z.read("word/document.xml").decode("utf-8")


# Real project: Branson Meadows Home2Suites (soft 27,825 / hard 11,795 /
# corridor 5,655). Money tokens supplied like the frontend's computeTokenValues.
def _gyp_vals(**over):
    v = {
        "work_type": "gyp",
        "project_name": "Branson Meadows Home2Suites", "job_name": "Branson Meadows Home2Suites",
        "city_state": "Branson, MO", "estimator_name": "Kyle",
        "bid_date_formatted": "7/10/26",
        "gyp_soft_sf": "27,825", "gyp_hard_sf": "11,795", "gyp_corridor_sf": "5,655",
        "base_bid_formatted": "$98,000.00", "material_tax_formatted": "$5,364.00",
        "tax_amount_formatted": "$0.00", "total_formatted": "$103,364.00",
    }
    v.update(over)
    return v


# ── (a) template annotation ─────────────────────────────────────────────
def test_gyp_template_has_every_token_exactly_twice():
    xml = _template_xml(_GYP_REL)
    for tok in _GYP_TOKENS:
        assert xml.count(tok) == 2, f"{tok} count={xml.count(tok)} (want 2)"


def test_gyp_template_has_no_leftover_placeholders():
    xml = _template_xml(_GYP_REL)
    for s in _GYP_LEFTOVERS:
        assert s not in xml, f"leftover placeholder still present: {s!r}"
    # No bare price placeholders (the "$x –" / "$  x" rows became money tokens).
    assert "$x" not in xml
    assert "$  x" not in xml


def test_gyp_template_opens_cleanly():
    d = docx.Document(str(pw.TEMPLATES_ROOT / _GYP_REL))
    assert len(d.paragraphs) >= 0  # parses without raising


# ── (b) pick_template ────────────────────────────────────────────────────
def test_pick_template_gyp_is_underlayment_for_all_audiences():
    for aud in ("Direct", "GC", None):
        p = pw.pick_template("gyp", aud)
        assert p.name == "xx TREADWELL UNDERLAYMENT PROPOSAL - xx.docx", aud


# ── (c) _ensure_value_aliases gyp defaults ───────────────────────────────
def test_gyp_narrative_and_token_defaults():
    v = {"work_type": "gyp"}
    main._ensure_value_aliases(v, "Direct")
    assert v["scope_notes"] == main._DEFAULT_SCOPE_GYP
    assert v["exclusions"] == main._DEFAULT_EXCLUSIONS_GYP
    assert v["schedule_notes"] == main._DEFAULT_SCHEDULE
    # Blank SFs -> "0"; thickness + mobilization defaults; forced work_description.
    assert v["gyp_soft_sf"] == "0"
    assert v["gyp_soft_thickness"] == '3/4"'
    assert v["gyp_hard_thickness"] == '1"'
    assert v["gyp_corridor_thickness"] == '3/4"'
    assert v["mobilizations_line"] == "1 Mobilization to Site."
    assert v["work_description"] == "per plans & specifications provided"


def test_gyp_defaults_same_for_gc_audience():
    # Gyp uses one template regardless of audience -> the gyp branch fires BEFORE
    # is_gc, so a GC gyp job still gets the gyp (not GC-Resinous) wording.
    v = {"work_type": "gyp"}
    main._ensure_value_aliases(v, "GC")
    assert v["scope_notes"] == main._DEFAULT_SCOPE_GYP
    assert v["exclusions"] == main._DEFAULT_EXCLUSIONS_GYP


def test_gyp_supplied_sf_passes_through_unformatted():
    # A non-blank SF the frontend already comma-formatted is NOT re-touched.
    v = {"work_type": "gyp", "gyp_soft_sf": "27,825"}
    main._ensure_value_aliases(v, "Direct")
    assert v["gyp_soft_sf"] == "27,825"


# ── (d) notes: gyp = mobilization-based, no phase bullet ──────────────────
def test_notes_for_gyp_has_mob_line_and_no_phase_placeholder():
    # _notes_for returns [{"text": ...}, ...]; pull the text for the assertions.
    texts = [n["text"] for n in main._notes_for("gyp", [])]
    assert len(texts) == 14
    assert not any("xxxx" in t for t in texts), "gyp notes must have no $xxxx placeholder"
    assert any("Add $8,600" in t and "Add $10,800" in t for t in texts)
    # No "additional phase" bullet — gyp prices mobilizations, not phases.
    assert not any("additional phase" in t for t in texts)


# ── (e) /api/default-notes ────────────────────────────────────────────────
def test_default_notes_endpoint_gyp():
    r = client.get("/api/default-notes?work_type=gyp")
    assert r.status_code == 200, r.text
    j = r.json()
    assert j["work_type"] == "gyp"
    assert len(j["notes"]) == 14
    assert any("Add $8,600" in n for n in j["notes"])


# ── (f) E2E generate ──────────────────────────────────────────────────────
def test_generate_gyp_end_to_end():
    r = client.post("/api/generate", json={"work_type": "gyp", "audience": "Direct",
                                           "values": _gyp_vals()})
    assert r.status_code == 200, r.text
    docx_bytes = client.get(r.json()["docx_download_url"]).content
    txt = _rendered(docx_bytes)
    xml = _doc_xml(docx_bytes)

    # The 3 SF buckets + job name + total render.
    assert "27,825" in txt
    assert "11,795" in txt
    assert "5,655" in txt
    assert "Branson Meadows Home2Suites" in txt
    assert "$103,364" in txt
    # Mobilization note present (gyp's phase-equivalent bullet).
    assert "Add $8,600" in txt
    # No raw tokens or bare price placeholders leaked to the customer doc.
    assert "{{" not in xml and "}}" not in xml
    assert "$x" not in xml


def test_gyp_base_bid_not_labeled_tax_included():
    """Gyp itemizes Base + Material Sales Tax + Kansas Remodel Tax + Total, so the
    base line must NOT claim "(material sales tax INCLUDED)" — that contradicts the
    separately-added tax lines (Base + taxes = Total). Regression: the gyp template
    hardcoded the phrase instead of the {{base_tax_phrase}} token the code blanks."""
    r = client.post("/api/generate", json={"work_type": "gyp", "audience": "Direct",
                                            "values": _gyp_vals()})
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    assert "Gypsum Underlayment System as described above" in txt   # base line present
    assert "material sales tax INCLUDED" not in txt                 # but NOT the contradictory label
    assert "Material Sales Tax" in txt                              # tax is itemized instead


def test_proposal_template_endpoint_gyp():
    r = client.get("/api/proposal-template?work_type=gyp&audience=Direct")
    assert r.status_code == 200, r.text
    # The editor gets ordered blocks; the notes region is a block group.
    body = r.json()
    assert isinstance(body.get("blocks"), list) and body["blocks"]
