"""The last locked WORDS on the proposal page, and a text box that grows instead of clipping.

Kyle, 2026-08-19 and again 2026-08-20 and again 2026-08-24, on the proposal document editor:
    "Some of the labels are not editable why not make it like a word document??"
    "Everything on that page must be editable like a word doc"
    "do not make them as subsections to edit but as a whole section you could edit"
    "I cant delete SF of epoxy flooring"

WHAT WAS ACTUALLY LOCKED, measured against the template files rather than assumed. Every label
in every proposal template is real docx text — only the page frame, the rotated WORK/PRICE rail
captions and the logo are pictures. "Scope:", "Schedule:", "Exclusions:" and "Notes:" were
already inside contenteditable paragraphs. The genuinely locked text is the three rows inside
the Direct epoxy template's `{{#system}}` region:

    {{system.prefix}}   {{system.name}}          ← "System:" / "Option N:", a COMPUTED token
    Texture:  {{system.texture}}                 ← static template text
    Area: ~{{system.sqft}} SF of epoxy flooring… ← static template text

The first answer gave each of those an editable ISLAND — the label got one, the value got one —
and that is what Kyle rejected. An island model leaves the words BETWEEN the islands dead: in the
Area row the "~", the phrase " SF of epoxy flooring" and the whole cove clause had no element to
put a caret in, so "I cant delete SF of epoxy flooring" was literally true. Nothing in either
half of the system could write that text.

WHAT IT IS NOW. One line, one editable region, every word typeable — the model the PRICE lines
(base bid, tax rows, options, alternate) already use. The whole row is stored as one string:
`name_line` / `texture_line` / `area_line`, on the same per-index `system_overrides` channel the
narrower fields used, and written by `proposal_writer._apply_system_row_line`. The row is found
by the TOKEN it carries, so a re-authored template still works. The older per-field keys
(`prefix`, `texture_label`, `area_label`, `name`, `texture`, `sqft`) are still honoured, because
a draft saved under the island editor carries them and throwing them away would delete text the
estimator typed.

THE REGION STILL CANNOT BE EDITED PARAGRAPH-BY-PARAGRAPH, which is why a per-row channel has to
exist at all: a `{{#block}}` region is expanded once per priced system at generate time, so its
paragraph ids stop describing anything the estimator saw, and `_apply_paragraph_overrides`
refuses any id with `in_block` set for exactly that reason.

UNTOUCHED TRACKS, TOUCHED FREEZES. A row with no stored line is still composed from the tokens,
so a changed square footage flows through. A row the estimator typed in prints verbatim and stops
following the sheet — a hand-written sentence has no slot to re-substitute a number into. That is
the same trade already accepted for every PRICE line including the base bid, which is money, and
the screen carries a ⚠ saying the line differs from the estimate.

THE ID SPACE IS UNTOUCHED, which is the load-bearing part. `iter_editable_blocks` yields the same
blocks in the same order as before, so every `paragraph_overrides` entry saved against a draft in
flight still lands on the paragraph it was captured from.

THE NUMBERING RULE. One system renders "System:", two or more render "Option 1:" / "Option 2:".
Rewriting one row does NOT switch numbering off for the rows the estimator did not touch: each
row's computed label is a default for that index only. Asserted below for one system and for
three.

THE BOX GROWTH. Over-long content used to shrink its own font and then be clipped behind a
"Too long for this box" badge. It now grows into whatever room the page has, and that height goes
to the writer through the same `box_overrides` channel a manual drag uses. Room is real geometry,
not optimism: on Kyle's Direct epoxy sheet PRICE starts 2.7pt ABOVE the bottom of WORK, so WORK
cannot grow by a single point and keeps the clip — with a badge that now says so.

The browser half RUNS in `js/doc-editor-labels-harness.js` rather than being read: whether the
whole row came out as ONE editable element with nothing editable nested inside it, whether an
untouched row still follows the estimate after the sheet moves, and whether a box may get taller
given where the other five boxes are, are all properties of executed code. The precedent is expensive — on 2026-08-12 an
unbound identifier shipped with every source-text assertion green and took the production board
down.
"""
import io
import json
import pathlib
import re
import shutil
import subprocess

import docx
import pytest

import main
import proposal_writer as pw

FRONTEND = pathlib.Path(__file__).resolve().parents[2] / "frontend"
HARNESS = pathlib.Path(__file__).resolve().parent / "js" / "doc-editor-labels-harness.js"
CSS = (FRONTEND / "styles.css").read_text(encoding="utf-8")

DIRECT_EPOXY = (pathlib.Path(__file__).resolve().parents[1] / "templates" / "Direct"
                / "XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx")

_A = "{%s}" % pw._A_NS
_WPS = "{%s}" % pw._WPS_NS
EMU_PER_PT = 12700.0


# ══ the fixture facts, re-derived from the template ═══════════════════════════
# Everything below is only meaningful if the template still looks like this. If Kyle re-authors
# it, these fail first and say what changed instead of leaving the feature quietly broken.
def _blocks(path):
    d = docx.Document(str(path))
    return [(i, in_block, text) for i, _k, _p, in_block, text, _t
            in pw.iter_editable_blocks(d)]


def test_the_system_labels_really_do_live_in_a_repeatable_region():
    """The whole reason a per-row channel had to exist. These three rows are inside
    `{{#system}}`, and `_apply_paragraph_overrides` skips anything with `in_block` set.

    The Area row's text is quoted in full on purpose: " SF of epoxy flooring" is the phrase
    Kyle could not delete, and it is template text sitting between two tokens — not a token,
    not a label, so neither the token channel nor the label channel could ever reach it."""
    rows = {text: in_block for _i, in_block, text in _blocks(DIRECT_EPOXY)}
    assert rows.get("{{system.prefix}}   {{system.name}}") == "system"
    assert rows.get("Texture:  {{system.texture}}") == "system"
    assert rows.get("Area: ~{{system.sqft}} SF of epoxy flooring{{system.lf_clause}}") == "system"


def test_the_other_labels_were_already_editable():
    """Kyle said "some of the labels", and this is which. Scope/Schedule/Exclusions sit in
    ordinary paragraphs (`in_block is None`), so they were always reachable through
    paragraph_overrides — the fix for those was an affordance, not a channel."""
    free = [text for _i, in_block, text in _blocks(DIRECT_EPOXY)
            if in_block is None and re.match(r"^(Scope|Schedule|Exclusions):", text.strip())]
    assert len(free) >= 3, "the free-text labels moved into a block: %r" % (free,)


def test_the_block_id_space_is_unchanged_by_this_feature():
    """A saved paragraph_override is keyed by a POSITION in this walk. The label channel had to
    avoid adding, removing or reordering a block, or every override on every draft in flight
    would land on the wrong paragraph. The walk is index-dense and the region still starts where
    it did, so an id captured before this change still means the same paragraph."""
    ids = [i for i, _b, _t in _blocks(DIRECT_EPOXY)]
    assert ids == list(range(len(ids)))
    starts = [i for i, b, t in _blocks(DIRECT_EPOXY) if t.strip() == "{{#system}}"]
    assert starts == [110], (
        "the {{#system}} region moved to %r — every saved paragraph_override id after it is now "
        "pointing at a different paragraph" % (starts,))


# ══ the Python half: what actually reaches the customer's document ═══════════
BASE_VALUES = {"job_name": "Test Job", "texture": "Light Broadcast"}


def _fill(systems, **kw):
    return docx.Document(io.BytesIO(pw.fill_proposal(
        work_type="epoxy", audience="Direct", values=dict(BASE_VALUES), systems=systems, **kw)))


def _texts(d):
    return [p.text for p in pw._iter_all_paragraphs(d) if p.text.strip()]


def _row(name="Broadcast Quartz", **extra):
    row = {"prefix": "System:", "name": name, "texture": "Light Broadcast",
           "sqft": "5,000", "lf_clause": ""}
    row.update(extra)
    return row


def test_a_renamed_system_label_reaches_the_document():
    got = _texts(_fill([_row(prefix="Base System:")]))
    assert any(t.startswith("Base System:") for t in got), got
    assert not any(t.startswith("System:") for t in got)


def test_a_renamed_label_keeps_its_bold_run():
    """The label is a bold run and the value is not. Rewriting the text must not flatten that —
    Kyle's page design is the bold lead-in."""
    d = _fill([_row(prefix="Base System:", texture_label="Surface texture:")])
    rows = {}
    for p in pw._iter_all_paragraphs(d):
        if p.text.startswith(("Base System:", "Surface texture:")):
            rows[p.text.split(":")[0]] = [(r.text, r.bold) for r in p.runs if r.text]
    assert rows["Base System"][0] == ("Base System:", True)
    assert rows["Surface texture"][0] == ("Surface texture:", True)
    assert rows["Surface texture"][-1][1] is False, "the value ran into the label's bold"


def test_the_static_texture_and_area_labels_can_be_renamed():
    """The older, narrower channel. Kept because a draft saved under the island editor still
    carries these keys, and dropping them would delete text the estimator typed."""
    got = _texts(_fill([_row(texture_label="Surface texture:", area_label="Coverage:")]))
    assert "Surface texture:  Light Broadcast" in got, got
    assert any(t.startswith("Coverage: ~5,000 SF") for t in got), got
    assert not any(t.startswith("Texture:") for t in got)
    assert not any(t.startswith("Area:") for t in got)


# ── THE COMPLAINT, at the end that writes the customer's document ────────────────
def test_the_static_words_around_a_token_can_be_deleted():
    """"I cant delete SF of epoxy flooring." This is that sentence as a test. The phrase is
    template text between {{system.sqft}} and {{system.lf_clause}}; under the island model no
    code path in either half of the system could write it."""
    got = _texts(_fill([_row(area_line="Area: ~5,000")]))
    assert "Area: ~5,000" in got, got
    assert not any("SF of epoxy flooring" in t for t in got), got
    assert not any("epoxy cove base" in t for t in got), got


def test_a_whole_line_prints_exactly_what_the_estimator_typed():
    """Every word, including ones the template never had, and including the label. Nothing
    re-composes the line from parts, because there are no parts."""
    line = "Coverage 5,000 square feet of quartz broadcast, cove included"
    got = _texts(_fill([_row(area_line=line)]))
    assert line in got, got


def test_each_row_has_its_own_whole_line_channel():
    got = _texts(_fill([_row(name_line="Base build:  Quartz broadcast",
                             texture_line="Finish:  matte",
                             area_line="Coverage:  5,000 sf")]))
    assert "Base build:  Quartz broadcast" in got, got
    assert "Finish:  matte" in got, got
    assert "Coverage:  5,000 sf" in got, got
    assert not any(t.startswith(("System:", "Texture:", "Area:")) for t in got), got


def test_a_whole_line_wins_over_the_label_and_the_value_on_the_same_row():
    """Both channels can be present on one row: the line is what the estimator was looking at
    when they typed, so it is what prints. Anything else would show them one thing on screen and
    send the customer another."""
    got = _texts(_fill([_row(area_label="Coverage:", sqft="1,234",
                             area_line="Area: ~5,000 SF")]))
    assert "Area: ~5,000 SF" in got, got
    assert not any(t.startswith("Coverage:") for t in got), got
    assert not any("1,234" in t for t in got), got


def test_an_edited_line_keeps_the_bold_lead_in():
    """Kyle's page design is the bold label. `_normalize_work_label_formatting` re-derives it
    from the rewritten text (bold through the first colon), and the on-screen line renders the
    same split — so the page and the file agree about the weight."""
    d = _fill([_row(name_line="Base System:   Quartz", area_line="Coverage: 5,000 SF")])
    rows = {}
    for p in pw._iter_all_paragraphs(d):
        if p.text.startswith(("Base System:", "Coverage:")):
            rows[p.text.split(":")[0]] = [(r.text, r.bold) for r in p.runs if r.text]
    assert rows["Base System"][0] == ("Base System:", True)
    assert rows["Base System"][-1][1] is False, "the value ran into the label's bold"
    assert rows["Coverage"][0] == ("Coverage:", True)
    assert rows["Coverage"][-1][1] is False


def test_a_line_with_no_colon_keeps_the_row_weight_the_page_shows():
    """Delete the colon and the normalizer stands down, so the row keeps its template run
    weight: bold for System and Area, normal for Texture. The preview's `workLabelHtml`
    boldFallback argument is exactly these three values — if they diverge the customer gets a
    weight the estimator never saw."""
    d = _fill([_row(name_line="Base build no colon",
                    texture_line="Finish matte no colon",
                    area_line="Coverage 5000 sq ft")])
    weight = {}
    for p in pw._iter_all_paragraphs(d):
        for lead in ("Base build", "Finish matte", "Coverage 5000"):
            if p.text.startswith(lead):
                weight[lead] = [r.bold for r in p.runs if r.text]
    assert weight["Base build"] == [True]
    assert weight["Finish matte"] == [None], "Texture's template run is not bold"
    assert weight["Coverage 5000"] == [True]
    js = (FRONTEND / "js" / "proposal-review.js").read_text(encoding="utf-8")
    # The subject is the boldFallback flag -- the LAST argument -- so that is what is matched.
    # It used to be pinned by quoting the whole call including a hardcoded inline margin
    # ("margin:0 0 4pt;"), which broke the moment those rows started taking their geometry from
    # the template instead of from hand-written numbers. The margin was never what this test
    # was about; matching it made an unrelated fidelity fix look like a regression here.
    calls = re.findall(r'lineRow\(i, "(\w+)"[^;]*?(true|false)\)', js)
    assert dict(calls) == {"name_line": "true", "texture_line": "false", "area_line": "true"}, (
        "renderSystemPreview's boldFallback arguments moved — re-derive them from the runs "
        "above: %r" % (calls,))


def test_a_rewritten_cove_only_row_is_not_re_edited_by_the_zero_sf_regex():
    """`_drop_zero_sf_prefix` tidies a line the ENGINE composed. A line a person typed is not
    the engine's to tidy: he can write "~0 SF of epoxy flooring and …" on purpose, and if he
    does that is what prints. Without the protect set this row silently became "Area: 240 LF …"
    on the way to the customer while the page went on showing what he typed."""
    typed = 'Area: ~0 SF of epoxy flooring and 240 LF of 6" epoxy cove base'
    got = _texts(_fill([_row(sqft="0", lf_clause=' and 240 LF of 6" epoxy cove base',
                             area_line=typed)]))
    assert typed in got, got


def test_an_untouched_cove_only_row_is_still_tidied():
    """The other half of the same rule — the protect set must not switch the tidy-up off for
    rows nobody edited."""
    got = _texts(_fill([_row(sqft="0", lf_clause=' and 240 LF of 6" epoxy cove base')]))
    assert any(t.startswith('Area: 240 LF of 6" epoxy cove base') for t in got), got
    assert not any("~0 SF of epoxy flooring" in t for t in got), got


def test_an_untouched_row_still_follows_the_estimate():
    """The freeze is per row and per line. Row 2 has no stored line, so its SF is still the one
    the estimate produced."""
    got = _texts(_fill([
        _row("Broadcast Quartz", area_line="Coverage: whatever I typed"),
        _row("Decorative Flake", prefix="Option 2:", sqft="1,800"),
    ]))
    assert "Coverage: whatever I typed" in got
    assert any(t.startswith("Area: ~1,800 SF of epoxy flooring") for t in got), got


@pytest.mark.parametrize("blank", ["", "   ", "{{system.sqft}}", "{{", "}}"])
def test_a_blank_or_brace_only_line_cannot_reach_the_document(blank):
    """Same defence in depth the labels get. A blank line means "revert", and a literal
    "{{token}}" printed to a customer is the exact failure the writer exists to prevent — this
    applier runs after the per-item substitution and before the flat pass, so an unstripped
    brace would survive all the way to the file."""
    got = _texts(_fill([_row(area_line=blank)]))
    body = "\n".join(got)
    # Same caveat test_no_work_row_carries_a_raw_token gives: `values` is deliberately minimal
    # here, so the front-page {{job_name}}-style tokens ARE still present. What must never
    # survive is a {{system.…}} token, which is what an unstripped brace would leave.
    assert "{{system." not in body, body
    assert not any("{{" in t or "}}" in t
                   for t in got if t.startswith(("Area:", "system.sqft"))), got
    if blank.strip() in ("", "{{", "}}"):
        assert any(t.startswith("Area: ~5,000 SF") for t in got), got
    else:
        assert "system.sqft" in body, body


def test_a_rename_is_per_row():
    """Row 2 keeps the template's wording and its own number. A label channel that leaked across
    rows would rewrite a row the estimator never opened, in a document a customer receives."""
    got = _texts(_fill([
        _row("Broadcast Quartz", prefix="Base System:", texture_label="Surface texture:",
             area_label="Coverage:"),
        _row("Decorative Flake", prefix="Option 2:", sqft="1,800"),
    ]))
    assert "Surface texture:  Light Broadcast" in got
    assert "Texture:  Light Broadcast" in got, "row 2 lost the template's own label"
    assert any(t.startswith("Coverage: ~5,000 SF") for t in got)
    assert any(t.startswith("Area: ~1,800 SF") for t in got)


def test_an_absent_label_leaves_the_template_wording():
    """"Emptied" is how the editor spells "revert": the input handler deletes the key, so the
    writer never sees it. Nothing here may emit a bare token or a lone colon."""
    got = _texts(_fill([_row()]))
    assert "Texture:  Light Broadcast" in got
    assert any(t.startswith("Area: ~5,000 SF") for t in got)


@pytest.mark.parametrize("blank", ["", "   ", "{{system.texture}}", "{{", "}}"])
def test_a_blank_or_brace_only_label_cannot_reach_the_document(blank):  # noqa: D401
    """Defence in depth. main._sanitize_system_overrides already drops a blank, but
    proposal_writer is also called directly (the To-Dropbox reconstruction path, and tests), and
    a literal "{{token}}" printed to a customer is the exact failure the writer exists to
    prevent. A brace-only override is stripped to nothing and therefore reverts."""
    got = _texts(_fill([_row(texture_label=blank, area_label=blank)]))
    body = "\n".join(got)
    assert "{{system." not in body, body
    # A brace-only override strips to nothing, so the template's own wording comes back. The
    # "{{system.texture}}" case is the interesting one: stripped it reads "system.texture", which
    # is not a token and cannot be mistaken for one by a reader of the printed page.
    if blank.strip() in ("", "{{", "}}"):
        assert "Texture:  Light Broadcast" in got, got
    else:
        assert "system.texture  Light Broadcast" in got, got


def test_no_work_row_carries_a_raw_token():
    """Whatever the estimator typed into a label, the WORK rows the customer reads must not show
    a `{{token}}`. Every `{{system.…}}` in the document has to be consumed by the expansion —
    checked document-wide, so the VML fallback copies are included. (The bare `{{job_name}}`-style
    front-page tokens ARE still present here: `values` is deliberately minimal in this file, and
    main._ensure_value_aliases is what backfills them on the live path.)"""
    d = _fill([_row(prefix="Base System:", texture_label="Surface texture:",
                    area_label="Coverage:", lf_clause=" and 240 LF of 6\" epoxy cove base")])
    body = "\n".join(_texts(d))
    assert "{{system." not in body and "{{#system" not in body and "{{/system" not in body
    rows = [t for t in body.split("\n")
            if t.startswith(("Base System:", "Surface texture:", "Coverage:"))]
    assert len(rows) == 6, rows          # three rows x (shape + VML fallback twin)
    assert not [t for t in rows if "{{" in t or "}}" in t]


def test_both_the_modern_and_the_legacy_copy_of_a_row_are_rewritten():
    """A floating text box is stored twice — the DrawingML shape and its VML `mc:Fallback` twin.
    Which one a renderer believes is version-dependent, so a label rewritten in only one of them
    would make Word and the customer's PDF disagree about what the proposal says."""
    d = _fill([_row(texture_label="Surface texture:")])
    hits = [p.text for p in pw._iter_all_paragraphs(d) if p.text.startswith("Surface texture:")]
    assert len(hits) == 2, "expected the shape and its fallback twin, got %r" % (hits,)


# ── the numbering rule, at both ends of the range ─────────────────────────────
def test_one_system_is_labelled_system():
    rows = main._build_epoxy_systems({}, dict(BASE_VALUES),
                                     [{"name": "Broadcast Quartz", "sf": 5000, "lf": 0}])
    assert [r["prefix"] for r in rows] == ["System:"]


def test_three_systems_are_numbered():
    rows = main._build_epoxy_systems({}, dict(BASE_VALUES), [
        {"name": "Quartz", "sf": 5000, "lf": 0},
        {"name": "Flake", "sf": 1800, "lf": 0},
        {"name": "Urethane", "sf": 900, "lf": 0},
    ])
    assert [r["prefix"] for r in rows] == ["Option 1:", "Option 2:", "Option 3:"]


def test_renaming_row_one_leaves_the_other_rows_numbered():
    """THE RULE, stated as a test. The computed prefix is a per-index default, so an override on
    row 1 changes row 1. The alternative — one manual label suppressing numbering for the whole
    list — would silently rewrite rows nobody edited."""
    rows = main._build_epoxy_systems({}, dict(BASE_VALUES), [
        {"name": "Quartz", "sf": 5000, "lf": 0},
        {"name": "Flake", "sf": 1800, "lf": 0},
        {"name": "Urethane", "sf": 900, "lf": 0},
    ])
    for i, ov in enumerate(main._sanitize_system_overrides([{"prefix": "Base System:"}])):
        rows[i].update(ov)
    assert [r["prefix"] for r in rows] == ["Base System:", "Option 2:", "Option 3:"]
    got = _texts(_fill(rows))
    assert any(t.startswith("Base System:") for t in got)
    assert any(t.startswith("Option 2:") for t in got)
    assert any(t.startswith("Option 3:") for t in got)


# ── the payload gate in main.py ───────────────────────────────────────────────
def test_the_line_and_label_fields_are_on_the_override_whitelist():
    """`_sanitize_system_overrides` drops anything not named here, so the whitelist IS the
    channel. A field the editor writes and the sanitizer drops is worse than no field: the draft
    remembers the edit and shows it back, and the customer's document does not have it."""
    for field in ("name_line", "texture_line", "area_line",
                  "prefix", "texture_label", "area_label"):
        assert field in main._SYSTEM_OVERRIDE_FIELDS, field


def test_the_browser_and_the_sanitizer_agree_on_the_whole_line_field_names():
    """Two files, one vocabulary. The preview writes these keys and the sanitizer is the only
    thing that lets them through, so a rename in one place has to fail here."""
    js = (FRONTEND / "js" / "proposal-review.js").read_text(encoding="utf-8")
    m = re.search(r"const _SYS_ROW_LINE_FIELDS = \[([^\]]*)\]", js)
    assert m, "the preview no longer declares _SYS_ROW_LINE_FIELDS"
    browser = re.findall(r'"([\w]+)"', m.group(1))
    assert browser == ["name_line", "texture_line", "area_line"], browser
    for field in browser:
        assert field in main._SYSTEM_OVERRIDE_FIELDS, field
    assert [k for k, _tok in pw._SYSTEM_ROW_LINES] == browser, (
        "the writer's row order/names drifted from the browser's")


def test_the_sanitizer_keeps_what_was_typed_and_drops_a_blank_field():
    """SPACES SURVIVE. This sanitizer used to `.strip()` every field, so a space typed at
    either end of a WORK row was stored in the draft, shown back after a reload, and then
    thrown away server-side. Kyle: "everything in the Proposals when editing should refelect 1
    to 1 in the customer side." Whitespace-only still means "revert"."""
    got = main._sanitize_system_overrides([
        {"area_line": "  Area:  ~5,000 SF  ", "texture_line": "   ",
         "prefix": " Base System: ", "junk": "dropped"},
        {},
    ])
    assert got == [{"area_line": "  Area:  ~5,000 SF  ", "prefix": " Base System: "}, {}]


def test_the_sanitizer_stays_index_preserving():
    """The list is positional. Dropping a malformed entry would shift every later override onto
    the wrong system — a renamed row 2 landing on row 3's label."""
    got = main._sanitize_system_overrides(["junk", {"name": "Decorative Flake"}])
    assert got == [{}, {"name": "Decorative Flake"}]


# ══ the grown height, once it leaves the browser ══════════════════════════════
def _box_heights(d, idx):
    """(anchor extent pt, shape transform pt, VML twin pt) for one box — the three places a
    box's size is recorded, in two unit systems."""
    txbx = list(pw._iter_txbx(d))[idx]
    anchor = pw._txbx_anchor(txbx)
    ext = anchor.find(pw.qn("wp:extent")) if anchor is not None else None
    extent = round(int(ext.get("cy")) / EMU_PER_PT, 2) if ext is not None else None
    xfrm = None
    for wsp in txbx.iterancestors(_WPS + "wsp"):
        x = next(iter(wsp.iter(_A + "xfrm")), None)
        e = x.find(_A + "ext") if x is not None else None
        if e is not None:
            xfrm = round(int(e.get("cy")) / EMU_PER_PT, 2)
        break
    vml = None
    for shape in pw._txbx_vml_twins(txbx):
        for part in (shape.get("style") or "").split(";"):
            k, _, v = part.partition(":")
            if k.strip().lower() == "height":
                vml = pw._vml_len_pt(v)
        break
    return extent, xfrm, vml


def test_an_auto_grown_height_reaches_all_three_places_in_the_docx():
    """The height the preview grew to is written through the ordinary box_overrides path, so the
    generated file has to agree with the screen in every branch a renderer might read."""
    grown = _fill([_row()], box_overrides={"3": {"h_pt": 219.75}})
    assert _box_heights(grown, 3) == (219.75, 219.75, 219.75)


def test_a_box_that_fits_is_left_byte_identical():
    """The constraint that protects every existing job: growth writes nothing for a box whose
    text already fits, so its bytes are the bytes it always had."""
    plain = pw.fill_proposal(work_type="epoxy", audience="Direct", values=dict(BASE_VALUES),
                             systems=[_row()])
    empty = pw.fill_proposal(work_type="epoxy", audience="Direct", values=dict(BASE_VALUES),
                             systems=[_row()], box_overrides={})
    assert _box_heights(docx.Document(io.BytesIO(plain)), 3) == \
           _box_heights(docx.Document(io.BytesIO(empty)), 3)
    assert _box_heights(docx.Document(io.BytesIO(plain)), 3)[0] == 162.0, (
        "the NOTES box is no longer 162pt — the growth-room numbers below need re-deriving")


def test_the_page_geometry_the_growth_rule_is_built_on():
    """The measurement that makes the WORK box's clip honest rather than lazy. PRICE's top is
    ABOVE the bottom of WORK, so WORK has 168.3pt of room for a 171pt shape: it cannot grow at
    all, and the code that decides that must not be reading "is it below my bottom edge"."""
    geo = pw.template_geometry(docx.Document(str(DIRECT_EPOXY)))
    by_id = {b["id"]: b for b in geo["boxes"]}
    work, price, notes = by_id[2], by_id[4], by_id[3]
    assert round(work["y_pt"] + work["h_pt"], 2) == 323.65
    assert round(price["y_pt"], 2) == 320.95, "the overlap this rule exists for is gone"
    assert round(price["y_pt"] - work["y_pt"], 2) == 168.3 < work["h_pt"]
    assert round(notes["y_pt"] + notes["h_pt"], 2) == 656.6
    assert geo["page"]["h_pt"] - geo["page"]["margin"]["bottom"] == 720.0


# ══ the browser half, executed ════════════════════════════════════════════════
@pytest.fixture(scope="module")
def ran():
    if shutil.which("node") is None:
        pytest.skip("node is not installed")
    proc = subprocess.run(["node", str(HARNESS), str(FRONTEND)],
                          capture_output=True, text=True, encoding="utf-8", timeout=180)
    assert proc.returncode == 0, (
        "the harness itself failed — read this before assuming a product bug:\n" + proc.stderr)
    return json.loads(proc.stdout.strip().splitlines()[-1])


def _fields(rows):
    return [(r["i"], r["field"]) for r in rows]


def test_every_work_row_is_editable_through_the_box_and_declares_no_host_of_its_own(ran):
    """THE ANSWER TO THE COMPLAINT, executed — and rewritten, because the complaint moved on.

    It used to assert that each row declared `contenteditable` ITSELF. That was the right shape
    while the argument was about islands: the words between two editable spans belonged to no
    editable element, so " SF of epoxy flooring" could not be typed in, and one editable element
    per row fixed it.

    Hanz, 2026-08-26: "just make every Major section of the proposal textbox just one BIG TEXT
    BOX not One big and some smaller textboxes. That is not how the Word DOc textboxes work."
    An element with its own `contenteditable` is its own EDITING HOST, and a browser selection
    cannot cross a host boundary — so three editable rows meant three places a drag could stop.
    The host moved up to the box, and these rows now inherit from it.

    So the assertion inverts, and what it protects does not: every row still takes a caret
    anywhere in it, and nothing inside a row is separately editable. `islands` is unchanged and
    still zero — a nested editable span would re-create both the island model AND a host boundary
    in the middle of a line."""
    got = ran["workGeometry"]["rows"]
    assert _fields(got) == [(0, "name_line"), (0, "texture_line"), (0, "area_line")]
    assert all(r["host"] == "tw-txbx" for r in got), (
        "a WORK row is not editable through its text box: %r" % [r["host"] for r in got])
    assert not any(r["ownHost"] for r in got), (
        "a WORK row declared contenteditable itself — that is a second editing host inside the "
        "box, and a drag through the section will stop at it again")
    assert all(r["wholeLine"] for r in got), (
        "a WORK row stopped using the same whole-line model as the base bid")
    assert [r["islands"] for r in got] == [0, 0, 0], (
        "something editable is nested inside a row again — that is the island model returning")


def test_every_price_row_is_editable_through_the_box_and_declares_no_host_of_its_own(ran):
    """THE SAME CLAIM AS THE WORK ROWS ABOVE, for the family that never had it — and the absence
    is why this shipped broken.

    Hanz, 2026-08-26: "why do we still have subboxes for the main text box?" / "Also remove the sub
    textboxes the subsections." Six PRICE rows in proposal-review.html still hard-coded
    `contenteditable="true"`: the Base Bid heading, the base line, the two tax rows, the total and
    the Options heading. Each one is its own EDITING HOST, and a browser selection cannot cross
    one — so Ctrl+A stopped at a row, a drag stopped at a row, and clicking from one price line to
    the next fired a focusout whose normalizer re-rendered every other line in the same box under
    the estimator's caret. The WORK rows were checked for exactly this from the day the box became
    the host; nothing checked these, and the markup sat there.

    These rows are STATIC MARKUP rather than the output of a render function, so the harness reads
    them out of the page and mounts them the way `mountRegionPreviews` does — into a real
    `.tw-txbx` — and then asks `hostOf`, the same walk the WORK rows answer with."""
    got = ran["priceRows"]
    assert got["keys"] == ["heading_base", "base", "sales_tax", "remodel", "total",
                           "heading_options"], (
        "the staged price rows changed — re-derive what this test is about: %r" % (got["keys"],))
    assert not any(got["ownHost"]), (
        "a PRICE row declares contenteditable itself: that is a second editing host inside the "
        "box, and every box-wide gesture stops at it")
    assert all(h == "tw-txbx" for h in got["host"]), (
        "a PRICE row is not editable through its text box at all: %r" % (got["host"],))
    assert all(got["wholeLine"]), (
        "a PRICE row stopped using the whole-line model the base bid established")
    assert got["islands"] == [0] * len(got["keys"]), (
        "something editable is nested inside a price row — the island model returning")
    assert got["anyHostBelowTheBox"] == 0, (
        "something in the staging block still declares contenteditable below the box")


def test_typing_in_notes_leaves_another_boxs_expanded_state_alone(ran):
    """Hanz, 2026-08-26, on the editor being clunky between sections.

    `syncNotesFromDom` ends by re-fitting, and it used to call `fitNotesBox()` — which loops EVERY
    `.tw-txbx` on the page and hands each one to `fitTxbx`, which resets fontSize, transform,
    maxHeight, overflow and zIndex and removes `tw-notes-open`. So one character typed in a notes
    bullet re-ran the shrink ladder on WORK and PRICE and folded shut any box the estimator had
    expanded to read. Only the notes box's content changed, so only the notes box is re-measured.

    The WORK box here is genuinely over capacity (400pt of content in Kyle's 171pt box) and
    genuinely expanded through its own button, so there is something real to destroy."""
    got = ran["notesRefitScope"]
    assert got["before"]["open"] is True, "the fixture never expanded the other box"
    assert got["after"] == got["before"], (
        "typing in a notes bullet changed another box: %r -> %r" % (got["before"], got["after"]))
    assert got["notesFitted"], "the notes box itself was not re-measured, which is the point of it"
    assert got["textarea"] == ("Price includes one mobilization.\n"
                              "Anchor bolts by others, plus layout."), (
        "the bullets no longer reach the textarea that is their source of truth: %r"
        % got["textarea"])


def test_the_whole_line_the_estimator_edits_carries_every_static_word(ran):
    """The static words are IN the editable element's own text, not beside it. Nothing else has
    to be true for him to be able to delete them."""
    area = next(r for r in ran["workGeometry"]["rows"] if r["field"] == "area_line")
    assert area["text"] == (
        'Area: ~5,000 SF of epoxy flooring and 240 LF of 6" epoxy cove base')
    for phrase in ("Area:", "~", " SF of epoxy flooring", 'LF of 6" epoxy cove base'):
        assert phrase in area["text"], phrase


def test_the_rendered_rows_still_read_like_the_template(ran):
    """Making the rows editable must not change what the page says. These are the template's own
    three lines, with the estimate's values in them."""
    assert ran["oneSystem"]["lines"] == [
        "System:   Broadcast Quartz",
        "Texture:  Light Broadcast",
        'Area: ~5,000 SF of epoxy flooring and 240 LF of 6" epoxy cove base']


def test_the_bold_lead_in_matches_what_the_writer_will_print(ran):
    """One <strong> per row, ending at the first colon — the same split
    `_normalize_work_label_formatting` applies to the .docx. If these two ever disagree the
    estimator approves one weight and the customer receives another."""
    assert [r["bold"] for r in ran["workGeometry"]["rows"]] == [
        ["System:"], ["Texture:"], ["Area:"]]


def test_two_systems_number_themselves(ran):
    labels = [r["text"].split(":")[0] + ":" for r in ran["twoSystems"]["rows"]
              if r["field"] == "name_line"]
    assert labels == ["Option 1:", "Option 2:"]


def test_rewriting_one_row_leaves_the_next_row_numbered(ran):
    """The browser half of the rule. The store holds one entry, for row 1 only, and the re-render
    shows row 2 still carrying its own number."""
    assert ran["renamedRow1"]["stored"] == [{"name_line": "Base System:   Broadcast Quartz"}]
    assert ran["renamedRow1"]["persisted"] == [{"name_line": "Base System:   Broadcast Quartz"}]
    assert ran["renamedRow1"]["lines"][0] == "Base System:   Broadcast Quartz"
    assert ran["renamedRow1"]["lines"][3] == "Option 2:   Decorative Flake"
    row1 = next(r for r in ran["renamedRow1"]["rows"]
                if r["i"] == 0 and r["field"] == "name_line")
    assert (row1["text"], row1["computed"]) == (
        "Base System:   Broadcast Quartz", "Option 1:   Broadcast Quartz"), (
        "the line lost the computed text it reverts to")


def test_emptying_a_line_reverts_it_instead_of_printing_a_token(ran):
    """The one outcome that would be visible to a customer. Clearing the line deletes the
    override, so the computed text comes back — never a bare "{{system.prefix}}", never a lone
    colon left where the label was."""
    assert ran["emptiedLine"]["stored"] == [{}]
    assert ran["emptiedLine"]["lines"][0] == "Option 1:   Broadcast Quartz"
    body = " ".join(ran["emptiedLine"]["lines"])
    assert "{{" not in body and "}}" not in body


def test_the_static_words_can_be_deleted_on_screen_and_only_on_that_row(ran):
    """The Area line loses " SF of epoxy flooring" AND the cove clause, and the label with them.
    Row 2, which nobody touched, still reads the template's wording with the estimate's SF."""
    assert ran["staticWordsDeleted"]["stored"] == [
        {"texture_line": "Surface texture:  Light Broadcast", "area_line": "Coverage: 5,000"}]
    assert ran["staticWordsDeleted"]["lines"][1] == "Surface texture:  Light Broadcast"
    assert ran["staticWordsDeleted"]["lines"][2] == "Coverage: 5,000"
    assert ran["staticWordsDeleted"]["lines"][4] == "Texture:  Light Broadcast"
    assert ran["staticWordsDeleted"]["lines"][5].startswith("Area: ~1,800 SF")
    edited = next(r for r in ran["staticWordsDeleted"]["rows"]
                  if r["i"] == 0 and r["field"] == "area_line")
    assert "SF of epoxy flooring" not in edited["text"]
    assert "SF of epoxy flooring" in edited["computed"], (
        "the row forgot the computed line, so clearing it could not revert")


def test_an_untouched_line_follows_the_estimate_and_an_edited_one_stops(ran):
    """CONSTRAINT 2, and the honest cost of the model Kyle asked for. The sheet moves from
    5,000/1,800 to 7,777/2,222 AFTER the edit. Row 2 (untouched) picks the new figure up. Row 1
    keeps his sentence and stops tracking — a hand-written line has no slot to re-substitute a
    number into. The ⚠ and the remembered computed text are what make that visible and
    reversible."""
    lines = ran["estimateMoved"]["lines"]
    assert lines[2] == "Coverage: 5,000 SF, cove included", "his words were overwritten"
    assert lines[5] == "Area: ~2,222 SF of epoxy flooring", (
        "an untouched row stopped following the estimate")
    frozen = next(r for r in ran["estimateMoved"]["rows"]
                  if r["i"] == 0 and r["field"] == "area_line")
    assert frozen["warned"] is True
    assert "~7,777" in frozen["computed"], (
        "the frozen row must still remember the current estimate line to revert to")


def test_a_reworded_line_is_not_reported_as_a_re_priced_one(ran):
    """The ⚠ marker means "this differs from the estimate", which is a review risk for a NUMBER.
    Digits are compared AFTER the label, so renaming "System:" is reported as a rewording and an
    SF typed off the sheet is reported as a pricing edit. One visual state either way — two
    would be the island model again in another costume."""
    by_field = {r["field"]: r for r in ran["warnings"]}
    assert by_field["name_line"]["warned"] is True
    assert "Reworded" in by_field["name_line"]["title"]
    assert by_field["area_line"]["warned"] is True
    assert "differs from the computed estimate" in by_field["area_line"]["title"]
    assert by_field["texture_line"]["warned"] is False
    assert by_field["texture_line"]["title"] is None


def test_a_colon_less_line_keeps_the_rows_template_weight_on_screen(ran):
    """The other half of test_a_line_with_no_colon_keeps_the_row_weight_the_page_shows. Delete
    the colon and there is no label to bold, so the page must fall back to the weight the writer
    will actually produce — the row's first run: bold on System and Area, normal on Texture. A
    fallback that always returned plain text would show him a light line and print a bold one."""
    by_field = {r["field"]: r for r in ran["noColon"]}
    assert by_field["name_line"]["bold"] == ["Base build no colon"]
    assert by_field["texture_line"]["bold"] == []
    assert by_field["area_line"]["bold"] == ["Coverage 5000 sq ft"]


def test_the_browser_stores_the_line_with_its_spaces(ran):
    """The estimator now types at both ends of a whole line. The browser keeps what he typed;
    test_the_sanitizer_keeps_what_was_typed_and_drops_a_blank_field is the server half."""
    assert ran["spacesKept"] == [{"area_line": "  Area:  ~5,000 SF  "}]


# ── growth, run against the template's real geometry ─────────────────────────
def test_the_growth_room_matches_the_page(ran):
    """Room is the distance to the next box below with a horizontal overlap — and ZERO when there
    is no such box. Restated here independently of the harness, off the real template geometry:

        box 0 header  36.00 → box 2 at 152.65   = 116.65
        box 1 date    36.00 → box 5 at 501.95   = 465.95
        box 2 WORK   152.65 → box 4 at 320.95   = 168.30   (its shape is 171pt: no room at all)
        box 3 NOTES  494.60 → NOTHING BELOW IT  =   0.00
        box 4 PRICE  320.95 → box 3 at 494.60   = 173.65
        box 5 logo   501.95 → NOTHING BELOW IT  =   0.00   (NOTES is beside it, not below it)

    THE TWO ZEROS ARE THE POINT, and they used to be 225.40 and 218.05 — the distance to the
    page's bottom margin. The margin is not a measurement of empty space. What sits under NOTES
    is the ACCEPTANCE and signature frame, and under the logo the rest of the letterhead, both
    baked into a full-page PNG: there is no element, so there is no rect, so the DOM cannot see
    them. Bounding by the margin let NOTES grow its bottom edge from 656.6pt to 714.35pt, over
    that frame — and a grown box disarms the server-side shrink, so nothing downstream caught it
    and the customer received a proposal with the terms printed across the artwork
    (review 2026-08-20).

    The honest consequence, which this test also pins: on this template growth is available on
    four boxes and not on two. Reading these zeros as a bug and restoring the margin fallback
    would reintroduce exactly that. Growing into artwork needs the artwork measured — see
    measureTermsBand for the precedent — and that is a separate piece of work.
    """
    assert ran["roomBoxCount"] == 6
    assert ran["room"] == {"0": 116.65, "1": 465.95, "2": 168.3,
                           "3": 0, "4": 173.65, "5": 0}


def test_the_last_box_on_the_page_is_never_offered_growth(ran):
    """The artwork case, end to end rather than as an arithmetic claim: NOTES overflows by 58pt,
    has 225pt of blank page beneath it, and is still refused — no button, no geometry, and a
    warning that says why. This is the test that fails if anyone reinstates the margin fallback."""
    got = ran["artBlocked"]
    assert got["offered"] is False, "offered to grow a box over the baked letterhead artwork"
    assert got["payload"] == {}, "wrote geometry for a box it must not resize"
    assert got["geom"]["boxHPt"] == "162", "the box left the template's height"
    assert got["geom"]["overflow"] is True and got["geom"]["blocked"] is True
    assert "part of the letterhead picture" in got["geom"]["title"], (
        "the estimator is told it cannot grow but not that artwork is the reason")


def test_a_box_that_fits_is_not_touched(ran):
    """Byte-identical output for every job that was already fine — the payload stays empty, so
    the writer has nothing to apply and nothing to persist."""
    got = ran["fitsUntouched"]
    assert got["before"] == got["after"]
    assert got["payload"] == {}
    assert got["persisted"] == 0, "a box that fits triggered a draft save"
    assert got["after"]["minHeight"] == "162pt" and got["after"]["boxHPt"] == "162"
    assert got["after"]["overflow"] is False and got["after"]["grown"] is False


def test_an_overflowing_box_with_room_grows_and_the_height_is_persisted(ran):
    """170pt of content in the 164.5pt PRICE box, and the estimator PRESSES Fit to text. PRICE has
    NOTES below it, so its room (173.65pt) is bounded by a real box and growing it is provably
    safe — which is why this scenario is PRICE and not NOTES. The box becomes 170.25pt (the
    content height rounded up to the template's own 2dp) and that number is in the box_overrides
    the generate payload carries, which is what makes the .docx match the preview.

    A press, not a repaint: growth used to happen inside fitNotesBox, i.e. on first paint and on
    every keystroke, writing document geometry with no gesture behind it."""
    got = ran["grows"]
    assert got["geom"]["minHeight"] == "170.25pt"
    assert got["geom"]["boxHPt"] == "170.25"
    assert got["geom"]["overflow"] is False, "it grew and still claims its text is cut off"
    assert got["geom"]["fontSize"] == "", "it grew AND shrank the type"
    assert got["geom"]["grown"] is True and got["geom"]["moved"] is True
    assert got["payload"] == {"4": {"h_pt": 170.25}}
    assert got["stored"] == {"4": {"h_pt": 170.25}}
    assert got["persistCalls"] == 1
    assert got["autoGrown"] is True


def test_a_box_with_nowhere_to_go_still_clips_and_says_so(ran):
    """WORK is 171pt and has 168.3pt of room, so growing it would print over the PRICE frame —
    which is baked into the letterhead PNG and cannot move. Clipping is the honest outcome, and
    the tooltip has to explain why the box was not simply made bigger."""
    got = ran["blocked"]
    assert got["payload"] == {}, "it moved a box it had just decided it could not move"
    assert got["geom"]["minHeight"] == "171pt" and got["geom"]["boxHPt"] == "171"
    assert got["geom"]["overflow"] is True
    assert got["geom"]["blocked"] is True
    assert "the next box on the page starts where this one ends" in got["geom"]["title"]


def test_the_blocked_badge_is_on_screen_and_not_only_in_a_tooltip():
    """A tooltip is not a warning if nobody hovers. The clipped-and-cannot-grow badge is its own
    CSS rule, more specific than the generic one so it wins wherever they are written, and it
    stands down while the box is expanded (that state has its own message)."""
    sel = r"\.tw-txbx\.tw-notes-overflow\.tw-grow-blocked:not\(\.tw-notes-open\)::after"
    m = re.search(r"(?m)^" + sel + r"\s*\{([^}]*)\}", CSS)
    assert m, "the grow-blocked badge has no top-level rule in styles.css"
    assert "cannot grow" in m.group(1)


def test_the_grown_note_says_what_happened(ran):
    """The geometry of a customer-facing document changed, so the page has to admit it. The note
    is a labelled word in the tools layer — not a grip, which is what confused Kyle already — and
    adding it must not have displaced the resize handles."""
    got = ran["grownNote"]
    assert got["present"] and got["label"] == "Grown to fit"
    assert got["isNotAGrip"], "the note reads as a drag handle"
    assert "Reset box puts it back" in got["title"]
    grips = [c.split()[-1] for c in got["order"] if c.startswith("tw-grip ")]
    assert grips == ["tw-grip-move", "tw-grip-e", "tw-grip-s", "tw-grip-se"], (
        "adding the note reordered the grips test_box_drag_ui.py pins")


def test_the_grown_note_adds_no_height_to_the_box():
    """fitTxbx decides what overflows from the box's offsetHeight, so a control in the normal
    flow would make every box measure taller than its text — i.e. would break the very notice
    this feature reads."""
    m = re.search(r"(?m)^\.tw-box-grown-note\s*\{([^}]*)\}", CSS)
    assert m, ".tw-box-grown-note has no top-level rule in styles.css"
    assert "position: absolute" in m.group(1)


def test_trimming_the_text_gives_the_space_back(ran):
    """The height we added is recomputed from the template each pass, not accumulated. Otherwise
    one long paste would enlarge a box permanently and the estimator would have to know that
    "Reset box" was the cure."""
    got = ran["trimGivesItBack"]
    assert got["grown"]["boxHPt"] == "170.25"
    assert got["after"]["boxHPt"] == "164.5"
    assert got["after"]["grown"] is False and got["after"]["moved"] is False
    assert got["payload"] == {}


def test_a_height_the_estimator_dragged_is_never_auto_grown_over(ran):
    """They dragged the box SHORTER than its text needs. That is a deliberate instruction, so it
    stands — with the overflow warning, which is the honest report — rather than being quietly
    undone by the auto-fit."""
    got = ran["manualHeightWins"]
    assert got["dragged"]["boxHPt"] == "122"
    assert got["after"]["boxHPt"] == "122", "auto-grow overrode a deliberate resize"
    assert got["after"]["overflow"] is True
    assert got["payload"] == {"3": {"h_pt": 122}}


def test_reset_box_stays_reset(ran):
    """"Put this box back where the template has it" has to survive the next repaint. Without the
    suppression the auto-fit would re-grow it immediately and the button would look broken."""
    got = ran["resetSticks"]
    assert got["grown"]["boxHPt"] == "170.25"
    assert got["reset"]["boxHPt"] == "164.5"
    assert got["afterRefit"]["boxHPt"] == "164.5", "the box re-grew itself after Reset"
    # And the text is still handled honestly at that size rather than silently spilling: 170pt of
    # content in a 164.5pt box is only 3% over, so fitTxbx's first step — stepping the type down —
    # absorbs it, which is exactly what it is for. (The older version of this test asserted the
    # overflow BADGE instead, using a box 36% over capacity where shrinking cannot save it. That
    # box was NOTES, which is now correctly refused growth altogether, so the scenario moved to
    # PRICE and the observable moved with it.)
    assert got["afterRefit"]["fontSize"] == "95%", (
        "reset to a size its text does not fit, and neither shrank the type nor warned")
    assert got["payload"] == {}

# ══ the editor's geometry against the template's own numbers ══════════════════
# Hanz, 2026-08-25, over a side-by-side of the editor and the generated PDF: "is this really the
# spacing format for the epoxy? Because it doesnt follow the exact font size and spacing on the
# editor in which it should." He was right, and the PDF was the honest one.


def _work_paras():
    """The epoxy WORK box's paragraphs, straight out of the template, keyed by the token they
    carry. This is the ground truth both the editor and the generated document answer to."""
    import glob
    from docx.oxml.ns import qn
    import docx as _docx
    tpl = pathlib.Path(__file__).resolve().parents[1] / "templates" / "Direct"
    path = [p for p in glob.glob(str(tpl / "*.docx")) if "EPOXY" in p.upper()][0]
    d = _docx.Document(path)
    seen = 0
    for tx in d.element.body.iter(qn("w:txbxContent")):
        seen += 1
        if seen != 5:
            continue
        out = {}
        for p in tx.iter(qn("w:p")):
            txt = "".join(x.text or "" for x in p.iter(qn("w:t")))
            for tok in ("system.name", "system.texture", "system.sqft"):
                if tok in txt:
                    out[tok] = pw.para_props(d, p)
        return d, out
    return d, {}


def test_the_editor_places_the_work_rows_where_the_template_does(ran):
    """The fake nesting, pinned shut.

    Every bulleted WORK row in the file is `left=288 hanging=288` — text at 14.4pt, marker hard
    against the margin. Word puts the TEXT at `left` and the MARKER at `left - hanging`, so an
    editor reading `left` alone draws the bullet where the text belongs and pushes the whole row in
    by the hanging distance. That is what produced two different indents for rows that are
    identical in the document.

    Asserted as margin-left = left − hanging and padding-left = hanging, because that is the pair
    that reproduces Word's own arithmetic. `.tw-li::before` sits at `left: 0` inside the padding,
    which is what puts the red square exactly where it prints."""
    _, paras = _work_paras()
    assert paras, "the epoxy WORK box was not found in the template"
    name, area = paras["system.name"], paras["system.sqft"]
    for label, p in (("System", name), ("Area", area)):
        assert p["bullet"] is True, label
        assert p["indent"] == 288 and p["hanging"] == 288, (label, p)
    got = {r["field"]: r["style"] for r in ran["workGeometry"]["rows"]}
    for field in ("name_line", "area_line"):
        st = got[field]
        assert "padding-left:14.4pt" in st.replace(" ", ""), (field, st)
        # left − hanging = 0: the marker belongs AT the margin, not 14pt inside it.
        assert "0pt" in st, (field, st)


def test_the_texture_row_carries_the_deep_indent_the_file_gives_it(ran):
    """The outlier, and the one the editor was hiding rather than inventing.

    Texture is the only WORK row with a real indent of its own — `left=1008 firstLine=72`, so its
    text sits at 54pt with NO bullet. The editor rendered it at the same shallow offset as the
    rows above it, so the one row that IS indented in the document was the one row not indented on
    screen."""
    _, paras = _work_paras()
    tex = paras["system.texture"]
    assert tex["bullet"] is False, "Texture is not a bulleted row in the template"
    assert tex["indent"] == 1008 and tex["first_line"] == 72, tex
    st = {r["field"]: r["style"] for r in ran["workGeometry"]["rows"]}["texture_line"]
    flat = st.replace(" ", "")
    assert "margin:0pt00pt50.4pt" in flat, st          # 1008tw = 50.4pt, no hanging
    assert "text-indent:3.6pt" in flat, st             # firstLine 72tw
    assert "padding-left:0pt" in flat, st              # no marker, so no gap for one


def test_the_synthesized_rows_carry_the_templates_own_font_size(ran):
    """The other half of what Hanz saw. `workLabelHtml` emitted no font-size, so these three rows
    inherited `.tw-page { font-size: 9pt }` while every real block beside them carried an explicit
    8pt from its runs — 12.5% larger, on the rows at the top of the box.

    It is also what the box-overflow shrink was silently singling out: `fitTxbx` sets a PERCENTAGE
    font-size, and a percentage only scales INHERITED sizes, so the shrink hit exactly these three
    rows and left the rest alone. Giving them an explicit size closes both."""
    for r in ran["workGeometry"]["rows"]:
        assert "font-size:8pt" in r["style"].replace(" ", ""), (r["field"], r["style"])


def test_the_editor_invents_no_gaps_between_the_work_rows(ran):
    """The file has no `before` or `after` spacing anywhere in the WORK box — the rows sit flush,
    and `contextualSpacing` from ListParagraph is why. The editor was adding 1pt, 1pt and 4pt of
    its own, which is the "tighter spacing" half of the sub-group illusion."""
    _, paras = _work_paras()
    for tok, p in paras.items():
        assert p["spacing"]["before"] is None and p["spacing"]["after"] is None, (tok, p["spacing"])
    for r in ran["workGeometry"]["rows"]:
        flat = r["style"].replace(" ", "")
        assert "margin:0pt0" in flat, (
            "the editor invents a gap the document does not have: %r" % (r["style"],))


def test_the_line_spacing_comes_from_the_file_not_a_constant(ran):
    """The editor used one flat `line-height: 1.32` for a box whose rows are genuinely 1.15 and
    1.25 — looser than both, and erasing the distinction between them. `line` is 240ths of a line
    under `lineRule="auto"`, so 276 is 1.15 and 300 is 1.25; the RULE has to travel with the number
    because the same field is twips under `exact`."""
    _, paras = _work_paras()
    assert paras["system.name"]["spacing"]["line"] == 276, paras["system.name"]["spacing"]
    assert paras["system.sqft"]["spacing"]["line"] == 300, paras["system.sqft"]["spacing"]
    got = {r["field"]: r["style"].replace(" ", "") for r in ran["workGeometry"]["rows"]}
    assert "line-height:1.15" in got["name_line"], got["name_line"]
    assert "line-height:1.25" in got["area_line"], got["area_line"]
