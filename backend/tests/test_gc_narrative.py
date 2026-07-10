"""Audience-aware Scope / Schedule / Exclusions for the GC proposal templates.

The 3 GC templates historically shipped their Scope/Schedule/Exclusions as
STATIC text (no token), so the estimator's sidebar edits never reached the GC
doc and a GC job showed Direct wording that was never used. annotate_templates.py
now tokenizes them ({{scope_notes}}/{{schedule_notes}}/{{exclusions}}) and
main._ensure_value_aliases seeds GC-specific defaults keyed by work_type.

Covers:
  (a) each GC template's document.xml has the 3 tokens exactly 2x + zero leftover
      static scope/schedule/exclusions step strings;
  (b) fill_proposal(epoxy, GC) renders the multi-line scope as <w:br/>-separated
      steps under a still-bold "Scope:" label;
  (c) POST /api/generate (GC, blank narrative) yields GC wording, never the Direct
      schedule/exclusions boilerplate;
  (d) epoxy+GC and combo+GC -> Resinous wording, polish+GC -> Polish, sealer+GC
      -> Sealer (mirrors proposal_writer.pick_template);
  (e) the template_version guard drops stale paragraph_overrides but applies them
      when the version matches (or is empty/legacy);
  (f) Direct parity: Direct narrative fallbacks unchanged AND the 4 Direct
      template files are byte-identical (checksum-guarded).
"""
import hashlib
import io
import re
import zipfile

import docx
from docx import Document
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_WP = re.compile(r"<w:p\b[^>]*>.*?</w:p>", re.DOTALL)
_WT = re.compile(r"<w:t\b[^>]*>([^<]*)</w:t>")
_WR = re.compile(r"<w:r\b[^>]*>.*?</w:r>", re.DOTALL)

_GC_TEMPLATES = {
    "resinous": "GC/xx TREADWELL RESINOUS PROPOSAL - xx.docx",
    "polish":   "GC/xx TREADWELL POLISH PROPOSAL - xx.docx",
    "sealer":   "GC/xx TREADWELL SEALER PROPOSAL - xx.docx",
}

_DIRECT_TEMPLATES = {
    "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx":       "be9051b8c52bc62a",
    "Direct/xx.xx TREADWELL BUDGET PRICING.docx":                    "684e11d2ac9763a4",
    "Direct/xx.xx TREADWELL POLISH PROPOSAL - NewDirect.docx":       "5953fbe86f3963fa",
    "Direct/xx.xx.xx TREADWELL COMBO PROPOSAL - CUSTMOER NAME.docx": "cd6e90ecd4c161d1",
}

# Static step strings that MUST be gone from each annotated GC template (escaped
# form for &). If any remain, tokenization left the wording behind twice.
_LEFTOVER = {
    "resinous": [
        "Perform relative humidity test on concrete slab prior to installation",
        "Prepare substrate surface profile utilizing mechanical means",
        "Prep substrate (includes patch of minor substrate defects",
        "Install Resinous System",
        "Assumes installation over: clean, sound",
        "1week to complete full scope",
        "Epoxy Paint Walls",
    ],
    "polish": [
        "Grind and polish concrete with successive passes",
        "Apply hardener/densifier",
        "Apply joint filler",
        "Assumes polish over: clean, sound",
        "Scope:Prep substrate",
        "1week to complete full scope",
        "Cove Base, Dye, Demo",
    ],
    "sealer": [
        "Clean Concrete; -or- Perform 1-2 passes",
        "Apply [1 coat -or- up to 2 coats",
        "Assumes sealer over: clean, sound",
        "Scope:Prep substrate",
        "1week to complete full scope",
        "Patching, Grinding, Joint Filler",
    ],
}


def _rendered(docx_bytes):
    """Text of the paragraphs Word renders (mc:Choice), skipping the VML
    mc:Fallback duplicate — mirrors the rest of the suite."""
    d = Document(io.BytesIO(docx_bytes))
    out = []
    for p in d.element.xpath("//w:p"):
        if any(True for _ in p.iterancestors(f"{_MC}Fallback")):
            continue
        t = "".join(x.text or "" for x in p.xpath(".//w:t")).strip()
        if t:
            out.append(t)
    return "\n".join(out)


def _doc_xml(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        return z.read("word/document.xml").decode("utf-8")


def _template_xml(rel_path):
    with zipfile.ZipFile(str(pw.TEMPLATES_ROOT / rel_path)) as z:
        return z.read("word/document.xml").decode("utf-8")


def _paras_starting(xml, prefix):
    return [p for p in _WP.findall(xml)
            if "".join(_WT.findall(p)).strip().startswith(prefix)]


def _first_run_bold(p_xml):
    m = _WR.search(p_xml)
    if not m:
        return False
    rpr = re.search(r"<w:rPr\b[^>]*>.*?</w:rPr>", m.group(0), re.DOTALL)
    body = rpr.group(0) if rpr else ""
    return "<w:b/>" in body or "<w:b " in body or "<w:b>" in body


# Minimal generate payload values (blank scope/schedule/exclusions so the backend
# backfills them). Extra tokens the GC template doesn't get render as literals —
# harmless for these assertions, which only look at the narrative fields.
def _gc_vals(**over):
    v = {
        "project_name": "Test GC Job", "job_name": "Test GC Job",
        "city_state": "Olathe, KS", "sqft": "12,000", "cove_lf": "250",
        "estimator_name": "Greg", "bid_date_formatted": "6/15/26",
        "base_bid_formatted": "$58,000.00", "material_tax_formatted": "$0.00",
        "tax_amount_formatted": "$0.00", "total_formatted": "$58,000.00",
        "texture": "OP",
    }
    v.update(over)
    return v


# ── (a) template annotation ─────────────────────────────────────────────
def test_gc_templates_have_three_tokens_exactly_twice():
    for _key, rel in _GC_TEMPLATES.items():
        xml = _template_xml(rel)
        for tok in ("{{scope_notes}}", "{{schedule_notes}}", "{{exclusions}}"):
            assert xml.count(tok) == 2, f"{rel}: {tok} count={xml.count(tok)} (want 2)"


def test_gc_templates_have_no_leftover_static_steps():
    for key, rel in _GC_TEMPLATES.items():
        xml = _template_xml(rel)
        for s in _LEFTOVER[key]:
            assert s not in xml, f"{rel}: leftover static step still present: {s!r}"


def test_gc_labels_still_bold_and_read_label_colon_token():
    want = [("Scope:", "{{scope_notes}}"), ("Schedule:", "{{schedule_notes}}"),
            ("Exclusions:", "{{exclusions}}")]
    for _key, rel in _GC_TEMPLATES.items():
        xml = _template_xml(rel)
        for label, token in want:
            paras = [p for p in _WP.findall(xml)
                     if (_WT.search(p) and _WT.search(p).group(1).strip() == label)]
            assert len(paras) == 2, f"{rel}: {label} appears {len(paras)}x (want 2)"
            for p in paras:
                assert _first_run_bold(p), f"{rel}: {label} label lost bold"
                assert "".join(_WT.findall(p)).strip() == f"{label} {token}"


def test_gc_annotated_templates_open_cleanly():
    for _key, rel in _GC_TEMPLATES.items():
        d = docx.Document(str(pw.TEMPLATES_ROOT / rel))
        assert len(d.paragraphs) >= 0  # parses without raising


# ── (b) multi-line scope rendering + bold label ─────────────────────────
def test_fill_gc_epoxy_scope_is_multiline_under_bold_label():
    vals = _gc_vals(scope_notes=main._DEFAULT_SCOPE_GC_RESINOUS)
    out = pw.fill_proposal(work_type="epoxy", audience="GC", values=vals)
    xml = _doc_xml(out)

    scope_paras = _paras_starting(xml, "Scope:")
    assert scope_paras, "no Scope: paragraph in output"
    p = scope_paras[0]
    # 5 scope steps -> 4 line breaks separating them within the value run.
    n_br = len(re.findall(r"<w:br\b", p))
    assert n_br == 4, f"expected 4 <w:br/> between the 5 scope steps, got {n_br}"
    # Bold "Scope:" label survives; first + last steps present.
    assert _first_run_bold(p), "Scope: label lost its bold weight"
    joined = "".join(_WT.findall(p))
    assert joined.startswith("Scope: Perform relative humidity test")
    assert joined.endswith("Assumes installation over: clean, sound &amp; solid concrete substrate")
    # No raw token left behind.
    assert "{{scope_notes}}" not in xml


# ── (c) /api/generate GC blank narrative -> GC wording, not Direct ──────
def test_generate_gc_epoxy_uses_gc_narrative_not_direct():
    body = {"work_type": "epoxy", "audience": "GC", "values": _gc_vals()}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)

    # GC schedule ("1week", no space) present; Direct schedule ("1 week") absent.
    assert "1week to complete full scope" in txt
    assert "approx. 1 week to complete full scope" not in txt
    # GC Resinous exclusions present; Direct exclusions absent.
    assert "Epoxy Paint Walls" in txt
    assert "Multiple layers of floor to be removed" not in txt
    # GC Resinous scope present.
    assert "Perform relative humidity test on concrete slab" in txt


# ── (d) work_type -> template wording mapping (mirrors pick_template) ────
def test_gc_scope_defaults_by_work_type():
    cases = {
        "epoxy":  main._DEFAULT_SCOPE_GC_RESINOUS,
        "combo":  main._DEFAULT_SCOPE_GC_RESINOUS,   # GC combo uses the Resinous doc
        "polish": main._DEFAULT_SCOPE_GC_POLISH,
        "sealer": main._DEFAULT_SCOPE_GC_SEALER,
    }
    for wt, expected_scope in cases.items():
        v = {"work_type": wt}
        main._ensure_value_aliases(v, "GC")
        assert v["scope_notes"] == expected_scope, wt
        assert v["schedule_notes"] == main._DEFAULT_SCHEDULE_GC, wt


def test_gc_exclusions_defaults_by_work_type():
    cases = {
        "epoxy":  main._DEFAULT_EXCLUSIONS_GC_RESINOUS,
        "combo":  main._DEFAULT_EXCLUSIONS_GC_RESINOUS,
        "polish": main._DEFAULT_EXCLUSIONS_GC_POLISH,
        "sealer": main._DEFAULT_EXCLUSIONS_GC_SEALER,
    }
    for wt, expected in cases.items():
        v = {"work_type": wt}
        main._ensure_value_aliases(v, "GC")
        assert v["exclusions"] == expected, wt


def test_generate_gc_polish_uses_polish_wording():
    body = {"work_type": "polish", "audience": "GC", "values": _gc_vals()}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    txt = _rendered(client.get(r.json()["docx_download_url"]).content)
    # GC Polish scope step unique to the polish template.
    assert "Apply joint filler" in txt
    # GC Polish exclusions, not the Resinous ones.
    assert "Cove Base, Dye" in txt
    assert "Epoxy Paint Walls" not in txt


# ── (e) template_version guard ──────────────────────────────────────────
_DIRECT_BASE_VALS = {
    "lump_sum_formatted": "$61,162.00", "tax_amount_formatted": "$2,639.00",
    "state_name": "Kansas", "total_formatted": "$63,801.00",
    "total_label": "$63,801.00 – Total", "system_name": "MACRO", "texture": "OP",
    "epoxy_sf": "12,000", "cove_lf": "250", "bid_date_formatted": "6/15/26",
    "job_name": "J", "city_state": "C", "area_description": "x", "disposal": "d",
    "site_visit_date": "6/15", "work_description": "w",
    "base_bid_formatted": "$58,523.00", "material_tax_formatted": "$2,639.00",
    "base_tax_phrase": "(material sales tax INCLUDED)",
    "site_visit_phrase": "per site visit on 6/15/26",
}
_EPOXY_TEMPLATE = "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx"


def _scope_override_id():
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    for idx, _k, _p, in_block, text, _tb in pw.iter_editable_blocks(d):
        if in_block is None and text.strip().startswith("Scope:"):
            return idx
    raise AssertionError("no editable Scope: paragraph")


def test_version_guard_drops_stale_but_applies_matching_and_empty():
    sid = _scope_override_id()
    ov = [{"id": sid, "text": "Scope:  GUARD-TEST-OVERRIDE"}]
    base = {"work_type": "epoxy", "audience": "Direct",
            "values": dict(_DIRECT_BASE_VALS), "paragraph_overrides": ov}

    # Stale version -> overrides dropped (custom text does NOT appear).
    r = client.post("/api/generate", json={**base, "template_version": "STALE-NOPE"})
    assert r.status_code == 200, r.text
    assert "GUARD-TEST-OVERRIDE" not in _rendered(client.get(r.json()["docx_download_url"]).content)

    # Empty version (legacy caller) -> applied.
    r = client.post("/api/generate", json={**base, "template_version": ""})
    assert r.status_code == 200, r.text
    assert "GUARD-TEST-OVERRIDE" in _rendered(client.get(r.json()["docx_download_url"]).content)

    # Matching version -> applied.
    tv = client.get("/api/proposal-template?work_type=epoxy&audience=Direct").json()["template_version"]
    r = client.post("/api/generate", json={**base, "template_version": tv})
    assert r.status_code == 200, r.text
    assert "GUARD-TEST-OVERRIDE" in _rendered(client.get(r.json()["docx_download_url"]).content)


# ── (f) Direct parity + checksum guard ──────────────────────────────────
def test_direct_narrative_defaults_unchanged():
    v = {"work_type": "epoxy"}
    main._ensure_value_aliases(v)                       # no audience -> Direct
    assert v["scope_notes"] == main._DEFAULT_SCOPE_EPOXY
    assert v["schedule_notes"] == main._DEFAULT_SCHEDULE
    assert v["exclusions"] == main._DEFAULT_EXCLUSIONS

    v = {"work_type": "polish"}
    main._ensure_value_aliases(v, "Direct")
    assert v["scope_notes"] == main._DEFAULT_SCOPE_POLISH
    assert v["schedule_notes"] == main._DEFAULT_SCHEDULE
    assert v["exclusions"] == main._DEFAULT_EXCLUSIONS


def test_direct_templates_are_byte_identical():
    for rel, want16 in _DIRECT_TEMPLATES.items():
        data = (pw.TEMPLATES_ROOT / rel).read_bytes()
        got16 = hashlib.sha256(data).hexdigest()[:16]
        assert got16 == want16, f"{rel}: SHA256 changed ({got16} != {want16})"
