"""Per-room estimate + proposal feature.

- proposal_writer `{{#room}}` block (heading + price + stacked notes), the
  `{{#single_bid}}` toggle, and the \\n -> <w:br/> note rendering.
- main._build_options (tax phrase, system line + signed difference toggles, notes).
- estimate_writer.fill_estimate(tab_copies=…) duplicating a worksheet.

Block-engine tests run on synthetic docs (no template dependency); real-template
rendering is covered in test_price_lines.py once the Direct templates carry the
{{#room}} / {{#single_bid}} markers.
"""
import io

from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from openpyxl import load_workbook

import estimate_writer as ew
import main
import proposal_writer as pw

client = TestClient(main.app)


def _doc(lines):
    d = Document()
    for line in lines:
        d.add_paragraph(line)
    return d


def _texts(d):
    return [p.text for p in d.paragraphs]


# ── {{#room}} block ────────────────────────────────────────────────────
def test_room_block_expands_heading_price_notes():
    for n in (0, 1, 3):
        rooms = [{"heading": f"Room {i}:", "price_formatted": f"${i},000",
                  "price_desc": "Epoxy flooring as described above (material sales tax INCLUDED)",
                  "notes_joined": 'Includes 6" Cove Base'} for i in range(1, n + 1)]
        d = _doc(["{{#room}}", "{{room.heading}}",
                  "{{room.price_formatted}} – {{room.price_desc}}",
                  "{{room.notes_joined}}", "{{/room}}"])
        pw._expand_all_blocks(d, {"room": rooms})
        txt = "\n".join(_texts(d))
        assert "{{#room}}" not in txt and "{{/room}}" not in txt
        assert "room." not in txt                       # all per-item tokens resolved
        for i in range(1, n + 1):
            assert f"Room {i}:" in txt
            assert f"${i},000 – Epoxy flooring as described above (material sales tax INCLUDED)" in txt
        if n:
            assert 'Includes 6" Cove Base' in txt


def test_room_notes_render_as_linebreaks():
    rooms = [{"heading": "Grooming:", "price_formatted": "$8,310",
              "price_desc": "Epoxy flooring (material sales tax INCLUDED)",
              "notes_joined": 'Includes 6" Cove Base\nTo be completed in separate mobilization'}]
    d = _doc(["{{#room}}", "{{room.heading}}", "{{room.notes_joined}}", "{{/room}}"])
    pw._expand_all_blocks(d, {"room": rooms})
    txt = "\n".join(_texts(d))
    assert 'Includes 6" Cove Base' in txt
    assert "To be completed in separate mobilization" in txt
    # stacked notes share one paragraph/bullet but split on a real <w:br/> line break
    assert d.element.body.findall(".//" + qn("w:br"))


def test_single_bid_block_toggles():
    for items, shown in (([{}], True), ([], False)):
        d = _doc(["before", "{{#single_bid}}", "Base Bid",
                  "{{total_formatted}} – Total", "{{/single_bid}}", "after"])
        pw._expand_all_blocks(d, {"single_bid": items})
        txt = "\n".join(_texts(d))
        assert "before" in txt and "after" in txt
        assert "{{#single_bid}}" not in txt and "{{/single_bid}}" not in txt
        assert ("Base Bid" in txt) is shown


# ── main._build_options (total / deduct modes, show gate, base excluded) ──
def test_build_options_total_mode():
    rooms_in = [
        {"is_base": True, "bid": {"total": 50000}},                # base -> single_bid, excluded here
        {"name": "Epoxy copy", "is_base": False, "base_total": 50000,
         "price_mode": "total", "option_desc": "Treadwell MACRO Flake",
         "bid": {"total": 8310, "remodel": 0},
         "notes_auto": ['Includes 6" Cove Base'], "notes_manual": ["separate mobilization"]},
        {"name": "Remodel opt", "is_base": False, "base_total": 50000,
         "price_mode": "total", "bid": {"total": 15035, "remodel": 120}},
        {"name": "Empty", "is_base": False, "bid": {"total": 0}},  # no total -> skipped
    ]
    out = main._build_options(rooms_in, {"state_name": "Kansas"}, "epoxy")
    assert len(out) == 2                                           # base + empty excluded
    a, b = out
    assert a["heading"] == ""                                      # no per-option heading now
    assert a["price_formatted"] == "$8,310"
    assert a["price_desc"] == "Treadwell MACRO Flake as described above (material sales tax INCLUDED)"
    assert a["notes_joined"] == 'Includes 6" Cove Base\nseparate mobilization'   # auto before manual
    assert "Remodel Tax AND material sales tax INCLUDED" in b["price_desc"]
    assert "Kansas" not in b["price_desc"]                         # state name dropped


def test_build_options_deduct_mode():
    # Auto add/deduct: diff = option − base = −6,000 → a Deduct, word in the AMOUNT.
    opt = {"name": "Grind & Seal", "is_base": False, "base_total": 50000,
           "price_mode": "deduct", "option_desc": "Grind & Seal", "base_desc": "Polished Concrete",
           "bid": {"total": 44000, "remodel": 0}}
    o = main._build_options([opt], {}, "polish")[0]
    assert o["price_formatted"] == "Deduct ($6,000)"
    assert o["price_desc"] == "VE for Grind & Seal, in lieu of Polished Concrete."


def test_build_options_deduct_costlier_becomes_add():
    # Option costs MORE than the base → diff = +11,000 → an ADD line (Will's
    # spec: option − base, positive = add), no silent fall-back to its own total.
    opt = {"name": "Premium", "is_base": False, "base_total": 50000,
           "price_mode": "deduct", "option_desc": "Premium System",
           "bid": {"total": 61000, "remodel": 0}}
    o = main._build_options([opt], {}, "epoxy")[0]
    assert o["price_formatted"] == "Add $11,000"
    assert o["price_desc"] == "Premium System"


def test_build_options_deduct_equal_totals_is_add_zero():
    # Same price as the base → diff = 0 → renders as "Add $0" (edge, but defined).
    opt = {"name": "Same", "is_base": False, "base_total": 50000,
           "price_mode": "deduct", "option_desc": "Same System",
           "bid": {"total": 50000, "remodel": 0}}
    o = main._build_options([opt], {}, "epoxy")[0]
    assert o["price_formatted"] == "Add $0"


def test_build_options_show_gate_and_base_excluded():
    rooms = [
        {"is_base": True, "bid": {"total": 50000}},
        {"name": "Hidden", "is_base": False, "base_total": 50000, "show": False,
         "price_mode": "total", "bid": {"total": 8000}},
        {"name": "Shown", "is_base": False, "base_total": 50000, "show": True,
         "price_mode": "total", "bid": {"total": 9000}},
    ]
    out = main._build_options(rooms, {}, "epoxy")
    assert len(out) == 1                                           # base + hidden excluded
    assert out[0]["price_formatted"] == "$9,000"


def test_build_options_empty():
    assert main._build_options([], {}, "epoxy") == []
    assert main._build_options(None, {}, "epoxy") == []


def test_flooring_noun_and_fmt_usd_parens():
    assert main._flooring_noun("epoxy") == "Epoxy flooring"
    assert main._flooring_noun("combo") == "Epoxy flooring"
    assert main._flooring_noun("polish") == "Polished Concrete Flooring"
    assert main._fmt_usd(6000, parens=True) == "($6,000)"
    assert main._fmt_usd(6000.5, parens=True) == "($6,000.50)"
    assert main._fmt_usd(6000) == "$6,000"


# ── estimate_writer.fill_estimate(tab_copies=…) — duplicated worksheets ─
def test_fill_estimate_duplicates_tab():
    data = ew.fill_estimate({"project_name": "X"},
                            cell_values={"Copy1!E20": 4000},
                            tab_copies=[{"id": "Copy1", "source": "Epoxy"}])
    wb = load_workbook(io.BytesIO(data))
    assert "Copy1" in wb.sheetnames                   # copied worksheet created
    assert "Epoxy" in wb.sheetnames                   # source intact
    ws = wb["Copy1"]
    assert ws["E20"].value == 4000                    # copy's own cell write landed
    assert str(ws["D88"].value or "").startswith("=")            # bid formula copied (self-ref)
    assert ws["B1"].value == wb["Epoxy"]["B1"].value             # project info mirror copied


def test_fill_estimate_no_copies_unchanged():
    data = ew.fill_estimate({"project_name": "X"})
    wb = load_workbook(io.BytesIO(data))
    assert "Copy1" not in wb.sheetnames


# ── real Direct templates (epoxy + combo) render the {{#room}} block ────
_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

_VALS = {
    "job_name": "J", "city_state": "C", "bid_date_formatted": "6/19/26",
    "base_bid_formatted": "$58,523.00", "material_tax_formatted": "$2,639.00",
    "state_name": "Kansas", "total_formatted": "$63,801.00",
    "system_name": "MACRO", "texture": "OP", "epoxy_sf": "12,000",
    "scope_notes": "demo + install", "schedule_notes": "~5 days",
    "work_description": "w", "site_visit_date": "6/19", "disposal": "d",
    "site_visit_phrase": "per site visit on 6/19/26",
    "base_tax_phrase": "(material sales tax INCLUDED)",
    "exclusions": "standard exclusions",
}
def _rendered(docx_bytes):
    """Text of paragraphs Word actually renders (mc:Choice copy), excluding the
    legacy mc:Fallback duplicate — same helper as test_price_lines.py."""
    d = Document(io.BytesIO(docx_bytes))
    out = []
    for p in d.element.xpath("//w:p"):
        if any(True for _ in p.iterancestors(f"{_MC}Fallback")):
            continue
        t = "".join(x.text or "" for x in p.xpath(".//w:t")).strip()
        if t:
            out.append(t)
    return "\n".join(out)


def test_notes_block_editable_with_default_fallback():
    # Fix #6: NOTES are an editable {{#notes}} block. Custom notes render;
    # _notes_for falls back to the standard per-work-type boilerplate when empty.
    blob = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=_VALS,
                                      notes=[{"text": "Custom note A"}, {"text": "Custom note B"}]))
    assert "Custom note A" in blob and "Custom note B" in blob
    assert "{{notes" not in blob and "{{#notes}}" not in blob
    items = main._notes_for("epoxy", [])             # empty -> standard boilerplate
    assert items and "Excludes saw cutting" in items[0]["text"]
    assert main._notes_for("epoxy", ["only this"]) == [{"text": "only this"}]
    blob2 = _rendered(pw.fill_proposal(work_type="polish", audience="Direct", values=_VALS,
                                       notes=main._notes_for("polish", [])))
    assert "Excludes saw cutting" in blob2


def test_site_visit_phrase_and_base_tax_tokens_render():
    # Fix #3/#4: epoxy template uses {{site_visit_phrase}} + {{base_tax_phrase}};
    # both must fill (never render literally) and the old hardcoded phrasing is gone.
    vals = {**_VALS, "site_visit_phrase": "per plans and specifications provided",
            "base_tax_phrase": "(tax exempt)"}
    blob = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=vals))
    assert "per plans and specifications provided" in blob
    assert "(tax exempt)" in blob
    assert "{{site_visit_phrase}}" not in blob and "{{base_tax_phrase}}" not in blob
    assert "per site visit on N/A" not in blob


def test_exclusions_token_carries_to_doc():
    # Fix #1: the Exclusions box must drive the doc (was hardcoded boilerplate).
    for wt in ("epoxy", "combo", "polish"):
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct",
                                          values={**_VALS, "exclusions": "Existing floor demo only"}))
        assert "Existing floor demo only" in blob, f"{wt}: edited exclusions did not carry"
        assert "Multiple layers of floor to be removed" not in blob, f"{wt}: stale hardcoded exclusions"


def test_remodel_tax_label_has_no_state_name():
    # Fix #2: the remodel-tax line reads "Remodel Tax", never "Kansas/Missouri Remodel Tax".
    import re
    for wt in ("epoxy", "combo"):
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct", values=_VALS,
                                          remodel=[{"amount_formatted": "$1,200"}]))
        assert "Remodel Tax" in blob, f"{wt}: remodel line missing"
        assert not re.search(r"(Kansas|Missouri)\s+Remodel Tax", blob), f"{wt}: state name still on remodel tax"


def test_options_render_via_price_lines():
    # Options now ride the {{#price_line}} block (under the "Options:" heading, after
    # the base bid); the base itself shows via {{#single_bid}}. Total + deduct lines
    # render with Treadwell's "$x – …" / "($x) – Deduct VE …" phrasing.
    import re
    price_lines = [
        {"amount_formatted": "$8,310",
         "label": "Treadwell MACRO Flake as described above (material sales tax INCLUDED)"},
        {"amount_formatted": "($6,000)",
         "label": "Deduct VE for Grind & Seal, in lieu of Epoxy flooring."},
    ]
    for wt in ("epoxy", "combo"):
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct", values=_VALS,
                                          price_lines=price_lines, has_options=True,
                                          single_bid=None, tax_breakout=True))
        assert "Base Bid" in blob, f"{wt}: base bid missing"
        assert "$63,801.00 – Total" in blob, f"{wt}: base total missing"
        assert ("$8,310 – Treadwell MACRO Flake as described above "
                "(material sales tax INCLUDED)") in blob, f"{wt}: option total line"
        assert "($6,000) – Deduct VE for Grind & Seal, in lieu of Epoxy flooring." in blob, f"{wt}: deduct line"
        assert not re.search(r"\{\{[#/]", blob), f"{wt}: leftover block marker"


def test_no_options_keeps_single_bid_clean():
    import re
    for wt in ("epoxy", "combo"):
        # tax_breakout=True so the itemized Total line renders; base always shows.
        blob = _rendered(pw.fill_proposal(work_type=wt, audience="Direct", values=_VALS,
                                          tax_breakout=True))
        assert "Base Bid" in blob, f"{wt}: single-bid Base Bid missing"
        assert "$63,801.00 – Total" in blob, f"{wt}: single-bid Total missing"
        assert "Deduct VE" not in blob, f"{wt}: stray option content"
        assert not re.search(r"\{\{[#/]", blob), f"{wt}: leftover block marker"


def test_combo_per_option_breakout_render():
    # Combo: Option 1 (Epoxy) + Option 2 (Polish), each with its own flooring /
    # Kansas Remodel Tax / Total, lead the PRICE section (via {{#price_line}}); the
    # combined single-bid line is suppressed (single_bid=[]).
    import re
    combo_lines = [
        {"amount_formatted": "$28,400", "label": "Option 1: Epoxy flooring as described above (material sales tax INCLUDED)"},
        {"amount_formatted": "$1,200",  "label": "Kansas Remodel Tax"},
        {"amount_formatted": "$29,600", "label": "Total"},
        {"amount_formatted": "$17,200", "label": "Option 2: Polished Concrete flooring as described above (material sales tax INCLUDED)"},
        {"amount_formatted": "$17,200", "label": "Total"},
    ]
    blob = _rendered(pw.fill_proposal(work_type="combo", audience="Direct", values=_VALS,
                                      price_lines=combo_lines, has_options=True, single_bid=[]))
    assert "$28,400 – Option 1: Epoxy flooring as described above (material sales tax INCLUDED)" in blob
    assert "$17,200 – Option 2: Polished Concrete flooring as described above (material sales tax INCLUDED)" in blob
    assert "$1,200 – Kansas Remodel Tax" in blob
    assert "Epoxy & Polished Concrete flooring" not in blob   # combined single-bid line suppressed
    assert not re.search(r"\{\{[#/]", blob)


# ── main.api_generate — combo_options sanitization (malformed shapes) ──────
def test_combo_options_malformed_shapes_do_not_500():
    # A client sending {"label": 123} (or a dict/list/bool value for label /
    # amount_formatted) used to blow up main.api_generate: the old filter called
    # `.strip()` on the raw value BEFORE coercing to str(), and .strip() doesn't
    # exist on int/dict/list/bool -> unhandled AttributeError -> 500. Each shape
    # below IS a dict (so it wasn't caught by the old isinstance(c, dict) guard);
    # only the field VALUE is the wrong type. Drives the real /api/generate route
    # (no X-Project-Id => no DB write) so this proves the sanitization landed on
    # the actual request path, not just in isolation.
    malformed = [
        {"label": 123},
        {"label": {"a": 1}},
        {"amount_formatted": ["$1"]},
        {"label": True},
    ]
    body = {"work_type": "combo", "audience": "Direct", "values": dict(_VALS),
            "combo_options": malformed}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text


def test_combo_options_list_is_capped():
    # No accepted-list cap meant N client-supplied entries -> N deep-copied
    # paragraph sets in the generated docx. 200 fake entries must not all render;
    # only the first 50 are honored. Unique "CapLabelN" tokens so the assertion
    # can't collide with unrelated doc text. (Note: _rendered's raw text repeats
    # each real paragraph a fixed number of times because of how the Direct
    # templates duplicate the price/textbox content — that's a pre-existing
    # quirk of the fixture, not this fix — so this checks presence/absence of
    # specific labels rather than counting occurrences.)
    many = [{"label": f"CapLabel{i}", "amount_formatted": f"${i}"} for i in range(200)]
    body = {"work_type": "combo", "audience": "Direct", "values": dict(_VALS),
            "combo_options": many}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    f = client.get(r.json()["docx_download_url"])
    assert f.status_code == 200, f.text
    text = _rendered(f.content)
    assert "CapLabel0" in text and "CapLabel49" in text
    assert "CapLabel50" not in text and "CapLabel199" not in text


# ── Combo breakout + trailing options: restored "Options:" separator ──────
# Bug: the template's {{#has_options}}Options:{{/has_options}} label is
# nested INSIDE {{#single_bid}}. The combo breakout suppresses single_bid
# (single_bid=[]) so its own Option 1/2 totals stand in for the base price —
# but that also deletes the nested "Options:" heading, so a trailing manual
# "Add for …" price line rendered right after the combo block with no visual
# separator from the base price. Fixed in two places: main.py restores the
# heading as a label-only price_line row between the combo lines and any
# trailing price lines; proposal_writer.py strips the template's hardcoded
# " – " separator off that row since its amount is empty.
_COMBO_LINES = [
    {"amount_formatted": "$28,400",
     "label": "Option 1: Epoxy flooring as described above (material sales tax INCLUDED)"},
    {"amount_formatted": "$1,200", "label": "Kansas Remodel Tax"},
    {"amount_formatted": "$29,600", "label": "Total"},
    {"amount_formatted": "$17,200",
     "label": "Option 2: Polished Concrete flooring as described above (material sales tax INCLUDED)"},
    {"amount_formatted": "$17,200", "label": "Total"},
]


def test_combo_breakout_with_trailing_line_restores_options_separator():
    body = {"work_type": "combo", "audience": "Direct", "values": dict(_VALS),
            "combo_options": _COMBO_LINES,
            "price_lines": [{"label": "Add for moisture mitigation", "amount": 500}]}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    f = client.get(r.json()["docx_download_url"])
    text = _rendered(f.content)
    # Clean label-only heading, no leading dash/en-dash slop.
    assert "Options:" in text
    assert "– Options:" not in text and "-Options:" not in text
    # Sits right after the LAST combo "Total" line and right before the
    # trailing manual price line — not floating loose, not merged with either.
    assert ("$17,200 – Total\nOptions:\n$500 – Add for moisture mitigation") in text
    # The suppressed combined single-bid line never reappears.
    assert "Epoxy & Polished Concrete flooring" not in text


def test_combo_breakout_alone_has_no_options_label():
    # No trailing option/manual price lines after the combo breakout -> nothing
    # for a separator to introduce, so "Options:" must not appear at all.
    body = {"work_type": "combo", "audience": "Direct", "values": dict(_VALS),
            "combo_options": _COMBO_LINES}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    f = client.get(r.json()["docx_download_url"])
    text = _rendered(f.content)
    assert "Options:" not in text
    assert "$17,200 – Total" in text                      # combo lines still render


def test_non_combo_options_label_unaffected_by_separator_fix():
    # Non-combo path: "Options:" comes from the {{#has_options}} block (a
    # plain-text label paragraph), NOT from a price_line row — the
    # proposal_writer dash-strip only fires for a price_line row with an
    # empty amount_formatted, so this path is untouched by either fix.
    price_lines = [{"amount_formatted": "$8,310",
                    "label": "Treadwell MACRO Flake as described above (material sales tax INCLUDED)"}]
    blob = _rendered(pw.fill_proposal(work_type="epoxy", audience="Direct", values=_VALS,
                                      price_lines=price_lines, has_options=True,
                                      single_bid=None, tax_breakout=True))
    assert "Options:" in blob
    assert "– Options:" not in blob and "-Options:" not in blob


# ── Proposal Review PRICE preview parity (backend contract the fix mirrors) ─
# The Proposal Review PAGE preview used to duplicate the base bid + hide the
# "Options:" label. The generated .docx has always been correct: ONE base bid
# (via {{#single_bid}}), then the "Options:" label, then the priced option lines
# (via {{#price_line}}). These guard that backend contract so the fixed preview
# has something stable to mirror. (Fixture quirk: the templates ALSO concatenate
# the whole price block into one anchor paragraph, so we match lines that ARE
# EXACTLY a heading — the anchor line contains them only as substrings.)
def test_generate_rooms_render_single_base_and_options_label():
    rooms = [
        {"id": "Epoxy", "is_base": True, "name": "Base",
         "system_desc": "MACRO Flake Single Broadcast", "bid": {"total": 36763, "remodel": 0}},
        {"id": "Copy1", "is_base": False, "show": True, "base_total": 36763,
         "price_mode": "total", "option_desc": "Quartz Double Broadcast",
         "system_desc": "Quartz Double Broadcast", "bid": {"total": 11126, "remodel": 0},
         "notes_auto": ["Includes cove base"], "notes_manual": []},
        # Add/deduct option that costs MORE than the base → an ADD line:
        # "Add $12,014 – Premium System" (option − base; Will's spec — no more
        # silent fall-back to the option's own total).
        {"id": "Copy2", "is_base": False, "show": True, "base_total": 36763,
         "price_mode": "deduct", "option_desc": "Premium System",
         "base_desc": "MACRO Flake Single Broadcast", "bid": {"total": 48777, "remodel": 0}},
    ]
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_VALS), "rooms": rooms}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    lines = _rendered(client.get(r.json()["docx_download_url"]).content).split("\n")
    assert lines.count("Base Bid") == 1                # exactly one standalone base heading
    assert "Options:" in lines                          # the label renders
    base_i = lines.index("Base Bid")
    opt_i = lines.index("Options:")
    assert base_i < opt_i                               # base first, then the Options label
    after = lines[opt_i + 1:]                            # the real option lines follow it
    quartz = next(l for l in after if "Quartz Double Broadcast" in l)
    premium = next(l for l in after if "Premium System" in l)
    assert " — Includes cove base" in quartz       # auto note folded inline (main.py " — ")
    assert "Add $12,014" in premium                      # 48,777 − 36,763 → an Add line
    assert "Deduct" not in "\n".join(after)


def test_generate_no_options_price_section_unchanged():
    rooms = [{"id": "Epoxy", "is_base": True, "name": "Base",
              "system_desc": "MACRO Flake Single Broadcast", "bid": {"total": 36763, "remodel": 0}}]
    body = {"work_type": "epoxy", "audience": "Direct", "values": dict(_VALS), "rooms": rooms}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    lines = _rendered(client.get(r.json()["docx_download_url"]).content).split("\n")
    assert lines.count("Base Bid") == 1                 # base still shows via single_bid
    assert "Options:" not in lines                       # no options -> label stripped
    assert "Deduct VE" not in "\n".join(lines)
