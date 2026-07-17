"""The Terms & Conditions section must start on its own page. Kyle's templates
ship without a forced break, so a short body (notably the combo) let the T&C
heading + its letterhead spill onto the bottom of page 1, over the ACCEPTANCE
frame. proposal_writer._force_terms_on_new_page adds a <w:pageBreakBefore/>."""
import io

import proposal_writer as pw
import main
from docx import Document
from docx.oxml.ns import qn

VALS = {
    "job_name": "QA", "city_state": "Olathe, KS", "bid_date_formatted": "7/22/26",
    "system_name": "X", "texture": "Smooth", "epoxy_sf": "4,000", "cove_lf": "180", "polish_sf": "3,000",
    "epoxy_system_name": "Epoxy System", "site_visit_phrase": "per site visit",
    "schedule_notes": main._DEFAULT_SCHEDULE, "exclusions": main._DEFAULT_EXCLUSIONS, "work_notes": "",
    "base_tax_phrase": "(material sales tax INCLUDED)", "lump_sum_formatted": "$32,985.00",
    "total_formatted": "$32,985.00", "base_bid_formatted": "$8,967.00", "estimator_name": "Hanz",
    "scope_notes": "scope", "disposal": "dumpster", "state_name": "Kansas",
    "material_tax_formatted": "$100.00", "tax_amount_formatted": "$100.00",
}


def _terms_forced_to_new_page(doc) -> bool:
    """True if a <w:pageBreakBefore/> sits on the T&C heading or the letterhead
    paragraph just before it (so the whole section starts a fresh page)."""
    tops = [c for c in doc.element.body if c.tag == qn("w:p")]
    h = next((i for i, p in enumerate(tops)
              if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip().upper() == "TERMS AND CONDITIONS"), None)
    assert h is not None, "T&C heading not found"
    for j in range(h, max(-1, h - 4), -1):
        ppr = tops[j].find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
            return True
    return False


def test_combo_terms_start_on_own_page():
    doc = Document(io.BytesIO(pw.fill_proposal(work_type="combo", audience="Direct", values=dict(VALS))))
    assert _terms_forced_to_new_page(doc)


def test_epoxy_terms_start_on_own_page():
    doc = Document(io.BytesIO(pw.fill_proposal(work_type="epoxy", audience="Direct", values=dict(VALS))))
    assert _terms_forced_to_new_page(doc)


def test_polish_terms_start_on_own_page():
    doc = Document(io.BytesIO(pw.fill_proposal(work_type="polish", audience="Direct", values=dict(VALS))))
    assert _terms_forced_to_new_page(doc)


# ── gyp NOTES box top inset (first bullet clears its frame border) ──────────
_NOTES = [{"text": "Excludes Union Labor/Prevailing Wage Labor bond and liquidated damages"},
          {"text": "Excludes hoisting or lifting of equipment to elevated slabs"}]


def _notes_box_tins(doc):
    for txbx in pw._iter_txbx(doc):
        if "Excludes Union Labor" in "".join(t.text or "" for t in txbx.iter(qn("w:t"))):
            shape = pw._shape_of_txbx(txbx)
            for bp in (shape.iter() if shape is not None else []):
                if bp.tag.endswith("}bodyPr"):
                    try:
                        return int(bp.get("tIns") or 0)
                    except (TypeError, ValueError):
                        return 0
    return None


def test_gyp_notes_box_gets_top_inset():
    doc = Document(io.BytesIO(pw.fill_proposal(work_type="gyp", audience="Direct", values=dict(VALS), notes=_NOTES)))
    assert _notes_box_tins(doc) == pw._GYP_NOTES_TOP_INSET_EMU


def test_non_gyp_notes_box_not_padded():
    doc = Document(io.BytesIO(pw.fill_proposal(work_type="polish", audience="Direct", values=dict(VALS), notes=_NOTES)))
    tins = _notes_box_tins(doc)
    assert tins is None or tins < pw._GYP_NOTES_TOP_INSET_EMU   # polish NOTES left as designed
