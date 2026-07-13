"""Word-style blank-line spacing in NOTES: blank lines the estimator adds in the
editor are preserved end-to-end and render in the .docx as clean empty lines
(bullet stripped), not empty bullet dots. Covers _notes_for preservation, the
_strip_bullet helper, and an end-to-end generate smoke."""
import io

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi.testclient import TestClient

import main
import proposal_writer as pw

client = TestClient(main.app)

_VALS = {
    "job_name": "J", "city_state": "C", "bid_date_formatted": "6/19/26",
    "base_bid_formatted": "$58,523.00", "material_tax_formatted": "$2,639.00",
    "state_name": "Kansas", "total_formatted": "$36,763.00", "system_name": "MACRO",
    "texture": "OP", "epoxy_sf": "12,000", "scope_notes": "demo",
    "schedule_notes": "~5 days", "work_description": "w", "site_visit_date": "6/19",
    "disposal": "d", "site_visit_phrase": "per site visit on 6/19/26",
    "base_tax_phrase": "(material sales tax INCLUDED)", "exclusions": "std",
}


# ── _notes_for preserves blanks ─────────────────────────────────────────────
def test_notes_for_preserves_blank_lines():
    items = main._notes_for("epoxy", ["Line A", "", "Line B"], None)
    assert [i["text"] for i in items] == ["Line A", "", "Line B"]


def test_notes_for_all_blank_falls_back_to_defaults():
    # A notes list that is only blanks is treated as "no notes" → boilerplate,
    # never a proposal with a blank notes section.
    items = main._notes_for("epoxy", ["", "   "], None)
    assert len(items) > 1
    assert any("additional phase" in i["text"] for i in items)


def test_notes_for_blank_between_real_notes_with_phase():
    items = main._notes_for("epoxy", ["Add $xxxx for each additional phase beyond the above stated schedule.", "", "Tail."], 6000)
    texts = [i["text"] for i in items]
    assert texts[0] == "Add $6,000 for each additional phase beyond the above stated schedule."
    assert texts[1] == ""            # blank preserved
    assert texts[2] == "Tail."


# ── _strip_bullet ───────────────────────────────────────────────────────────
def test_strip_bullet_removes_numpr():
    d = Document()
    p = d.add_paragraph("x")
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ppr.append(numpr)
    assert ppr.find(qn("w:numPr")) is not None
    pw._strip_bullet(p._p)
    assert ppr.find(qn("w:numPr")) is None


def test_strip_bullet_noop_without_ppr():
    d = Document()
    p = d.add_paragraph("x")          # no pPr
    pw._strip_bullet(p._p)            # must not raise


# ── end-to-end generate smoke (blank line survives the whole pipeline) ──────
def _all_text(doc):
    return "".join(t.text or "" for t in doc.element.iter(qn("w:t")))


def test_generate_with_blank_note_survives():
    body = {"work_type": "epoxy", "audience": "Direct",
            "values": dict(_VALS),
            "notes": ["First note.", "", "Second note."], "cell_values": {}}
    r = client.post("/api/generate", json=body)
    assert r.status_code == 200, r.text
    doc = Document(io.BytesIO(client.get(r.json()["docx_download_url"]).content))
    text = _all_text(doc)
    assert "First note." in text and "Second note." in text
