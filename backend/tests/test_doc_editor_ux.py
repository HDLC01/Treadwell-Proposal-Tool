"""Kyle's four complaints about the Proposal Review document editor, 2026-08-19.

Verbatim:
  (a) "Some of the labels are not editable why not make it like a word document??"
  (b) "when he pressed enter to add spacing it did not generate in the proposal"
  (c) "the textboxes are clunky"
  (d) "He is confused on how to get out of that Textbox view can you improve that as well"

WHAT THE INVESTIGATION FOUND, because the fixes only make sense against it.

(a) is TWO different things wearing one sentence.

    "Epoxy Flooring:", "Scope:", "Schedule:", "Exclusions:" and "Notes:" are real text in real
    editable paragraphs — blocks 109 and 115-118 of the Direct epoxy template, `in_block` None,
    each rendered as a whole contenteditable `.tw-block`. They were ALREADY editable; nothing
    on screen said so, because only the `{{token}}` values carried a cue. That is complaint
    (c), and it is fixed as (c). `test_a_work_label_is_ordinary_editable_text_that_stays_bold`
    proves the label round-trips and keeps its bold;
    `test_an_emptied_label_leaves_no_stray_token_and_no_lone_colon` proves clearing it does not
    expose the token it sat in front of.

    "System:" / "Option N:", "Texture:" and "Area:" are real docx text too (blocks 111-113) —
    NOT baked into the letterhead artwork — but they live inside the `{{#system}}` repeat
    region, which the editor collapses into one read-only `.tw-priced-region` and replaces with
    a preview whose labels are frontend `<strong>` text (proposal-review.js renderSystemPreview).
    `proposal_writer._apply_paragraph_overrides` refuses any id whose `in_block` is not None on
    purpose. So they are genuinely not editable, and making them so is a feature — a new
    override channel through main.py, not a CSS fix.
    `test_the_region_labels_are_real_docx_text_not_artwork` pins the evidence so the next person
    does not have to re-derive it.

(b) had a cause at each end.

    The BROWSER end: what Enter does to a contenteditable is not one thing. Depending on the
    engine and on `white-space` it inserts a `<br>`, a bare "\\n", or a wrapper `<div>` with its
    own placeholder `<br>` — and `serializeBlock` reads that last shape as TWO newlines. So one
    Enter could become a blank line and two could become three.

    The SERVER end, which is where the blank line actually died:
    `_normalize_work_label_formatting` re-bolds a WORK row up to its first colon by splitting
    the run at that character. It measured the run with `"".join(t.text …)`, which is BLIND to
    `<w:br/>`, then handed that string to `_set_direct_run_text`, which clears the run's
    w:t/w:br children and rewrites them from what it was given. Measured before the fix, a
    plain-text override of "Scope:  line one\\n\\nline two" came out of the generator as
    `[t "Scope:", t "  line oneline two"]` — both breaks gone, silently, in a customer document.
    Only the WORK box is affected, and the WORK box is where Scope / Schedule / Exclusions /
    Notes live, which is exactly what Kyle was typing into.

(d) was diagnosed and confirmed: `wireOverflowExpand` toggled on a click on the box and
    deliberately ignored clicks on `.tw-block` / `.tw-line-edit` / `[contenteditable=true]` so a
    click meant for a paragraph puts a caret in it. An OPEN box is nearly all editable content,
    so there was frequently no pixel left that would close it again.

WHY THE FRONTEND HALF RUNS UNDER NODE. Every claim above is a behaviour: which of four clicks
reaches which branch, and whether three walkers agree on one character. A source-text assertion
cannot see either, and this repo has already paid for that lesson — on 2026-08-12 `STAGE_CREATED`
shipped unbound with every source assertion green and took the production board down. So
`js/doc-editor-harness.js` lifts the shipped functions out of proposal-review.js, gives them the
smallest DOM they touch, and fires real events at them.
"""
import io
import json
import pathlib
import re
import shutil
import subprocess

import docx
import pytest
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

import proposal_writer as pw

FRONTEND = pathlib.Path(__file__).resolve().parents[2] / "frontend"
HARNESS = pathlib.Path(__file__).resolve().parent / "js" / "doc-editor-harness.js"
CSS = (FRONTEND / "styles.css").read_text(encoding="utf-8")

_EPOXY_TEMPLATE = "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx"


@pytest.fixture(scope="module")
def ran():
    if shutil.which("node") is None:
        pytest.skip("node is not installed")
    proc = subprocess.run(["node", str(HARNESS), str(FRONTEND)],
                          capture_output=True, text=True, encoding="utf-8", timeout=180)
    assert proc.returncode == 0, (
        "the harness itself failed — read this before assuming a product bug:\n" + proc.stderr)
    return json.loads(proc.stdout.strip().splitlines()[-1])


# ═══ shared docx helpers ═════════════════════════════════════════════════════
def _blocks(template_rel_path=_EPOXY_TEMPLATE):
    """The same walk the editor's ids come from, so a paragraph is found by its text rather
    than by a magic number that a template re-annotation would silently invalidate."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / template_rel_path))
    return [{"id": idx, "text": text, "in_block": in_block, "in_txbx": in_txbx}
            for idx, _kind, _p, in_block, text, in_txbx in pw.iter_editable_blocks(d)]


def _block_id(prefix, blocks=None):
    hits = [b for b in (blocks or _blocks()) if b["text"].startswith(prefix)]
    assert hits, "no template paragraph starts with %r any more" % prefix
    return hits[0]["id"]


def _generate(overrides):
    return pw.fill_proposal(
        work_type="epoxy", audience="Direct",
        values={"job_name": "Cedar Ridge Distribution Center", "scope_notes": "SCOPE"},
        paragraph_overrides=overrides,
    )


def _read_run(run_elem):
    """A run's text with each `<w:br/>` shown as a newline.

    Deliberately its OWN reader rather than `pw._run_text_with_breaks`: that function is half
    the fix under test, and measuring the output with it would make the assertion circular —
    a reader that lost the breaks would agree with a writer that lost them."""
    bits = []
    for el in run_elem.iter():
        if el.tag == qn("w:br"):
            bits.append("\n")
        elif el.tag == qn("w:t"):
            bits.append(el.text or "")
    return "".join(bits)


def _paragraph_with(docx_bytes, needle):
    """The generated paragraph whose OWN text contains `needle`, plus its run structure with
    the line breaks made visible.

    `_own_text` — which is what the editor reports and what the test's `needle` matches — sees
    no breaks at all, which is exactly why a br-blind reading of the paragraph cannot prove the
    break survived."""
    d = docx.Document(io.BytesIO(docx_bytes))
    for _idx, _kind, p_elem, _in_block, text, _txbx in pw.iter_editable_blocks(d):
        if needle in text:
            runs = p_elem.findall(qn("w:r"))
            return {
                "own_text": text,
                "with_breaks": "".join(_read_run(r) for r in runs),
                "breaks": sum(1 for r in runs for _ in r.iter(qn("w:br"))),
                "bold_runs": [_read_run(r) for r in runs if _is_bold(r) and _read_run(r)],
            }
    raise AssertionError("no generated paragraph contains %r" % needle)


def _is_bold(run_elem):
    rpr = run_elem.find(qn("w:rPr"))
    if rpr is None:
        return False
    b = rpr.find(qn("w:b"))
    return b is not None and b.get(qn("w:val")) not in ("0", "false")


# ═══ (b) the blank line, server end ══════════════════════════════════════════
def test_a_typed_blank_line_survives_the_work_label_normalizer():
    """THE bug behind "when he pressed enter to add spacing it did not generate in the
    proposal". A plain-text override arrives as ONE run holding w:t / w:br / w:t / w:br / w:t.
    `_normalize_work_label_formatting` splits that run at the label's colon, and it used to
    measure it with a `<w:br/>`-blind join before handing the result to
    `_set_direct_run_text`, which rewrites the run's children from the string it is given.

    Measured before the fix, this exact input generated `[t "Scope:", t "  line oneline two"]`.
    """
    sent = "Scope:  line one\n\nline two"
    got = _paragraph_with(_generate([{"id": _block_id("Scope:"), "text": sent}]), "line one")
    assert got["with_breaks"] == sent, (
        "the generated paragraph reads back as %r — the estimator's line breaks were joined "
        "away between the override and the file" % got["with_breaks"])
    assert got["breaks"] == 2, (
        "a blank line is TWO line breaks; got %d, so the gap the estimator typed is not in the "
        "document" % got["breaks"])
    # Stated as lines rather than as breaks, because "one blank line" is the thing Kyle asked
    # for: re-parsing must give back exactly one empty line, not two and not none.
    lines = got["with_breaks"].split("\n")
    assert lines == ["Scope:  line one", "", "line two"]
    assert lines.count("") == 1


def test_the_blank_line_is_a_real_break_not_a_literal_newline_in_a_run():
    """A "\\n" left inside a `<w:t>` is not a line break to Word — it is whitespace, and Word
    normalizes it away. The break has to be a real `<w:br/>` element, which is what
    `_write_t_text` emits and what this asserts on the file rather than on the payload."""
    d = docx.Document(io.BytesIO(
        _generate([{"id": _block_id("Schedule:"), "text": "Schedule:  a\n\nb"}])))
    hits = [t for t in d.element.body.iter(qn("w:t")) if "\n" in (t.text or "")]
    assert not hits, (
        "%d <w:t> node(s) still carry a literal newline: %r"
        % (len(hits), [t.text for t in hits[:3]]))


@pytest.mark.parametrize("label", ["Scope:", "Schedule:", "Exclusions:", "Notes:"])
def test_every_work_row_the_estimator_types_into_keeps_its_breaks(label):
    """All four WORK rows go through the same normalizer, so all four had the same bug.

    The needle is a phrase that occurs nowhere else in the template — "first" also appears in
    the Terms boilerplate, and matching that paragraph made this look like a product failure."""
    sent = label + "  Sawcut control joints\n\nProtect adjacent finishes"
    got = _paragraph_with(_generate([{"id": _block_id(label), "text": sent}]),
                          "Sawcut control joints")
    assert got["with_breaks"] == sent
    assert got["breaks"] == 2


def test_the_label_is_still_bolded_through_its_colon():
    """The normalizer exists to bold the label and un-bold the value; the break fix must not
    cost that. The label half of the split keeps the bold, the value half does not."""
    got = _paragraph_with(
        _generate([{"id": _block_id("Scope:"), "text": "Scope:  line one\n\nline two"}]),
        "line one")
    assert got["bold_runs"], "the WORK label lost its bold"
    assert any(r.startswith("Scope:") for r in got["bold_runs"])
    assert not any("line two" in r for r in got["bold_runs"]), (
        "the value was bolded along with the label")


def test_a_run_that_is_only_a_break_is_left_alone():
    """The visible-text guard: a run with no `<w:t>` text is skipped, exactly as it was before
    the fix, so a bare `<w:br/>` run cannot be rewritten into nothing."""
    assert pw._run_text_with_breaks(_only_break_run()) == "\n"


def _only_break_run():
    from docx.oxml import OxmlElement
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:br"))
    return r


@pytest.mark.parametrize("s,n,expect", [
    ("Scope:  line one\nline two", 7, ("Scope: ", " line one\nline two")),
    ("Scope:", 6, ("Scope:", "")),
    ("a\n\nb", 1, ("a", "\n\nb")),
    ("a\n\nb", 2, ("a\n\nb", "")),
    ("", 0, ("", "")),
    ("abc", 99, ("abc", "")),
])
def test_the_split_counts_visible_characters_only(s, n, expect):
    """The label/colon offsets are measured in the document's TEXT (`_own_text` sees no
    breaks), so converting an index in that coordinate system into a cut in a string that
    carries the breaks is the whole of the fix. Off by one here re-bolds the wrong character."""
    assert pw._split_after_visible(s, n) == expect


# ═══ (b) the blank line, browser end ═════════════════════════════════════════
def test_one_enter_is_one_newline(ran):
    """Driven through the page's real keydown handler with a caret it can read. One press, one
    "\\n"; two presses, one blank line. Not two, not zero."""
    assert ran["enter"]["once"] == "Scope:  Grind and coat.\n"
    assert ran["enter"]["twice"] == "Scope:  Grind and coat.\n\n"
    assert ran["enter"]["defaultPrevented"] is True, (
        "the browser's own Enter still ran alongside ours, so the block gets both")
    # The caret moves past the break it just inserted, and the edit is marked as text (not as
    # formatting — `tw-fmt` would push every Enter onto the richer runs payload).
    assert ran["enter"]["caretsPlaced"] == [[24, 24], [25, 25]]
    assert ran["enter"]["dirtied"] == [["115", False], ["115", False]]


def test_the_caret_ends_up_after_the_break_it_just_typed(ran):
    """The regression that intercepting Enter would otherwise introduce. `placeSelection` can
    only build a Range inside a TEXT node, and `pointAt` has to skip a `<br>` because there is
    no text position in it — so rendering the break as `<br>` left the caret at the end of the
    PREVIOUS line and the next character typed went in above the break. `.tw-block` is
    `white-space: pre-wrap`, so the break is rendered as a real newline character instead,
    which is also what Blink's own editor inserts in this exact CSS context."""
    for key in ("caretLanding", "caretLandingAfterRerender"):
        landing = ran["enter"][key]
        assert landing is not None, "%s: there is no caret position at the end of the block" % key
        assert landing["isText"] is True, "%s: the caret would land on an element, not in text" % key
        assert landing["after"] == "\n", (
            "%s: the character before the caret is %r, so the caret is on the wrong side of the "
            "break" % (key, landing["after"]))


def test_the_blank_line_survives_a_re_render_and_a_reload(ran):
    """Three walkers have to agree on the same character: `serializeBlock` (what is sent),
    `renderRuns` (what applyFormat and paste rebuild the block with) and the reload path, where
    `restoreSavedOverrides` writes the stored text back in as textContent. A disagreement shows
    up as the blank line doubling or vanishing on the next sidebar change."""
    assert ran["enter"]["afterRerender"] == ran["enter"]["twice"]
    assert ran["enter"]["afterReload"] == ran["enter"]["twice"]


def test_the_harness_fixture_is_the_real_template_block(ran):
    """The (b) tests are only worth their assertions if the block they type into is the block
    the editor gets. `/api/proposal-template` reports the WORK rows AFTER running
    `_normalize_work_label_formatting` over the pristine template (main.py does that so the
    preview's metadata matches the generated file), so these are the runs the page renders."""
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    pw._normalize_work_label_formatting(d)
    real = None
    for _idx, _kind, p_elem, _in_block, text, _txbx in pw.iter_editable_blocks(d):
        if text.startswith("Scope:"):
            real = pw._block_runs(p_elem, Paragraph(p_elem, d))
            break
    assert real is not None, "the Scope: row is gone from the template"
    assert ran["fixtureRuns"] == real, (
        "the harness's block fixture has drifted from the template:\nfixture %r\nreal    %r"
        % (ran["fixtureRuns"], real))


def test_the_break_gets_its_own_run_and_leaves_the_label_alone(ran):
    """What actually ships: `serializeRuns`. The label keeps `bold: True`, the value keeps the
    explicit `bold: False` the template carries (absent would mean "inherit", which on a bold
    paragraph style is the opposite instruction), and the break arrives un-styled so the writer
    gives it the template's own base run properties rather than pinning anything."""
    label = {"text": "Scope:", "bold": True, "size_pt": 8}
    value = {"text": "  Grind and coat.", "bold": False, "size_pt": 8}
    assert ran["pristine"]["runs"] == [label, value]
    assert ran["enter"]["runs"] == [label, value, {"text": "\n\n"}]


def test_the_browsers_own_enter_would_have_doubled_it(ran):
    """WHY the handler intercepts Enter at all, stated as a measurement rather than a belief:
    the wrapper-div-plus-placeholder-br shape a contenteditable is left in reads as TWO
    newlines. That is one keypress becoming a blank line."""
    assert ran["browserDefaultShape"]["before"] == "Scope:  Grind and coat."
    assert ran["browserDefaultShape"]["after"] == "Scope:  Grind and coat.\n\n"


@pytest.mark.parametrize("case", ["ctrl", "meta", "alt", "composing", "other"])
def test_enter_is_only_intercepted_when_it_means_a_line_break(ran, case):
    """Ctrl/Cmd/Alt+Enter are other people's shortcuts, an in-flight IME composition is a
    candidate being committed rather than a new line, and any other key is not ours."""
    assert ran["enterGuards"][case] is True, "%s+Enter inserted a line break" % case


def test_enter_is_the_editors_on_every_line_and_refused_when_the_caret_is_unreadable(ran):
    """REVERSED on both counts, 2026-08-26, by the same structural change.

    It used to read: a price line and a notes bullet handle their own Enter, and a caret the page
    cannot read is a caret it must not guess at. Both halves rested on each line being its own
    contenteditable, where the worst a browser Enter could do was leave a stray wrapper element
    inside one line.

    With one editing host per text box, the browser's Enter SPLITS the paragraph into two
    elements. On a computed line that means a second `<p data-sys-line="area">` — a row the writer
    has no channel for, so half of what the estimator typed reaches the customer and half of it
    disappears. On a template paragraph it means a `.tw-block` whose id nothing owns. One break
    inside one element is the only shape this editor can send, so:

      * A computed line's Enter is taken, becomes one break, and the line's own save channel is
        told about it (its channels are delegated `input` listeners — being told is how it saves).
      * A caret the page cannot read means the page will not let ANYTHING happen. Handing the key
        to the browser as a fallback is the one thing that could split the paragraph, so
        "I don't know where the caret is" has to be a refusal, not a delegation. The paragraph is
        left exactly as it was, which is what the estimator sees either way.

    A node that is not an editable line at all — the box's own tools layer, the page behind it —
    still belongs to the browser. That is the guard that stayed."""
    g = ran["enterGuards"]
    assert g["computedLineTaken"] is True, "a computed line's Enter fell through and split its <p>"
    assert g["computedLineText"] == "\n", "the break did not land in the line"
    assert g["computedLineTold"] == 1, "the line's own save channel was never told"
    assert g["notALine"] is True, "Enter on chrome was hijacked"
    assert ran["enterNoCaret"]["defaultPrevented"] is True
    assert ran["enterNoCaret"]["text"] == "Scope:  Grind and coat."


def test_a_break_inside_a_token_value_keeps_the_value_a_token(ran):
    """A `.tw-fill` span is a live estimate value. Splitting it must leave two fills, not
    dissolve it into hand-typed text — that is how a computed figure gets frozen."""
    assert ran["midBreak"]["text"] == "Scope:  Grind and \ncoat."
    assert ran["midBreak"]["fills"] == ["scope_notes", "scope_notes"]
    # And the two halves of the split value keep the template's own formatting, so the .docx
    # does not come back with half a line at a different weight or size.
    assert ran["midBreak"]["runs"] == [
        {"text": "Scope:", "bold": True, "size_pt": 8},
        {"text": "  Grind and ", "bold": False, "size_pt": 8},
        {"text": "\n"},
        {"text": "coat.", "bold": False, "size_pt": 8}]


def test_the_whole_frontend_to_docx_round_trip(ran):
    """The two halves joined: the text the harness's real serializer produced, sent through the
    real writer, read back out of the real file. Exactly one blank line, end to end."""
    sent = ran["enter"]["twice"]
    got = _paragraph_with(_generate([{"id": _block_id("Scope:"), "text": sent}]),
                          "Grind and coat.")
    assert got["with_breaks"] == sent
    assert got["breaks"] == 2
    assert got["with_breaks"].split("\n") == ["Scope:  Grind and coat.", "", ""], (
        "two Enters at the end of the paragraph did not come back as one blank line")
    # And through the richer `runs` shape, which is what a formatted paragraph sends.
    got_runs = _paragraph_with(
        _generate([{"id": _block_id("Scope:"), "runs": ran["enter"]["runs"]}]),
        "Grind and coat.")
    assert got_runs["with_breaks"] == sent
    assert got_runs["breaks"] == 2


# ═══ (d) the way out of an expanded text box ══════════════════════════════════
CLIPPED = "246px"        # round(183.75pt * 96/72 + 1) — the design height fitTxbx clips to


def test_the_fixture_box_really_is_over_capacity(ran):
    """Every assertion below is worthless against a box that fits. 400pt of content in Kyle's
    183.75pt GC Resinous box is the real complaint (a long WORK scope)."""
    assert ran["clipped"] == {"open": False, "overflow": True, "maxHeight": CLIPPED,
                             "overflowStyle": "hidden", "zIndex": ""}


def test_an_expanded_box_carries_a_labelled_way_out(ran):
    """Not a grip. Somebody who cannot find the exit will not find it by hovering a 7px square,
    and the grips are for resizing — a control that shares their look would read as one."""
    t = ran["tools"]
    assert t["hasCollapse"] and t["label"] == "Collapse"
    assert t["inToolsLayer"], "the button is inside the editable content, where clicks are ignored"
    assert t["isNotAGrip"], "the way out is a drag handle, which is what confused Kyle already"
    assert "Esc" in t["title"] and "outside" in t["title"], (
        "the tooltip does not mention the other two ways out: %r" % t["title"])
    # Adding a control to the tools layer must not reorder the grips — test_box_drag_ui.py
    # asserts they come out as move / e / s / se, in that order.
    grips = [c.split()[-1] for c in t["order"] if c.startswith("tw-grip ")]
    assert grips == ["tw-grip-move", "tw-grip-e", "tw-grip-s", "tw-grip-se"]


def test_the_collapse_button_clips_the_box_again(ran):
    """The point: it must restore the clipped maxHeight fitTxbx set, not merely drop the
    class — a box left at `max-height: none` looks expanded with no way to say so."""
    assert ran["collapseButton"]["opened"]["open"] is True
    assert ran["collapseButton"]["opened"]["maxHeight"] == "none"
    assert ran["collapseButton"]["closed"] == {
        "open": False, "overflow": True, "maxHeight": CLIPPED,
        "overflowStyle": "hidden", "zIndex": ""}


def test_clicking_the_text_of_an_open_box_does_not_close_it(ran):
    """THE reason the other exits had to exist, kept as a test so nobody "simplifies" it away:
    a click on a paragraph is a click that puts a caret in the paragraph. If it also collapsed
    the box, the estimator would lose the view every time they went to type."""
    assert ran["typingKeepsItOpen"]["afterBlockClick"]["open"] is True
    assert ran["typingKeepsItOpen"]["afterFillClick"]["open"] is True, (
        "a click on a .tw-fill island inside the paragraph collapsed the box")


def test_escape_collapses_the_box_and_blurs_the_caret_first(ran):
    """Order matters. A collapsed box is `overflow: hidden`, so a caret left in the clipped
    part makes the browser scroll the box back to it — which reads as the collapse not working.
    Blurring loses nothing: the edit is already in the DOM and already marked dirty."""
    e = ran["escape"]
    assert e["closed"]["open"] is False
    assert e["closed"]["maxHeight"] == CLIPPED
    assert e["blurred"] is True, "the caret was left inside the box it just collapsed"
    assert e["defaultPrevented"] is True


def test_escape_with_nothing_open_is_not_swallowed(ran):
    """Escape means other things on this page. A handler that always preventDefaults steals
    them from whatever else is listening."""
    assert ran["escapeWhenClosed"] == {"open": False, "defaultPrevented": False}


def test_escape_does_not_blur_a_field_outside_the_box(ran):
    """The sidebar is full of inputs. Escape has to close the box the estimator is looking at
    without yanking the focus out of whatever they were typing in."""
    assert ran["escapeFromElsewhere"]["closed"]["open"] is False
    assert ran["escapeFromElsewhere"]["stillFocused"] is True


def test_clicking_the_page_outside_the_box_collapses_it(ran):
    """The third exit, and the one most people try first."""
    assert ran["outsideClick"]["open"] is False
    assert ran["outsideClick"]["maxHeight"] == CLIPPED


def test_the_formatting_toolbar_is_not_the_outside_of_the_box(ran):
    """ensureFmtBar mounts the B / I / U bar in the page's top chrome (`#fmt-ribbon`) — it has to
    escape the box's own clipping to be visible at all — so in the DOM it IS outside the box.
    Treating it as such would close the box the moment the estimator bolded a word inside it,
    taking their selection with it. Since 2026-08-24 the bar is a static ribbon and never sits
    over a box at all, which makes this exclusion permanent rather than positional."""
    assert ran["formatBarClick"]["open"] is True, (
        "clicking the formatting toolbar collapsed the box being formatted")
    assert ran["formatBarClick"]["maxHeight"] == "none"


def test_no_box_is_left_expanded_behind_another(ran):
    """Two boxes can be opened at once (WORK and NOTES, say). One click off them puts both
    back, or the page keeps a broken layout the estimator has stopped looking at."""
    assert ran["manyBoxes"]["bothOpen"] == [True, True], "the fixture never opened both"
    assert [b["open"] for b in ran["manyBoxes"]["afterOutside"]] == [False, False]
    assert [b["maxHeight"] for b in ran["manyBoxes"]["afterOutside"]] == [CLIPPED, CLIPPED]


def test_a_refit_puts_an_open_box_back(ran):
    """fitTxbx re-runs after every edit and every repagination. It already cleared the class;
    what matters is that it also restores the clip, so an open box cannot survive a render with
    `max-height: none` still on it."""
    assert ran["refitCollapses"] == {"open": False, "overflow": True, "maxHeight": CLIPPED,
                                    "overflowStyle": "hidden", "zIndex": ""}


# ═══ (c) clunky: what is editable, and what is being edited ═══════════════════
def _css_rule(selector):
    """The declarations of EVERY top-level rule with exactly this selector, concatenated.

    Scoped to the rule rather than to a slice of the file, for the reason test_box_drag_ui.py
    records: a nearby rule can carry the declaration you are looking for and pass an assertion
    you had already broken. All of them rather than the first, because styles.css legitimately
    declares `.tw-txbx.tw-notes-overflow` twice — once for the clipped box's fade and once for
    its cursor — and reading only the first is how a real declaration reads as missing."""
    found = [m.group(1) for m in
             re.finditer(r"(?m)^" + re.escape(selector) + r"\s*\{([^}]*)\}", CSS)]
    assert found, "%s has no top-level rule in styles.css" % selector
    return "\n".join(found)


def test_an_editable_paragraph_says_so_with_the_pointer():
    """Kyle read the un-highlighted half of a paragraph as not editable, and the pointer was
    agreeing with him: `.tw-block` set no cursor, and `cursor` is inherited — so inside a
    clipped box the caret pointer over the text was the box's `zoom-in` magnifier."""
    assert "cursor: text" in _css_rule(".tw-block")
    assert "cursor: text" in _css_rule(".tw-block .tw-fill"), (
        "a token value inside an editable paragraph still claims to be read-only")
    assert "cursor: zoom-in" in _css_rule(".tw-txbx.tw-notes-overflow"), (
        "the box's own peek affordance was removed rather than overridden on the text")


def test_hover_and_focus_are_visible_without_moving_the_page():
    """This surface is a to-scale preview of a printed page and each box is registered against
    baked page artwork, so an affordance that changed the geometry would be a worse bug than
    the clunkiness. Every cue is `background` or `box-shadow`."""
    for selector in (".tw-block:hover", ".tw-block:focus", ".tw-block.tw-dirty:focus"):
        decls = _css_rule(selector)
        assert "box-shadow" in decls or "background" in decls
        for banned in ("border", "padding", "margin", "font-size", "letter-spacing", "outline-width"):
            assert not re.search(r"(?m)^\s*%s\s*:" % banned, decls), (
                "%s changes the layout (%s), which shifts the text off the artwork"
                % (selector, banned))


def test_the_paragraph_being_edited_keeps_its_focus_ring_once_it_is_dirty():
    """`.tw-block.tw-dirty` and `.tw-block:focus` have the same specificity, so while the focus
    rule came first the hand-edited paragraph — the one most likely to still be being typed in —
    lost its ring to the dirty bar. The combined rule has to come last and carry both shadows."""
    combined = _css_rule(".tw-block.tw-dirty:focus")
    assert "inset 2px 0 0" in combined, "the dirty bar disappears while the block has focus"
    assert combined.count("rgba(158, 0, 31") >= 2, "the focus ring is missing from the combination"
    assert CSS.index(".tw-block.tw-dirty {") < CSS.index(".tw-block.tw-dirty:focus"), (
        "the combined rule is declared before the plain dirty rule, which overrides it")


def test_the_collapse_button_adds_no_height_to_the_box():
    """fitTxbx decides what overflows from the box's offsetHeight, so a control in the normal
    flow would make every box measure taller than its text — i.e. would break the overflow
    notice for the sake of the button that exists to answer it."""
    assert "position: absolute" in _css_rule(".tw-box-collapse")
    assert "display: none" in _css_rule(".tw-box-collapse"), (
        "the button shows on every box, not only on the expanded one")
    assert "display: block" in _css_rule(".tw-txbx.tw-notes-open .tw-box-collapse")


def test_an_expanded_box_still_says_the_text_does_not_fit():
    """Expanding must not make the overflow look solved. Nothing is clipped ON SCREEN any more,
    but the .docx is unchanged and Word's normAutofit will cramp it just the same, so the badge
    keeps naming both facts. The old copy — "Showing all of it — click to collapse" — retired
    the warning at exactly the moment the box looks fine."""
    badge = _css_rule(".tw-txbx.tw-notes-open::after")
    assert "content:" in badge
    assert "longer" in badge.lower(), "the expanded badge no longer warns about the box: %r" % badge
    assert "Word" in badge, "the badge does not say the generated document is cramped too"
    assert "Esc" in badge, "the badge does not name a way out"


# ═══ (a) the labels: what is real text, and what is locked ════════════════════
def test_a_work_label_is_ordinary_editable_text_that_stays_bold(ran):
    """Retyping "Scope:" as "Scope of work:" in place — which is what typing inside the bold
    span does; the browser edits that text node, it does not restyle it — keeps the run bold and
    at the template's own size, all the way into the file."""
    assert ran["labelRetyped"]["text"] == "Scope of work:  Grind and coat."
    assert ran["labelRetyped"]["runs"] == [
        {"text": "Scope of work:", "bold": True, "size_pt": 8},
        {"text": "  Grind and coat.", "bold": False, "size_pt": 8}]
    got = _paragraph_with(
        _generate([{"id": _block_id("Scope:"), "runs": ran["labelRetyped"]["runs"]}]),
        "Scope of work:")
    assert any(r.startswith("Scope of work:") for r in got["bold_runs"]), (
        "the retyped label reached the document un-bolded")


def test_an_emptied_label_leaves_no_stray_token_and_no_lone_colon(ran):
    """Emptying the label must not expose the `{{scope_notes}}` token it sat in front of, and
    must not leave a colon on its own. The value is what is left, and nothing else."""
    assert ran["labelEmptied"]["text"] == "  Grind and coat."
    assert ran["labelEmptied"]["runs"] == [
        {"text": "  Grind and coat.", "bold": False, "size_pt": 8}]
    got = _paragraph_with(
        _generate([{"id": _block_id("Scope:"), "text": ran["labelEmptied"]["text"]}]),
        "Grind and coat.")
    assert "{{" not in got["own_text"], "an emptied label re-exposed a raw token"
    assert ":" not in got["own_text"], "an emptied label left a colon behind"
    assert got["own_text"].strip() == "Grind and coat."


def test_the_region_labels_are_real_docx_text_not_artwork():
    """The evidence behind the honest answer to "why not make it like a word document":
    "System:" / "Texture:" / "Area:" are ordinary paragraph text in the template — the reason
    they cannot be edited is that they sit inside the `{{#system}}` repeat region, which the
    editor collapses into a read-only preview and which `_apply_paragraph_overrides` refuses by
    design. Nothing here is baked into a PNG, so this is a missing feature and not a limit of
    the template artwork; whoever adds it needs a new override channel, not a CSS change."""
    blocks = _blocks()
    labelled = {b["text"].split(":")[0]: b for b in blocks if ":" in b["text"]}
    for name in ("Texture", "Area"):
        assert name in labelled, "%s: is gone from the template" % name
        assert labelled[name]["in_block"] == "system", (
            "%s: is no longer inside the {{#system}} region — if it became a free paragraph it "
            "is now editable and this test should become the round-trip one" % name)
    # And the paragraphs Kyle can already edit are free paragraphs in the same text box.
    for prefix in ("Scope:", "Schedule:", "Exclusions:", "Notes:", "Epoxy Flooring:"):
        b = [x for x in blocks if x["text"].startswith(prefix)][0]
        assert b["in_block"] is None, "%s moved into a region and stopped being editable" % prefix
        assert b["in_txbx"] == labelled["Texture"]["in_txbx"], (
            "%s is no longer in the WORK text box, so the label normalizer no longer covers it"
            % prefix)


def test_a_region_paragraph_override_is_still_refused():
    """The refusal is what makes the region labels un-editable, so it is stated here rather than
    left as a comment: a client that sends one anyway must be ignored, not obeyed."""
    system_id = [b["id"] for b in _blocks() if b["in_block"] == "system"][1]
    d = docx.Document(str(pw.TEMPLATES_ROOT / _EPOXY_TEMPLATE))
    assert pw._apply_paragraph_overrides(d, [{"id": system_id, "text": "Widget:  x"}]) == 0
