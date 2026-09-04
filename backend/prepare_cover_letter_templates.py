"""One-shot: build the Cover Letter templates from Treadwell's letterhead.

Run it when the letterhead OR the draft copy below changes, then commit the
.docx files it writes:

    python backend/prepare_cover_letter_templates.py

WHY A GENERATOR AND NOT A FOLDER OF HAND-MADE FILES.

The cover letter is Treadwell's branded letterhead (`docs/Cover Letter/Treadwell
Letterhead.docx`) with a short body typed onto it — one variant per
(work type, audience), files that differ only in ~10 paragraphs of copy.
Hand-editing them keeps them out of step: the 2026-08 first draft already needs a
copy pass from Hanz (see `templates/CoverLetter/README.md`), and re-typing his
wording into seven documents by hand is how three of them end up with last week's
sentence.

AUDIENCE, BECAUSE THE PROPOSAL HAS ONE.

`proposal_writer.TEMPLATE_PICKER` is keyed on `(work_type, audience)` and its
files live in audience-first folders — `Direct/` (4), `GC/` (3), `Gyp/` (1) — and
it is deliberately NOT a clean grid: gyp ignores audience entirely (`audience`
is `None`), and GC combo reuses the GC resinous file because Kyle never made a
GC combo document. The cover letters mirror that shape:

    CoverLetter/Direct/{Epoxy,Polish,Combo}.docx
    CoverLetter/GC/{Epoxy,Polish,Combo}.docx
    CoverLetter/Gyp/Gyp.docx                     (audience-agnostic)

Two deliberate departures from the proposal picker, both because these files are
GENERATED and Kyle's are not:

  * GC gets its OWN Combo letter instead of reusing the GC epoxy one. The
    proposal shares that file because no GC combo document exists; a LETTER that
    shares it would tell a combo customer, in prose, that the pages behind it are
    an epoxy proposal — and half the scope would simply be missing. Costing
    nothing to generate, it is written.
  * There is no sealer or budget letter. Those work types have a proposal but no
    cover-letter copy in Hanz's source PDF, so `cover_letter_writer.pick_template`
    falls back and `has_template()` reports False — see that module.

The GC and Direct bodies currently carry the SAME copy. The source PDF has no GC
variant, and inventing contractor-flavoured sentences would put wording nobody
approved in front of a customer. Separate files exist so the copy pass can
diverge them without a code change; the README says so.

The letterhead itself is COPIED BYTE-FOR-BYTE and then appended to. Its page-one
artwork is a full-page PNG anchored inside a `w:sdt` (a "Cover Pages" content
control), and everything about that — the anchor offsets, the two-section page
setup, the theme fonts — is left exactly as Kyle's file has it. We add
paragraphs to the BODY after the `w:sdt`, one decimal numbering definition
copied out of Hanz's example letter, and one anchored text box (the date).

Body paragraphs outside the `w:sdt`, deliberately:
`proposal_writer.iter_editable_blocks` (which the doc editor and
`/api/coverletter-template` walk for their block ids) yields the body's DIRECT
`w:p` children. Paragraphs nested inside a `w:sdt` are invisible to it — which is
exactly why Hanz's own example letter shows up as two blocks instead of
twenty-two. Anything the estimator must be able to edit therefore has to be a
direct body child.

THERE ARE NO TEXT BOXES, AND THERE IS NO TITLE LINE.

Both used to be here. Kyle's example letter floats the date in a small
anchored box over the artwork and opens with a red underlined heading, and
these templates copied both faithfully. On 2026-09-04 Hanz asked for both to
come off EVERY format: the proposal behind the letter already carries the job
name and the date, and repeating them on the cover page is two more things to
keep true. So a generated letter is now nothing but flowing body paragraphs,
which is the simplest shape this file can produce and the one the rule above
wants: every block is a direct body child, `template_geometry`'s `boxes` is
empty, and the editor takes the no-box layout branch it already had.
`{{proposal_date_short}}` is still resolved server-side for the proposal; no
letter prints it.

Source files (Hanz, 2026-08-28), kept OUT of the image — they are reference
material, not runtime inputs:
    docs/Cover Letter/Treadwell Letterhead.docx        — branding, page setup
    docs/Cover Letter/Treadwell Cover Letter - Example1.docx
                                                       — the finished-page model,
                                                         and the source of the
                                                         decimal list
    docs/Cover Letter/1 Treadwell Proposal Cover Letter Templates.pdf
                                                       — the copy, as 6 EMAIL drafts

The PDF's templates were written as outbound emails. Templates 1-4 (Epoxy,
Polish, Combo, Gyp) are adapted here; 5 (a gyp addendum note) and 6
(BuildingConnected) are out of scope. The email framing is stripped — no
"Subject Line:", no cc note, and "I've attached our proposal to this email"
becomes a line that introduces the proposal the customer is about to read in
the portal, because this is a portal DOCUMENT PAGE and there is no email and no
attachment.
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
REF_DIR = REPO_ROOT / "docs" / "Cover Letter"
LETTERHEAD = REF_DIR / "Treadwell Letterhead.docx"
EXAMPLE = REF_DIR / "Treadwell Cover Letter - Example1.docx"
OUT_DIR = Path(__file__).resolve().parent / "templates" / "CoverLetter"

# ── House formatting, read off Hanz's example letter ──────────────────────────
# Body: Zetta Serif Book, 40% grey, with a 1.625" right indent so the column
# stops short of the page edge (the example's `w:ind w:right="2340"`).
# Signature: Arial 10pt.
# The heading face (Century Gothic, Treadwell red A91120) is gone with the
# title line it set — see `build`.
BODY_FONT = "Zetta Serif Book"
BODY_GREY = RGBColor(0x40, 0x40, 0x40)
SIG_FONT = "Arial"
SIG_GREY = RGBColor(0x59, 0x59, 0x59)

# 1.625" = 117pt, Kyle's own `w:ind w:right="2340"` — present on every body
# paragraph of `Treadwell Cover Letter - Example1.docx` and absent from the bare
# letterhead, so it is his design decision and not an artefact of this generator.
#
# KNOWN, DELIBERATELY NOT RECONCILED: it means the real text measure is 405pt
# while `proposal_writer._page_metrics` reports this section's column as 522pt.
# That is not a bug in either — the metric describes the SECTION and the indent
# describes the PARAGRAPH — but a future Letterhead Editor that lays text out
# from `geometry.page` alone will draw the lines 117pt too wide. No consumer
# exists yet; the fix belongs with the one that does, which will know whether it
# wants the section column or the paragraph measure. Guessing at it now is how
# the last two geometry bugs got in.
RIGHT_INDENT = Inches(1.625)

# 11pt, where the example is 12pt. The Combo letter carries two full system
# sections (six numbered items, because Hanz's Template 3 repeats Schedule and
# Options under both headings) and runs onto a second page at 12pt — and page two
# has no letterhead artwork, because the art is anchored to page one.
#
# The point is NOT what buys the page back; measured, it is worth about 30pt of a
# ~69pt overrun. The spacing below does the rest — see GROUP_SPACE_BEFORE and
# `_set_mark_size`. There is still no one-page GUARANTEE for Combo: a long
# {{job_name}} (the title wraps) or a multi-line {{schedule_notes}} can put it
# back over. Flagged in the CoverLetter README so the copy pass can decide
# whether to shorten Combo instead.
BODY_PT = Pt(11)
SIG_PT = Pt(10)

# Air above a group heading, above the first numbered item of an unheaded group,
# and above the closing lines. Kyle's example gets ~14pt there from Word's
# `beforeAutospacing`/`afterAutospacing` on the list paragraphs — which is an
# HTML-compatibility feature whose exact value is the RENDERER's, and production
# renders with LibreOffice while this was verified in Word. An explicit measure
# lands in the same place on both. Correspondingly there is no space BETWEEN
# items: Kyle's run together (measured at a 0.02pt gap), because auto-spacing
# collapses between adjacent paragraphs of the same kind.
GROUP_SPACE_BEFORE = Pt(12)

# Numbering: the letterhead ships 32 bullet lists and no decimal one, so the
# "1. / 2. / 3." list definition is lifted out of the example letter (its
# abstractNum 12, `%1.`) and re-homed on ids that are free in the letterhead.
#
# Two `w:num` instances point at the one abstract definition — and pointing at it
# is NOT enough to restart the count. The counter lives on the abstractNum, so
# Combo's second list continued 4./5./6. under "Polished Concrete:" (seen in the
# render, 2026-08-29). `_install_decimal_numbering` therefore gives every `w:num`
# an explicit `w:lvlOverride/w:startOverride val="1"`, which is the only thing in
# OOXML that resets a shared counter for one instance.
NUM_IDS = (33, 34)
_ABSTRACT_ID = 32


def _seg(text, *, bold=False, italic=False):
    return {"text": text, "bold": bold, "italic": italic}


def _ph(text):
    """A PLACEHOLDER segment — copy nobody has signed off yet, in italics so it
    is impossible to miss on the page. Square brackets, never `{{braces}}`: a
    `{{token}}` would be silently substituted-or-left by the writer, whereas
    "[THICKNESS - pick one: ...]" prints as itself and reads as an instruction.
    Every one of these is listed in templates/CoverLetter/README.md."""
    return _seg(text, italic=True)


# ── The copy, per work type ───────────────────────────────────────────────────
# Each entry: title, the line that introduces the proposal, and one or more
# numbered GROUPS (heading, items). Items are lists of runs so a bold lead-in
# label and an italic placeholder can sit in the same sentence.
GREETING = "Hello,"
THANKS = "Thanks for the opportunity to bid on this {{job_name}} project!"
CLOSE_1 = "Feel free to reach out if you have any questions."
CLOSE_2 = "Looking forward to working with you!"

_SCHEDULE_ITEM = [
    _seg("Schedule: ", bold=True),
    _seg("{{schedule_notes}} "),
    _ph("[SCHEDULE - assumes 1 mobilization per phase, with all areas available "
        "at one time; edit if this job phases differently.]"),
]

_OPTIONS_EPOXY = [
    _seg("Options: ", bold=True),
    _ph("[OPTIONS - keep the lines that apply: add for an onsite / in-place "
        "mockup; add for temporary protection after install; moisture-mitigation "
        "unit price if slab RH is over 75%; add for generator use if onsite power "
        "is unavailable.]"),
]

_OPTIONS_POLISH = [
    _seg("Options: ", bold=True),
    _ph("[OPTIONS - keep the lines that apply: add for an onsite / in-place "
        "mockup; add for temporary protection after install; add for generator "
        "use if onsite power is unavailable.]"),
]

_EPOXY_SYSTEM_ITEM = [
    _seg("Materials / System: ", bold=True),
    _seg("{{system_name}}, {{texture}} texture, over approximately {{epoxy_sf}} SF, "
         "with {{cove_lf}} LF of integral cove base. "),
    _ph('[THICKNESS - pick one: 1/8" / 3/16" / 1/4" nominal thickness with '
        'urethane topcoat.] [COVE HEIGHT - pick one: 4" / 6" / 8".]'),
]

_POLISH_SYSTEM_ITEM = [
    _seg("System: ", bold=True),
    _seg("Polished concrete over approximately {{polish_sf}} SF. "),
    _ph("[AGGREGATE EXPOSURE - pick one: Class A cream finish (no exposure) / "
        "Class B salt & pepper / Class C coarse, full exposure.] "
        "[SHEEN - pick one: Level 2 (400 grit) / Level 3 (800 grit).]"),
]

COPY = {
    "epoxy": {
        "title": "Epoxy / Resinous Flooring Proposal - {{job_name}}",
        "intro": "The pages that follow are our Epoxy / Resinous Flooring proposal "
                 "for this project. A few things to note:",
        "groups": [
            (None, [_EPOXY_SYSTEM_ITEM, _SCHEDULE_ITEM, _OPTIONS_EPOXY]),
        ],
    },
    "polish": {
        "title": "Polished Concrete Proposal - {{job_name}}",
        "intro": "The pages that follow are our Polished Concrete proposal for this "
                 "project. A few things to note:",
        "groups": [
            (None, [
                [
                    _seg("System: ", bold=True),
                    _seg("{{system_name}} over approximately {{polish_sf}} SF. "),
                    _ph("[AGGREGATE EXPOSURE - pick one: Class A cream finish "
                        "(no exposure) / Class B salt & pepper / Class C coarse, "
                        "full exposure.] [SHEEN - pick one: Level 2 (400 grit) / "
                        "Level 3 (800 grit).]"),
                ],
                _SCHEDULE_ITEM,
                _OPTIONS_POLISH,
            ]),
        ],
    },
    "combo": {
        "title": "Epoxy / Resinous Flooring & Polished Concrete Proposal - {{job_name}}",
        "intro": "The pages that follow are our Epoxy / Resinous Flooring and "
                 "Polished Concrete proposals for this project. A few things to note:",
        "groups": [
            ("Epoxy / Resinous Flooring:", [
                [
                    _seg("Materials / System: ", bold=True),
                    _seg("{{epoxy_system_name}}, {{texture}} texture, over "
                         "approximately {{epoxy_sf}} SF, with {{cove_lf}} LF of "
                         "integral cove base. "),
                    _ph('[THICKNESS - pick one: 1/8" / 3/16" / 1/4" nominal '
                        'thickness with urethane topcoat.] [COVE HEIGHT - pick '
                        'one: 4" / 6" / 8".]'),
                ],
                _SCHEDULE_ITEM,
                _OPTIONS_EPOXY,
            ]),
            ("Polished Concrete:", [
                _POLISH_SYSTEM_ITEM,
                _SCHEDULE_ITEM,
                _OPTIONS_POLISH,
            ]),
        ],
    },
    "gyp": {
        "title": "Gypsum Underlayment Proposal - {{job_name}}",
        "intro": "The pages that follow are our Gypsum Underlayment proposal for "
                 "this project. A few things to note:",
        "groups": [
            (None, [
                [
                    _seg("System: ", bold=True),
                    _seg("{{gyp_soft_thickness}} gypsum underlayment over "
                         "{{gyp_soft_sf}} SF of soft-surface area and "
                         "{{gyp_hard_thickness}} over {{gyp_hard_sf}} SF of "
                         "hard-surface area. "),
                    _ph("[SOUND MAT - state the mat thickness and where it goes, "
                        "e.g. an 1/8\" sound mat at the hard-surface areas.]"),
                ],
                _SCHEDULE_ITEM,
                [
                    _seg("Options: ", bold=True),
                    _ph("[OPTIONS - keep the lines that apply: add for offsite "
                        "storage if onsite storage is not available; add for gyp "
                        "sealer material.]"),
                ],
            ]),
        ],
        # Gyp's source template carries sub-notes under item 1 (spec conflicts,
        # STC/IIC ratings, excluded mat, GC-provided storage). They are job-
        # specific to the point of being unwritable as boilerplate, so they ship
        # as one instruction line rather than four invented sentences.
        "sub_notes": [
            "[NOTES - add the ones this job needs: a spec thickness conflict "
            "(3/4\" gyp over a 1/4\" mat cracks; a 3/4\" pour needs a 1/8\" mat, "
            "3/16\" maximum); which assembly meets the specified STC & IIC "
            "ratings; any area where the sound mat is excluded; whether the GC "
            "provides covered storage.]",
        ],
    },
}

# ── Direct: Will Buchanan's wording, 2026-09-03 ──────────────────────────────
# He sent the Direct letter's text verbatim ("In addition to what Greg sent you
# for GC projects please use the text template below for direct projects. The
# highlighted text should be pulled from the intake form") and this is it, to the
# comma. What is DELIBERATELY not here:
#
#   * GC KEEPS THE COPY ABOVE. Will's text is for Direct; Greg's GC text has not
#     reached this repo. The module header's rule stands -- inventing
#     contractor-flavoured sentences puts wording nobody approved in front of a
#     customer -- so the two diverge here rather than both moving. This is the
#     divergence the separate files were created for.
#   * POLISH KEEPS ITS OWN SYSTEM SENTENCE. Will's Materials / System line is
#     epoxy ('1/4" Flake Floor broadcast with 6" Integral Cove Base'); he wrote
#     nothing about polished concrete, and the aggregate-exposure and sheen
#     choices are real spec decisions. Polish takes his STRUCTURE -- the Area
#     line, the fixed Schedule sentence, the closings -- and keeps its own
#     system wording and its own placeholders.
#   * THE THICKNESS AND SCHEDULE PLACEHOLDERS ARE GONE from Direct. They existed
#     because nothing captured a thickness and the mobilization assumption was
#     unmodelled; the intake form now asks for the thickness and Will has
#     written the assumption out as a sentence. [OPTIONS] stays: it is Kyle's
#     add-alternate checklist, not a hole.
#   * THE "pages that follow" SENTENCE IS GONE from Direct. Will's text opens
#     the list with "A few things to note:" alone. That sentence was added when
#     this letter stopped being an email and became a page in front of the
#     proposal; the red underlined title one line above already says which
#     proposal follows, so it was saying it twice.
DIRECT_GREETING = "{{greeting}}"
DIRECT_THANKS = "Thank you for the opportunity to provide a quote for this project."
DIRECT_INTRO = "A few things to note:"
DIRECT_CLOSE_1 = "Feel free to reach out, if you have questions."

# Fixed boilerplate, not {{schedule_notes}} -- Hanz's call, and Will spelled the
# whole sentence out. It states the assumption the price was built on and invites
# the customer to correct it, which a filled-in note cannot do. The proposal
# behind the letter still carries {{schedule_notes}}, and the estimator can edit
# this paragraph in the document editor on any job that phases differently.
_DIRECT_SCHEDULE_ITEM = [
    _seg("Schedule: ", bold=True),
    _seg("This price is based on all work taking place in 1 phase/mobilization. "
         "If this needs to be split into multiple phases and/or over weekends, "
         "please let me know."),
]

# "Area: Warehouse expansion and 4 offices" -- the estimator's own words, off the
# intake form. Falls back to the SF line when the box is empty (see
# `cover_letter_writer._ensure_cover_letter_values`), so this never prints a bare
# label and the square footage is what a job with no typed area shows.
_DIRECT_AREA_ITEM = [
    _seg("Area: ", bold=True),
    _seg("{{work_areas}}"),
]

# '1/4" MACRO Flake Single Broadcast with 6" Integral Cove Base'. One token, not
# four: the thickness is skipped when Kyle's system name already states one (three
# of his fifteen do), and the cove clause is dropped entirely on a job with no
# cove rather than printing "with 0 LF of integral cove base" the way the
# proposal body still does. That composition cannot be expressed in template
# text, so it lives in the writer -- resolved identically on both sides.
_DIRECT_EPOXY_SYSTEM_ITEM = [
    _seg("Materials / System: ", bold=True),
    _seg("{{cover_system_line}}"),
]

DIRECT_COPY = {
    "epoxy": {
        "title": COPY["epoxy"]["title"],
        "intro": DIRECT_INTRO,
        "greeting": DIRECT_GREETING,
        "thanks": DIRECT_THANKS,
        "close_1": DIRECT_CLOSE_1,
        "groups": [
            (None, [_DIRECT_EPOXY_SYSTEM_ITEM, _DIRECT_AREA_ITEM,
                    _DIRECT_SCHEDULE_ITEM, _OPTIONS_EPOXY]),
        ],
    },
    "polish": {
        "title": COPY["polish"]["title"],
        "intro": DIRECT_INTRO,
        "greeting": DIRECT_GREETING,
        "thanks": DIRECT_THANKS,
        "close_1": DIRECT_CLOSE_1,
        "groups": [
            (None, [
                [
                    _seg("System: ", bold=True),
                    _seg("{{system_name}} over approximately {{polish_sf}} SF. "),
                    _ph("[AGGREGATE EXPOSURE - pick one: Class A cream finish "
                        "(no exposure) / Class B salt & pepper / Class C coarse, "
                        "full exposure.] [SHEEN - pick one: Level 2 (400 grit) / "
                        "Level 3 (800 grit).]"),
                ],
                _DIRECT_AREA_ITEM,
                _DIRECT_SCHEDULE_ITEM,
                _OPTIONS_POLISH,
            ]),
        ],
    },
    "combo": {
        "title": COPY["combo"]["title"],
        "intro": DIRECT_INTRO,
        "greeting": DIRECT_GREETING,
        "thanks": DIRECT_THANKS,
        "close_1": DIRECT_CLOSE_1,
        "groups": [
            # The Area line sits under the FIRST heading only. It describes what
            # the floor covers across the whole job, so repeating it verbatim
            # under "Polished Concrete:" would read as a second, different area.
            ("Epoxy / Resinous Flooring:", [
                [
                    _seg("Materials / System: ", bold=True),
                    _seg("{{cover_system_line}}"),
                ],
                _DIRECT_AREA_ITEM,
                _DIRECT_SCHEDULE_ITEM,
                _OPTIONS_EPOXY,
            ]),
            ("Polished Concrete:", [
                _POLISH_SYSTEM_ITEM,
                _DIRECT_SCHEDULE_ITEM,
                _OPTIONS_POLISH,
            ]),
        ],
    },
    # Gyp has no audience -- one file for both readers -- so there is no Direct
    # variant of it to write. `spec_for` falls back to COPY for anything absent.
}


SIGNATURE = [
    [_seg("--")],
    [_seg("{{estimator_name}}", bold=True), _seg(" | Estimator")],
    [_ph("[ESTIMATOR EMAIL]"), _seg(" | wetreadwell.com")],
    [],
    [_seg("TREADWELL", bold=True),
     _seg(" | 913.396.6216 | 1707 E. 123rd Ter, Olathe, KS 66061")],
    [_seg("Epoxy Flooring + Polished Concrete + Gypsum Underlayments")],
]


# ── docx plumbing ─────────────────────────────────────────────────────────────
def _install_decimal_numbering(d) -> None:
    """Copy the example letter's decimal list definition into `d`'s numbering
    part, as abstractNum 32 + num 33/34 (all free in the letterhead).

    Schema order matters: every `w:abstractNum` must precede every `w:num`, so
    the abstract goes after the last abstract and the instances after the last
    instance. Getting this wrong yields a file Word refuses to open."""
    src = docx.Document(str(EXAMPLE))
    src_numbering = src.part.numbering_part.element
    abstract = None
    for a in src_numbering.findall(qn("w:abstractNum")):
        lvl = a.find(qn("w:lvl"))
        fmt = lvl.find(qn("w:numFmt")) if lvl is not None else None
        if fmt is not None and fmt.get(qn("w:val")) == "decimal":
            abstract = copy.deepcopy(a)
            break
    if abstract is None:
        raise SystemExit(
            "No decimal numbering definition in %s — the '1. / 2. / 3.' list in "
            "the cover letters is copied from it, so it cannot be built without "
            "one." % EXAMPLE.name)
    abstract.set(qn("w:abstractNumId"), str(_ABSTRACT_ID))

    numbering = d.part.numbering_part.element
    taken_abstract = {a.get(qn("w:abstractNumId")) for a in numbering.findall(qn("w:abstractNum"))}
    taken_num = {n.get(qn("w:numId")) for n in numbering.findall(qn("w:num"))}
    clash = ({str(_ABSTRACT_ID)} & taken_abstract) | ({str(i) for i in NUM_IDS} & taken_num)
    if clash:
        raise SystemExit(
            "Numbering ids %s are already used by the letterhead; pick free ones "
            "in prepare_cover_letter_templates.py." % sorted(clash))

    abstracts = numbering.findall(qn("w:abstractNum"))
    abstracts[-1].addnext(abstract)
    anchor = abstract
    for num_id in NUM_IDS:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(_ABSTRACT_ID))
        num.append(ref)
        # Restart at 1. The count belongs to the abstractNum, so two `w:num`
        # sharing one abstract share one counter and Combo's second system
        # numbered 4./5./6. `w:startOverride` is the reset, and it is applied to
        # BOTH instances: the first list must also start at 1 in a document where
        # something else has already used this abstract.
        # Schema order inside `w:num` is `w:abstractNumId` then `w:lvlOverride`*.
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        num.append(override)
        existing = numbering.findall(qn("w:num"))
        if existing:
            existing[-1].addnext(num)
        else:
            anchor.addnext(num)
        anchor = num


# THE DATE BOX IS GONE, AND ON PURPOSE. Kyle's example letter floats the date in a
# small anchored text box over the letterhead artwork, and these templates copied it
# faithfully until 2026-09-04, when Hanz asked for it to come off every format: a
# proposal already carries its own dates, and a second one on the letter is one more
# thing to be wrong. Removing it also takes the ONLY text box out of these documents,
# so a generated letter is now pure flowing body -- `template_geometry`'s `boxes` is
# empty, every block is a direct body child, and the editor renders the no-box layout
# it already had for exactly this case. `{{proposal_date_short}}` still resolves in
# `cover_letter_writer` (the proposal uses it); nothing in a letter prints it.


def _set_numbering(p, num_id: int) -> None:
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num)
    ppr.insert(0, numpr)


def _set_mark_size(p, font, size) -> None:
    """Size the PARAGRAPH MARK, not just the runs.

    Word gives a line the height of the tallest thing on it, and the invisible
    `¶` at the end of a paragraph counts. `add_paragraph()` leaves that mark at
    the style default — 12pt here — so in an 11pt letter every paragraph's LAST
    line was 14.06pt tall against 12.94pt for its wrapped ones, and a "blank
    line" between 10pt signature lines was a 12pt blank line. Measured in the
    render, 2026-08-29: ~22pt of a 792pt page, on Combo alone.

    Schema order inside `w:pPr` puts `w:rPr` after the spacing/indent/jc group
    and before `w:sectPr`/`w:pPrChange`. python-docx has no accessor for it
    (`CT_PPr` models the properties it edits, not the mark), so it is placed by
    hand — before those two if they are ever present, appended otherwise."""
    ppr = p._p.get_or_add_pPr()
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        tail = [ppr.find(qn(t)) for t in ("w:sectPr", "w:pPrChange")]
        tail = [el for el in tail if el is not None]
        if tail:
            tail[0].addprevious(rpr)
        else:
            ppr.append(rpr)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    rpr.append(fonts)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(round(size.pt * 2))))     # half-points
        rpr.append(el)


def _add(d, segments, *, font=BODY_FONT, size=BODY_PT, color=BODY_GREY,
         underline=False, num_id=None, left_indent=None, right_indent=RIGHT_INDENT,
         space_before=None, space_after=None):
    p = d.add_paragraph()
    pf = p.paragraph_format
    if right_indent is not None:
        pf.right_indent = right_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if num_id is not None:
        _set_numbering(p, num_id)
    _set_mark_size(p, font, size)
    for seg in segments:
        r = p.add_run(seg["text"])
        r.font.name = font
        r.font.size = size
        r.font.color.rgb = color
        r.bold = bool(seg.get("bold"))
        r.italic = bool(seg.get("italic"))
        if underline:
            r.underline = True
    return p


def spec_for(work_type: str, audience) -> dict:
    """The copy for one variant. Direct diverged on 2026-09-03 (see DIRECT_COPY);
    everything else -- GC, and gyp for either reader -- takes the shared copy."""
    if audience == "Direct" and work_type in DIRECT_COPY:
        return DIRECT_COPY[work_type]
    return COPY[work_type]


def build(work_type: str, spec: dict, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(LETTERHEAD, out_path)       # branding, byte-for-byte
    d = docx.Document(str(out_path))
    _install_decimal_numbering(d)
    # The paragraph sequence is Example1's, measured off its render rather than
    # eyeballed (see GROUP_SPACE_BEFORE). Where Kyle has a blank LINE there is a
    # blank paragraph here; where his air comes from paragraph spacing, so does
    # ours. Specifically: he runs THANKS straight into the intro line, runs the
    # numbered items together with no gap, and runs the last item straight into
    # "Feel free…". Blank paragraphs in those three places were this generator's
    # invention, cost ~42pt, and are what pushed Combo's sign-off onto a second,
    # letterhead-less page.
    # NO TITLE LINE. Every letter used to open with a red underlined heading —
    # "Epoxy / Resinous Flooring Proposal - {{job_name}}" and its siblings —
    # because Kyle's own example letter opens with one. Hanz took it off every
    # format on 2026-09-04: the proposal stapled behind the letter says which
    # system it is and whose job it is, on its own front page, and a heading that
    # repeats it is a second place for the job name to go stale. The greeting is
    # now the first line on the page.
    #
    # `spec["title"]` is deliberately LEFT IN the copy tables. It is the one
    # human-readable name each variant has, it costs nothing, and putting it back
    # is one `_add` call rather than seven rewritten dictionaries.
    _add(d, [_seg(spec.get("greeting", GREETING))])
    _add(d, [])
    _add(d, [_seg(spec.get("thanks", THANKS))])
    # No air: Kyle runs the thank-you straight into the line that introduces the
    # proposal (measured at a 0.02pt gap in his render). They read as one pair.
    _add(d, [_seg(spec["intro"])])

    for gi, (heading, items) in enumerate(spec["groups"]):
        # The air goes on whichever paragraph opens the group, so an unheaded
        # group (epoxy/polish/gyp) is spaced off the intro exactly as a headed
        # one (combo) is spaced off the group above it.
        lead = GROUP_SPACE_BEFORE
        if heading:
            _add(d, [_seg(heading, bold=True)], space_before=lead)
            lead = None
        for item in items:
            _add(d, item, num_id=NUM_IDS[gi], space_before=lead)
            lead = None
        if gi == 0:
            for note in spec.get("sub_notes") or []:
                _add(d, [_ph(note)], left_indent=Inches(1.0))

    _add(d, [_seg(spec.get("close_1", CLOSE_1))], space_before=GROUP_SPACE_BEFORE)
    _add(d, [_seg(CLOSE_2)])
    # Blank line at the SIGNATURE's size, not the body's: an empty paragraph is
    # as tall as its paragraph mark, and a 12pt gap over a 10pt block reads wrong.
    _add(d, [], font=SIG_FONT, size=SIG_PT)
    for line in SIGNATURE:
        _add(d, line, font=SIG_FONT, size=SIG_PT, color=SIG_GREY)

    d.save(str(out_path))
    print("wrote %s (%d body paragraphs, no text boxes)"
          % (out_path.relative_to(OUT_DIR).as_posix(), len(d.paragraphs)))


# (work_type, audience) -> path under templates/CoverLetter/, mirroring
# `proposal_writer.TEMPLATE_PICKER`'s audience-first folders. `None` is gyp's
# audience there and here: a gypsum-underlayment bid reads the same to either
# reader, so there is one file and no Direct/GC split to keep in step.
#
# Kept in step with `cover_letter_writer.TEMPLATE_PICKER` by
# `backend/tests/test_cover_letter.py`, which asserts the two tables name the
# same files and that every one of them exists — a generator that stops writing
# a file the writer still picks is a 500 on a real send.
VARIANTS = [
    ("epoxy",  "Direct", "Direct/Epoxy.docx"),
    ("epoxy",  "GC",     "GC/Epoxy.docx"),
    ("polish", "Direct", "Direct/Polish.docx"),
    ("polish", "GC",     "GC/Polish.docx"),
    ("combo",  "Direct", "Direct/Combo.docx"),
    ("combo",  "GC",     "GC/Combo.docx"),
    ("gyp",    None,     "Gyp/Gyp.docx"),
]


def main() -> int:
    if not LETTERHEAD.is_file() or not EXAMPLE.is_file():
        print("Missing reference file(s) under %s.\n"
              "They are Hanz's originals and are NOT committed — copy them back "
              "in before regenerating." % REF_DIR, file=sys.stderr)
        return 2
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for work_type, audience, rel in VARIANTS:
        build(work_type, spec_for(work_type, audience), OUT_DIR / rel)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
