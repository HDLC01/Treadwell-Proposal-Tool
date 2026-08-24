"""
Proposal Word-doc writer.

Takes a dict of values + a work_type + an audience, picks the right
Treadwell proposal template, runs `{{token}}` Jinja-style substitution,
returns the filled .docx as bytes.

Kyle's templates in `templates/Direct/`, `templates/GC/`,
`templates/Gyp/` were copied straight from his Numbers 5.7.26 folder.
For v1 they need to be **annotated** with `{{token}}` placeholders
before this writer can fill them — see TEMPLATE_PREP.md (or the
"template prep" Phase in the plan file).

If a template has zero `{{tokens}}`, the writer still returns the file
unchanged with a logged warning — the user gets a usable starter
document; they just have to copy values manually for that template
until tokens are added.

Multi-system support (v2, added 2026-06):
    `fill_proposal` accepts an optional `systems` list. Each item is a
    per-system dict (e.g. system_name, texture, scope_notes, sqft,
    lump_sum). When a template contains a repeatable BLOCK delimited by
    a `{{#system}}` paragraph and a `{{/system}}` paragraph that are
    SIBLINGS in the same container (body, one table cell, or one text
    box), the writer clones the paragraphs between those markers once
    per system, substituting `{{system.field}}` (or bare `{{field}}`)
    tokens against each system dict.

    This is 100% backward-compatible: a template with no `{{#system}}`
    marker, called with `systems=None` (the default), behaves exactly
    like v1 — flat `{{token}}` substitution against `values` only.
    See docs/MULTI-SYSTEM-PROPOSAL.md for the annotation workflow.
"""
from __future__ import annotations

import copy
import io
import logging
import math
import re
from pathlib import Path
from typing import Any, Mapping

import docx
from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run


log = logging.getLogger("proposal_tool.proposal_writer")

TEMPLATES_ROOT = Path(__file__).parent / "templates"


# ─── Template selection ───────────────────────────────────────────────
# (work_type, audience) → relative template path. None audience means
# the template is audience-agnostic (e.g. gypsum, budget).
TEMPLATE_PICKER: dict[tuple[str, str | None], str] = {
    ("epoxy",   "Direct"): "Direct/XX.XX TREADWELL EPOXY PROPOSAL - New Direct.docx",
    ("epoxy",   "GC"):     "GC/xx TREADWELL RESINOUS PROPOSAL - xx.docx",
    ("polish",  "Direct"): "Direct/xx.xx TREADWELL POLISH PROPOSAL - NewDirect.docx",
    ("polish",  "GC"):     "GC/xx TREADWELL POLISH PROPOSAL - xx.docx",
    ("combo",   "Direct"): "Direct/xx.xx.xx TREADWELL COMBO PROPOSAL - CUSTMOER NAME.docx",
    # No dedicated GC combo template — use the GC Resinous (covers the
    # epoxy/resinous side in GC format) instead of falling back to a Direct doc.
    ("combo",   "GC"):     "GC/xx TREADWELL RESINOUS PROPOSAL - xx.docx",
    ("sealer",  "GC"):     "GC/xx TREADWELL SEALER PROPOSAL - xx.docx",
    ("gyp",     None):     "Gyp/xx TREADWELL UNDERLAYMENT PROPOSAL - xx.docx",
    ("budget",  "Direct"): "Direct/xx.xx TREADWELL BUDGET PRICING.docx",
}


def pick_template(work_type: str, audience: str | None) -> Path:
    """Resolve (work_type, audience) → absolute template path.

    Falls back to ('epoxy', 'Direct') if the combination isn't mapped,
    so the tool never hard-fails on an unmapped audience.
    """
    key = (work_type, audience)
    if key not in TEMPLATE_PICKER:
        # Try audience-agnostic fallback (e.g. gyp ignores audience).
        if (work_type, None) in TEMPLATE_PICKER:
            key = (work_type, None)
        else:
            log.warning(
                "No template for (%s, %s); falling back to (epoxy, Direct)",
                work_type, audience,
            )
            key = ("epoxy", "Direct")
    return TEMPLATES_ROOT / TEMPLATE_PICKER[key]


# ─── Token substitution ───────────────────────────────────────────────
TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _replace_in_paragraph(p: Paragraph, values: Mapping[str, Any]) -> int:
    """Replace `{{token}}` in a paragraph, preserving EACH run's formatting.

    Word splits text across multiple <w:r> runs whenever formatting changes —
    e.g. a BOLD "Scope:" label run followed by a NORMAL-weight
    "{{scope_notes}}" value run. The substituted value must keep its OWN run's
    formatting (font, size, bold), not inherit the leading run's. The old code
    collapsed the whole paragraph into run[0], which made every value bold like
    its label. We now rewrite the token's text in place across only the runs it
    actually spans (see `_sub_runs_preserving`).
    """
    if "{{" not in p.text:
        return 0
    return _sub_runs_preserving(
        p._p, TOKEN_RE,
        lambda m: str(values[m.group(1)]) if m.group(1) in values else None,
    )


def _iter_all_paragraphs(d: Document):
    """Yield every paragraph in the doc — body + tables + headers/footers + text boxes.

    python-docx's default `d.paragraphs` skips text in tables, headers,
    footers, and floating text boxes (shapes). For proposal templates
    where the project info often lives in a text box at the top of the
    page, we need to walk the document XML and yield every <w:p>.
    """
    yield from d.paragraphs

    # Tables (recursive — table cells can contain nested tables)
    def walk_table(t):
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in d.tables:
        yield from walk_table(table)

    # Headers / footers
    for section in d.sections:
        for hp in section.header.paragraphs:
            yield hp
        for fp in section.footer.paragraphs:
            yield fp

    # Text boxes / shapes — these live in w:txbxContent inside the body XML.
    # Wrap each <w:p> we find there as a Paragraph object.
    body = d.element.body
    for txbx in body.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            yield Paragraph(p_elem, d)


# ─── Repeatable per-system blocks ─────────────────────────────────────
# A block is delimited by two marker paragraphs that are SIBLINGS in the
# same parent <w:txbxContent> / <w:tc> / <w:body>:
#
#     {{#system}}      ← start marker paragraph (whole paragraph is removed)
#     ... template ... ← cloned once per system
#     {{/system}}      ← end marker paragraph (whole paragraph is removed)
#
# Inside a block, `{{system.field}}` (and bare `{{field}}` as a fallback)
# resolve against each system dict; any other `{{token}}` is left alone
# here and picked up later by the normal flat pass against `values`.
#
# Cloning operates on plain <w:p> elements only — never on the enclosing
# drawing/shape — so it is safe inside floating text boxes (no drawing-id
# or VML-fallback duplication problems). The whole block must therefore
# live inside ONE container (one text box, one table cell, or the body).
# Name-capturing block markers — `{{#<name>}}` / `{{/<name>}}` — so any named
# list can drive a repeatable block (`system`, `price_line`, `alternate`, …).
_WORK_ANCHOR_RE = re.compile(r"^\s*(?:scope|schedule|exclusions)\s*:", re.I)


def _set_run_bold(run_elem, value: bool) -> None:
    """Set an explicit bold value without disturbing the run's other styling."""
    rpr = run_elem.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_elem.insert(0, rpr)
    bold = rpr.find(qn("w:b"))
    if bold is None:
        bold = OxmlElement("w:b")
        rpr.append(bold)
    if value:
        bold.attrib.pop(qn("w:val"), None)
    else:
        bold.set(qn("w:val"), "0")


def _set_direct_run_text(run_elem, text: str) -> None:
    """Replace one run's visible text/break children, retaining its rPr."""
    for child in list(run_elem):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab")):
            run_elem.remove(child)
    t = OxmlElement("w:t")
    run_elem.append(t)
    _write_t_text(t, text)


def _run_text_with_breaks(run_elem) -> str:
    """One run's text with every `<w:br/>` rendered back as a newline.

    `"".join(t.text …)` is BLIND to line breaks, and that blindness is where a line break
    dies: `_set_direct_run_text` clears the run's `<w:t>`/`<w:br>`/`<w:tab>` children and
    rewrites them from the string it is handed, so a caller that measured the run with a
    br-blind join hands back a string with the breaks already gone.

    Kyle, 2026-08-19: "when he pressed enter to add spacing it did not generate in the
    proposal." The WORK box is exactly the box `_normalize_work_label_formatting` splits, and
    a plain-text `paragraph_override` arrives as ONE run holding `<w:t>`/`<w:br>`/`<w:t>` —
    so every blank line typed into Scope / Schedule / Exclusions / Notes was joined away
    there. `_write_t_text` turns the newlines back into `<w:br/>` on the way out, which is
    why round-tripping through a string with them in it is enough to fix it.

    `<w:tab>` is deliberately NOT represented: nothing in these templates writes one into a
    WORK row, and inventing a character for it would shift the label/colon offsets below.
    """
    parts = []
    for el in run_elem.iter(qn("w:t"), qn("w:br")):
        parts.append("\n" if el.tag == qn("w:br") else (el.text or ""))
    return "".join(parts)


def _split_after_visible(s: str, n: int) -> tuple[str, str]:
    """Split `s` after `n` VISIBLE characters, where a newline counts for none.

    The label/colon arithmetic in `_normalize_work_label_formatting` measures offsets in the
    document's TEXT, which is what `_own_text` reports and which contains no breaks. This
    converts an index in that coordinate system into a cut in a string that also carries the
    breaks, so the split lands on the same character it always did and the breaks travel with
    whichever half they were in.
    """
    seen = 0
    for i, ch in enumerate(s):
        if seen == n:
            return s[:i], s[i:]
        if ch != "\n":
            seen += 1
    return s, ""


def _normalize_work_label_formatting(d: Document) -> int:
    """Make WORK-box labels bold through their first colon, values normal.

    SKIPS any run whose weight the estimator stated by hand (`_user_bolded_runs`). Without that
    exemption this pass un-bolded every run after the colon unconditionally, which silently
    made the format toolbar's most-used button a no-op in the WORK box — the area estimators
    edit most. Bold on a phrase inside a Scope / Schedule / Exclusions value reached this pass
    intact and left it as `w:b val="0"`.

    A skipped run is not split either. The split exists ONLY to give the tail its own weight,
    so with the weight left alone there is nothing to split for, and the row keeps the run
    structure the estimator's own edit produced.

    The label still gets bolded automatically in every case where nobody said otherwise: a
    plain-text override registers no runs at all (so a row Kyle retyped is normalized exactly
    as before), and a runs override carries the template's bold back on the label because the
    preview renders it as `font-weight:700` and `serializeRuns` reads it back. What changes is
    only that an EXPLICIT weight now wins over this pass instead of losing to it.
    """
    changed = 0
    hand_bolded = _user_bolded_runs(d)
    for txbx in d.element.body.iter(qn("w:txbxContent")):
        paragraphs = list(txbx.iter(qn("w:p")))
        if not any(_WORK_ANCHOR_RE.match(_own_text(p).strip()) for p in paragraphs):
            continue
        for p_elem in paragraphs:
            text = _own_text(p_elem)
            colon = text.find(":")
            if colon < 0:
                continue
            # A colon occurring later in prose (for example the Gyp terms'
            # "following: access to …") is not a label/value row.
            label = text[:colon].strip()
            if not label or len(label) > 48 or any(ch in label for ch in ".?!"):
                continue
            offset = 0
            passed_colon = False
            for run_elem in list(p_elem.findall(qn("w:r"))):
                # Drawing/object runs only anchor artwork or nested text boxes.
                # `raw` keeps the run's line breaks; `run_text` is the visible text only, which
                # is the coordinate system `colon` and `offset` are measured in (`_own_text`
                # sees no breaks either). A run that is nothing BUT a break has no visible text
                # and is skipped, exactly as it was before — so its break survives untouched.
                raw = _run_text_with_breaks(run_elem)
                run_text = raw.replace("\n", "")
                if not run_text:
                    continue
                start, end = offset, offset + len(run_text)
                offset = end
                # THE ESTIMATOR'S OWN WEIGHT OUTRANKS THIS PASS. The check sits inside each
                # branch rather than short-circuiting the whole run, so `passed_colon` and
                # `offset` stay maintained by the one piece of code that has always owned them.
                hand = id(run_elem) in hand_bolded
                if passed_colon or start > colon:
                    if not hand:
                        _set_run_bold(run_elem, False)
                        changed += 1
                    continue
                if start <= colon < end:
                    passed_colon = True
                    # No split when the weight is the estimator's: the split exists ONLY to
                    # give the tail its own weight, so there is nothing to split for, and the
                    # row keeps the run structure their edit produced.
                    if hand:
                        continue
                    split_at = colon - start + 1
                    if split_at < len(run_text):
                        head, tail = _split_after_visible(raw, split_at)
                        suffix = copy.deepcopy(run_elem)
                        _set_direct_run_text(run_elem, head)
                        _set_direct_run_text(suffix, tail)
                        _set_run_bold(suffix, False)
                        run_elem.addnext(suffix)
                    _set_run_bold(run_elem, True)
                    changed += 1
                elif not hand:
                    _set_run_bold(run_elem, True)
                    changed += 1
    return changed


BLOCK_START_RE = re.compile(r"\{\{\s*#\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
BLOCK_END_RE = re.compile(r"\{\{\s*/\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _dotted_token_re(name: str) -> "re.Pattern":
    """`{{<name>.field}}` — dotted per-item token for a given block name."""
    return re.compile(r"\{\{\s*" + re.escape(name) + r"\.([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _p_text(p_elem) -> str:
    """Joined text of a raw <w:p> element (across all its <w:t> runs)."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _own_text(p_elem) -> str:
    """Like `_p_text`, but STOPS at a nested `<w:txbxContent>` — a paragraph
    that merely anchors a floating text box (the drawing lives in one of its
    runs) must not report that box's entire contents as its own text, since
    the box's inner paragraphs are walked and reported independently (see
    `_iter_body_editable`). Without this, a top-level anchor paragraph's
    `_p_text` would recurse into every nested `<w:t>` — including the whole
    text box's worth of paragraphs concatenated into one string — corrupting
    both the editor's `text` field for that paragraph and any block-marker
    detection run against it. A paragraph with no nested text box behaves
    identically to `_p_text`.
    """
    out = []
    txbx_tag = qn("w:txbxContent")
    for t in p_elem.iter(qn("w:t")):
        nested = False
        anc = t.getparent()
        while anc is not None and anc is not p_elem:
            if anc.tag == txbx_tag:
                nested = True
                break
            anc = anc.getparent()
        if not nested:
            out.append(t.text or "")
    return "".join(out)


def _set_t_multiline(t, text: str) -> None:
    """Write `text` into a <w:t>, rendering embedded newlines as <w:br/> line
    breaks within the same run (so stacked per-item notes share one bullet).
    The raw-lxml item path doesn't go through python-docx's run.text setter,
    which would otherwise convert \\n → <w:br/> for us — so we do it here."""
    parts = text.split("\n")
    t.text = parts[0]
    t.set(qn("xml:space"), "preserve")
    anchor = t
    for part in parts[1:]:
        br = OxmlElement("w:br")
        anchor.addnext(br)
        nt = OxmlElement("w:t")
        nt.set(qn("xml:space"), "preserve")
        nt.text = part
        br.addnext(nt)
        anchor = nt


def _write_t_text(t, text: str) -> None:
    """Set a <w:t>'s text, preserving whitespace and rendering \\n as <w:br/>."""
    if "\n" in text:
        _set_t_multiline(t, text)
    else:
        t.text = text
        t.set(qn("xml:space"), "preserve")


def _sub_runs_preserving(p_elem, pattern, repl, require_braces: bool = True) -> int:
    """Substitute `pattern` matches across a paragraph's runs WITHOUT collapsing
    run formatting.

    The replacement text lands in the run where the match STARTS, and any text
    before/after the match stays in its own run — so a normal-weight value run
    keeps its weight even when an earlier label run is bold (the fix for values
    inheriting the bold "Scope:" label). `repl(match) -> str | None`; returning
    None leaves that match untouched (e.g. a token not in this scope's values),
    so later/known tokens still resolve.

    Works on any element with <w:t> descendants (a python-docx paragraph's `_p`
    or a raw cloned block <w:p>), so both substitution phases share one engine.

    `require_braces` (default True) short-circuits paragraphs with no `{{` — the
    right optimization for the `{{token}}` passes. Pass False to substitute a
    plain-text pattern that isn't a token (e.g. rewriting a hardcoded amount that
    spans runs); the caller's `repl` must then return None once the match already
    equals the replacement, or the loop would rewrite it forever.
    """
    n = 0
    guard = 0
    while guard < 2000:
        guard += 1
        tnodes = list(p_elem.iter(qn("w:t")))
        if not tnodes:
            break
        texts = [(t.text or "") for t in tnodes]
        joined = "".join(texts)
        if require_braces and "{{" not in joined:
            break
        chosen = None
        for m in pattern.finditer(joined):
            r = repl(m)
            if r is not None:
                chosen = (m, r)
                break
        if chosen is None:
            break
        m, value = chosen
        s, e = m.start(), m.end()
        spans = []
        pos = 0
        for txt in texts:
            spans.append((pos, pos + len(txt)))
            pos += len(txt)
        si = so = ei = eo = None
        for i, (a, b) in enumerate(spans):
            if si is None and a <= s < b:
                si, so = i, s - a
            if a < e <= b:
                ei, eo = i, e - a
        if si is None:
            break
        if ei is None:                      # match runs to the very end
            ei, eo = len(tnodes) - 1, len(texts[-1])
        before, after = texts[si][:so], texts[ei][eo:]
        if si == ei:
            _write_t_text(tnodes[si], before + value + after)
        else:
            _write_t_text(tnodes[si], before + value)   # value keeps si's format
            for j in range(si + 1, ei):
                tnodes[j].text = ""
            tnodes[ei].text = after                      # 'after' keeps ei's format
            tnodes[ei].set(qn("xml:space"), "preserve")
        n += 1
    return n


def _substitute_item_tokens(p_elem, item: Mapping[str, Any], block_name: str) -> None:
    """Replace `{{<block>.field}}` / bare `{{field}}` in one cloned <w:p>,
    preserving each run's formatting (see `_sub_runs_preserving`).

    `{{<block>.field}}` always resolves against `item`; bare `{{field}}`
    resolves against `item` ONLY when the key exists there — any other
    `{{token}}` (e.g. {{state_name}}) is left for the flat pass.
    """
    if "{{" not in _p_text(p_elem):
        return
    dotted = _dotted_token_re(block_name)
    _sub_runs_preserving(
        p_elem, dotted,
        lambda m: str(item[m.group(1)]) if m.group(1) in item else None,
    )
    _sub_runs_preserving(
        p_elem, TOKEN_RE,
        lambda m: str(item[m.group(1)]) if m.group(1) in item else None,
    )


# A `{{#price_line}}` row whose amount is empty is a label-only heading row
# (e.g. the combo breakout's restored "Options:" separator — see main.py's
# `_combo_lines` handling). The template paragraph hardcodes the separator as
# literal text between the two tokens —
# `{{price_line.amount_formatted}} – {{price_line.label}}` — not a token
# itself, so once `amount_formatted` substitutes to "" the rendered text
# starts with a bare "– " before the label. Match hyphen, en dash, or em
# dash so this isn't brittle to which one a given template uses.
_LEADING_SEP_RE = re.compile(r"^\s*[-–—]\s*")


def _strip_leading_separator(p_elem) -> None:
    """Strip a leading `<amount> <dash> ` separator off an already-substituted
    price_line paragraph whose amount was empty.

    Operates on the rendered text across all of the paragraph's `<w:t>` runs
    (the separator may land in the same run as the tokens, as it currently
    does, or in a run of its own if a template is authored differently) and
    trims exactly the matched leading characters off the front run(s), so any
    remaining text keeps its own run/formatting untouched. No-op if the
    paragraph doesn't start with a separator (e.g. a normal priced row).
    """
    tnodes = list(p_elem.iter(qn("w:t")))
    if not tnodes:
        return
    joined = "".join(t.text or "" for t in tnodes)
    m = _LEADING_SEP_RE.match(joined)
    if not m or m.end() == 0:
        return
    remaining = m.end()
    for t in tnodes:
        if remaining <= 0:
            break
        cur = t.text or ""
        if len(cur) <= remaining:
            remaining -= len(cur)
            t.text = ""
        else:
            t.text = cur[remaining:]
            t.set(qn("xml:space"), "preserve")
            remaining = 0


# ── per-row WORK labels ("Texture:", "Area:") ────────────────────────────────
# Kyle, 2026-08-19, on the document editor: "Everything on that page must be editable like a
# word doc." Everything on the page IS a contenteditable paragraph — except the {{#system}}
# region, which the editor replaces with a synthesized preview because the region is expanded
# per priced system and its paragraph ids stop meaning anything once it is. "System:" was
# already the computed {{system.prefix}} token, so it only needed whitelisting. "Texture:" and
# "Area:" are STATIC template text inside those rows, and static text cannot be reached by
# paragraph_overrides (_apply_paragraph_overrides skips anything with in_block set — deliberately,
# because that content is engine-owned).
#
# So the label rides the row's OWN item dict, on the per-index `system_overrides` channel the
# values (name/texture/sqft) already use — `texture_label` / `area_label`. Nothing about the id
# space changes: iter_editable_blocks yields the same blocks in the same order, so every
# paragraph_override saved against a draft in flight still lands where it did before.
#
# Anchored on the row's TOKEN, not on the label's own wording: the Texture row is the one holding
# {{system.texture}} and the Area row the one holding {{system.sqft}}, which is true of a
# re-authored template too. Runs at expansion time, BEFORE _substitute_item_tokens consumes the
# tokens (see _expand_named_block).
_SYSTEM_ROW_LABELS = (
    ("texture_label", "texture"),
    ("area_label", "sqft"),
)

# The static label is everything up to the row's first token, trimmed to the last colon:
# "Texture:  " → "Texture:", "Area: ~" → "Area:". The greedy `.*:` is what makes the "~" (and
# any other separator a template puts between the label and the number) survive untouched.
_ROW_LABEL_RE = re.compile(r"^(\s*)(.*:)", re.DOTALL)


def _splice_t_range(tnodes, start: int, end: int, repl: str) -> bool:
    """Replace joined-text characters [start, end) across a paragraph's <w:t> nodes with `repl`.

    `repl` lands in the run that owned the FIRST replaced character, so a label keeps the bold /
    size / colour the template gave it; later runs in the range only lose their share of the old
    text. Same technique as `_strip_leading_separator`, generalized to an arbitrary span.
    """
    pos = 0
    placed = False
    for t in tnodes:
        cur = t.text or ""
        a, b = pos, pos + len(cur)
        pos = b
        if b <= start or a >= end:
            continue
        lo = max(0, start - a)
        hi = min(len(cur), end - a)
        t.text = cur[:lo] + ("" if placed else repl) + cur[hi:]
        t.set(qn("xml:space"), "preserve")
        placed = True
    return placed


def _apply_system_row_labels(p_elem, item: Mapping[str, Any]) -> int:
    """Rewrite one expanded {{#system}} row's static label from the item's `*_label` override.

    No-op unless the item carries the override for the token this paragraph holds, and no-op if
    the paragraph has no static label before its first token. Braces are stripped out of the
    override so a pasted "{{token}}" can never reach a customer-facing document as literal text.
    """
    n = 0
    for key, token in _SYSTEM_ROW_LABELS:
        raw = item.get(key)
        if raw is None:
            continue
        label = str(raw).replace("{", "").replace("}", "").strip()
        if not label:
            continue
        tnodes = list(p_elem.iter(qn("w:t")))
        if not tnodes:
            continue
        joined = "".join(t.text or "" for t in tnodes)
        if not re.search(r"\{\{\s*system\." + re.escape(token) + r"\s*\}\}", joined):
            continue
        head_end = joined.find("{{")
        if head_end <= 0:
            continue
        m = _ROW_LABEL_RE.match(joined[:head_end])
        if not m:
            continue
        start = len(m.group(1))
        if _splice_t_range(tnodes, start, start + len(m.group(2)), label):
            n += 1
    return n


# ── per-row WHOLE-LINE WORK overrides ────────────────────────────────────────
# Kyle, for the third time, 2026-08-24: "every line in the proposal must be editable AS ONE
# LINE, the way the base bid is. No token islands with untouchable words around them ... I
# cannot delete 'SF of epoxy flooring' from 'Area: ~2,305 SF of epoxy flooring'."
#
# He is right about the mechanism. The template's Area row is
#     Area: ~{{system.sqft}} SF of epoxy flooring{{system.lf_clause}}
# and until now the only two things this file could rewrite in that row were the label
# ("Area:", above) and the token's value. The literal "~", the words " SF of epoxy flooring"
# and the whole cove clause had NO channel at all — not the label one, not
# paragraph_overrides, which refuses any paragraph with `in_block` set because a {{#system}}
# region is expanded once per priced system and its paragraph ids stop describing anything the
# estimator saw.
#
# So the WHOLE LINE rides the same per-index `system_overrides` channel the labels and values
# already use — `name_line` / `texture_line` / `area_line`. The row is found by the TOKEN it
# carries, exactly as the labels are, which is true of a re-authored template too. That has to
# happen BEFORE substitution consumes the tokens; the text is written AFTER it, so the diff the
# estimator saw on screen (a fully resolved line) is the text that lands.
#
# THE TRADE, stated once. A row nobody has touched is still rendered from the tokens, so a
# changed square footage flows through. A row with a stored line prints that line verbatim and
# stops tracking the sheet, because a hand-written sentence has no slot to re-substitute a
# number into. That is the same trade already accepted for every PRICE line including the base
# bid, which is money; the on-screen ⚠ says the line differs from the estimate, and clearing
# the line brings the computed text back.
_SYSTEM_ROW_LINES = (
    ("name_line", "name"),
    ("texture_line", "texture"),
    ("area_line", "sqft"),
)


def _system_row_line_key(p_elem, item: Mapping[str, Any]) -> str | None:
    """Which `*_line` override belongs to this NOT-YET-SUBSTITUTED {{#system}} row, if any.

    Anchored on the row's own token (`{{system.name}}` / `.texture` / `.sqft`) because that is
    the only thing that survives a re-authored template — the label wording does not. Must be
    called before `_substitute_item_tokens`, which is what removes those tokens.
    """
    joined = _p_text(p_elem)
    for key, token in _SYSTEM_ROW_LINES:
        raw = item.get(key)
        if raw is None or not str(raw).strip():
            continue
        if re.search(r"\{\{\s*system\." + re.escape(token) + r"\s*\}\}", joined):
            return key
    return None


def _apply_system_row_line(p_elem, text: Any) -> bool:
    """Replace one expanded {{#system}} row's ENTIRE visible text with the estimator's line.

    Same writer the whole-line PRICE overrides use (`_set_paragraph_text`), so the paragraph
    keeps its first run's font/size/colour, any anchored drawing in the paragraph survives, and
    a newline the estimator typed becomes a real `<w:br/>`. The bold lead-in is re-derived
    afterwards by `_normalize_work_label_formatting` (bold through the first colon), which is
    also what the on-screen line renders — so screen and document agree on the weight even
    after the row has been rewritten.

    Braces are stripped, for the reason `_apply_system_row_labels` gives: this runs after the
    per-item substitution and before the flat pass, so a pasted "{{system.sqft}}" would
    otherwise reach a customer-facing document as literal text.

    Returns True when the paragraph was rewritten (the caller uses that to keep the row out of
    `_drop_zero_sf_prefix`'s reach — a hand-typed line must not be re-edited by a regex).
    """
    line = str(text).replace("{", "").replace("}", "")
    if not line.strip():
        return False
    _set_paragraph_text(p_elem, line)
    return True


def _strip_bullet(p_elem) -> None:
    """Remove list/bullet formatting (`<w:numPr>`) from a paragraph so a blank
    NOTES item renders as clean vertical spacing — a genuinely empty line —
    rather than a lone empty bullet dot. No-op if the paragraph isn't a list
    item. (The empty `<w:t>` from the substitution already makes the line
    blank; this just drops its bullet glyph.)"""
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is None:
        return
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


# ─── Word PARAGRAPH properties the doc editor can change ─────────────────────
# Kyle, 2026-08-20, on the proposal document editor:
#     "I cant dletet the bullet points"
#     "There is indentation in this but I cant remove tat if I want to to be
#      aligned on the polished concrete?"
#
# Both complaints are about `w:pPr`, not about text. The editor has always been able to rewrite
# a paragraph's RUNS (text, bold/italic/underline, size); it could never reach the paragraph's
# own numbering (`w:numPr`) or indentation (`w:ind`), which is where a bullet and a hanging
# indent live. These helpers are that reach, and they are shared by BOTH override channels:
# `paragraph_overrides` (the plain WORK rows — Scope / Schedule / Exclusions / Notes) and the
# per-index `system_overrides` rows (System / Texture / Area, which live inside the
# `{{#system}}` region and so can never be addressed by a paragraph id — see
# `_apply_system_row_labels` for the same split).
#
# TWO RULES, both read off these templates rather than assumed:
#
#  1. REMOVING A BULLET MUST NOT MOVE THE TEXT. A WORK row carries no `w:ind` of its own; its
#     indentation comes from the numbering level (numId 4 -> `w:ind w:left=288 w:hanging=288`,
#     so the square prints at 0 and the text at 288). Drop `w:numPr` alone and the level's
#     indent goes with it, the paragraph falls back to the style chain, and the line jumps.
#     Kyle's own template shows the target state on the Texture row: pStyle=ListParagraph, NO
#     numPr, and an explicit `w:ind`. So the level's left indent is copied onto the paragraph
#     before the numbering reference is removed.
#
#  2. AN ORDERED LIST IS OFF LIMITS. numId 5 is the numbered TERMS AND CONDITIONS list. Those
#     paragraphs are plain body paragraphs with `in_block=None`, so the paragraph-override
#     channel CAN address them — and removing one item from a decimal list renumbers every
#     clause after it, silently, in legal boilerplate. `para_props()` reports such a paragraph
#     as `locked` (the editor hides the controls) and `apply_para_props()` refuses it outright,
#     so neither half can renumber the Terms.
#
# Indents are absolute twips, not "levels": the client sends the left indent it wants, this
# clamps it, and 0 really is flush left — which is the whole point of Kyle's second complaint.
_INDENT_STEP_TW = 288       # one step = the WORK/NOTES list level's own indent in these templates
_INDENT_MAX_TW = 2880       # 2 inches, far past anything that still fits a text box
# Only used when a paragraph is on a list whose definition cannot be read (no numbering part, a
# dangling numId). Every Kyle template's bullet levels indent by 288, so a de-bulleted row lands
# where its neighbours are instead of at the margin.
_BULLET_FALLBACK_LEFT_TW = 288


def _tw_or_none(raw):
    """A twip measurement off an OOXML attribute as a non-negative int, or None when it is
    absent or unparseable. Word writes these as decimal strings; a corrupt one must not raise
    out of a paragraph-property read."""
    if raw is None:
        return None
    try:
        return max(0, int(float(raw)))
    except (TypeError, ValueError):
        return None


def _get_or_make_ppr(p_elem):
    """The paragraph's `w:pPr`, created (as the FIRST child, which OOXML requires) if absent."""
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_elem.insert(0, ppr)
    return ppr


def _numbering_levels(d) -> dict:
    """`{(numId, ilvl): {"fmt", "text", "start", "ind"}}` for this document.

    `fmt`   — the `w:numFmt` ("bullet", "decimal", ...).
    `text`  — the `w:lvlText`, i.e. what the level actually PRINTS in front of the paragraph:
              "%1." for the numbered Terms and Conditions clauses, a Wingdings glyph (U+F0A7)
              for every bullet list in Kyle's templates. `fmt` answers "decimal or bullet";
              this answers "what does the reader SEE", which is the question the document
              editor was getting wrong — it drew a red square in front of all 27 clauses that
              print "1." to "27." in the signed contract.
    `start` — the `w:start`, the first ordinal of the level (1 in every template here).
    `ind`   — the level's own `w:ind` attributes.

    Read once per Document and cached on it (same technique as `_user_sized_paragraphs` — a
    custom attribute in the XML would be invalid OOXML). A document with no numbering part
    caches an empty map rather than re-raising on every paragraph."""
    got = getattr(d, "_tw_num_levels", None)
    if got is not None:
        return got
    levels: dict = {}
    try:
        root = d.part.numbering_part.element
    except Exception:  # noqa: BLE001 — no numbering part, or an unreadable one
        root = None
    if root is not None:
        abstract: dict = {}
        for anum in root.iter(qn("w:abstractNum")):
            aid = anum.get(qn("w:abstractNumId"))
            for lvl in anum.findall(qn("w:lvl")):
                fmt_el = lvl.find(qn("w:numFmt"))
                txt_el = lvl.find(qn("w:lvlText"))
                start_el = lvl.find(qn("w:start"))
                lppr = lvl.find(qn("w:pPr"))
                ind = lppr.find(qn("w:ind")) if lppr is not None else None
                abstract[(aid, lvl.get(qn("w:ilvl")))] = {
                    "fmt": fmt_el.get(qn("w:val")) if fmt_el is not None else None,
                    "text": txt_el.get(qn("w:val")) if txt_el is not None else None,
                    "start": _tw_or_none(start_el.get(qn("w:val"))) if start_el is not None else None,
                    "ind": ({k.split("}")[-1]: v for k, v in ind.attrib.items()}
                            if ind is not None else {}),
                }
        for num in root.iter(qn("w:num")):
            nid = num.get(qn("w:numId"))
            aref = num.find(qn("w:abstractNumId"))
            aid = aref.get(qn("w:val")) if aref is not None else None
            if nid is None or aid is None:
                continue
            for (a, ilvl), info in abstract.items():
                if a == aid:
                    levels[(nid, ilvl)] = info
    try:
        d._tw_num_levels = levels
    except AttributeError:      # pragma: no cover — a Document always accepts attributes
        pass
    return levels


def _para_num_ref(p_elem):
    """`(numId, ilvl)` as STRINGS for a paragraph carrying `w:numPr`, else None.

    A `w:numPr` without an explicit `w:ilvl` means level 0, which is how Kyle's rows are
    authored."""
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    nid = numpr.find(qn("w:numId"))
    if nid is None or nid.get(qn("w:val")) is None:
        return None
    ilvl = numpr.find(qn("w:ilvl"))
    return nid.get(qn("w:val")), ((ilvl.get(qn("w:val")) if ilvl is not None else None) or "0")


def _num_fmt(d, p_elem):
    """The `w:numFmt` of the list level this paragraph is on ("bullet", "decimal", ...), or None
    when the paragraph is not on a list / the definition cannot be read."""
    ref = _para_num_ref(p_elem)
    if ref is None:
        return None
    info = _numbering_levels(d).get(ref)
    return (info or {}).get("fmt")


def _para_ordered_list(d, p_elem) -> bool:
    """True for a paragraph on a NUMBERED (non-bullet) list — the Terms and Conditions clauses.

    Rule 2 above: this is the one class of paragraph the controls refuse, because dropping its
    numbering renumbers every clause below it. An unreadable definition counts as ordered —
    refusing to touch a list we cannot identify is the safe direction when the alternative is
    renumbering a contract."""
    if _para_num_ref(p_elem) is None:
        return False
    return _num_fmt(d, p_elem) != "bullet"


# ─── WHAT AN ORDERED LEVEL ACTUALLY PRINTS ───────────────────────────────────
# The document editor drew a red square in front of all 27 numbered TERMS AND CONDITIONS
# clauses, because `/api/proposal-template` told it `list: True` (the paragraph carries
# `w:numPr`) and the renderer read that as "bulleted". The flag was never wrong; it was
# answering a different question. The question the editor needs answered is what the level
# PRINTS in front of this paragraph, and for numId 5 that is "1." to "27." — the numbers the
# signed contract shows and a bulleted preview does not.
#
# `w:numFmt` alone cannot answer it either: it says "decimal", not "1." — the trailing period,
# and any prefix, live in `w:lvlText` ("%1."). So the marker is `w:lvlText` with each `%N`
# replaced by that level's running count, which is exactly how Word renders one.
_MAX_ROMAN = 3999
_ROMAN_PARTS = ((1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"), (90, "xc"),
                (50, "l"), (40, "xl"), (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i"))


def _roman(n: int) -> str:
    """`n` as a lowercase Roman numeral. Empty for anything Roman cannot spell (<1, >3999)."""
    if n < 1 or n > _MAX_ROMAN:
        return ""
    out = []
    for val, sym in _ROMAN_PARTS:
        while n >= val:
            out.append(sym)
            n -= val
    return "".join(out)


def _alpha(n: int) -> str:
    """`n` as Word's letter sequence: 1→a, 26→z, 27→aa. Empty below 1."""
    out = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        out = chr(ord("a") + rem) + out
    return out


def _format_ordinal(n: int, fmt: str | None) -> str:
    """One ordinal in a `w:numFmt`. Anything unrecognised (or unspellable) falls back to the
    decimal, which is what every ordered level in Kyle's templates uses at the level that
    matters and is never a worse answer than showing nothing."""
    if fmt == "lowerLetter":
        return _alpha(n) or str(n)
    if fmt == "upperLetter":
        return (_alpha(n) or str(n)).upper()
    if fmt == "lowerRoman":
        return _roman(n) or str(n)
    if fmt == "upperRoman":
        return (_roman(n) or str(n)).upper()
    return str(n)


def _render_lvl_text(raw: str, num_id: str, counters: dict, levels: dict) -> str:
    """`w:lvlText` with every `%N` replaced by level N-1's current count.

    A level the walk has not reached yet contributes its own `w:start` rather than nothing, so a
    nested "%1.%2" reads as "1.1" rather than ".1". Literal text around the placeholders (the
    "." in "%1.") is the template's and is kept verbatim."""
    def one(m):
        lvl = str(int(m.group(1)) - 1)
        info = levels.get((num_id, lvl)) or {}
        n = counters.get((num_id, lvl))
        if n is None:
            n = info.get("start")
            n = 1 if n is None else n
        return _format_ordinal(n, info.get("fmt"))
    return re.sub(r"%(\d)", one, raw)


def _ordered_markers(d) -> dict:
    """`{id(w:p): (w:p, marker)}` for every paragraph on a NUMBERED list in this document.

    Document order, one counter per `(numId, ilvl)` starting at the level's `w:start`, which is
    how Word numbers a plain list. `w:lvlOverride` / `w:startOverride` restarts are NOT modelled:
    no template here uses one, and a marker this cannot resolve comes back as "" (see
    `para_props`) rather than as a guess.

    THE VALUE PINS THE ELEMENT, and it has to — same reason `_hand_formatted` gives. lxml frees
    an element proxy as soon as the last Python reference goes and hands out a brand new one, at
    a possibly REUSED address, on the next access. Holding the element keeps its proxy alive,
    which keeps `id()` both stable and unique, and means a caller who looks up the paragraph it
    is iterating gets that same proxy back.

    Cached per document and keyed on the paragraph COUNT: this map is a read-path answer about
    document order, so it must not survive a pass that inserts or removes a paragraph (block
    expansion does exactly that). The count is the cheapest honest way to notice."""
    paras = list(d.element.body.iter(qn("w:p")))
    got = getattr(d, "_tw_ordered_markers", None)
    if got is not None and got[0] == len(paras):
        return got[1]
    levels = _numbering_levels(d)
    counters: dict = {}
    out: dict = {}
    for p in paras:
        ref = _para_num_ref(p)
        if ref is None:
            continue
        info = levels.get(ref)
        fmt = (info or {}).get("fmt")
        # An unreadable level, a bullet level and Word's "no marker at all" level are all
        # "this paragraph prints no number". Only a real ordered level gets a counter.
        if not info or fmt is None or fmt in ("bullet", "none"):
            continue
        prev = counters.get(ref)
        start = info.get("start")
        counters[ref] = prev + 1 if prev is not None else (1 if start is None else start)
        raw = info.get("text")
        out[id(p)] = (p, _render_lvl_text(raw, ref[0], counters, levels) if raw
                      else _format_ordinal(counters[ref], fmt) + ".")
    try:
        d._tw_ordered_markers = (len(paras), out)
    except AttributeError:      # pragma: no cover — a Document always accepts attributes
        pass
    return out


def _para_marker(d, p_elem) -> str:
    """What an ORDERED list level prints in front of this paragraph ("1.", "27.", "a."), or "".

    Empty for a bullet row (its square is the preview's own CSS, and a Wingdings private-use
    glyph is not a character a browser can render), for a plain paragraph, and for an ordered
    level whose definition cannot be read — in that last case the caller keeps whatever it did
    before rather than being told a number that might be wrong."""
    hit = _ordered_markers(d).get(id(p_elem))
    return hit[1] if hit else ""


def _effective_left_tw(d, p_elem) -> int:
    """The paragraph's current left indent in twips, as the renderer resolves it.

    Direct `w:ind` on the paragraph wins; otherwise the numbering level's indent; otherwise 0.
    The STYLE chain's own indent is deliberately not walked: a `w:ind` written by this feature
    overrides it anyway, so folding it in here would make the toolbar's readout disagree with
    what an outdent actually produces."""
    ppr = p_elem.find(qn("w:pPr"))
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    if ind is not None:
        for attr in ("w:start", "w:left"):
            got = _tw_or_none(ind.get(qn(attr)))
            if got is not None:
                return got
    ref = _para_num_ref(p_elem)
    if ref is not None:
        lvl_ind = (_numbering_levels(d).get(ref) or {}).get("ind") or {}
        got = _tw_or_none(lvl_ind.get("start", lvl_ind.get("left")))
        if got is not None:
            return got
    return 0


def para_props(d, p_elem) -> dict:
    """This paragraph's editable paragraph properties, for `/api/proposal-template`.

    `bullet`  — does it currently render a bullet?
    `indent`  — its resolved left indent in twips (see `_effective_left_tw`).
    `locked`  — an ordered list; the editor must not offer the controls (rule 2).
    `marker`  — what an ORDERED level prints in front of it ("1.", "27."), else "".

    `bullet` and `marker` are the two halves of "what does this paragraph show in front of its
    text", and they are never both set. `locked` is a different kind of answer — a POLICY, about
    what the editor may change — and reading one as the other is the bug this field ends: the
    editor rendered every numbered contract clause as a red square because it trusted `list`,
    which is True for a bullet row and a numbered clause alike.
    """
    return {
        "bullet": _num_fmt(d, p_elem) == "bullet",
        "indent": _effective_left_tw(d, p_elem),
        "locked": _para_ordered_list(d, p_elem),
        "marker": _para_marker(d, p_elem),
    }


def _write_left_indent(ppr, left_tw: int, hanging_tw: int | None = None) -> None:
    """Pin the paragraph's left edge at `left_tw` twips.

    `w:left` AND `w:start` are both written — they are the same property in two schema
    generations and Word/LibreOffice disagree about which one to read.

    `hanging_tw=None` (the un-bulleted case) clears `w:hanging` / `w:firstLine`, because a
    hanging indent with no bullet in front of it prints as a first line that starts further
    left than the rest of its own paragraph.

    A paragraph that KEEPS its bullet passes the hanging it needs instead: the square prints at
    `left - hanging`, so dropping the hanging would print it inline with the words rather than
    ahead of them. The caller clamps it to `left` — a hanging bigger than the left indent puts
    the square at a negative position, i.e. out in the margin or off the page."""
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), str(int(left_tw)))
    ind.set(qn("w:start"), str(int(left_tw)))
    for gone in ("w:hanging", "w:firstLine", "w:startChars", "w:leftChars", "w:firstLineChars"):
        if ind.get(qn(gone)) is not None:
            del ind.attrib[qn(gone)]
    # `is not None`, not truthiness: a hanging of ZERO is a real request ("square and text at the
    # same place") and used to be silently dropped, which is indistinguishable in the output from
    # "no hanging wanted" and was half of the vanishing-square bug above. The caller no longer
    # asks for zero on a bulleted paragraph, but the writer must not be the thing that decides.
    if hanging_tw is not None:
        ind.set(qn("w:hanging"), str(int(hanging_tw)))


def _drop_left_indent(p_elem) -> bool:
    """Remove the paragraph's OWN `w:left`/`w:start`, so whatever it inherits governs again.

    Used when the indent being asked for is already exactly what the paragraph's bullet level
    provides: keeping an explicit `w:ind` there would be inert at best, and (because the
    explicit one has no `w:hanging`) actually moves the square. An emptied `w:ind` element is
    removed with the attributes, so a round trip leaves the XML as it found it."""
    ppr = p_elem.find(qn("w:pPr"))
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    if ind is None:
        return False
    hit = False
    for gone in ("w:left", "w:start", "w:leftChars", "w:startChars"):
        if ind.get(qn(gone)) is not None:
            del ind.attrib[qn(gone)]
            hit = True
    if not len(ind.attrib):
        ppr.remove(ind)
    return hit


def _remove_bullet_keep_indent(d, p_elem) -> bool:
    """Drop the paragraph's bullet WITHOUT letting its text move (rule 1).

    The numbering level's left indent is copied onto the paragraph as an explicit `w:ind`
    first, so the line stays where the list had it. Returns False when there was no bullet."""
    ref = _para_num_ref(p_elem)
    if ref is None:
        return False
    lvl_ind = (_numbering_levels(d).get(ref) or {}).get("ind") or {}
    left = _tw_or_none(lvl_ind.get("start", lvl_ind.get("left")))
    if left is None:
        left = _BULLET_FALLBACK_LEFT_TW
    ppr = _get_or_make_ppr(p_elem)
    existing = ppr.find(qn("w:ind"))
    # A paragraph that already states its own left indent keeps it — that value is what the
    # renderer was using, and the numbering level's was not.
    if existing is None or (existing.get(qn("w:left")) is None
                            and existing.get(qn("w:start")) is None):
        _write_left_indent(ppr, max(0, left))
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)
    return True


def _sibling_bullet_ref(d, p_elem):
    """The `(numId, ilvl)` of the BULLET list this paragraph should join when its bullet is
    switched back on: the one its siblings in the same container already use.

    Sibling-scoped on purpose. Picking "any bullet list in the document" would put a WORK row on
    the NOTES list, whose level indents differently, and there is no honest way to invent a
    numbering definition the template's own styles agree with. Returns None when no sibling is
    on a bullet list, in which case `apply_para_props` leaves the paragraph alone rather than
    guessing — `para_props()['bullet']` then keeps reporting False, so the toolbar and the
    document still agree about what happened."""
    parent = p_elem.getparent()
    if parent is None:
        return None
    levels = _numbering_levels(d)
    for sib in parent.iterchildren(qn("w:p")):
        if sib is p_elem:
            continue
        ref = _para_num_ref(sib)
        if ref is not None and (levels.get(ref) or {}).get("fmt") == "bullet":
            return ref
    return None


def _add_bullet(d, p_elem) -> bool:
    """Put the paragraph back on its siblings' bullet list.

    Any explicit `w:ind` a previous de-bullet wrote is removed again, so the numbering level's
    own indent (square at the hanging position, text after it) governs once more — leaving it
    behind is what would print the square in the middle of the line."""
    if _para_num_ref(p_elem) is not None:
        return False                    # already a list item
    ref = _sibling_bullet_ref(d, p_elem)
    if ref is None:
        return False
    num_id, ilvl = ref
    ppr = _get_or_make_ppr(p_elem)
    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        ppr.remove(ind)
    numpr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_el = OxmlElement("w:numId")
    num_el.set(qn("w:val"), str(num_id))
    numpr.append(ilvl_el)
    numpr.append(num_el)
    # w:numPr belongs near the top of w:pPr (right after w:pStyle) per the schema's sequence.
    style = ppr.find(qn("w:pStyle"))
    if style is not None:
        style.addnext(numpr)
    else:
        ppr.insert(0, numpr)
    return True


def sanitize_para_props(raw) -> dict:
    """Coerce one client-supplied paragraph-property dict to `{bullet?: bool, indent?: int}`.

    Defensive like every other override sanitizer here: anything unrecognised is dropped, an
    out-of-range indent is CLAMPED rather than rejected (a clamp still does what the estimator
    asked for, as far as the page can go), and an empty result means "no paragraph change" so
    the caller can skip the work. Never raises."""
    if not isinstance(raw, Mapping):
        return {}
    out: dict = {}
    b = raw.get("bullet")
    if isinstance(b, bool):
        out["bullet"] = b
    ind = raw.get("indent")
    if isinstance(ind, bool):       # bool is an int subclass; True would mean 1 twip
        ind = None
    if isinstance(ind, (int, float, str)):
        # `round(float("nan"))` raises ValueError and `float("inf")` clamps to the max but
        # `round` on it raises too, so the conversion is guarded rather than type-gated. A
        # JSON body really can carry NaN (Python's own json module emits and accepts it), and
        # this function's whole contract is that it never raises.
        try:
            out["indent"] = int(max(0, min(_INDENT_MAX_TW,
                                           round(float(ind.strip() if isinstance(ind, str) else ind)))))
        except (TypeError, ValueError, OverflowError):
            pass
    return out


def apply_para_props(d, p_elem, props) -> int:
    """Apply `{bullet?, indent?}` to one paragraph. Returns how many properties changed.

    REFUSES an ordered-list paragraph outright (rule 2 — the Terms and Conditions clauses),
    which is also what `para_props()['locked']` tells the editor so it never offers the controls
    there.

    Order matters: the bullet change runs FIRST, because turning a bullet off writes the indent
    it is preserving and turning one on hands the indent back to the list level. An `indent` in
    the same request is then resolved against whichever of the two the paragraph ended up on —
    which is why it is read after, not before.

    "Resolved", not "written": on a paragraph that is still bulleted at an indent its own list
    level already provides, the correct action is to state NOTHING and let the level govern.
    Stating it would drop the level's `w:hanging` and print the square inline with the words.
    """
    clean = sanitize_para_props(props)
    if not clean:
        return 0
    if _para_ordered_list(d, p_elem):
        return 0
    n = 0
    if "bullet" in clean:
        had = _num_fmt(d, p_elem) == "bullet"
        if clean["bullet"] and not had:
            if _add_bullet(d, p_elem):
                n += 1
        elif not clean["bullet"] and _para_num_ref(p_elem) is not None:
            if _remove_bullet_keep_indent(d, p_elem):
                n += 1
    if "indent" in clean:
        want = clean["indent"]
        # Read the bullet state AFTER the branch above: it may have just changed, and what a
        # correct indent looks like differs entirely between the two cases.
        lvl_ind = {}
        if _num_fmt(d, p_elem) == "bullet":
            lvl_ind = (_numbering_levels(d).get(_para_num_ref(p_elem)) or {}).get("ind") or {}
        lvl_left = _tw_or_none(lvl_ind.get("start", lvl_ind.get("left")))
        if lvl_ind and lvl_left == want:
            # STILL BULLETED, and the list level already puts the text exactly where this asks
            # for it. Writing our own `w:left` here would come with no `w:hanging` (the level's
            # is not ours to restate) and so would move the square from in front of the words
            # to inline with them — which is what happened to every paragraph the editor sent
            # its own unchanged state for, including a bullet switched off and straight back on.
            # Letting the level govern is both correct and byte-identical to the template.
            if _drop_left_indent(p_elem):
                n += 1
        else:
            ppr_now = p_elem.find(qn("w:pPr"))
            ind_now = ppr_now.find(qn("w:ind")) if ppr_now is not None else None
            stated = ind_now is not None and (ind_now.get(qn("w:left")) is not None
                                             or ind_now.get(qn("w:start")) is not None)
            # Written whenever the paragraph does not already STATE this indent itself. Comparing
            # only the resolved value would skip the write on a paragraph that inherits the same
            # number from its list level — and then a later bullet change would take the indent
            # with it, which is the bug this feature exists to fix.
            if not stated or _effective_left_tw(d, p_elem) != want:
                # A PARAGRAPH THAT KEPT ITS BULLET CANNOT PUT ITS TEXT AT THE MARGIN, because
                # the square has to go somewhere and it goes `hanging` twips to the LEFT of the
                # text. `indent` means the text's left edge everywhere in this feature (see
                # `_effective_left_tw`, which reads w:left), so an indent below the hanging asks
                # for a square at a negative position - out in the page margin, where Word simply
                # does not draw it.
                #
                # It used to ask for exactly that. `min(hang, want)` with want=0 gave a hanging of
                # 0, and `_write_left_indent` wrote the hanging under a plain truthiness test, so
                # a zero hanging was dropped and the paragraph got `<w:ind w:left="0"/>` while
                # still carrying `w:numPr`. Measured: no U+25AA anywhere on that line in the
                # rendered PDF. One press of outdent on a WORK row - the exact flow this feature
                # was built for - silently deleted the row's red square from the customer's
                # document, while the editor kept drawing it and `para_props()` kept reporting
                # `bullet: True`. Three readers, three different answers.
                #
                # So the floor for a still-bulleted row is its own hanging: the square lands at
                # `left - hanging` = 0, flush with the margin and in line with the squares of the
                # rows above and below it, and the text sits one hanging further in. That is as
                # far left as a bulleted row can travel and still be a bulleted row. An estimator
                # who wants the text itself at the margin turns the bullet off, which is what the
                # third control in the same toolbar is for.
                #
                # Nothing changes for want >= hang: left is want and the hanging is the level's,
                # exactly as before, since min(hang, want) was already hang there.
                hang = _tw_or_none(lvl_ind.get("hanging")) if lvl_ind else None
                if hang is None:
                    _write_left_indent(_get_or_make_ppr(p_elem), want)
                else:
                    _write_left_indent(_get_or_make_ppr(p_elem), max(want, hang), hang)
                n += 1
    return n


_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
# Default text-box insets when <a:bodyPr> omits them (OOXML defaults): 0.1" L/R,
# 0.05" T/B. In points.
_TXBX_INSET_LR_PT = 0.1 * 72 * 2      # left + right
_TXBX_INSET_TB_PT = 0.05 * 72 * 2     # top + bottom
# Rough proportional-font metrics for Carlito/Calibri body text: average glyph
# advance ≈ 0.5·fontSize, single line height ≈ 1.2·fontSize. Biased slightly
# toward OVER-estimating height (wider glyph, taller line) so we err on the side
# of shrinking a hair MORE rather than clipping. The floor mirrors the editor's
# on-screen fitTxbx (0.60).
_TXBX_GLYPH_W = 0.50
_TXBX_LINE_H = 1.20
_TXBX_SCALE_FLOOR = 0.60


def _estimate_txbx_scale(txbx, box: dict | None) -> float:
    """Estimate the font scale (0.60–1.0) needed for a text box's content to fit
    its fixed design height. Returns 1.0 when it already fits or geometry is
    unknown. Pure estimate (no renderer) — see the metric constants above."""
    if not box:
        return 1.0
    w_pt, h_pt = box.get("w_pt"), box.get("h_pt")
    if not w_pt or not h_pt or w_pt <= 0 or h_pt <= 0:
        return 1.0
    lIns, rIns, tIns, bIns = _txbx_insets(txbx)   # actual box insets (padding must reduce usable height)
    usable_w = w_pt - (lIns + rIns) / _EMU_PER_PT
    usable_h = h_pt - (tIns + bIns) / _EMU_PER_PT
    if usable_w <= 0 or usable_h <= 0:
        return 1.0
    content_h = 0.0
    for p in txbx.iter(qn("w:p")):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        sz = p.find(".//" + qn("w:sz"))
        try:
            font_pt = int(sz.get(qn("w:val"))) / 2.0 if sz is not None else 9.0
        except (TypeError, ValueError):
            font_pt = 9.0
        if font_pt <= 0:
            font_pt = 9.0
        chars_per_line = max(1.0, usable_w / (_TXBX_GLYPH_W * font_pt))
        lines = max(1, math.ceil(len(text) / chars_per_line))   # empty para → 1 line of height
        content_h += lines * _TXBX_LINE_H * font_pt
    if content_h <= usable_h or content_h <= 0:
        return 1.0
    return max(_TXBX_SCALE_FLOOR, usable_h / content_h)


def _shape_of_txbx(txbx):
    el = txbx.getparent()
    for _ in range(4):
        if el is None:
            return None
        if el.tag.endswith("}wsp"):
            return el
        el = el.getparent()
    return None


# OOXML default text-box insets (EMU) when <a:bodyPr> omits them.
_DEF_TXBX_INS = {"lIns": 91440, "rIns": 91440, "tIns": 45720, "bIns": 45720}


def _txbx_insets(txbx):
    """(lIns, rIns, tIns, bIns) in EMU for a text box, reading its <bodyPr> and
    falling back to the OOXML defaults."""
    ins = dict(_DEF_TXBX_INS)
    shape = _shape_of_txbx(txbx)
    if shape is not None:
        for bp in shape.iter():
            if bp.tag.endswith("}bodyPr"):
                for k in ins:
                    v = bp.get(k)
                    if v is not None:
                        try:
                            ins[k] = int(v)
                        except (TypeError, ValueError):
                            pass
                break
    return ins["lIns"], ins["rIns"], ins["tIns"], ins["bIns"]


def _scale_txbx_runs(txbx, scale: float, exempt: dict | None = None) -> None:
    """Directly shrink every run's font size in a text box by `scale`.

    `exempt` is `_user_sized_paragraphs`' register, keyed by the id() of paragraphs whose sizes
    the ESTIMATOR chose (it pins the elements; see `_hand_formatted`). Those are skipped:
    an automatic shrink that silently overrides a deliberate size is worse than a box that
    overflows, because the person who set it has no way to see what happened. Measured before
    this existed: an edited GC NOTES line came out of fill_proposal at 4.5pt.

    Why not autofit: LibreOffice-headless (our docx→PDF engine) does NOT apply
    DrawingML text autofit — neither an empty <a:normAutofit/> nor one with an
    explicit fontScale shrinks the render (verified on staging: the last WORK line
    still clipped). It DOES always honor explicit run sizes (<w:sz>), so we scale
    those directly. Size-less runs inherit — we give them the box's most common
    size so they shrink too. Floored at 4pt so nothing vanishes."""
    sizes = [int(v) for sz in txbx.iter(qn("w:sz"))
             if (v := sz.get(qn("w:val"))) and v.isdigit()]
    default_hp = max(set(sizes), key=sizes.count) if sizes else 18   # half-points; 18 = 9pt
    for r in txbx.iter(qn("w:r")):
        if exempt:
            # Walk up to the run's paragraph; if the estimator sized that paragraph, leave it.
            par = r.getparent()
            while par is not None and not par.tag.endswith("}p"):
                par = par.getparent()
            if par is not None and id(par) in exempt:
                continue
        rpr = r.find(qn("w:rPr"))
        cur = None
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            v = sz.get(qn("w:val")) if sz is not None else None
            if v and v.isdigit():
                cur = int(v)
        new_hp = max(8, int(round((cur if cur is not None else default_hp) * scale)))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), str(new_hp))


def _shrink_overflowing_text_boxes(d: Document) -> int:
    """Keep long text-box content from spilling past its fixed box (over the next
    box / the baked page-frame art) — e.g. a combo's two options + exclusions,
    whose last line ("*Assumes installation over…") was overdrawn by the PRICE
    frame (the "cut-off last line" bug).

    Kyle's boxes are fixed-size (<a:noAutofit/>). The obvious fix — flip to
    <a:normAutofit/> "shrink text on overflow" — is a NO-OP under LibreOffice-
    headless (it doesn't compute/apply DrawingML autofit, with or without an
    explicit fontScale). So for boxes we estimate to overflow, we shrink the RUN
    sizes directly (which LibreOffice always honors), mirroring the editor's
    on-screen `fitTxbx` so preview == generated doc. We still flip noAutofit→
    normAutofit (harmless; lets Word re-fit if the doc is opened there). Boxes
    that already fit are untouched (byte-identical output)."""
    NO, NORM = f"{{{_A_NS}}}noAutofit", f"{{{_A_NS}}}normAutofit"
    try:
        boxes = template_geometry(d).get("boxes", [])
    except Exception:                       # geometry is best-effort; never block generation
        boxes = []
    n = 0
    for i, txbx in enumerate(_iter_txbx(d)):
        shape = _shape_of_txbx(txbx)
        af = None
        if shape is not None:
            af = shape.find(f".//{NO}")
            if af is None:
                af = shape.find(f".//{NORM}")
        if af is None:
            continue
        af.tag = NORM
        af.attrib.pop("fontScale", None)    # empty normAutofit; we shrink runs directly below
        af.attrib.pop("lnSpcReduction", None)
        scale = _estimate_txbx_scale(txbx, boxes[i] if i < len(boxes) else None)
        if scale < 0.999:
            _scale_txbx_runs(txbx, scale, _user_sized_paragraphs(d))
        n += 1
    # Straggler noAutofit not paired to a geometry box: preserve the old intent.
    for na in list(d.element.iter(NO)):
        na.tag = NORM
        n += 1
    return n


def _force_terms_on_new_page(d: Document) -> bool:
    """Make the "TERMS AND CONDITIONS" section start on a fresh page.

    Kyle's templates have NO forced break before the T&C heading — they rely on
    the body flowing onto a later page, which fails for the combo (its body is
    short): the heading + its terms-page letterhead land on the bottom of page 1,
    over the ACCEPTANCE frame. We set <w:pageBreakBefore/> on the terms
    letterhead's host paragraph — the empty paragraph that anchors the terms-page
    PNG (positionV relative to that paragraph), immediately before the heading —
    so the letterhead AND heading move to the new page together. pageBreakBefore
    never inserts a blank page, so templates whose T&C already starts a page are
    unaffected. Budget pricing has no T&C section → no-op."""
    tops = [c for c in d.element.body if c.tag == qn("w:p")]
    h = None
    for i, p in enumerate(tops):
        if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip().upper() == "TERMS AND CONDITIONS":
            h = i
            break
    if h is None:
        return False
    target = tops[h]
    for j in range(h, max(-1, h - 4), -1):          # heading + up to 3 paras before it
        if list(tops[j].iter(qn("wp:anchor"))):      # the terms-page letterhead's host paragraph
            target = tops[j]
            break
    ppr = target.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        target.insert(0, ppr)
    if ppr.find(qn("w:pageBreakBefore")) is None:
        ppr.insert(0, OxmlElement("w:pageBreakBefore"))
    return True


# Some templates position a framed box's top edge at/above its red frame border
# with zero top-inset, so the first line hugs / rides over the border. Which boxes
# are affected varies PER TEMPLATE (Kyle positioned each individually; verified by
# rendering): the NOTES box crosses on the Polish + Gyp templates (Combo + Epoxy
# are fine); Gyp additionally has its WORK + PRICE boxes touching ("Base Bid" on
# the PRICE border). Give the affected boxes a top inset so the first line clears
# the border. EMU (1pt = 12700). MUST run BEFORE _shrink_overflowing_text_boxes so
# the shrink estimate (which reads the actual inset) accounts for the reduced
# usable height and can't push the WORK box into overflow.
_FRAME_BOX_TOP_INSET_EMU = 114300   # ~9pt
# The gyp template's NOTES content box sits ~0.54" further LEFT than its WORK/PRICE
# boxes (column-relative posH 0.451" vs 0.991"/1.000"), all with a zero left inset,
# so its bullets render on top of the baked-in rotated red "NOTES" gutter label
# (the labels are fixed raster art in the page PNGs — only the content box can move).
# Left-inset the gyp NOTES box so its text clears the label and left-aligns with the
# WORK/PRICE text. EMU (1pt=12700); 39pt ~= the 0.54" posH delta. ONLY the gyp NOTES
# box needs it — polish's NOTES box and gyp's WORK/PRICE already start at the right x.
_GYP_NOTES_LEFT_INSET_EMU = 495300   # ~39pt (~0.54")


def _pad_frame_boxes(d: Document, notes, work_type) -> int:
    """Inset the framed boxes whose text rides over the baked frame art. TOP-inset
    (first line hugs the border): NOTES on polish + gyp, WORK + PRICE additionally on
    gyp. LEFT-inset (bullets overlap the "NOTES" gutter label): the gyp NOTES box
    only. Boxes identified by content markers so the DATE/JOB-NAME/estimator header
    boxes are never touched; the left inset is guarded to the NOTES box (a note
    marker present, no WORK/PRICE marker) so WORK/PRICE are never shifted."""
    wt = str(work_type or "").lower()
    pad_notes = wt in ("polish", "gyp")     # these templates' NOTES box crosses its border
    pad_work_price = wt == "gyp"            # gyp also has WORK + PRICE touching
    left_inset_notes = wt == "gyp"          # gyp NOTES box also overlaps the left gutter label
    if not (pad_notes or pad_work_price):
        return 0
    note_markers, work_price_markers = [], []
    if pad_notes:
        note_keys = [str((n or {}).get("text") or "").strip()[:20] for n in (notes or [])]
        note_markers = [k for k in note_keys if len(k) >= 8][:4]
    if pad_work_price:
        # WORK has "Exclusions:"/"Assumptions:"/"per plans"; PRICE has "Base Bid".
        work_price_markers = ["Base Bid", "Exclusions", "Assumptions", "per plans"]
    if not (note_markers or work_price_markers):
        return 0
    n = 0
    for txbx in _iter_txbx(d):
        txt = "".join(t.text or "" for t in txbx.iter(qn("w:t")))
        is_notes = any(m in txt for m in note_markers)
        is_work_price = any(m in txt for m in work_price_markers)
        if not (is_notes or is_work_price):
            continue
        shape = _shape_of_txbx(txbx)
        if shape is None:
            continue
        for bp in shape.iter():
            if bp.tag.endswith("}bodyPr"):
                try:
                    cur_t = int(bp.get("tIns") or 0)
                except (TypeError, ValueError):
                    cur_t = 0
                if cur_t < _FRAME_BOX_TOP_INSET_EMU:
                    bp.set("tIns", str(_FRAME_BOX_TOP_INSET_EMU))
                    n += 1
                if left_inset_notes and is_notes and not is_work_price:
                    try:
                        cur_l = int(bp.get("lIns") or 0)
                    except (TypeError, ValueError):
                        cur_l = 0
                    if cur_l < _GYP_NOTES_LEFT_INSET_EMU:
                        bp.set("lIns", str(_GYP_NOTES_LEFT_INSET_EMU))
                        n += 1
                break
    return n


def _flatten_price_bullets(d: Document) -> int:
    """Remove list/bullet formatting from the PRICE section so amounts read as
    clean flush-left lines (Kyle: no bullet points in the pricing). Every price
    template puts its PRICE rows — base bid, Material Sales Tax, Remodel, Total,
    {{#price_line}} options, {{#room}}, {{#alternate}} — on list numId=3 (verified
    across all Direct/GC/Gyp templates); NOTES (numId 1), the WORK section
    (numId 4) and Terms (numId 5) keep their bullets. Runs AFTER block expansion
    (so cloned option/room/tax rows are covered) over body + text-box paragraphs.
    Supersedes the older per-row _zero_list_indent hide-the-bullet trick."""
    n = 0
    for p in d.element.body.iter(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            continue
        numpr = ppr.find(qn("w:numPr"))
        if numpr is None:
            continue
        numid = numpr.find(qn("w:numId"))
        if numid is None or numid.get(qn("w:val")) != "3":
            continue
        ppr.remove(numpr)
        # No longer a list item — pin flush-left so no orphaned hanging indent remains.
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:start"), "0")
        n += 1
    return n


def _space_before_options(d: Document, n: int = 2) -> int:
    """Insert `n` blank paragraphs before the PRICE "Options" heading so the
    base-bid Total isn't cramped against the Options section (Kyle: double
    spacing after the Total). Runs AFTER block expansion + substitution over
    body + text-box paragraphs; targets the first standalone "Options" heading.
    No-op for a bid with no options (no heading to anchor to). Blank paragraphs
    inherit the document default height — enough to read as clean line breaks."""
    # Insert before EVERY standalone "Options" heading — a floating text box is
    # duplicated across mc:Choice (DrawingML) + mc:Fallback (VML), and different
    # renderers (Word vs LibreOffice→PDF) pick different copies, so both need the
    # spacing to stay consistent.
    targets = [p for p in d.element.body.iter(qn("w:p"))
               if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip() == "Options"]
    for target in targets:
        for _ in range(n):
            target.addprevious(OxmlElement("w:p"))
    return len(targets)


def _is_total_row(p_elem) -> bool:
    """True for the PRICE block's Total row — the `{{#tax_breakout}}` paragraph
    carrying the `{{total_label}}` / `{{total_formatted}}` token (as opposed to
    the sibling Material Sales Tax row that shares the same block name)."""
    txt = _p_text(p_elem)
    return "{{total_label}}" in txt or "{{total_formatted}}" in txt


def _zero_list_indent(p_elem) -> None:
    """Zero a list paragraph's left indent (`<w:ind w:left="0" w:start="0"/>`) so
    the numbering level's hanging bullet tucks into the margin and doesn't print —
    the same trick Kyle's other PRICE rows use to hide their bullet while keeping
    the text flush left. Keeps `<w:numPr>` intact so spacing/style are unchanged.
    Appended last (after `<w:rPr>`) to match the sibling rows' element order.
    No-op without a pPr; idempotent where the indent is already zeroed."""
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is None:
        return
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:start"), "0")


# Base-bid line: `{{base_bid_formatted}} – <description> as described above {{base_tax_phrase}}`.
# The <description> is static text in every template (each work type / audience
# has its own wording, e.g. "Epoxy flooring", "Polished Concrete & Joint Filler"),
# NOT a token — so a display override swaps just the text BETWEEN the two tokens,
# leaving the amount + tax-phrase tokens (and each template's default wording when
# there's no override) untouched. Matches hyphen / en dash / em dash separators.
_BASE_DESC_RE = re.compile(
    r"(\{\{\s*base_bid_formatted\s*\}\}\s*[-–—]\s*).*?(\s*\{\{\s*base_tax_phrase\s*\}\})",
    re.DOTALL,
)


def _apply_base_desc_override(d: Document, desc: str) -> int:
    """Replace the base-bid line's description with `desc`, preserving the
    `{{base_bid_formatted}}`/`{{base_tax_phrase}}` tokens + their separators.

    Operates per `<w:t>` run across the whole document (body + text boxes,
    including the VML-fallback duplicate). The base line is authored as a single
    run in every template, so a template whose base line were split across runs
    simply wouldn't match — a safe no-op that keeps the default wording. Runs
    BEFORE the flat `{{token}}` pass so the anchor tokens are still present.
    """
    if not desc:
        return 0
    n = 0
    for t in d.element.body.iter(qn("w:t")):
        txt = t.text or ""
        if "base_bid_formatted" in txt and "base_tax_phrase" in txt:
            new = _BASE_DESC_RE.sub(lambda m: m.group(1) + desc + m.group(2), txt)
            if new != txt:
                t.text = new
                t.set(qn("xml:space"), "preserve")
                n += 1
    return n


# The GC proposal templates (Resinous/Polish/Sealer) hardcode an additional-phase
# surcharge amount in their Clarifications text — e.g. "Add $5,000 for each
# additional required phase beyond above stated schedule." — each with its OWN
# native default ($5,000 Resinous, $2,300 Polish/Sealer). It is static body text
# in a text box, NOT a {{token}}, and the digits span many single-char runs. When
# the estimator changes the "Add for additional phase" estimate cell, main.py sets
# `_phase_price_override`; we then rewrite JUST the digits in place, keeping "Add $"
# + the clause wording + every run's formatting. Absent → each template keeps its
# own literal default (mirrors `_base_desc_override`: no per-template default to
# drift). Matches only the phase clause (the trailing lookahead), never the "$500"
# mobilization figures elsewhere in the same paragraph.
_GC_PHASE_RE = re.compile(r"(?<=Add \$)[\d,]+(?= for each additional required phase)")
_TXBX_CONTENT = qn("w:txbxContent")


def _apply_gc_phase_override(d: Document, amount: str) -> int:
    """Replace the GC additional-phase amount with `amount` (a bare number string
    like "5,200") across runs, everywhere the clause appears (incl. the VML
    fallback duplicate). No-op if `amount` is falsy or no clause matches.
    """
    if not amount:
        return 0
    n = 0
    for p in d.element.body.iter(qn("w:p")):
        # Skip anchor paragraphs that merely CONTAIN a text box — the box's own
        # <w:p> children (where the clause text actually lives) are visited
        # separately, so processing the anchor too would rewrite runs across the
        # nesting boundary. Leaf paragraphs (incl. the ones inside the box) match.
        if p.find(".//" + _TXBX_CONTENT) is not None:
            continue
        n += _sub_runs_preserving(
            p, _GC_PHASE_RE,
            lambda m: None if m.group(0) == amount else amount,
            require_braces=False,
        )
    return n


# PRICE tax-row + ALTERNATE labels ("Material Sales Tax", "Remodel Tax", "Total",
# "Flooring as described above (…)") are STATIC text trailing their amount token,
# not tokens themselves. When the estimator overrides a label in the doc editor,
# main.py sets a private `_*_label_override` value; we rewrite the text that
# trails the anchor token IN PLACE, preserving runs. Each anchor maps a private
# key to the {{token}} that immediately precedes the label. MUST run in Phase 0.5
# (before block expansion) so the {{#remodel}} / {{#alternate}} ITEM tokens are
# still present as anchors — expansion consumes them.
_PRICE_LABEL_ANCHORS = (
    ("_sales_tax_label_override",    "material_tax_formatted"),
    ("_remodel_label_override",      "remodel.amount_formatted"),
    ("_total_label_override",        "total_formatted"),
    ("_alt_flooring_label_override", "alternate.lump_sum_formatted"),
    ("_alt_remodel_label_override",  "alternate.remodel_tax"),
    ("_alt_total_label_override",    "alternate.total_formatted"),
)


def _apply_price_label_overrides(d: Document, values) -> int:
    """Rewrite each PRICE/ALTERNATE row's static LABEL (the text after its amount
    token) to the estimator's override, preserving runs. Anchored on the amount
    token + separator so the token stays for the flat pass to fill. No-op for any
    label not set; never touches a row whose anchor token isn't present."""
    n = 0
    for key, anchor in _PRICE_LABEL_ANCHORS:
        label = values.get(key)
        if not label:
            continue
        label = str(label)
        # "{{ anchor }} <sep> <rest-of-line>" → keep the token+separator, replace the
        # trailing label. repl returns None once the tail already equals the label
        # (stops the loop) — the token keeps `{{`, so require_braces stays satisfied.
        pat = re.compile(r"(\{\{\s*" + re.escape(anchor) + r"\s*\}\}\s*[–—-]\s*)(.*)$", re.DOTALL)
        def _repl(m, _lbl=label):
            return None if m.group(2) == _lbl else m.group(1) + _lbl
        for p in d.element.body.iter(qn("w:p")):
            # Skip text-box anchor paragraphs (their inner <w:p> are visited on
            # their own) — same nesting guard as _apply_gc_phase_override.
            if p.find(".//" + _TXBX_CONTENT) is not None:
                continue
            n += _sub_runs_preserving(p, pat, _repl)
    return n


# WHOLE-LINE overrides: the estimator rewrote an ENTIRE price/heading/alternate
# line in the doc editor. main.py passes the full text in a private `_line_*`
# value; we find the paragraph that carries the line's anchor and REPLACE the
# whole paragraph text with it (preserving spaces via _set_paragraph_text →
# xml:space="preserve"). Token anchors ({{...}}) match pre-expansion (item tokens
# like {{#remodel}}/{{#alternate}} are consumed by expansion); heading anchors
# match the exact paragraph text. First match per anchor wins; text-box anchor
# paragraphs are skipped. (`is_token`: True = {{token}} regex, False = literal text.)
_LINE_ANCHORS = (
    ("_line_base",            "base_bid_formatted",           True),
    ("_line_sales_tax",       "material_tax_formatted",       True),
    ("_line_remodel",         "remodel.amount_formatted",     True),
    ("_line_total",           "total_formatted",              True),
    ("_line_alt_name",        "alternate.system_name",        True),
    ("_line_alt_flooring",    "alternate.lump_sum_formatted", True),
    ("_line_alt_remodel",     "alternate.remodel_tax",        True),
    ("_line_alt_total",       "alternate.total_formatted",    True),
    ("_line_heading_base",    "Base Bid",                     False),
    ("_line_heading_options", "Options:",                     False),
)


def _apply_line_overrides(d: Document, values) -> int:
    """Replace an ENTIRE price/heading/alternate line with the estimator's whole-line
    override. Runs in Phase 0.5 (before block expansion). No-op for any line not set
    or whose anchor isn't present. Rewrites EVERY matching paragraph — both the
    mc:Choice and the VML mc:Fallback copies — so no un-replaced copy keeps the
    token and renders the computed line (same reach as _apply_price_label_overrides)."""
    n = 0
    for key, anchor, is_token in _LINE_ANCHORS:
        text = values.get(key)
        if text is None:
            continue
        text = str(text)
        pat = re.compile(r"\{\{\s*" + re.escape(anchor) + r"\s*\}\}") if is_token else None
        for p in d.element.body.iter(qn("w:p")):
            # Skip anchor paragraphs that merely CONTAIN a text box (their inner
            # <w:p> are visited on their own) — same nesting guard as the others.
            if p.find(".//" + _TXBX_CONTENT) is not None:
                continue
            ptext = "".join(t.text or "" for t in p.iter(qn("w:t")))
            hit = pat.search(ptext) if is_token else (ptext.strip() == anchor)
            if hit:
                _set_paragraph_text(p, text)
                n += 1
    return n


# Cove-only WORK rows: after the flat {{token}} fill, an epoxy system with 0 SF
# but a cove clause reads "Area: ~0 SF of epoxy flooring and <n> LF …". Drop the
# meaningless "~0 SF of epoxy flooring and " prefix so it reads "Area: <n> LF …"
# (mirrors the on-screen renderSystemPreview). A 0-SF row with NO cove has no
# " and " after "flooring", so the pattern can't match — it keeps today's line.
_AREA_ZERO_RE = re.compile(r"Area:\s*~0 SF of epoxy flooring and ")


def _drop_zero_sf_prefix(d: Document, protect=None) -> int:
    """`protect` is the set of Area rows an estimator REWROTE (see `_apply_system_row_line`).

    Those are skipped. This function exists to tidy a line the engine composed; a line a person
    typed is not the engine's to tidy, and silently re-editing it would be the exact 1:1
    violation the whole-line channel was built to close — he can type "Area: ~0 SF of epoxy
    flooring and 240 LF …" on purpose, and if he does, that is what prints.

    Membership is tested by element identity. `protect` holds the paragraph objects themselves,
    which keeps lxml's proxies alive, so the objects this walk yields for those nodes are the
    same objects that went in.
    """
    n = 0
    for p in d.element.body.iter(qn("w:p")):
        # Skip text-box anchor paragraphs (their <w:p> children are visited on
        # their own) — same nesting guard as _apply_gc_phase_override.
        if p.find(".//" + _TXBX_CONTENT) is not None:
            continue
        if protect and any(p is q for q in protect):
            continue
        # repl returns a fixed "Area: " (never equal to the matched span, and the
        # result no longer contains the pattern) so the require_braces=False loop
        # can't rewrite forever.
        n += _sub_runs_preserving(p, _AREA_ZERO_RE, lambda m: "Area: ", require_braces=False)
    return n


def _expand_named_block(container, block_name: str, items: list[Mapping[str, Any]],
                        protect: set | None = None) -> int:
    """Expand EVERY `{{#<block_name>}}…{{/<block_name>}}` block in `container`.

    `container` is any element whose direct <w:p> children may hold the markers
    (a <w:body>, <w:tc>, or <w:txbxContent>). One container may hold several
    blocks of different names (e.g. the PRICE cell has {{#price_line}} AND
    {{#alternate}}) — this expands only the blocks whose name matches
    `block_name`, re-scanning after each so element indices stay valid.
    `items==[]` still removes the markers + template body (renders zero rows).
    Returns how many blocks were expanded.

    `protect`, when given, collects every paragraph this expansion rewrote from a
    whole-line override, so a later doc-wide tidy-up pass can leave the estimator's
    own words alone (see `_drop_zero_sf_prefix`).
    """
    expanded = 0
    while True:
        children = list(container)
        start_idx = end_idx = None
        for i, child in enumerate(children):
            if child.tag != qn("w:p"):
                continue
            txt = _p_text(child)
            if start_idx is None:
                m = BLOCK_START_RE.search(txt)
                if m and m.group(1) == block_name:
                    start_idx = i
            else:
                m = BLOCK_END_RE.search(txt)
                if m and m.group(1) == block_name:
                    end_idx = i
                    break
        if start_idx is None or end_idx is None:
            break

        # Template paragraphs strictly between the two markers.
        template_elems = children[start_idx + 1:end_idx]
        start_elem = children[start_idx]
        end_elem = children[end_idx]

        # For each item, a fresh deep copy of every template paragraph with
        # per-item tokens substituted.
        new_elems = []
        for item in items:
            for tmpl in template_elems:
                clone = copy.deepcopy(tmpl)
                # BEFORE the token substitution: both the label rewrite and the whole-line
                # rewrite find their row by the token that row carries ({{system.texture}} /
                # {{system.sqft}} / {{system.name}}), and substitution is what removes those
                # tokens. See _apply_system_row_labels / _system_row_line_key.
                line_key = None
                if block_name == "system":
                    _apply_system_row_labels(clone, item)
                    line_key = _system_row_line_key(clone, item)
                _substitute_item_tokens(clone, item, block_name)
                # AFTER it: a whole-line override replaces the row's fully resolved text, which
                # is what the estimator was looking at when they typed. It therefore also wins
                # over the label rewrite above on the same row, which is right — the line the
                # estimator wrote includes whatever label they wanted.
                if line_key and _apply_system_row_line(clone, item[line_key]):
                    if protect is not None:
                        protect.add(clone)
                # Label-only price_line row (empty amount) — drop the now-bare
                # leading "– " separator so it reads as just the label. Scoped
                # to price_line/empty-amount only; every other row/block is
                # untouched.
                if block_name == "price_line" and not str(item.get("amount_formatted") or "").strip():
                    _strip_leading_separator(clone)
                # Blank NOTES line — estimator's Word-style spacing. Drop the
                # bullet so it renders as an empty line, not an empty bullet dot.
                if block_name == "notes" and not str(item.get("text") or "").strip():
                    _strip_bullet(clone)
                # PRICE Total row: the sibling rows (base bid / Material Sales Tax /
                # Remodel) zero their list indent so the numbering's hanging bullet
                # tucks into the margin and doesn't print; the Polish template's
                # Total row was missed and so shows a lone stray bullet. Match the
                # siblings so the whole PRICE block formats consistently (no-op where
                # the template already zeros it — e.g. the Epoxy template).
                if block_name == "tax_breakout" and _is_total_row(clone):
                    _zero_list_indent(clone)
                new_elems.append(clone)

        for clone in new_elems:
            start_elem.addprevious(clone)
        for stale in [start_elem, end_elem, *template_elems]:
            container.remove(stale)
        expanded += 1
    return expanded


def _expand_all_blocks(d: Document, block_lists: Mapping[str, list],
                       protect: set | None = None) -> int:
    """Expand every named block in `block_lists` across the whole document.

    Walks the body, every table cell, and every text box (<w:txbxContent>,
    including the VML-fallback duplicate) for each block name — so a block
    authored in any of those locations expands consistently. A block whose
    list is empty is still processed: its markers + template body are stripped
    (zero rows) rather than left as literal `{{#name}}` text in the output.
    """
    total = 0
    body = d.element.body
    containers = [body]
    containers += list(body.iter(qn("w:tc")))
    containers += list(body.iter(qn("w:txbxContent")))
    for block_name, items in block_lists.items():
        for container in containers:
            total += _expand_named_block(container, block_name, list(items or []), protect)
    return total


# ─── Paragraph-editor id mapping (Proposal Review's document editor) ──────
# The web editor shows the estimator the REAL template — every paragraph, in
# document order, as an editable block — instead of the old hand-built HTML
# approximation. `iter_editable_blocks` is the ONE walk shared by:
#   1. `GET /api/proposal-template` (main.py)      — builds the JSON the
#      editor renders.
#   2. `_apply_paragraph_overrides` (below)         — maps an edited block's
#      `id` back to its paragraph when generating.
# Both MUST see the exact same ids for the exact same document, or an edit
# could silently land on the wrong paragraph. That's only guaranteed if both
# walk the PRISTINE template (before Phase 1's block expansion inserts/
# removes paragraphs and shifts every id after it) — see the call site in
# `fill_proposal` for where overrides are applied for that reason.
_MC_FALLBACK_TAG = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"


def _is_fallback_paragraph(p_elem) -> bool:
    """True for a <w:p> living inside the legacy VML `mc:Fallback` branch of a
    floating text box/shape — a byte-for-byte duplicate of the modern
    DrawingML version that `_iter_all_paragraphs` also visits (so old-Word/
    VML readers get filled tokens too). The editor must show ONE copy of each
    paragraph, not two, so every id-based walk below skips these."""
    return any(True for _ in p_elem.iterancestors(_MC_FALLBACK_TAG))


def _iter_body_editable(d: Document):
    """Yield `(p_elem, kind, txbx_idx)` for every REAL (non-Fallback)
    paragraph in the document BODY — top-level paragraphs (`kind="p"`),
    table-cell paragraphs (`kind="cell"`, recursing into nested tables), and
    floating text-box paragraphs (`kind="p"`, `txbx_idx` = the 0-based index
    of the enclosing text box in this walk's box order; `None` outside a
    box). Headers/footers are intentionally excluded (the editor is scoped
    to the body; none of Kyle's templates put tokens there today — see
    `_iter_all_paragraphs`, which still covers them for the flat fill pass).

    Text boxes carry almost all of the customer-facing proposal copy (job
    name, WORK, PRICE, NOTES, SIGN) — Kyle's templates lay the whole front
    page out as floating shapes over blank body paragraphs — so skipping
    them would leave the editor showing nothing but the Terms & Conditions
    boilerplate at the bottom of the document. `txbx_idx` pairs each block
    with its box's page geometry (`template_geometry` enumerates the SAME
    non-Fallback boxes in the SAME order), so the editor can place the
    content exactly where the printed page puts it — even though this walk,
    which defines the ids and therefore can never be reordered, visits body
    paragraphs before text boxes.
    """
    for p in d.paragraphs:
        if not _is_fallback_paragraph(p._p):
            yield p._p, "p", None

    def walk_table(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not _is_fallback_paragraph(p._p):
                        yield p._p, "cell", None
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in d.tables:
        yield from walk_table(table)

    for bi, txbx in enumerate(_iter_txbx(d)):
        for p_elem in txbx.iter(qn("w:p")):
            if not _is_fallback_paragraph(p_elem):
                yield p_elem, "p", bi


def _iter_txbx(d: Document):
    """The document body's REAL (non-Fallback) text boxes, in the one
    canonical order shared by `_iter_body_editable` (block → box pairing)
    and `template_geometry` (box → page position)."""
    for txbx in d.element.body.iter(qn("w:txbxContent")):
        if not _is_fallback_paragraph(txbx):
            yield txbx


def iter_editable_blocks(d: Document):
    """THE shared id-mapping walk (see module note above). Yields
    `(id, kind, p_elem, in_block, text, in_txbx)` for every editable
    paragraph in `d`:

      - `id`       — 0-based index in THIS walk's order (stable as long as
                     the document hasn't been mutated by block expansion).
      - `kind`     — "p" or "cell" (see `_iter_body_editable`).
      - `p_elem`   — the raw `<w:p>` lxml element.
      - `in_block` — the name of the innermost `{{#name}}…{{/name}}` region
                     this paragraph currently sits in (the start/end marker
                     paragraphs themselves count as "in" that block), else
                     `None`. Blocks nest lexically as flat marker pairs among
                     SIBLING paragraphs within one container (see the block-
                     engine docstring above `_expand_named_block`) — e.g.
                     `{{#tax_breakout}}`/`{{#remodel}}` sit inside
                     `{{#single_bid}}` — so a simple stack reproduces it.
      - `text`     — this paragraph's OWN text (`_own_text`, NOT `_p_text` —
                     see that helper for why a naive recursive join would
                     duplicate a nested text box's content onto its anchor
                     paragraph). Block-marker detection uses this same value,
                     so it's computed once and handed to the caller instead
                     of making every caller re-derive it (and risk using the
                     wrong helper).
      - `txbx_idx` — index of the enclosing floating text box (pairs the
                     block with `template_geometry`'s box positions), `None`
                     for plain-body/table paragraphs. Display placement
                     only, never id math.
    """
    stack: list[str] = []
    idx = 0
    for p_elem, kind, txbx_idx in _iter_body_editable(d):
        txt = _own_text(p_elem)
        start_m = BLOCK_START_RE.search(txt)
        if start_m:
            stack.append(start_m.group(1))
        in_block = stack[-1] if stack else None
        end_m = BLOCK_END_RE.search(txt)
        if end_m and stack and stack[-1] == end_m.group(1):
            stack.pop()
        yield idx, kind, p_elem, in_block, txt, txbx_idx
        idx += 1


# ─── Formatting + page-geometry extraction (fidelity rendering) ──────────
# The editor renders the REAL page: run-level formatting (bold lead-ins,
# Zetta Serif sizes/colors), true bullet flags, and the floating text boxes
# placed on the page over the template's baked-in letterhead artwork (Kyle's
# templates draw the DATE:/JOB NAME: labels, the buffalo logo, the red
# PROPOSAL stamp and the bordered WORK/PRICE/NOTES/ACCEPTANCE frame as
# full-page background PNGs — word/media/image1.png etc.).

# Empty body paragraphs are the vertical ruler Word hangs the floating
# anchors off ('paragraph'-relative positionV). Their rendered line height
# isn't in the XML (it's a layout result), so we use a constant calibrated
# against the Direct Epoxy artwork: with 14pt/line the WORK box lands at
# y≈153pt (art: ≈152pt), PRICE at ≈321pt (art: ≈318pt), NOTES at ≈495pt
# (art: ≈490pt). The spec accepts approximate anchoring.
_ANCHOR_LINE_H_PT = 14.0
_EMU_PER_PT = 12700.0

# ── Resizing a floating text box ──────────────────────────────────────────────
#
# A box's size is stored in THREE places that all have to agree, in two different unit
# systems, and which one a renderer believes depends on the renderer:
#
#   1. wp:extent/@cx,@cy                     — EMU, the DrawingML anchor's own size
#   2. wps:spPr/a:xfrm/a:ext/@cx,@cy         — EMU, the shape's transform
#   3. the mc:Fallback VML v:shape/@style    — POINTS (or inches) in a CSS-ish string
#
# Writing only the DrawingML pair leaves the legacy VML twin claiming the old size, and
# which branch gets read is Word-version and LibreOffice-version dependent — so Word and
# the customer's PDF could disagree about the same box. All three, every time.
#
# Two traps that cost time to find, both confirmed against
# `GC/xx TREADWELL RESINOUS PROPOSAL - xx.docx`:
#
#   * `wsp.iter(a:ext)` matches TWO elements — the real one under `a:xfrm` and a second
#     under `a:extLst` that carries no cx/cy at all. Only the `a:xfrm` child is geometry.
#   * VML lengths are not all in points. Box 0 reads `width:324.8pt;height:99pt` while box 1
#     reads `width:1in;height:18pt`, so parsing has to be unit-aware. (Both agree with their
#     DrawingML twins: 4124960 EMU = 324.8pt, 914400 EMU = 72pt = 1in.)
#
# Sizes are set EXPLICITLY and never by asking the renderer to autofit: project experience
# is that LibreOffice ignores DrawingML autofit outright, which is why
# `_shrink_overflowing_text_boxes` rewrites run sizes directly instead.
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_V_NS = "urn:schemas-microsoft-com:vml"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

_VML_UNIT_PT = {"pt": 1.0, "in": 72.0, "pc": 12.0, "cm": 72.0 / 2.54,
                "mm": 7.2 / 2.54, "px": 0.75, "": 1.0}


def _vml_len_pt(raw):
    """A VML length ("99pt", "1in", "324.8") to points, or None if unparseable."""
    m = re.match(r"^\s*(-?[\d.]+)\s*([a-z%]*)\s*$", str(raw or ""), re.I)
    if not m:
        return None
    try:
        n = float(m.group(1))
    except ValueError:
        return None
    factor = _VML_UNIT_PT.get(m.group(2).lower())
    return None if factor is None else n * factor


def _fmt_vml_pt(v):
    """Points as VML writes them: no trailing ".0", because the templates don't."""
    return ("%g" % round(float(v), 2)) + "pt"


def _set_vml_style(shape, wanted: dict) -> bool:
    """Rewrite named declarations in a VML @style, leaving every other one alone.

    `wanted` maps a declaration name to points, e.g. `{"width": 300.0, "margin-top": -36.0}`;
    a None value means "leave that one as it is". Order is preserved rather than rebuilt: the
    style also carries position, z-index and visibility, and reordering those is a needless
    diff in a customer-facing template.

    Matching is on the WHOLE declaration name, which is what keeps `margin-left` away from the
    bare `left`. `Invoice_Deposit.docx` ships `position:absolute;left:0;text-align:left;
    margin-left:-2.1pt;…` — so a substring match, or writing `left` when you meant
    `margin-left`, would edit a decoy and leave the real offset untouched. That is the same
    class of mistake as the `a:extLst/a:ext` decoy on the DrawingML side.
    """
    style = shape.get("style") or ""
    parts = [p for p in style.split(";") if p.strip()]
    live = {k: v for k, v in wanted.items() if v is not None}
    out, seen = [], set()
    for part in parts:
        key = part.partition(":")[0].strip().lower()
        if key in live:
            out.append(key + ":" + _fmt_vml_pt(live[key]))
            seen.add(key)
        else:
            out.append(part)
    for key, val in live.items():
        if key not in seen:
            out.append(key + ":" + _fmt_vml_pt(val))
    new = ";".join(out)
    if new == style:
        return False
    shape.set("style", new)
    return True


def _set_vml_size(shape, w_pt, h_pt) -> bool:
    """The size half of `_set_vml_style`."""
    return _set_vml_style(shape, {"width": w_pt, "height": h_pt})


def _set_vml_pos(shape, x_off_pt, y_off_pt) -> bool:
    """The position half of `_set_vml_style`: `margin-left` / `margin-top`.

    These carry the SAME numbers as the DrawingML `wp:posOffset` pair, because the twin also
    carries `mso-position-horizontal-relative:text` / `-vertical-relative:text` — VML's name
    for the `column`/`paragraph` frames the DrawingML anchor uses. Verified against every
    template: box 0 of the GC Resinous file is `positionH posOffset 447040` (= 35.2pt) and
    `margin-left:35.2pt`, `positionV posOffset -457200` (= -36pt) and `margin-top:-36pt`.

    So what goes here is the OFFSET, never the absolute page coordinate. Writing an absolute
    coordinate into a text-relative frame is how Word and the PDF end up disagreeing about a
    box the estimator can see on screen.
    """
    return _set_vml_style(shape, {"margin-left": x_off_pt, "margin-top": y_off_pt})


def _txbx_anchor(txbx):
    """The `wp:anchor`/`wp:inline` that positions this text box, or None."""
    wanted = (qn("wp:anchor"), qn("wp:inline"))
    for anc in txbx.iterancestors():
        if anc.tag in wanted:
            return anc
    return None


def _txbx_vml_twins(txbx):
    """The VML shapes in the `mc:Fallback` branch paired with this box.

    `_iter_txbx` deliberately skips Fallback content, so the twin has to be reached through
    the shared `mc:AlternateContent` ancestor.
    """
    for ac in txbx.iterancestors("{%s}AlternateContent" % _MC_NS):
        fb = ac.find("{%s}Fallback" % _MC_NS)
        if fb is None:
            return []
        out = []
        for tag in ("shape", "rect", "roundrect", "oval"):
            out.extend(fb.iter("{%s}%s" % (_V_NS, tag)))
        return out
    return []


def _resize_txbx(txbx, w_pt=None, h_pt=None) -> int:
    """Set one text box's size everywhere it is recorded. Returns the number of sites written.

    Either dimension may be None to leave it alone. Never raises: a template whose shape is
    built differently simply gets fewer sites written, which is a box at its design size —
    not a 500 on /api/generate.
    """
    wrote = 0
    emu = lambda v: str(int(round(float(v) * _EMU_PER_PT)))   # noqa: E731

    anchor = _txbx_anchor(txbx)
    if anchor is not None:
        ext = anchor.find(qn("wp:extent"))
        if ext is not None:
            if w_pt is not None:
                ext.set("cx", emu(w_pt))
            if h_pt is not None:
                ext.set("cy", emu(h_pt))
            wrote += 1

    for wsp in txbx.iterancestors("{%s}wsp" % _WPS_NS):
        # `a:xfrm/a:ext` ONLY — a bare iter() would also match the cx/cy-less a:extLst/a:ext.
        xfrm = next(iter(wsp.iter("{%s}xfrm" % _A_NS)), None)
        if xfrm is not None:
            e = xfrm.find("{%s}ext" % _A_NS)
            if e is not None:
                if w_pt is not None:
                    e.set("cx", emu(w_pt))
                if h_pt is not None:
                    e.set("cy", emu(h_pt))
                wrote += 1
        break

    for shape in _txbx_vml_twins(txbx):
        if _set_vml_size(shape, w_pt, h_pt):
            wrote += 1

    return wrote


# ── Moving a floating text box ────────────────────────────────────────────────
#
# Hanz, 2026-08-13: "Allow me to drag and resize the text box for the proposal please."
#
# Resizing already existed; nothing had ever WRITTEN a position. A box's position lives in two
# places, the same split as its size:
#
#   1. wp:positionH/wp:posOffset + wp:positionV/wp:posOffset  — EMU, the DrawingML anchor
#   2. mc:Fallback v:shape/@style margin-left / margin-top     — points, the legacy VML twin
#
# Both hold an OFFSET FROM A FRAME, not a page coordinate. Every box in every template anchors
# `positionH relativeFrom="column"` and `positionV relativeFrom="paragraph"` (measured across all
# nine .docx files), and the VML twins say `mso-position-*-relative:text`, which agrees.
#
# THE PARAGRAPH-RELATIVE PROBLEM, and why this writer never needs to solve it. The y the editor
# shows is a CALIBRATED ESTIMATE — `_pos_of_anchor` resolves a paragraph-relative anchor as
# `margin.top + paragraph_index * _ANCHOR_LINE_H_PT + posOffset`, because the rendered height of
# the empty anchor paragraphs is a layout result that is not in the XML. Round-tripping an
# absolute y through that formula would bake the calibration error into the document.
#
# So the move is applied as a DIFFERENCE against the very same estimate the estimator dragged
# from: read where `_pos_of_anchor` thinks the box is now, and shift the offset by
# `wanted - current`. Whatever the estimate gets wrong appears on both sides and cancels, exactly,
# leaving `posOffset + (wanted - current)`. Drag a box down 40pt and it moves 40pt, whether or not
# the anchor paragraphs are really 14pt tall.
#
# Which also answers the tempting alternative: rewriting `relativeFrom` to "page" or "margin" and
# writing an absolute offset would make the arithmetic trivial and change the DOCUMENT's
# behaviour — a page-relative box stops travelling with the text above it, so a template edit that
# adds a line would slide Kyle's frame art out from under the box it belongs to. The frame is
# baked into the letterhead PNG, so that is not a cosmetic difference. The anchor frame is left
# exactly as Kyle authored it.
def _set_anchor_offset(anchor, tag: str, off_pt) -> bool:
    """Write one `wp:positionH`/`wp:positionV` offset, in points. True when written."""
    p = anchor.find(qn("wp:" + tag))
    if p is None:
        return False
    o = p.find(qn("wp:posOffset"))
    if o is None:
        # `wp:align` and `wp:posOffset` are the two arms of ONE xsd:choice, so a box positioned by
        # keyword ("align: left") has no offset to shift. Replace the keyword instead of adding a
        # sibling: both arms present at once is invalid OOXML, and Word's answer to that is a
        # repair prompt on a customer-facing file. No template of Kyle's uses align — every box
        # carries a posOffset — but declining to move a box the estimator dragged would look like
        # the drag was broken, and the offset we write is measured from the same place
        # `_pos_of_anchor` reads an absent one as (zero), so the box lands where the screen said.
        al = p.find(qn("wp:align"))
        if al is not None:
            p.remove(al)
        o = OxmlElement("wp:posOffset")
        p.append(o)
    o.text = str(int(round(float(off_pt) * _EMU_PER_PT)))
    return True


def _move_txbx(txbx, x_off_pt=None, y_off_pt=None) -> int:
    """Set one text box's anchor OFFSETS everywhere they are recorded.

    Offsets, not page coordinates — see the module note above. Either may be None to leave that
    axis alone. Never raises: a template whose shape is built differently gets fewer sites
    written, which is a box at its design position rather than a 500 on /api/generate.
    """
    wrote = 0
    anchor = _txbx_anchor(txbx)
    if anchor is not None:
        if x_off_pt is not None and _set_anchor_offset(anchor, "positionH", x_off_pt):
            wrote += 1
        if y_off_pt is not None and _set_anchor_offset(anchor, "positionV", y_off_pt):
            wrote += 1
    # Both branches or neither. A drawing whose anchor carries no position at all — a `wp:inline`
    # box, which no template of Kyle's uses — has nothing for the fallback to AGREE with, and a
    # fallback written on its own is precisely the half-write this pair exists to avoid.
    if wrote:
        for shape in _txbx_vml_twins(txbx):
            if _set_vml_pos(shape, x_off_pt, y_off_pt):
                wrote += 1
    return wrote


# Kyle's page, for when no document is to hand; `_apply_box_overrides` measures the real one
# instead. US Letter (612 x 792pt) with 1.25in side margins and 1in top/bottom, so the printable
# area is 612-90-90 = 432 by 792-72-72 = 648. The three constants have to keep describing ONE
# sheet — `_page_metrics` asserts as much in test_the_fallback_page_and_the_fallback_size_limit
# _describe_one_sheet, because a mismatch here would let the size bound and the position bound
# disagree about the same template.
_DEFAULT_MAX_BOX_PT = (432.0, 648.0)
_DEFAULT_PAGE_PT = (612.0, 792.0)
_MIN_BOX_PT = 12.0


def box_size_limits(d: Document) -> tuple:
    """The biggest a text box may be made, `(w_pt, h_pt)` — the printable area of the page.

    A box taller than the printable height cannot fit from ANY starting position, so accepting one
    guarantees text outside the margins. The old ceiling was 1600pt, about two pages. Measured in
    the container against the real LibreOffice render, filling one box with 60 numbered lines:

        design 184pt   7/60 lines   lowest text 134pt
        limit  648pt  46/60 lines   lowest text 684pt   inside the 720pt bottom margin
               700pt  50/60 lines   lowest text 740pt   into the margin
              1600pt  56/60 lines   lowest text 789pt   3pt from the sheet edge, 4 lines GONE

    So the old ceiling did not push text off the sheet — LibreOffice clips it instead, which is
    worse: the last lines vanish silently and what survives prints into the unprintable margin.
    Nothing errors, and the customer receives a proposal missing text.

    This bounds the impossible, not every overflow: `wp:positionV` in these templates is
    `relativeFrom="paragraph"`, so where a box actually lands depends on the text flowing above
    it, which needs a layout engine rather than a geometry read. A box anchored low can still be
    given a height that fits the page and overflows anyway. Catching that is the render check's
    job, not this function's.
    """
    try:
        sec = d.sections[0]
        w = float(sec.page_width.pt) - float(sec.left_margin.pt) - float(sec.right_margin.pt)
        h = float(sec.page_height.pt) - float(sec.top_margin.pt) - float(sec.bottom_margin.pt)
    except Exception:  # noqa: BLE001 — a template with no usable sectPr falls back to Letter
        return _DEFAULT_MAX_BOX_PT
    # A section with absurd margins would otherwise produce a limit under the minimum, which
    # would refuse every resize on that template.
    if not (math.isfinite(w) and math.isfinite(h)) or w <= _MIN_BOX_PT or h <= _MIN_BOX_PT:
        return _DEFAULT_MAX_BOX_PT
    return (w, h)


def page_size(d: Document) -> tuple:
    """The sheet, `(w_pt, h_pt)` — the bound on where a box may be MOVED to.

    Deliberately the paper and not the printable area, which is the bound on how big a box may
    be MADE. The two differ because Kyle designs into the margins: measured across all nine
    templates, EVERY box sits outside the printable area — the DATE/JOB NAME header box lands at
    y=36pt against a 72pt top margin, the buffalo logo at x=27pt against a 90pt left margin — and
    not one sits off the paper. So a printable-area bound would refuse to move any box in any
    template, and a refused drag reads to the estimator as a broken drag.

    What the paper bound does buy is the thing that actually goes wrong: a box cannot be pushed
    off the sheet, where its text would be silently gone from the customer's PDF.
    """
    try:
        w = float(d.sections[0].page_width.pt)
        h = float(d.sections[0].page_height.pt)
    except Exception:  # noqa: BLE001 — a template with no usable sectPr falls back to Letter
        return _DEFAULT_PAGE_PT
    if not (math.isfinite(w) and math.isfinite(h)) or w <= 0 or h <= 0:
        return _DEFAULT_PAGE_PT
    return (w, h)


def _sanitize_box_overrides(raw, limits=None, page=None) -> dict:
    """`{"<box id>": {h_pt?, w_pt?, x_pt?, y_pt?}}`, coerced and bounded. Never raises.

    A dict keyed by id rather than a list, for the reason the paragraph-override sanitizer
    already documents: a list's positions shift, and a stale draft would then resize a
    different box than the estimator dragged.

    `h_pt`/`w_pt` are a size; `x_pt`/`y_pt` are the box's top-left corner in PAGE POINTS, the
    same coordinate system `template_geometry` reports and therefore the same one the drag
    handle works in. All four are independent: a box may be moved without being resized.

    Size bounds are the printable area (see `box_size_limits`); position bounds are the sheet
    (see `page_size`). Out-of-range values are REFUSED, not clamped: the drag handle stops at
    the same limits, so anything past them arrived from a stale draft or a hand-built request,
    and leaving the design geometry gives a document that still reads correctly. Nothing useful
    is under 12pt either way.

    Only the CORNER is bounded here. Whether the far edge stays on the paper depends on the
    box's size, which is a property of the document rather than of the request, so that half of
    the rule lives in `_apply_box_overrides` where the document is to hand.
    """
    out = {}
    if not isinstance(raw, dict):
        return out
    max_w, max_h = limits or _DEFAULT_MAX_BOX_PT
    page_w, page_h = page or _DEFAULT_PAGE_PT
    for key, spec in list(raw.items())[:200]:
        try:
            box_id = int(key)
        except (TypeError, ValueError):
            continue
        if box_id < 0 or not isinstance(spec, dict):
            continue
        one = {}
        for field, lo, hi in (("h_pt", _MIN_BOX_PT, max_h), ("w_pt", _MIN_BOX_PT, max_w),
                              ("x_pt", 0.0, page_w), ("y_pt", 0.0, page_h)):
            v = spec.get(field)
            if isinstance(v, bool) or not isinstance(v, (int, float)):
                continue
            v = float(v)
            if not math.isfinite(v) or not (lo <= v <= hi):
                continue
            one[field] = v
        if one:
            out[str(box_id)] = one
    return out


def _apply_box_overrides(d: Document, raw) -> int:
    """Move and resize the boxes the estimator dragged. Returns the number of boxes changed.

    MUST run before `_pad_frame_boxes`/`_shrink_overflowing_text_boxes`, because the shrink
    re-reads `template_geometry(d)` to decide what overflows. Applying the resize first is
    what lets the shrink stand down by itself on a box that has been made big enough — which
    is the entire point of the feature: a box the estimator enlarged should show its text at
    full size, not get its runs scaled to 4.5pt anyway.
    """
    # Bounded by THIS template's own page, not by constants: the fallbacks are Kyle's Letter
    # sheet, and a template with different margins should be held to its own.
    page = _page_metrics(d)
    limits = (page["max_box"]["w_pt"], page["max_box"]["h_pt"])
    boxes = _sanitize_box_overrides(raw, limits, (page["w_pt"], page["h_pt"]))
    if not boxes:
        return 0
    body = d.element.body
    top_ps = [c for c in body if c.tag == qn("w:p")]
    changed = 0
    for idx, txbx in enumerate(_iter_txbx(d)):
        spec = boxes.get(str(idx))
        if not spec:
            continue
        wrote = 0
        want_x, want_y = spec.get("x_pt"), spec.get("y_pt")
        anchor = _txbx_anchor(txbx) if (want_x is not None or want_y is not None) else None
        # Read the box's CURRENT place before touching it — the move is a difference against
        # this reading, which is what makes the paragraph-relative estimate cancel out (see the
        # "Moving a floating text box" note above).
        cur = _pos_of_anchor(anchor, page, top_ps, body) if anchor is not None else None
        # Asked for at all, before asked to do anything: `_resize_txbx` counts the sites it
        # VISITED, so calling it with two Nones reports a write that never happened — and this
        # count is what tells the caller (and the log line) whether the estimator's drag landed.
        if ("w_pt" in spec or "h_pt" in spec) and _resize_txbx(
                txbx, w_pt=spec.get("w_pt"), h_pt=spec.get("h_pt")):
            wrote += 1
        if cur is not None:
            cur_x, cur_y, cur_w, cur_h = cur
            new_w = spec.get("w_pt", cur_w) or cur_w
            new_h = spec.get("h_pt", cur_h) or cur_h
            tx = cur_x if want_x is None else want_x
            ty = cur_y if want_y is None else want_y
            # The far-edge half of the paper bound. A corner on the sheet with the box hanging
            # off the right or the bottom loses text: LibreOffice CLIPS rather than spilling, so
            # nothing errors and the customer's proposal is simply missing a paragraph.
            #
            # Half a point of slack, because the handle clamps to exactly `page - size` and then
            # rounds to 2dp: without it, the furthest position the estimator can actually drag to
            # would sometimes be refused by a rounding step of 0.005pt. Half a point is 1/144 of
            # an inch, well inside the unprintable edge of any real printer.
            if (tx + new_w <= page["w_pt"] + 0.5) and (ty + new_h <= page["h_pt"] + 0.5):
                ox, _rfx = _anchor_offset(anchor, "positionH")
                oy, _rfy = _anchor_offset(anchor, "positionV")
                wrote += _move_txbx(
                    txbx,
                    x_off_pt=None if want_x is None else ox + (tx - cur_x),
                    y_off_pt=None if want_y is None else oy + (ty - cur_y),
                )
        if wrote:
            changed += 1
    return changed


def _fmt_of_run(run: Run, para: Paragraph) -> dict:
    """Resolved character formatting for one run: the run's own font first,
    then up the paragraph-style chain (style.font → base_style.font …, max 4
    hops — cheap, no full Word style resolution). `None` = unresolved; the
    frontend falls back to the page default (Zetta Serif 9pt #404040)."""
    fonts = [run.font]
    st = para.style
    hops = 0
    while st is not None and hops < 4:
        try:
            fonts.append(st.font)
        except Exception:  # noqa: BLE001
            break
        st = getattr(st, "base_style", None)
        hops += 1

    def resolve(attr):
        for f in fonts:
            try:
                v = getattr(f, attr)
            except Exception:  # noqa: BLE001
                v = None
            if v is not None:
                return v
        return None

    color = None
    for f in fonts:
        try:
            c = f.color
            if c is not None and c.type is not None and c.rgb is not None:
                color = str(c.rgb)
                break
        except Exception:  # noqa: BLE001
            pass

    bold, italic, under = resolve("bold"), resolve("italic"), resolve("underline")
    size = resolve("size")
    return {
        "bold": bool(bold) if bold is not None else None,
        "italic": bool(italic) if italic is not None else None,
        "underline": bool(under) if under is not None else None,
        "size_pt": size.pt if size is not None else None,
        "font": resolve("name"),
        "color": color,
    }


def _block_runs(p_elem, para: Paragraph) -> list:
    """The paragraph's own text as formatted segments
    `[{text, bold, italic, underline, size_pt, font, color}]`, with two
    invariants the editor depends on:

      1. `"".join(seg.text) == _own_text(p_elem)` — the frontend verifies
         this and falls back to flat rendering if it ever doesn't hold
         (e.g. hyperlink runs, which aren't direct <w:r> children).
      2. No flat `{{token}}` straddles a segment boundary: each token is its
         own segment carrying the formatting of the run where the match
         STARTS — the same rule `_sub_runs_preserving` applies when actually
         filling the docx, so the preview shows a value with the exact
         formatting the generated document will give it.
    """
    txbx_tag = qn("w:txbxContent")

    def own_run_text(r_elem):
        out = []
        for t in r_elem.iter(qn("w:t")):
            anc, nested = t.getparent(), False
            while anc is not None and anc is not r_elem:
                if anc.tag == txbx_tag:
                    nested = True
                    break
                anc = anc.getparent()
            if not nested:
                out.append(t.text or "")
        return "".join(out)

    raw = []
    for r_elem in p_elem.findall(qn("w:r")):
        txt = own_run_text(r_elem)
        if txt:
            raw.append({"text": txt, **_fmt_of_run(Run(r_elem, para), para)})
    if not raw:
        return []

    joined = "".join(s["text"] for s in raw)
    spans = []
    pos = 0
    for s in raw:
        spans.append((pos, pos + len(s["text"]), s))
        pos += len(s["text"])

    def fmt_at(i):
        for a, b, s in spans:
            if a <= i < b:
                return s
        return spans[-1][2]

    fmt_keys = ("bold", "italic", "underline", "size_pt", "font", "color")

    def seg(a, b, src):
        return {"text": joined[a:b], **{k: src.get(k) for k in fmt_keys}}

    out = []
    cursor = 0
    for m in TOKEN_RE.finditer(joined):
        # non-token stretch before the match: split at run boundaries so a
        # bold lead-in ("Scope:") keeps its weight next to a normal value run
        for a, b, s in spans:
            lo, hi = max(a, cursor), min(b, m.start())
            if lo < hi:
                out.append(seg(lo, hi, s))
        out.append(seg(m.start(), m.end(), fmt_at(m.start())))
        cursor = m.end()
    for a, b, s in spans:
        lo, hi = max(a, cursor), min(b, len(joined))
        if lo < hi:
            out.append(seg(lo, hi, s))
    return out


_ALIGN_NAMES = {0: "left", 1: "center", 2: "right", 3: "justify"}


def _para_align(para: Paragraph):
    """Paragraph alignment as a CSS-friendly name, or None (inherit/left)."""
    try:
        a = para.alignment
        return _ALIGN_NAMES.get(int(a)) if a is not None else None
    except (TypeError, ValueError):
        return None


def _para_is_list(p_elem) -> bool:
    """True when the paragraph carries real Word numbering (<w:numPr>) — the
    template's bullet rows. Style name alone ("List Paragraph") is NOT enough:
    Kyle uses it for indent-only headings like "Base Bid" too."""
    ppr = p_elem.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def _para_price_list(p_elem) -> bool:
    """True when the paragraph is on the PRICE list (numId=3) that
    _flatten_price_bullets strips at generate time. The on-screen document
    editor uses this to render those rows flush/bullet-less so the preview
    matches the generated .docx (Kyle: no bullet points in the pricing)."""
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is None:
        return False
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return False
    numid = numpr.find(qn("w:numId"))
    return numid is not None and numid.get(qn("w:val")) == "3"


def _anchor_offset(anchor, tag: str) -> tuple:
    """`(offset_pt, relativeFrom)` for one `wp:positionH`/`wp:positionV`.

    Its own function so the READER (`_pos_of_anchor`) and the WRITER
    (`_apply_box_overrides`, which shifts the offset by a difference) cannot drift apart about
    what an absent or keyword-positioned anchor counts as.
    """
    p = anchor.find(qn("wp:" + tag))
    if p is None:
        return 0.0, "page"
    o = p.find(qn("wp:posOffset"))
    return (int(o.text) / _EMU_PER_PT if o is not None and o.text else 0.0,
            p.get("relativeFrom") or "page")


def _pos_of_anchor(anchor, page: dict, top_ps: list, body) -> tuple:
    """(x_pt, y_pt, w_pt, h_pt) of a floating drawing on its page.

    Word stores positionH/positionV relative to page/margin/column/paragraph;
    the paragraph-relative vertical (what these templates use) is resolved
    against the anchor's enclosing top-level paragraph index at the
    calibrated `_ANCHOR_LINE_H_PT` per empty line (see the constant's note).

    The y this returns is therefore an ESTIMATE, and `_apply_box_overrides` depends on it being
    the SAME estimate the editor was given rather than on it being right — read the
    "Moving a floating text box" note before changing the arithmetic here.
    """
    ext = anchor.find(qn("wp:extent"))
    w = int(ext.get("cx")) / _EMU_PER_PT if ext is not None else 0.0
    h = int(ext.get("cy")) / _EMU_PER_PT if ext is not None else 0.0

    ox, rfx = _anchor_offset(anchor, "positionH")
    oy, rfy = _anchor_offset(anchor, "positionV")

    x = ox + (page["margin"]["left"] if rfx in ("column", "margin") else 0.0)

    anc = anchor
    while anc is not None and anc.getparent() is not body:
        anc = anc.getparent()
    try:
        pidx = top_ps.index(anc)
    except ValueError:
        pidx = 0
    if rfy in ("paragraph", "line"):
        y = page["margin"]["top"] + pidx * _ANCHOR_LINE_H_PT + oy
    elif rfy == "margin":
        y = page["margin"]["top"] + oy
    else:                                     # "page" and anything unmapped
        y = oy
    return x, y, w, h


def _page_metrics(d: Document) -> dict:
    """`{w_pt, h_pt, margin{top,left,right,bottom}, max_box{w_pt,h_pt,min_pt}}` from sectPr.

    THE one reader of the section's page setup, shared by `template_geometry` (which hands it to
    the editor) and `_apply_box_overrides` (which bounds a drag with it). Two independent
    derivations of the same numbers is how a handle ends up letting somebody drag to a place the
    server then throws away, which reads as the drag not having worked.
    """
    max_w, max_h = box_size_limits(d)
    page_w, page_h = page_size(d)
    try:
        sec = d.sections[0]
        margin = {"top": sec.top_margin.pt, "left": sec.left_margin.pt,
                  "right": sec.right_margin.pt, "bottom": sec.bottom_margin.pt}
    except Exception:  # noqa: BLE001 — Kyle's own page setup, so the three fallbacks agree
        # 612-90-90 = 432 and 792-72-72 = 648, i.e. exactly `_DEFAULT_MAX_BOX_PT`. Picking a
        # different guess here would make the size limit and the margins describe two different
        # sheets on the one template that ever reaches this branch.
        margin = {"top": 72.0, "left": 90.0, "right": 90.0, "bottom": 72.0}
    return {
        "w_pt": page_w, "h_pt": page_h, "margin": margin,
        # Stated rather than left for the editor to derive, so the drag handle stops exactly where
        # the sanitiser starts refusing. `max_box` is how BIG a box may be made (the printable
        # area); the page above is how far it may be MOVED (the sheet) — see `page_size` for why
        # those are two different rectangles.
        "max_box": {"w_pt": max_w, "h_pt": max_h, "min_pt": _MIN_BOX_PT},
    }


def template_geometry(d: Document) -> dict:
    """Page metrics + floating-object placement for the editor's page view:

      page   — {w_pt, h_pt, margin:{top,left,right,bottom}} from sectPr.
      boxes  — one {id, x_pt, y_pt, w_pt, h_pt} per REAL text box, in the
               SAME order `_iter_body_editable` numbers them (`txbx_idx`).
      images — the anchored artwork {name, x_pt, y_pt, w_pt, h_pt,
               para_index}; `name` is served by /api/proposal-template/media.
               For Kyle's templates these are the full-page letterhead PNGs
               (page 1's labeled/bordered art, then the plain terms-page
               letterhead — `para_index` orders them by where they anchor).
    """
    page = _page_metrics(d)
    body = d.element.body
    top_ps = [c for c in body if c.tag == qn("w:p")]

    def enclosing_anchor(el):
        anc = el.getparent()
        want = (qn("wp:anchor"), qn("wp:inline"))
        while anc is not None and anc.tag not in want:
            anc = anc.getparent()
        return anc

    boxes = []
    for bi, txbx in enumerate(_iter_txbx(d)):
        anchor = enclosing_anchor(txbx)
        if anchor is not None:
            x, y, w, h = _pos_of_anchor(anchor, page, top_ps, body)
        else:
            x = y = w = h = None
        boxes.append({"id": bi, "x_pt": x, "y_pt": y, "w_pt": w, "h_pt": h})

    images = []
    for anchor in body.iter(qn("wp:anchor"), qn("wp:inline")):
        if _is_fallback_paragraph(anchor):
            continue
        blip = anchor.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        rid = blip.get(qn("r:embed"))
        try:
            target = d.part.rels[rid].target_ref
        except (KeyError, AttributeError):
            continue
        x, y, w, h = _pos_of_anchor(anchor, page, top_ps, body)
        anc = anchor
        while anc is not None and anc.getparent() is not body:
            anc = anc.getparent()
        try:
            pidx = top_ps.index(anc)
        except ValueError:
            pidx = 0
        images.append({"name": target.rsplit("/", 1)[-1],
                       "x_pt": x, "y_pt": y, "w_pt": w, "h_pt": h,
                       "para_index": pidx})
    return {"page": page, "boxes": boxes, "images": images}


def _hand_formatted(d, attr: str) -> dict:
    """A per-document `{id(element): element}` register of what the ESTIMATOR chose by hand.

    Kept on the Document object rather than in the XML: a custom attribute on `w:p`/`w:r`
    would be invalid OOXML and Word may reject the file.

    THE VALUE IS THE ELEMENT, AND IT HAS TO BE. lxml builds an element proxy on demand and
    frees it the moment the last Python reference goes, then hands out a BRAND NEW proxy — at
    a different address — on the next access. Every walk in this module is a generator that
    keeps no references, so a bare `set()` of `id()`s goes stale as soon as the walk that
    filled it ends: measured on the GC Resinous template, 1 paragraph marked, 0 still matching
    one fresh `iter_editable_blocks` later. Worse than useless, because a freed address can be
    REUSED by an unrelated element and exempt the wrong thing. Holding the element keeps its
    proxy alive, which keeps its `id()` both stable and unique for as long as the register
    lives. Membership tests (`id(x) in register`) read exactly as they did against a set.
    """
    got = getattr(d, attr, None)
    if got is None:
        got = {}
        try:
            setattr(d, attr, got)
        except Exception:  # noqa: BLE001 — a read-only Document still works, just unexempted
            return {}
    return got


def _user_sized_paragraphs(d) -> dict:
    """Paragraphs whose run sizes the ESTIMATOR set, per document.

    `_shrink_overflowing_text_boxes` consults this so a deliberate size is not rewritten by the
    automatic shrink — which was measured rewriting an edited NOTES line down to 4.5pt, i.e.
    silently undoing the estimator on exactly the overflowing boxes they were fixing."""
    return _hand_formatted(d, "_tw_user_sized")


def _user_bolded_runs(d) -> dict:
    """Runs whose WEIGHT the estimator stated explicitly, per document.

    `_normalize_work_label_formatting` consults this. That pass makes a WORK row bold through
    its first colon and normal after it, which is right for the template's own text and was
    wrong for an override: bold applied to a phrase inside a Scope / Schedule / Exclusions /
    Notes value showed on screen, survived the reload, travelled in the payload, was rebuilt
    faithfully by `_set_paragraph_runs` — and was flattened one pass later. Measured on Direct
    block 115 with bold on "3-coat system": the run split survived, `w:b` came out `val="0"`.

    Per RUN, not per paragraph, because the normalization is per run. Exempting the whole
    paragraph would also drop the automatic label bold, and the label is exactly what should
    keep it: the browser renders the template's own bold as `font-weight:700` and
    `serializeRuns` reads it straight back, so a runs override arrives with `bold: True`
    already on the label — the estimator only ever states the opposite deliberately."""
    return _hand_formatted(d, "_tw_user_bolded")


def _set_paragraph_runs(p_elem, runs, bold_marks: dict | None = None) -> bool:
    """Replace a paragraph's text with `runs`, KEEPING each run's own formatting.

    `_set_paragraph_text` (below) collapses a paragraph to run[0]'s rPr. That is fine for a
    plain text edit and destructive for a formatted one: Kyle's templates genuinely mix
    formats inside one paragraph — a GC label row is bold+underlined 9pt followed by 8pt body
    with an italic aside — and an override flattened all of it to the first run's look.

    Each run is {text, bold?, italic?, underline?, size_pt?}. An ABSENT key means "inherit",
    which is not the same as False: absent leaves the template's own rPr alone, False writes an
    explicit off. That distinction is what lets somebody bold one phrase without pinning the
    size of everything around it.

    The first template run's rPr is the base for every new run, so whatever the estimator did
    NOT touch — font, colour, and the size when they set none — still comes from Kyle's design.
    Media runs are never removed, for the same reason as `_set_paragraph_text`: they anchor the
    letterhead and every floating text box.

    Returns True when any run carries an explicit size, so the caller can exempt this
    paragraph from the overflow shrink (which would otherwise rewrite it — measured at 4.5pt
    on a real GC NOTES line). `bold_marks`, when given, collects the runs whose `bold` the
    estimator STATED (True or False alike — both are a choice, absent is not), so
    `_normalize_work_label_formatting` can leave those alone; see `_user_bolded_runs`.
    """
    _MEDIA_TAGS = (qn("w:drawing"), qn("w:pict"), qn("w:object"))
    all_runs = p_elem.findall(qn("w:r"))
    text_runs = [r for r in all_runs
                 if not any(next(r.iter(tag), None) is not None for tag in _MEDIA_TAGS)]

    base_rpr = None
    for r in (text_runs or all_runs):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            base_rpr = copy.deepcopy(rpr)
            break

    # Where the text used to start, so the new runs land in the same place relative to any
    # media runs (an anchored text box in the same paragraph must stay put).
    children = list(p_elem)
    insert_at = children.index(text_runs[0]) if text_runs else len(children)
    for r in text_runs:
        p_elem.remove(r)

    user_sized = False
    for offset, spec in enumerate(runs):
        r = OxmlElement("w:r")
        rpr = copy.deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")

        def toggle(tag, on):
            """Word booleans: <w:b/> on, <w:b w:val="0"/> explicitly off, absent inherit —
            all three are meaningful."""
            if on is None:
                return
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), "1" if on else "0")

        toggle("w:b", spec.get("bold"))
        toggle("w:bCs", spec.get("bold"))
        toggle("w:i", spec.get("italic"))
        toggle("w:iCs", spec.get("italic"))

        u = spec.get("underline")
        if u is not None:
            el = rpr.find(qn("w:u"))
            if el is None:
                el = OxmlElement("w:u")
                rpr.append(el)
            el.set(qn("w:val"), "single" if u else "none")

        size_pt = spec.get("size_pt")
        if size_pt:
            # Half-points on BOTH w:sz and w:szCs — the idiom _scale_txbx_runs already uses,
            # and the one LibreOffice reliably honours (it ignores DrawingML autofit).
            hp = max(2, int(round(float(size_pt) * 2)))
            for tag in ("w:sz", "w:szCs"):
                el = rpr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rpr.append(el)
                el.set(qn("w:val"), str(hp))
            user_sized = True

        if len(rpr):
            r.append(rpr)
        t = OxmlElement("w:t")
        r.append(t)
        _write_t_text(t, str(spec.get("text", "")))
        p_elem.insert(insert_at + offset, r)
        # Registered only once the run is IN the tree — an element that is about to be
        # discarded must never end up pinned in the register.
        if bold_marks is not None and spec.get("bold") is not None:
            bold_marks[id(r)] = r

    return user_sized


def _set_paragraph_text(p_elem, text: str) -> None:
    """Replace a paragraph's visible text with `text` IN PLACE, preserving the
    paragraph's formatting by keeping its FIRST text run (and that run's
    `<w:rPr>` — font/bold/size/color) and writing the new text into it; every
    other text run is dropped. Embedded newlines render as `<w:br/>` (via
    `_write_t_text`, the same helper the block engine uses for multi-line
    item notes).

    Runs that carry a drawing/picture/object are NEVER removed — Kyle's
    templates anchor the page letterhead artwork AND every floating text box
    in runs of otherwise-blank body paragraphs, so dropping those runs on an
    override would silently delete the letterhead (or an entire text box)
    from the customer document.

    A paragraph with no text runs (a blank spacer line, or one holding only
    a drawing) gets a fresh run appended so non-empty override text still
    renders. An override that blanks a paragraph (`text == ""`) is honored —
    the paragraph keeps its (now textless) run so its formatting/paragraph-
    mark survives.
    """
    # Descendant (not direct-child) search: Word wraps floating drawings in
    # mc:AlternateContent inside the run, so w:drawing is a grandchild.
    _MEDIA_TAGS = (qn("w:drawing"), qn("w:pict"), qn("w:object"))
    runs = p_elem.findall(qn("w:r"))
    text_runs = [r for r in runs
                 if not any(next(r.iter(tag), None) is not None for tag in _MEDIA_TAGS)]
    if not text_runs:
        r = OxmlElement("w:r")
        # Match the paragraph's look: clone rPr off an existing (media) run.
        if runs:
            rpr = runs[0].find(qn("w:rPr"))
            if rpr is not None:
                r.append(copy.deepcopy(rpr))
        p_elem.append(r)
        text_runs = [r]
    first = text_runs[0]
    for extra in text_runs[1:]:
        p_elem.remove(extra)
    # A run can hold several <w:t>/<w:br>/<w:tab> children (e.g. a value we
    # previously wrote with embedded line breaks) — collapse to a single
    # fresh <w:t> so re-overriding a multi-line paragraph doesn't leave stale
    # break/text nodes behind.
    for child in list(first):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab")):
            first.remove(child)
    t = OxmlElement("w:t")
    first.append(t)
    _write_t_text(t, text)


def _override_is_blank(val) -> bool:
    """True when an override leaves the paragraph with no words at all.

    One definition for both shapes the editor sends — a plain string, or the `runs` list a
    formatted paragraph travels as (emptying a formatted paragraph sends `runs: []`, which
    reaches this as the `text: ""` that came with it). Whitespace-only counts as blank: a lone
    newline is what a browser leaves behind when the last character is deleted, and it prints
    exactly as nothing."""
    if isinstance(val, list):
        return not "".join(str(r.get("text") or "") for r in val).strip()
    return not str(val or "").strip()


def _apply_paragraph_overrides(d: Document, overrides: list) -> int:
    """Apply the web editor's `paragraph_overrides` to the PRISTINE template —
    i.e. this MUST run before Phase 1 (block expansion) in `fill_proposal`,
    because block expansion inserts/removes paragraphs and would shift every
    id after the touched block, desyncing them from what the editor showed.

    Each override's `text` is whatever the estimator left in that block on
    the page — already-resolved values, not `{{tokens}}` — EXCEPT any
    `{{token}}` they deliberately left in place, which still gets filled by
    the normal flat substitution pass that runs after this (Phase 2), since
    that pass re-scans every paragraph regardless of whether it was just
    overridden.

    Defensive by design — never raises on bad input, so a malformed payload
    can't 500 `/api/generate`:
      - non-dict entries, non-int ids, or non-str text are skipped;
      - an id that doesn't exist in this document is skipped (no-op);
      - an id whose paragraph is inside a repeatable block (`in_block` is not
        None) is skipped — that content is pricing-engine/template owned and
        is never user-overridable, regardless of what the client sends.

    Returns the number of overrides actually applied.
    """
    by_id: dict[int, object] = {}
    # PARAGRAPH properties (bullet on/off, left indent) ride the SAME entry as the text, under
    # an optional `para` key — additive, so an entry's identity is unchanged and every override
    # already saved against a draft in flight still means what it meant. They are collected
    # separately here because an entry may carry `para` and NO text at all: switching a bullet
    # off is not an edit to the words, and requiring a text field to come with it would make a
    # formatting-only change either impossible or a silent rewrite of the paragraph.
    para_by_id: dict[int, dict] = {}
    for o in overrides or []:
        if not isinstance(o, dict):
            continue
        pid = o.get("id")
        if isinstance(pid, bool) or not isinstance(pid, int):
            continue
        para = sanitize_para_props(o.get("para"))
        if para:
            para_by_id[pid] = para
        # `runs` is the richer shape: [{text, bold?, italic?, underline?, size_pt?}]. It only
        # appears when the estimator applied formatting; a plain edit still sends `text`, so the
        # common case takes the simpler path and the payload stays the size it always was.
        runs = o.get("runs")
        if isinstance(runs, list) and runs and all(isinstance(r, dict) for r in runs):
            clean = []
            for r in runs:
                t = r.get("text")
                if not isinstance(t, str):
                    continue
                one = {"text": t}
                for k in ("bold", "italic", "underline"):
                    v = r.get(k)
                    if isinstance(v, bool):
                        one[k] = v
                sz = r.get("size_pt")
                if isinstance(sz, (int, float)) and not isinstance(sz, bool) and 1 <= float(sz) <= 200:
                    one["size_pt"] = float(sz)
                clean.append(one)
            if clean:
                by_id[pid] = clean
                continue
        text = o.get("text")
        if not isinstance(text, str):
            continue
        by_id[pid] = text   # last one wins on a duplicate id

    if not by_id and not para_by_id:
        return 0

    applied = 0
    refused = 0
    for idx, _kind, p_elem, in_block, _text, _txbx in iter_editable_blocks(d):
        if in_block is not None or (idx not in by_id and idx not in para_by_id):
            continue
        text_refused = False
        if idx in by_id:
            val = by_id[idx]
            # A NUMBERED CLAUSE CANNOT BE EMPTIED. The renumbering guard below keeps the clause
            # count fixed by KEEPING the numbering on a paragraph whose text was just deleted —
            # so Word printed "1." followed by nothing, in a signed contract, and the estimator
            # had no way to see it coming (the paragraph controls are hidden on that row and no
            # warning was raised). Three ways out, and they are not equal:
            #
            #   * DROP THE PARAGRAPH — renumbers every clause below it. This is the exact
            #     regression the guard was written for: measured at 26 numbered clauses of 27.
            #   * KEEP THE TEXT SILENTLY — the screen says empty, the .docx says otherwise, which
            #     is the same class of lie as the red squares, pointed the other way.
            #   * REFUSE THE EDIT AND SAY SO — what "locked" already implies everywhere else in
            #     this feature (`apply_para_props` refuses an ordered paragraph outright), and the
            #     only option where screen, document and intent can all agree.
            #
            # So: refused here, and refused visibly in the editor (proposal-review.js
            # restoreEmptiedClause puts the clause back the moment it is emptied and says why).
            # This half is what protects a request the editor did not build — a stale draft, a
            # replayed payload — and it is deliberately the authoritative one.
            blank = _override_is_blank(val)
            if blank and _para_ordered_list(d, p_elem):
                log.warning("Refused a blank override on numbered clause block %s: emptying a "
                            "Terms and Conditions clause would print a bare clause number", idx)
                refused += 1
                text_refused = True
            else:
                if isinstance(val, list):
                    if _set_paragraph_runs(p_elem, val, _user_bolded_runs(d)):
                        # Remember the box so the overflow shrink leaves this paragraph's sizes alone.
                        _user_sized_paragraphs(d)[id(p_elem)] = p_elem
                else:
                    _set_paragraph_text(p_elem, val)
                # AN EMPTIED ROW KEEPS NO BULLET. `_strip_bullet` already did this for a blank
                # {{#notes}} item ("a lone empty bullet dot"), but a WORK row emptied by hand went
                # through this channel instead and kept its numbering: the .docx printed a red
                # square with nothing after it. Same rule, same reason, now on both channels — and
                # it is what lets the on-screen preview honestly suppress its own square
                # (styles.css .tw-block.tw-empty.tw-li::before).
                #
                # Reachable only for a BULLET row now: an ordered clause never gets here, because
                # emptying one is refused above. A bullet carries no meaning that outlives its
                # text; a clause number carries the identity of the clause, and whatever
                # references "Section 7" does not move with it.
                #
                # Skipped when the same entry states a `bullet` explicitly. `apply_para_props` runs
                # after this and would put a `bullet: True` back, but only by joining whichever list
                # a SIBLING is on — so not stripping in the first place is how the row keeps its own
                # numbering identity rather than being re-homed onto the neighbouring list.
                if blank and "bullet" not in para_by_id.get(idx, {}):
                    _strip_bullet(p_elem)
        if idx in para_by_id:
            # Phase 0 runs long before `_shrink_overflowing_text_boxes`, which is required:
            # the shrink re-reads the box geometry to decide what overflows, and an indent
            # applied after it would change how much text fits behind its back.
            apply_para_props(d, p_elem, para_by_id[idx])
        # A refused entry applied NOTHING, so it is not counted — the same answer this function
        # already gives for an id it skipped. An entry that also carried `para` still counts,
        # because that half went through the writer (which refuses it in its own right).
        if text_refused and idx not in para_by_id:
            continue
        applied += 1
    if refused:
        log.warning("Kept %d numbered clause(s) the payload asked to empty", refused)
    return applied


def fill_proposal(
    *,
    work_type: str,
    audience: str | None,
    values: Mapping[str, Any],
    systems: list[Mapping[str, Any]] | None = None,
    price_lines: list[Mapping[str, Any]] | None = None,
    alternates: list[Mapping[str, Any]] | None = None,
    remodel: list[Mapping[str, Any]] | None = None,
    rooms: list[Mapping[str, Any]] | None = None,
    single_bid: list[Mapping[str, Any]] | None = None,
    notes: list[Mapping[str, Any]] | None = None,
    tax_breakout: bool = False,
    has_options: bool = False,
    paragraph_overrides: list[Mapping[str, Any]] | None = None,
    box_overrides: Mapping[str, Any] | None = None,
) -> bytes:
    """Open the matching template, substitute tokens, return docx bytes.

    `values` is a flat dict keyed by token name (e.g. `job_name`,
    `lump_sum`, `scope_notes`). Tokens not present in `values` are left
    as-is in the doc, so Troy can see which fields were missing.

    Repeatable blocks (Phase 1), each cloned once per list item before the
    flat pass:
      - `systems`     → `{{#system}}…{{/system}}`     (only when supplied)
      - `price_lines` → `{{#price_line}}…{{/price_line}}` (option/unit-price lines)
      - `alternates`  → `{{#alternate}}…{{/alternate}}`   (0/1 recommended system)
    `price_line`/`alternate` always run so their markers are stripped (zero
    rows) when empty — never left as literal text. A template with no marker,
    and the default args, is 100% backward-compatible with v1 fills.

    A `systems` row may also carry `name_line` / `texture_line` / `area_line` — the doc
    editor's rewrite of that WHOLE row, static words and all, which is the only channel that
    can reach text like " SF of epoxy flooring" (see `_apply_system_row_line`). The older
    per-row `texture_label` / `area_label` keys rename just the label and are still honoured
    for drafts saved before the whole-line editor. Absent keys leave the template's own
    wording; a row's `*_line` wins over its `*_label` and over its token values.

    `paragraph_overrides` — free-text edits from the Proposal Review document
    editor (Phase 0, runs BEFORE block expansion — see `_apply_paragraph_overrides`
    for why ids must be resolved against the pristine template).
    """
    template_path = pick_template(work_type, audience)
    log.info("Filling proposal: work_type=%s audience=%s template=%s systems=%d price_lines=%d alt=%d",
             work_type, audience, template_path.name,
             len(systems) if systems else 0,
             len(price_lines) if price_lines else 0,
             len(alternates) if alternates else 0)

    if not template_path.exists():
        raise FileNotFoundError(f"Proposal template not found: {template_path}")

    d = docx.Document(str(template_path))

    # Phase 0 — apply the document editor's paragraph overrides FIRST, against
    # the pristine (just-opened, unexpanded) template — the same document
    # `iter_editable_blocks` walked to hand the editor its ids. Doing this
    # before Phase 1 is load-bearing: block expansion inserts/removes
    # paragraphs, which would shift ids computed afterward out from under the
    # editor's.
    if paragraph_overrides:
        n_over = _apply_paragraph_overrides(d, paragraph_overrides)
        if n_over:
            log.info("Applied %d paragraph override(s)", n_over)

    # Phase 1 — expand repeatable blocks. All three always run so their markers
    # are stripped (zero rows when empty) rather than left as literal {{#…}} text
    # in the output. A template with no marker for a block is unaffected (no-op),
    # so this stays byte-identical for templates that don't use a given block.
    block_lists: dict[str, list] = {
        "price_line": list(price_lines or []),
        "alternate": list(alternates or []),
        "system": list(systems or []),
        # {{#remodel}} line — present (1 row) only when a remodel tax applies,
        # stripped otherwise so the proposal hides "Kansas Remodel Tax" entirely.
        "remodel": list(remodel or []),
        # {{#tax_breakout}} — the itemized Material Sales Tax + Total lines. Shown
        # (1 row) only when the estimator chooses "sales tax broken out"; stripped
        # by DEFAULT so the price collapses to a single all-in line
        # ("$Total – … (material sales tax INCLUDED)") per Kyle's preferred layout.
        "tax_breakout": [{}] if tax_breakout else [],
        # {{#has_options}} — the "Options:" label. Shown only when there are
        # actual options (price lines or a recommended alternate); stripped
        # otherwise so an empty "Options:" never prints.
        "has_options": [{}] if has_options else [],
        # {{#room}} — per-room priced options (per-room jobs); stripped when empty.
        "room": list(rooms or []),
        # {{#single_bid}} — the single Base-Bid/Total layout. Shown by DEFAULT
        # (single_bid is None → one row) so existing callers are unaffected;
        # callers pass single_bid=[] to SUPPRESS it when room options replace it.
        "single_bid": [{}] if single_bid is None else list(single_bid),
        # {{#notes}} — editable boilerplate notes (one bullet per item).
        "notes": list(notes or []),
    }
    # Phase 0.5 — doc-editor display overrides, BEFORE block expansion so the
    # {{#remodel}} / {{#alternate}} item-token anchors still exist (expansion
    # consumes them). WHOLE-LINE overrides first (they replace the entire line and
    # drop the token), then per-field LABEL overrides (a no-op on any line already
    # whole-line-replaced, since its anchor token is gone). No-op unless a private
    # `_line_*` / `_*_label_override` key is set.
    _n_line = _apply_line_overrides(d, values)
    if _n_line:
        log.info("Applied %d whole-line PRICE override(s)", _n_line)
    _n_lbl = _apply_price_label_overrides(d, values)
    if _n_lbl:
        log.info("Applied %d PRICE/ALTERNATE label override(s)", _n_lbl)
    # Rows the estimator rewrote whole. Collected here and handed to _drop_zero_sf_prefix
    # below so a regex tidy-up can never re-edit a hand-typed line.
    _rewritten_rows: set = set()
    n_blocks = _expand_all_blocks(d, block_lists, _rewritten_rows)
    if n_blocks:
        log.info("Expanded %d repeatable block(s)", n_blocks)

    # Base-bid line DISPLAY override (single_bid.desc): swap the static
    # description noun between {{base_bid_formatted}} and {{base_tax_phrase}}
    # BEFORE the flat pass fills those tokens. No-op unless the caller set
    # `_base_desc_override` (private key — the flat pass never emits it).
    _bdo = values.get("_base_desc_override")
    if _bdo and _apply_base_desc_override(d, str(_bdo)):
        log.info("Applied base-bid description override")

    # GC additional-phase amount (static Clarifications text, not a token) — only
    # rewritten when the estimator changed the phase cell (main.py sets
    # `_phase_price_override`). Absent → each GC template keeps its native default.
    _ppo = values.get("_phase_price_override")
    if _ppo and _apply_gc_phase_override(d, str(_ppo)):
        log.info("Applied GC additional-phase override (%s)", _ppo)

    # Phase 2 — flat {{token}} substitution against `values`. This runs
    # unchanged from v1 and also fills any non-system tokens left inside
    # the expanded block paragraphs.
    total_subs = 0
    for p in _iter_all_paragraphs(d):
        total_subs += _replace_in_paragraph(p, values)

    log.info("Substituted %d tokens", total_subs)
    # Cove-only WORK rows: drop the "~0 SF of epoxy flooring and " prefix now that
    # the sqft/lf_clause tokens are filled (matches the on-screen preview).
    if _drop_zero_sf_prefix(d, _rewritten_rows):
        log.info("Dropped ~0 SF prefix on cove-only WORK row(s)")
    _n_work_format = _normalize_work_label_formatting(d)
    if _n_work_format:
        log.info("Normalized %d WORK label/value run(s)", _n_work_format)
    # PRICE section reads as clean flush-left lines — Kyle wants NO bullets in the
    # pricing (confirmed by Hanz 2026-07-16, reversing the earlier "keep the red
    # squares" read). _flatten_price_bullets strips the numId=3 list formatting off
    # every PRICE row (base bid, Material Sales Tax, Remodel, Total, {{#price_line}}
    # options, {{#room}}, {{#alternate}}) across all Direct/GC/Gyp templates; the
    # WORK (numId 4), NOTES (numId 1) and Terms (numId 5) lists keep their bullets.
    _n_flat = _flatten_price_bullets(d)
    if _n_flat:
        log.info("Flattened %d PRICE bullet row(s)", _n_flat)
    # Double spacing after the base-bid Total, before the Options section (Kyle).
    if _space_before_options(d, 2):
        log.info("Added double spacing before the PRICE Options heading")
    # Boxes the estimator dragged or resized, FIRST — before the padding and therefore before
    # the shrink, which re-reads template_geometry(d) to decide what overflows. Applying a
    # resize first is what lets the shrink stand down by itself on a box that is now big
    # enough: a box somebody enlarged precisely because its text was being cut off must show
    # that text at full size, not get its runs scaled down anyway.
    _laid_out = _apply_box_overrides(d, box_overrides)
    if _laid_out:
        log.info("Moved/resized %d text box(es) from the estimator's box overrides", _laid_out)
    # Pad affected framed boxes' top inset (so the first NOTES bullet / "Base Bid"
    # clear their red borders) BEFORE the shrink, so the shrink estimate sees the
    # reduced usable height and can't push the WORK box into overflow.
    _padded = _pad_frame_boxes(d, notes, work_type)
    if _padded:
        log.info("Padded %d framed box(es) top inset (clears the frame border)", _padded)
    # Shrink-to-fit: long content (esp. gyp's verbose WORK scope) would otherwise
    # overflow its fixed box and overlap the next box / frame art.
    _shrunk = _shrink_overflowing_text_boxes(d)
    if _shrunk:
        log.info("Set %d text box(es) to shrink-on-overflow (normAutofit)", _shrunk)
    # Force the Terms & Conditions onto their own page (templates ship without a
    # forced break, so a short body — e.g. combo — spills T&C over the acceptance).
    if _force_terms_on_new_page(d):
        log.info("Forced a page break before the Terms & Conditions section")
    if total_subs == 0 and not systems:
        log.warning(
            "Template has no {{tokens}}: %s. Returning unmodified.",
            template_path.name,
        )

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf.read()
